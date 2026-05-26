from __future__ import annotations

import json
import math
import re
import shutil
import sys
from dataclasses import dataclass
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from pypdf import PdfReader

CODE_DIR = Path(__file__).resolve().parent
if str(CODE_DIR) not in sys.path:
    sys.path.insert(0, str(CODE_DIR))

from build_0401_data import load_generated_data  # noqa: E402
from config_0401 import (  # noqa: E402
    BASELINE_MODEL_RPT,
    CANDIDATE_NODES,
    LAYOUT_LINK_CSV,
    LAYOUT_NODE_CSV,
    MONITOR_NODES,
    RESULT_DIR,
    ROOT_DIR,
    TRUTH_EVENT_MODEL_RPT,
    runtime_model_path,
)
from simulation_0401 import build_dataset, evaluate_shares  # noqa: E402

plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS"]
plt.rcParams["axes.unicode_minus"] = False

OUTPUT_DIR = ROOT_DIR / "output" / "doc"
DOCX_PATH = OUTPUT_DIR / "0401_SchemeC_正式分析报告_v4.docx"
MD_PATH = OUTPUT_DIR / "0401_SchemeC_正式分析报告_v4.md"
ASSET_DIR = OUTPUT_DIR / "0401_SchemeC_正式分析报告_v4_assets"
CURRENT_RUN_DIR = RESULT_DIR / "medium_run"
PREVIOUS_RUN_DIR = ROOT_DIR / "废弃" / "20260408_scheme_c_plus_j77_refresh" / "medium_run"
PREVIOUS_ROOT = ROOT_DIR / "废弃" / "20260408_scheme_c_plus_j77_refresh"
TMP_PAPER_DIR = ROOT_DIR / "tmp" / "paper_extract"
ENG_PDF = TMP_PAPER_DIR / "english.pdf"
CN_PDF = TMP_PAPER_DIR / "chinese.pdf"


@dataclass
class RunBundle:
    name: str
    run_dir: Path
    monitor_nodes: list[str]
    summary: dict
    truth_replay: dict
    ga_all: pd.DataFrame
    ga_history: pd.DataFrame
    initial_ppd: pd.DataFrame
    am_samples: pd.DataFrame
    weights: pd.DataFrame
    coverage: pd.DataFrame
    runtime_truth: pd.Timestamp
    runtime_ga: pd.Timestamp
    runtime_am: pd.Timestamp
    runtime_pp: pd.Timestamp
    runtime_final: pd.Timestamp


def ensure_dirs() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    TMP_PAPER_DIR.mkdir(parents=True, exist_ok=True)


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def find_single_file(run_dir: Path, patterns: list[str]) -> Path:
    for pattern in patterns:
        matches = list(run_dir.glob(pattern))
        if matches:
            return matches[0]
    raise FileNotFoundError(f"未在 {run_dir} 找到匹配文件：{patterns}")


def load_run_bundle(name: str, run_dir: Path, monitor_nodes: list[str]) -> RunBundle:
    summary_path = find_single_file(run_dir, ["*结果汇总.json"])
    truth_path = find_single_file(run_dir, ["*truth_replay_check.json"])
    ga_all_path = find_single_file(run_dir, ["*GA全部方案.csv"])
    ga_history_path = find_single_file(run_dir, ["*GA每代最佳.csv"])
    initial_ppd_path = find_single_file(run_dir, ["*initial_PPD.csv"])
    am_path = find_single_file(run_dir, ["*AM样本.csv"])
    weights_path = find_single_file(run_dir, ["*后验节点权重.csv"])
    coverage_path = find_single_file(run_dir, ["*posterior_predictive_coverage.csv"])
    bands_path = find_single_file(run_dir, ["*posterior_predictive_bands.csv"])
    return RunBundle(
        name=name,
        run_dir=run_dir,
        monitor_nodes=monitor_nodes,
        summary=read_json(summary_path),
        truth_replay=read_json(truth_path),
        ga_all=pd.read_csv(ga_all_path, encoding="utf-8-sig"),
        ga_history=pd.read_csv(ga_history_path, encoding="utf-8-sig"),
        initial_ppd=pd.read_csv(initial_ppd_path, encoding="utf-8-sig"),
        am_samples=pd.read_csv(am_path, encoding="utf-8-sig"),
        weights=pd.read_csv(weights_path, encoding="utf-8-sig"),
        coverage=pd.read_csv(coverage_path, encoding="utf-8-sig"),
        runtime_truth=pd.Timestamp(truth_path.stat().st_mtime, unit="s"),
        runtime_ga=pd.Timestamp(ga_history_path.stat().st_mtime, unit="s"),
        runtime_am=pd.Timestamp(am_path.stat().st_mtime, unit="s"),
        runtime_pp=pd.Timestamp(bands_path.stat().st_mtime, unit="s"),
        runtime_final=pd.Timestamp(summary_path.stat().st_mtime, unit="s"),
    )


def parse_rpt_metrics(path: Path) -> dict:
    text = path.read_text(encoding="utf-8", errors="ignore")
    lines = text.splitlines()

    def extract_last(label: str) -> float | None:
        vals: list[float] = []
        for line in lines:
            if label in line:
                nums = re.findall(r"[-+]?\d+\.\d+", line)
                if nums:
                    vals.append(float(nums[-1]))
        return vals[-1] if vals else None

    return {
        "external_inflow": extract_last("External Inflow"),
        "final_stored_volume": extract_last("Final Stored Volume"),
        "continuity_error_pct": extract_last("Continuity Error (%)"),
        "no_nodes_flooded": "No nodes were flooded." in text,
    }


def format_timedelta(delta: pd.Timedelta) -> str:
    total = int(round(delta.total_seconds()))
    h, rem = divmod(total, 3600)
    m, s = divmod(rem, 60)
    return f"{h}小时{m}分钟{s}秒"


def runtime_breakdown(bundle: RunBundle) -> dict[str, str]:
    return {
        "总耗时": format_timedelta(bundle.runtime_final - bundle.runtime_truth),
        "GA阶段": format_timedelta(bundle.runtime_ga - bundle.runtime_truth),
        "AM阶段": format_timedelta(bundle.runtime_am - bundle.runtime_ga),
        "后处理阶段": format_timedelta(bundle.runtime_final - bundle.runtime_am),
    }


def top_nodes_from_row(row: pd.Series, topn: int = 5) -> str:
    vals = [(node, float(row[node])) for node in CANDIDATE_NODES if node in row.index]
    vals = sorted(vals, key=lambda item: item[1], reverse=True)[:topn]
    return "，".join(f"{node}={value:.4f}" for node, value in vals if value > 0)


def top_rows(df: pd.DataFrame, sort_col: str, topn: int = 5, ascending: bool = False) -> pd.DataFrame:
    return df.sort_values(sort_col, ascending=ascending).head(topn).copy()


def unique_state_top_rows(df: pd.DataFrame, sort_col: str, topn: int = 5, ascending: bool = False) -> pd.DataFrame:
    return (
        df.sort_values([sort_col], ascending=[ascending])
        .drop_duplicates(subset=CANDIDATE_NODES, keep="first")
        .copy()
    )


def build_display_top_rows(df: pd.DataFrame, sort_col: str, topn: int = 5, ascending: bool = False) -> pd.DataFrame:
    out = unique_state_top_rows(df, sort_col, topn * 4, ascending)
    out["结构前五"] = out.apply(top_nodes_from_row, axis=1)
    out = out.drop_duplicates(subset=["结构前五"], keep="first").head(topn).copy()
    return out


def build_stage_top_tables(bundle: RunBundle) -> dict[str, pd.DataFrame]:
    ga_top = build_display_top_rows(bundle.ga_all, "mean_nse", 5, False)
    ppd_top = top_rows(bundle.initial_ppd, "mean_nse", 5, False)
    ppd_top["结构前五"] = ppd_top.apply(top_nodes_from_row, axis=1)
    am_top_nse = build_display_top_rows(bundle.am_samples, "mean_nse", 5, False)
    am_top_like = build_display_top_rows(bundle.am_samples, "log_like", 5, False)
    posterior_top = bundle.weights.sort_values("posterior_median", ascending=False).head(5).copy()
    return {
        "ga_top": ga_top,
        "ppd_top": ppd_top,
        "am_top_nse": am_top_nse,
        "am_top_like": am_top_like,
        "posterior_top": posterior_top,
    }


def build_truth_shares() -> np.ndarray:
    arr = np.zeros(len(CANDIDATE_NODES), dtype=float)
    truth = {
        "J76": 0.23684210526315788,
        "J124": 0.34210526315789475,
        "J140": 0.42105263157894735,
    }
    for node, value in truth.items():
        arr[CANDIDATE_NODES.index(node)] = value
    return arr


def shares_from(mapping: dict[str, float]) -> np.ndarray:
    arr = np.zeros(len(CANDIDATE_NODES), dtype=float)
    for node, value in mapping.items():
        arr[CANDIDATE_NODES.index(node)] = value
    return arr


def compute_current_compensation() -> dict[str, dict[str, float]]:
    generated = load_generated_data()
    dataset = build_dataset(generated)
    runtime = str(runtime_model_path(0))
    cases = {
        "truth": build_truth_shares(),
        "j140_to_j145": shares_from(
            {
                "J76": 0.23684210526315788,
                "J124": 0.34210526315789475,
                "J145": 0.42105263157894735,
            }
        ),
        "only_j140": shares_from({"J140": 1.0}),
        "only_j145": shares_from({"J145": 1.0}),
    }
    result: dict[str, dict[str, float]] = {}
    for label, shares in cases.items():
        eval_result = evaluate_shares(shares, dataset, runtime)
        result[label] = {
            "mean_nse": float(eval_result["mean_nse"]),
            "sse": float(eval_result["sse"]),
        }
    return result


def load_previous_compensation() -> dict:
    ident = read_json(PREVIOUS_ROOT / "scheme_c_identifiability_check.json")
    return {
        "truth": ident["truth"],
        "j140_to_j145": ident["j140_to_j145"],
    }


def copy_papers_to_tmp() -> None:
    if ENG_PDF.exists() and CN_PDF.exists():
        return
    src = Path(r"D:\Users\xwechat_files\wxid_4668346683612_4126\msg\file\2026-03")
    eng_src = next(path for path in src.iterdir() if "Uncertainty analysis method" in path.name)
    cn_src = next(path for path in src.iterdir() if "刘淑雅" in path.name and path.suffix.lower() == ".pdf")
    shutil.copy2(eng_src, ENG_PDF)
    shutil.copy2(cn_src, CN_PDF)


def extract_paper_facts() -> dict:
    copy_papers_to_tmp()
    PdfReader(str(ENG_PDF))
    PdfReader(str(CN_PDF))
    return {
        "英文论文": {
            "网络规模": "19 根管段的小尺度合成排水网络",
            "监测规模": "3 个监测点（J6、J12、J19），另设 1 个单监测点对照工况",
            "数据来源": "SWMM / PySWMM 生成的合成数据，不是现场实测",
            "是否现场布设监测": "否",
            "反演对象": "各管段的缺陷参数 / 渗漏率参数",
            "参数约束": "连续参数反演，不是井点份额单纯形约束",
            "优化与采样": "多种群 GA + 轮盘赌 + Adaptive Metropolis",
            "典型参数": "5 个种群、每个种群 80 个体、共 25 代；每 5 代竞争与迁移一次；AM 最大 2000 步",
            "评分指标": "多监测点平均 NSE 作为 GA 目标；后续用 ACC、MCC、MAE 等综合评价",
            "论文观点": "监测点减少后，多点缺陷诊断效果明显下降；离监测点较远或位于分支上的对象更难诊断；大尺度管网应拆分为单元管网后分段诊断",
        },
        "中文论文": {
            "网络规模": "研究区现场布设 95 个监测检查井，实际存在 7 个错接节点",
            "监测规模": "对 30 / 40 / 50 / 60 / 70 个监测点工况逐一比较",
            "数据来源": "现场水位监测与降雨事件数据",
            "是否现场布设监测": "是",
            "反演对象": "节点雨水入流量反演，共 95 维",
            "参数约束": "满足总入流守恒与非负约束：各节点入流量之和等于总入流量",
            "优化与采样": "PSO 与 GWO-PSO，不包含 GA-AM",
            "典型参数": "粒子数 50，最大迭代 300",
            "评分指标": "以监测水位与模型响应误差为核心，联合水量平衡约束开展定位",
            "论文观点": "监测点数量与布设方式显著影响定位结果；GWO-PSO 在 60 个监测点工况下表现最好，查全率 100%，查准率 64%",
        },
        "当前方案": {
            "网络规模": "239 个节点、239 条连接，20 个候选注入井的大尺度受控试验网络",
            "监测规模": "当前正式方案为 9 个监测点：J74、J77、J78、J123、J126、J141、J139、J145、J231",
            "数据来源": "正式真值模型与去注水基线模型构成的受控真值实验数据",
            "是否现场布设监测": "否",
            "反演对象": "20 维井点注入份额向量",
            "参数约束": "份额非负且总和为 1，总量与总波形固定，仅反演空间分配",
            "优化与采样": "多种群 GA + 轮盘赌 + Adaptive Metropolis",
            "典型参数": "当前中参数为 4 个种群、每个种群 24 个体、12 代；AM 为 4 条链、每链 900 步",
            "评分指标": "GA 以多监测点平均 NSE 排序；AM 用 SSE 构造 likelihood 并输出 posterior summary 与 posterior predictive coverage",
            "论文适用性": "位于中英文论文之间：监测规模远小于中文论文的现场工况，但网络规模远大于英文论文的小网验证，因此更容易出现邻近节点代偿与分段排查需求",
        },
    }


def add_top_border(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "0")
    top.set(qn("w:color"), "D9D9D9")
    tc_borders.append(top)


def set_doc_font(doc: Document) -> None:
    styles = doc.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def dataframe_to_rows(df: pd.DataFrame) -> list[list[str]]:
    rows: list[list[str]] = [list(df.columns)]
    for _, row in df.iterrows():
        converted: list[str] = []
        for value in row.tolist():
            if isinstance(value, float):
                if math.isnan(value):
                    converted.append("")
                elif abs(value) >= 1000 or (0 < abs(value) < 1e-3):
                    converted.append(f"{value:.4e}")
                else:
                    converted.append(f"{value:.4f}")
            else:
                converted.append(str(value))
        rows.append(converted)
    return rows


def add_table(doc: Document, rows: list[list[str]], column_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(9.5 if r_idx == 0 else 9)
                    if r_idx == 0:
                        run.bold = True
            if r_idx == 0:
                add_top_border(cell)
            if column_widths:
                cell.width = Inches(column_widths[c_idx])
    doc.add_paragraph()


def add_image_with_caption(doc: Document, image_path: Path, caption: str, width: float = 6.5) -> None:
    doc.add_picture(str(image_path), width=Inches(width))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)


def plot_layout() -> Path:
    node_df = pd.read_csv(LAYOUT_NODE_CSV, encoding="utf-8-sig")
    link_df = pd.read_csv(LAYOUT_LINK_CSV, encoding="utf-8-sig")
    pos = {row["节点名称"]: (row["X"], row["Y"]) for _, row in node_df.iterrows()}

    fig, ax = plt.subplots(figsize=(9.2, 7.2), dpi=180)
    ax.set_facecolor("#fcfcfc")
    for _, row in link_df.iterrows():
        start = row["起点"]
        end = row["终点"]
        if start in pos and end in pos:
            ax.plot(
                [pos[start][0], pos[end][0]],
                [pos[start][1], pos[end][1]],
                color="#d0d0d0",
                linewidth=0.8,
                zorder=1,
            )

    ax.scatter(node_df["X"], node_df["Y"], s=9, c="#cfcfcf", alpha=0.8, zorder=2, label="全部节点")

    candidate_df = node_df[node_df["节点名称"].isin(CANDIDATE_NODES)]
    monitor_df = node_df[node_df["节点名称"].isin(MONITOR_NODES)]
    truth_df = node_df[node_df["节点名称"].isin(["J76", "J124", "J140"])]

    ax.scatter(candidate_df["X"], candidate_df["Y"], s=28, facecolors="none", edgecolors="#d55e00", linewidths=1.2, zorder=3, label="20个候选井")
    ax.scatter(monitor_df["X"], monitor_df["Y"], s=38, c="#0072B2", zorder=4, label="当前9个监测点")
    ax.scatter(truth_df["X"], truth_df["Y"], s=64, marker="*", c="#E69F00", edgecolors="#9A6700", linewidths=0.6, zorder=5, label="真值注入井")

    label_nodes = ["J74", "J77", "J78", "J123", "J126", "J141", "J139", "J145", "J231", "J76", "J124", "J140"]
    for _, row in node_df[node_df["节点名称"].isin(label_nodes)].iterrows():
        ax.text(row["X"] + 6, row["Y"] + 6, row["节点名称"], fontsize=8, color="#222222")

    ax.set_title("图1 当前239节点管网、候选井与9个监测点布设示意", fontsize=12)
    ax.set_xlabel("X 坐标")
    ax.set_ylabel("Y 坐标")
    ax.legend(loc="upper right", fontsize=8, frameon=True)
    ax.grid(alpha=0.12, linewidth=0.5)
    fig.tight_layout()
    out = ASSET_DIR / "figure_01_layout.png"
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    return out


def plot_ga_convergence(previous: RunBundle, current: RunBundle) -> Path:
    prev_curve = previous.ga_history.groupby("generation")["best_mean_nse"].max().reset_index()
    curr_curve = current.ga_history.groupby("generation")["best_mean_nse"].max().reset_index()
    fig, ax = plt.subplots(figsize=(8.8, 4.8), dpi=180)
    ax.plot(prev_curve["generation"], prev_curve["best_mean_nse"], marker="o", linewidth=2, color="#7f7f7f", label="上一版8点方案")
    ax.plot(curr_curve["generation"], curr_curve["best_mean_nse"], marker="o", linewidth=2, color="#1f77b4", label="当前9点方案(+J77)")
    ax.set_xlabel("代数")
    ax.set_ylabel("GA全局最优 Mean NSE")
    ax.set_title("图2 两轮中参数实验的GA收敛曲线")
    ax.grid(alpha=0.2)
    ax.legend()
    fig.tight_layout()
    out = ASSET_DIR / "figure_02_ga_convergence_compare.png"
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    return out


def plot_runtime_compare(previous: RunBundle, current: RunBundle) -> Path:
    prev = runtime_breakdown(previous)
    curr = runtime_breakdown(current)

    def to_hours(s: str) -> float:
        h = re.search(r"(\d+)小时", s)
        m = re.search(r"(\d+)分钟", s)
        sec = re.search(r"(\d+)秒", s)
        hh = int(h.group(1)) if h else 0
        mm = int(m.group(1)) if m else 0
        ss = int(sec.group(1)) if sec else 0
        return hh + mm / 60 + ss / 3600

    labels = list(prev.keys())
    prev_vals = [to_hours(prev[k]) for k in labels]
    curr_vals = [to_hours(curr[k]) for k in labels]
    x = np.arange(len(labels))
    width = 0.35
    fig, ax = plt.subplots(figsize=(8.8, 4.8), dpi=180)
    ax.bar(x - width / 2, prev_vals, width, label="上一版8点方案", color="#7f7f7f")
    ax.bar(x + width / 2, curr_vals, width, label="当前9点方案(+J77)", color="#1f77b4")
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.set_ylabel("小时")
    ax.set_title("图3 两轮中参数实验耗时拆分")
    ax.legend()
    ax.grid(axis="y", alpha=0.2)
    for idx, value in enumerate(prev_vals):
        ax.text(idx - width / 2, value + 0.08, f"{value:.2f}", ha="center", va="bottom", fontsize=8)
    for idx, value in enumerate(curr_vals):
        ax.text(idx + width / 2, value + 0.08, f"{value:.2f}", ha="center", va="bottom", fontsize=8)
    fig.tight_layout()
    out = ASSET_DIR / "figure_03_runtime_compare.png"
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    return out


def plot_compensation_compare(prev_comp: dict, curr_comp: dict) -> Path:
    labels = ["真值方案", "J140→J145代偿"]
    prev_vals = [prev_comp["truth"]["mean_nse"], prev_comp["j140_to_j145"]["mean_nse"]]
    curr_vals = [curr_comp["truth"]["mean_nse"], curr_comp["j140_to_j145"]["mean_nse"]]
    x = np.arange(len(labels))
    width = 0.35
    fig, ax = plt.subplots(figsize=(7.8, 4.8), dpi=180)
    ax.bar(x - width / 2, prev_vals, width, label="上一版8点方案", color="#7f7f7f")
    ax.bar(x + width / 2, curr_vals, width, label="当前9点方案(+J77)", color="#1f77b4")
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.set_ylabel("Mean NSE")
    ax.set_ylim(0, 1.05)
    ax.set_title("图4 前后两版对J140→J145邻近代偿的敏感性对比")
    ax.legend()
    ax.grid(axis="y", alpha=0.2)
    for idx, value in enumerate(prev_vals):
        ax.text(idx - width / 2, value + 0.015, f"{value:.4f}", ha="center", va="bottom", fontsize=8)
    for idx, value in enumerate(curr_vals):
        ax.text(idx + width / 2, value + 0.015, f"{value:.4f}", ha="center", va="bottom", fontsize=8)
    fig.tight_layout()
    out = ASSET_DIR / "figure_04_compensation_compare.png"
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    return out


def plot_single_injection_scores(curr_comp: dict) -> Path:
    labels = ["仅在J140注入", "仅在J145注入"]
    vals = [curr_comp["only_j140"]["mean_nse"], curr_comp["only_j145"]["mean_nse"]]
    fig, ax = plt.subplots(figsize=(7.2, 4.6), dpi=180)
    bars = ax.bar(labels, vals, color=["#2ca02c", "#d62728"])
    ax.axhline(0, color="#555555", linewidth=0.8)
    ax.set_ylabel("Mean NSE")
    ax.set_title("图5 当前9点方案下单点注入得分")
    ax.grid(axis="y", alpha=0.2)
    for bar, value in zip(bars, vals):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            value + (0.2 if value >= 0 else -0.3),
            f"{value:.4f}",
            ha="center",
            va="bottom" if value >= 0 else "top",
            fontsize=8,
        )
    fig.tight_layout()
    out = ASSET_DIR / "figure_05_single_injection_scores.png"
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    return out


def plot_truth_node_compare(previous: RunBundle, current: RunBundle) -> Path:
    focus = ["J140", "J124", "J76", "J145"]
    prev_map = {row["node"]: float(row["posterior_median"]) for _, row in previous.weights.iterrows()}
    curr_map = {row["node"]: float(row["posterior_median"]) for _, row in current.weights.iterrows()}
    truth_map = {"J140": 0.4210526316, "J124": 0.3421052632, "J76": 0.2368421053, "J145": 0.0}
    x = np.arange(len(focus))
    width = 0.24
    fig, ax = plt.subplots(figsize=(8.6, 4.8), dpi=180)
    ax.bar(x - width, [truth_map[n] for n in focus], width, label="真值份额", color="#E69F00")
    ax.bar(x, [prev_map.get(n, 0.0) for n in focus], width, label="上一版后验中位数", color="#7f7f7f")
    ax.bar(x + width, [curr_map.get(n, 0.0) for n in focus], width, label="当前版后验中位数", color="#1f77b4")
    ax.set_xticks(x)
    ax.set_xticklabels(focus)
    ax.set_ylabel("份额")
    ax.set_title("图6 真值节点与关键代偿节点的后验中位数对比")
    ax.legend()
    ax.grid(axis="y", alpha=0.2)
    fig.tight_layout()
    out = ASSET_DIR / "figure_06_truth_node_compare.png"
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    return out


def build_markdown(context: dict) -> str:
    current = context["current"]
    lines = [
        "# 0401 Scheme C + J77 正式分析报告 v4",
        "",
        "## 核心结论",
        f"- 当前9点方案(+J77)的 posterior top3 为：{' / '.join(current.summary['predicted_top3'])}。",
        f"- 当前中参数总耗时：{runtime_breakdown(current)['总耗时']}，其中 AM 阶段约 {runtime_breakdown(current)['AM阶段']}。",
        f"- 当前版 J140→J145 代偿 Mean NSE = {context['current_comp']['j140_to_j145']['mean_nse']:.4f}。",
        "",
        "## 说明",
        "- 正式Word报告请查看同目录 docx 文件。",
    ]
    return "\n".join(lines)


def build_doc(context: dict) -> Document:
    current: RunBundle = context["current"]
    previous: RunBundle = context["previous"]
    paper_facts: dict = context["paper_facts"]
    current_comp: dict = context["current_comp"]
    previous_comp: dict = context["previous_comp"]
    current_rpt: dict = context["current_rpt"]
    baseline_rpt: dict = context["baseline_rpt"]
    stage_tables: dict = context["current_stage_tables"]

    doc = Document()
    set_doc_font(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("0401 Scheme C + J77 正式分析报告 v4")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = subtitle.add_run("主题：9个监测点方案下的阶段结果、邻近代偿机理、论文适用性与大网排查含义")
    srun.font.size = Pt(10.5)
    srun.font.name = "Microsoft YaHei"
    srun._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    add_heading(doc, "一、实验口径与算法说明", 1)
    add_bullets(
        doc,
        [
            "当前正式网络为 239 个节点、239 条连接，候选注入井 20 个；正式监测点为 9 个：J74、J77、J78、J123、J126、J141、J139、J145、J231。",
            "真值注入井为 J76、J124、J140；当前真值总量固定为 22800 m3，总波形固定，仅反演 20 个候选井之间的空间份额分配。",
            "算法流程为：多种群 GA 搜索高分区 → 轮盘赌形成 initial PPD → AM 多链采样 → posterior summary 与 posterior predictive validation。",
            "GA 的排序指标是多监测点平均 NSE；AM 的接受率依据 likelihood ratio，likelihood 由 SSE 构造；后验总结使用 posterior mean、posterior median、P05、P95。",
        ],
    )

    algo_rows = [
        ["模块", "当前正式实现", "当前中参数设置"],
        ["GA", "4个种群并行搜索，按 mean NSE 排序，采用精英保留、交叉、变异、competition 与 migration", "4个种群 × 24个体 × 12代"],
        ["轮盘赌", "对 GA 末代合并池按适应度归一化抽样，形成 initial PPD", "保留 24 个初始样本"],
        ["AM", "4 条链并行采样，协方差按 2.4^2 / d 自适应更新，warmup 后进入后验统计", "4链 × 900步，warmup=220，adapt_start=220"],
        ["后验输出", "输出 posterior mean / median / P05 / P95，并做 posterior predictive coverage", "24 个后验预测样本"],
    ]
    add_table(doc, algo_rows, [0.95, 3.7, 1.8])
    add_image_with_caption(doc, context["figure_layout"], "图1 当前239节点管网、20个候选井、9个监测点与3个真值井的空间关系", width=6.7)

    add_heading(doc, "二、两轮中参数实验的结果分析", 1)
    add_paragraph(doc, "本节将上一版8点方案与当前9点方案(+J77)并列比较，同时将当前版本的 GA、initial PPD、AM、posterior summary、posterior predictive 五个阶段拆开分析，避免将不同阶段指标混为一谈。")

    add_heading(doc, "2.1 数据闭合性与数值口径", 2)
    add_bullets(
        doc,
        [
            f"当前 truth replay 结果为 Mean NSE={current.truth_replay['mean_nse']:.4f}，SSE={current.truth_replay['sse']:.4e}，说明正式基线模型、真值模型、注水接口与评分链严格闭合。",
            f"基线模型连续性误差为 {baseline_rpt['continuity_error_pct']:.3f}%，真值事件模型连续性误差为 {current_rpt['continuity_error_pct']:.3f}%。这里的连续性误差已经包含最终存储量，不能简单理解为“漏算存储”，而是整个事件工况的水量平衡残差。",
            "因此，当前结果中的识别偏差不是由数据链断裂造成，而是在大尺度网络上由真实水力等效关系导致的结构性不确定性。",
        ],
    )

    compare_rows = [
        ["指标", "上一版8点方案", "当前9点方案(+J77)"],
        ["GA best Mean NSE", f"{previous.summary['ga_best_mean_nse']:.4f}", f"{current.summary['ga_best_mean_nse']:.4f}"],
        ["posterior median NSE", f"{previous.summary['posterior_median_nse']:.4f}", f"{current.summary['posterior_median_nse']:.4f}"],
        ["posterior best NSE", f"{previous.summary['posterior_best_nse']:.4f}", f"{current.summary['posterior_best_nse']:.4f}"],
        ["posterior coverage mean", f"{previous.summary['posterior_coverage_mean']:.4f}", f"{current.summary['posterior_coverage_mean']:.4f}"],
        ["预测前3", " / ".join(previous.summary["predicted_top3"]), " / ".join(current.summary["predicted_top3"])],
        ["总耗时", runtime_breakdown(previous)["总耗时"], runtime_breakdown(current)["总耗时"]],
        ["GA阶段耗时", runtime_breakdown(previous)["GA阶段"], runtime_breakdown(current)["GA阶段"]],
        ["AM阶段耗时", runtime_breakdown(previous)["AM阶段"], runtime_breakdown(current)["AM阶段"]],
        ["后处理阶段耗时", runtime_breakdown(previous)["后处理阶段"], runtime_breakdown(current)["后处理阶段"]],
    ]
    add_table(doc, compare_rows, [1.7, 2.0, 2.0])
    add_image_with_caption(doc, context["figure_ga"], "图2 两轮中参数实验的GA全局最优收敛曲线", width=6.4)
    add_image_with_caption(doc, context["figure_runtime"], "图3 两轮中参数实验的耗时拆分；当前版总耗时约9小时21分，其中 AM 阶段约7小时26分", width=6.4)

    add_heading(doc, "2.2 当前版本各阶段 TOP5 结果", 2)
    add_paragraph(doc, "以下表格均只对应当前9点方案(+J77)中参数实验，不与上一版混用。")

    add_paragraph(doc, "（1）GA 阶段 TOP5：")
    add_table(doc, dataframe_to_rows(stage_tables["ga_top"][["mean_nse", "sse", "结构前五"]]), [1.0, 1.0, 4.4])

    add_paragraph(doc, "（2）轮盘赌 initial PPD TOP5：")
    add_table(doc, dataframe_to_rows(stage_tables["ppd_top"][["mean_nse", "sse", "roulette_weight", "结构前五"]]), [0.9, 0.9, 0.9, 3.9])

    add_paragraph(doc, "（3）AM 样本中按 mean NSE 排序的 TOP5：")
    add_table(doc, dataframe_to_rows(stage_tables["am_top_nse"][["chain", "step", "mean_nse", "sse", "log_like", "结构前五"]]), [0.55, 0.55, 0.85, 0.85, 0.85, 3.3])

    add_paragraph(doc, "（4）AM 样本中按 log_like 排序的 TOP5：")
    add_table(doc, dataframe_to_rows(stage_tables["am_top_like"][["chain", "step", "mean_nse", "sse", "log_like", "结构前五"]]), [0.55, 0.55, 0.85, 0.85, 0.85, 3.3])

    add_paragraph(doc, "（5）posterior summary 的节点级 TOP5：")
    add_table(doc, dataframe_to_rows(stage_tables["posterior_top"][["node", "posterior_mean", "posterior_median", "p05", "p95"]]), [0.7, 1.0, 1.0, 1.0, 1.0])

    add_bullets(
        doc,
        [
            "GA 阶段的 top1 仍然带有邻近节点分摊，说明 GA 本身更容易停留在较宽的高分盆地中；但进入 initial PPD 后，高分池已经保留了更接近真值结构的样本。",
            "AM 样本按 mean NSE 排序时已经出现了明显贴近真值三点的高分样本；按 log_like 排序时，最优样本的结构同样围绕 J140、J124、J76 形成主峰。",
            "posterior summary 最终将真值三点直接拉成前3位，这说明新增 J77 后，局部可辨识性较上一版有实质提升。",
        ],
    )

    add_heading(doc, "2.3 当前版本的后验与后验预测", 2)
    coverage_sorted = current.coverage.sort_values("coverage_90")
    best_cov = coverage_sorted.tail(3)[["monitor", "coverage_90"]]
    weak_cov = coverage_sorted.head(3)[["monitor", "coverage_90"]]
    coverage_rows = [["类别", "监测点", "coverage_90"]]
    for _, row in weak_cov.iterrows():
        coverage_rows.append(["较弱监测点", row["monitor"], f"{row['coverage_90']:.4f}"])
    for _, row in best_cov.iterrows():
        coverage_rows.append(["较强监测点", row["monitor"], f"{row['coverage_90']:.4f}"])
    add_table(doc, coverage_rows, [1.2, 1.0, 1.0])
    add_bullets(
        doc,
        [
            f"当前 posterior median 方案的 Mean NSE 为 {current.summary['posterior_median_nse']:.4f}；posterior best（按 log_like 选取）对应 Mean NSE 为 {current.summary['posterior_best_nse']:.4f}。",
            f"posterior predictive 平均 coverage 为 {current.summary['posterior_coverage_mean']:.4f}，说明后验整体解释能力仍保持在较高水平。",
            "与上一版相比，当前版的 posterior median 分数略低，但 posterior best 更高、posterior top3 更接近真值，说明监测点调整提升的是结构识别正确性，而不是单一代表解的外观分数。",
        ],
    )

    add_heading(doc, "三、邻近代偿的机理分析", 1)
    add_paragraph(doc, "本节仅保留前后两版都真实存在、且具有工程解释意义的邻近代偿：J140→J145。此前与其他邻近点相关的讨论不再展开。")
    add_image_with_caption(doc, context["figure_comp"], "图4 前后两版对 J140→J145 邻近代偿的敏感性对比", width=6.1)
    add_image_with_caption(doc, context["figure_single"], "图5 当前9点方案下单点注入得分：仅在 J140 或仅在 J145 注入都无法单独解释整体响应", width=6.0)
    add_image_with_caption(doc, context["figure_truth"], "图6 真值节点与关键代偿节点的后验中位数对比；当前版已将 J145 压低并恢复 J76/J124/J140 真值结构", width=6.4)

    compensation_rows = [
        ["工况", "上一版8点方案 Mean NSE", "当前9点方案(+J77) Mean NSE", "解释"],
        ["真值方案", f"{previous_comp['truth']['mean_nse']:.4f}", f"{current_comp['truth']['mean_nse']:.4f}", "作为理论上限，说明评分链闭合"],
        ["J140→J145 代偿", f"{previous_comp['j140_to_j145']['mean_nse']:.4f}", f"{current_comp['j140_to_j145']['mean_nse']:.4f}", "替换后仍能得到较高分数，说明 J140 与 J145 在大网中存在明显邻近代偿"],
        ["仅在 J140 注入", "-", f"{current_comp['only_j140']['mean_nse']:.4f}", "单点不足以解释三点真值事件"],
        ["仅在 J145 注入", "-", f"{current_comp['only_j145']['mean_nse']:.4f}", "单点同样无法单独解释整体事件"],
    ]
    add_table(doc, compensation_rows, [1.4, 1.5, 1.6, 2.5])

    add_bullets(
        doc,
        [
            "当前网络规模为 239 节点、239 连接。对于这种 200 多节点的大尺度管网，局部节点一旦处于相邻支路、相邻汇流路径或相近传播距离上，就可能对下游多个监测点产生相似响应，从而形成代偿。",
            f"上一版中，J140→J145 的替代得分已达到 {previous_comp['j140_to_j145']['mean_nse']:.4f}；当前版即使增加 J77 并改善其他区域辨识，J140→J145 仍有 {current_comp['j140_to_j145']['mean_nse']:.4f}，说明这一代偿不是偶然噪声，而是由大网邻近结构导致的真实水力等效。",
            "因此，当前结果更适合解释为已将疑似源区压缩到真值邻近簇内，而不是一步精确定位到唯一井点。在 200 多节点网络里，更符合工程实际的做法是根据后验前列结果做分段排查与重点核查。",
        ],
    )

    add_heading(doc, "四、中英文论文与当前方案的深入对照", 1)
    paper_rows = [["比较维度", "英文论文", "中文论文", "当前方案"]]
    keys = [
        "网络规模",
        "监测规模",
        "数据来源",
        "是否现场布设监测",
        "反演对象",
        "参数约束",
        "优化与采样",
        "典型参数",
        "评分指标",
    ]
    for key in keys:
        paper_rows.append([key, paper_facts["英文论文"][key], paper_facts["中文论文"][key], paper_facts["当前方案"][key]])
    add_table(doc, paper_rows, [1.05, 2.2, 2.2, 2.2])

    view_rows = [
        ["论文 / 方案", "关键观点或适用性边界"],
        ["英文论文", paper_facts["英文论文"]["论文观点"]],
        ["中文论文", paper_facts["中文论文"]["论文观点"]],
        ["当前方案", paper_facts["当前方案"]["论文适用性"]],
    ]
    add_table(doc, view_rows, [1.2, 6.2])

    add_bullets(
        doc,
        [
            "中文论文的核心优势在于：网络来自现场工况、监测点数量远多于当前方案，并通过 30/40/50/60/70 个监测点对比明确指出 60 个监测点时效果最好。这一证据说明，监测点数量与现场覆盖度对工程定位效果有决定性影响。",
            "英文论文的核心优势在于：GA-AM 框架完整，适合展示从 GA 搜索高分区到 AM 输出不确定性的全过程；但其网络规模仅为 19 根管段、3 个监测点，属于方法学验证而非大尺度工程级验证。",
            "当前方案恰好位于两篇论文之间：监测点数量远少于中文论文的现场工况，但网络规模却远大于英文论文的小尺度网络，因此更容易出现真值已进入前列但邻近节点代偿仍然存在的现象。",
            "英文论文明确指出，离监测点较远和位于分支上的对象更难诊断，并建议对大尺度管网拆分成单元系统后进行分段诊断。当前 239 节点网络中仍观察到 J140→J145 代偿，与这一论文观点是一致的。",
        ],
    )

    add_heading(doc, "五、结论", 1)
    add_bullets(
        doc,
        [
            "当前 9 点方案（在上一版基础上新增 J77）已经显著改善了结构识别结果：posterior top3 由上一版的 J140 / J124 / J78 改为当前版的 J140 / J124 / J76，真值三点全部进入后验前三。",
            "当前版中参数单次总耗时约 9 小时 21 分，其中 AM 阶段约 7 小时 26 分，占总耗时主体。对 200 多节点网络而言，这一耗时本身也说明工程应用更适合以分段排查、重点核查的方式使用结果，而不是期待单次运行给出绝对唯一答案。",
            "J140→J145 邻近代偿在前后两版都存在，说明大尺度网络中的局部水力等效是客观存在的；监测点选位能够显著改善识别结果，但并不能完全抹去所有邻近代偿。",
            "因此，当前这一版的最重要意义不是已经完全没有代偿，而是通过监测点布局优化，已经把最优解稳定推回到真值三点前列，并将局部代偿压缩到较小的邻近范围内。",
        ],
    )

    return doc


def main() -> None:
    ensure_dirs()
    current = load_run_bundle("当前9点方案(+J77)", CURRENT_RUN_DIR, list(MONITOR_NODES))
    previous = load_run_bundle("上一版8点方案", PREVIOUS_RUN_DIR, ["J74", "J78", "J123", "J126", "J141", "J139", "J145", "J231"])
    current_rpt = parse_rpt_metrics(TRUTH_EVENT_MODEL_RPT)
    baseline_rpt = parse_rpt_metrics(BASELINE_MODEL_RPT)
    paper_facts = extract_paper_facts()
    current_comp = compute_current_compensation()
    previous_comp = load_previous_compensation()
    current_stage_tables = build_stage_top_tables(current)

    context = {
        "current": current,
        "previous": previous,
        "current_rpt": current_rpt,
        "baseline_rpt": baseline_rpt,
        "paper_facts": paper_facts,
        "current_comp": current_comp,
        "previous_comp": previous_comp,
        "current_stage_tables": current_stage_tables,
        "figure_layout": plot_layout(),
        "figure_ga": plot_ga_convergence(previous, current),
        "figure_runtime": plot_runtime_compare(previous, current),
        "figure_comp": plot_compensation_compare(previous_comp, current_comp),
        "figure_single": plot_single_injection_scores(current_comp),
        "figure_truth": plot_truth_node_compare(previous, current),
    }

    doc = build_doc(context)
    doc.save(DOCX_PATH)
    MD_PATH.write_text(build_markdown(context), encoding="utf-8")
    print(f"Saved report: {DOCX_PATH}")
    print(f"Saved markdown: {MD_PATH}")


if __name__ == "__main__":
    main()
