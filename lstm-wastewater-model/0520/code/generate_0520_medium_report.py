from __future__ import annotations

import json
from pathlib import Path

import numpy as np
import pandas as pd
import plotly.graph_objects as go
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
import matplotlib.pyplot as plt
from matplotlib.lines import Line2D
from matplotlib.patches import Patch


ROOT = Path(__file__).resolve().parents[1]
RESULT_DIR = ROOT / "results_0520" / "medium_run"
DATA_DIR = ROOT / "data" / "generated_0520"
ANALYSIS_DIR = ROOT / "analysis_0520"
OUT_DIR = RESULT_DIR / "analysis_report_0520_medium"
FIG_DIR = OUT_DIR / "figures"

CANDIDATE_NODES = [
    "10", "62", "124", "42", "178", "63", "103", "241", "273", "215",
    "216", "60", "308", "310", "312", "118", "304", "64", "91", "85",
]
MONITOR_NODES = ["286", "223", "239", "267", "8", "251", "252", "189", "37"]
TRUTH_NODES = ["103", "304", "10", "178", "42"]


def setup_matplotlib() -> None:
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "SimSun", "Arial Unicode MS"]
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.dpi"] = 150


def pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def fmt(value: float, digits: int = 4) -> str:
    return f"{value:.{digits}f}"


def load_inputs() -> dict[str, object]:
    return {
        "summary": json.loads((RESULT_DIR / "0520_summary.json").read_text(encoding="utf-8")),
        "data_summary": json.loads((DATA_DIR / "0520_data_summary.json").read_text(encoding="utf-8")),
        "scores": pd.read_csv(RESULT_DIR / "0520_solution_scores.csv", encoding="utf-8-sig"),
        "shares": pd.read_csv(RESULT_DIR / "0520_solution_shares.csv", encoding="utf-8-sig"),
        "ga_history": pd.read_csv(RESULT_DIR / "0520_GA_best_by_generation.csv", encoding="utf-8-sig"),
        "ga_all": pd.read_csv(RESULT_DIR / "0520_GA_all.csv", encoding="utf-8-sig"),
        "ga_last": pd.read_csv(RESULT_DIR / "0520_GA_last_generation.csv", encoding="utf-8-sig"),
        "initial_ppd": pd.read_csv(RESULT_DIR / "0520_initial_PPD.csv", encoding="utf-8-sig"),
        "am": pd.read_csv(RESULT_DIR / "0520_AM_samples.csv", encoding="utf-8-sig"),
        "posterior": pd.read_csv(RESULT_DIR / "0520_posterior_node_weights.csv", encoding="utf-8-sig"),
        "observed_delta": pd.read_csv(DATA_DIR / "0520_observed_delta_10min.csv", encoding="utf-8-sig"),
        "map_delta": pd.read_csv(RESULT_DIR / "0520_solution_posterior_best_map_delta.csv", encoding="utf-8-sig"),
        "ga_delta": pd.read_csv(RESULT_DIR / "0520_solution_ga_best_delta.csv", encoding="utf-8-sig"),
        "median_delta": pd.read_csv(RESULT_DIR / "0520_solution_posterior_median_summary_delta.csv", encoding="utf-8-sig"),
        "outlet": pd.read_csv(DATA_DIR / "0520_outlet_series_10min.csv", encoding="utf-8-sig"),
        "map_outlet": pd.read_csv(RESULT_DIR / "0520_solution_posterior_best_map_outlet.csv", encoding="utf-8-sig"),
        "nodes": pd.read_csv(ANALYSIS_DIR / "0520_nodes_classified.csv", encoding="utf-8-sig"),
        "links": pd.read_csv(ANALYSIS_DIR / "0520_links_classified.csv", encoding="utf-8-sig"),
        "plan": pd.read_csv(ANALYSIS_DIR / "0520_monitor_candidate_injection_plan.csv", encoding="utf-8-sig"),
    }


def solution_share_series(shares: pd.DataFrame, name: str) -> pd.Series:
    row = shares.loc[shares["solution"] == name].iloc[0]
    return row[CANDIDATE_NODES].astype(float)


def truth_share_series() -> pd.Series:
    s = pd.Series(0.0, index=CANDIDATE_NODES)
    for node in TRUTH_NODES:
        s[node] = 1.0 / len(TRUTH_NODES)
    return s


def compute_detection_table(shares: pd.DataFrame, threshold: float = 0.05) -> pd.DataFrame:
    rows = []
    truth = set(TRUTH_NODES)
    for name in ["ga_best", "posterior_best_map"]:
        s = solution_share_series(shares, name)
        active = set(s[s > threshold].index)
        rows.append(
            {
                "方案": name,
                "阈值": pct(threshold),
                "超过阈值点数": len(active),
                "命中真值点": "、".join([n for n in TRUTH_NODES if n in active]) or "无",
                "命中数量": len(active & truth),
                "漏掉真值点": "、".join([n for n in TRUTH_NODES if n not in active]) or "无",
                "主要代偿点": "、".join(sorted(active - truth, key=lambda n: float(s[n]), reverse=True)) or "无",
                "真值点总比例": s.loc[TRUTH_NODES].sum(),
                "非真值点总比例": 1.0 - s.loc[TRUTH_NODES].sum(),
            }
        )
    return pd.DataFrame(rows)


def nse(obs: np.ndarray, sim: np.ndarray) -> float:
    denom = float(np.sum((obs - obs.mean()) ** 2))
    if denom <= 1e-12:
        return float("nan")
    return 1.0 - float(np.sum((sim - obs) ** 2)) / denom


def monitor_metrics(observed: pd.DataFrame, simulated: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for node in MONITOR_NODES:
        obs = observed[node].to_numpy(dtype=float)
        sim = simulated[node].to_numpy(dtype=float)
        rows.append(
            {
                "监测点": node,
                "NSE": nse(obs, sim),
                "SSE": float(np.sum((sim - obs) ** 2)),
                "观测峰值": float(np.max(obs)),
                "模拟峰值": float(np.max(sim)),
                "峰值差": float(np.max(sim) - np.max(obs)),
            }
        )
    return pd.DataFrame(rows)


def save_score_bar(scores: pd.DataFrame) -> Path:
    path = FIG_DIR / "01_三类结果评分对比.png"
    scores = scores[scores["solution"].isin(["ga_best", "posterior_best_map"])].copy()
    labels = ["GA最佳", "AM MAP"]
    mean_nse = scores["mean_nse"].to_numpy(float)
    sse = scores["sse"].to_numpy(float)
    fig, ax1 = plt.subplots(figsize=(9, 4.8))
    x = np.arange(len(labels))
    bars = ax1.bar(x - 0.18, mean_nse, width=0.36, color="#2E75B6", label="mean NSE")
    ax1.set_ylabel("mean NSE，越高越好")
    ax1.set_ylim(min(-0.1, float(mean_nse.min()) - 0.1), 1.05)
    ax1.set_xticks(x)
    ax1.set_xticklabels(labels)
    ax1.grid(axis="y", alpha=0.25)
    ax2 = ax1.twinx()
    sse_bars = ax2.bar(x + 0.18, sse, width=0.36, color="#ED7D31", label="SSE")
    ax2.set_ylabel("SSE，越低越好")
    for bar, val in zip(bars, mean_nse):
        ax1.text(bar.get_x() + bar.get_width() / 2, val + 0.02, f"{val:.3f}", ha="center", fontsize=9)
    fig.suptitle("本次中参数运行的核心结果评分")
    ax1.legend([bars, sse_bars], ["mean NSE，越高越好", "SSE，越低越好"], loc="upper left")
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_ga_convergence(ga_history: pd.DataFrame) -> Path:
    path = FIG_DIR / "02_GA逐代收敛过程.png"
    best_by_gen = ga_history.groupby("generation")["best_mean_nse"].max()
    median_by_gen = ga_history.groupby("generation")["best_mean_nse"].median()
    fig, ax = plt.subplots(figsize=(9, 4.8))
    ax.plot(best_by_gen.index, best_by_gen.values, marker="o", color="#2E75B6", label="每代最优")
    ax.plot(median_by_gen.index, median_by_gen.values, marker="s", color="#70AD47", label="各种群最优的中位数")
    ax.set_xlabel("GA代数")
    ax.set_ylabel("mean NSE")
    ax.set_title("GA阶段：前期快速提升，后期进入平台")
    ax.grid(alpha=0.25)
    ax.legend()
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_share_comparison(shares: pd.DataFrame) -> Path:
    path = FIG_DIR / "03_真值与识别比例对比.png"
    truth = truth_share_series()
    data = pd.DataFrame({
        "真值": truth,
        "GA最佳": solution_share_series(shares, "ga_best"),
        "AM MAP": solution_share_series(shares, "posterior_best_map"),
    })
    x = np.arange(len(CANDIDATE_NODES))
    fig, ax = plt.subplots(figsize=(13, 5.2))
    width = 0.26
    colors = ["#A5A5A5", "#4472C4", "#ED7D31"]
    for i, col in enumerate(data.columns):
        ax.bar(x + (i - 1.0) * width, data[col].values, width=width, label=col, color=colors[i])
    ax.set_xticks(x)
    ax.set_xticklabels(CANDIDATE_NODES, rotation=45)
    ax.set_ylabel("注入比例")
    ax.set_title("20个候选点上的比例分配：拟合结果与真值对比")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(ncol=3, loc="upper right")
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_active_count_distribution(ga_all: pd.DataFrame, ga_last: pd.DataFrame, initial_ppd: pd.DataFrame) -> Path:
    path = FIG_DIR / "04_GA有效激活点数量分布.png"
    datasets = {"GA全部": ga_all, "GA最后一代": ga_last, "initial PPD": initial_ppd}
    fig, axes = plt.subplots(1, 3, figsize=(12, 3.8), sharey=True)
    for ax, (label, df) in zip(axes, datasets.items()):
        counts = (df[CANDIDATE_NODES].to_numpy(float) > 0.05).sum(axis=1)
        bins = np.arange(0.5, max(12, counts.max() + 1.5), 1)
        ax.hist(counts, bins=bins, color="#5B9BD5", edgecolor="white")
        ax.axvline(np.median(counts), color="#C00000", linestyle="--", label=f"中位数 {np.median(counts):.0f}")
        ax.set_title(label)
        ax.set_xlabel("比例超过5%的点数")
        ax.grid(axis="y", alpha=0.25)
        ax.legend(fontsize=8)
    axes[0].set_ylabel("解的数量")
    fig.suptitle("GA并未强制只保留5个点：实际有效点数分布")
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_monitor_fit(observed: pd.DataFrame, simulated: pd.DataFrame) -> Path:
    path = FIG_DIR / "05_AM_MAP各监测点拟合曲线.png"
    fig, axes = plt.subplots(3, 3, figsize=(13, 8.8), sharex=True)
    hours = observed["relative_hour"].to_numpy(float)
    for ax, node in zip(axes.flatten(), MONITOR_NODES):
        ax.plot(hours, observed[node], color="#1F4E79", linewidth=1.6, label="观测增量")
        ax.plot(hours, simulated[node], color="#ED7D31", linewidth=1.3, linestyle="--", label="AM MAP模拟")
        node_nse = nse(observed[node].to_numpy(float), simulated[node].to_numpy(float))
        nse_text = "近零响应" if np.isnan(node_nse) else f"NSE={node_nse:.3f}"
        ax.set_title(f"监测点 {node}，{nse_text}", fontsize=10)
        ax.grid(alpha=0.2)
    axes[-1, 0].set_xlabel("小时")
    axes[-1, 1].set_xlabel("小时")
    axes[-1, 2].set_xlabel("小时")
    axes[1, 0].set_ylabel("流量增量 m3/s")
    handles, labels = axes[0, 0].get_legend_handles_labels()
    fig.legend(handles, labels, loc="lower center", ncol=2, bbox_to_anchor=(0.5, -0.01))
    fig.suptitle("AM MAP方案在9个监测点上的拟合情况", y=0.995)
    fig.tight_layout(rect=[0, 0.035, 1, 0.965])
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_am_trace(am: pd.DataFrame) -> Path:
    path = FIG_DIR / "06_AM多链搜索过程.png"
    chain_best = am.groupby("chain").agg(min_sse=("sse", "min"), max_mean_nse=("mean_nse", "max"), accept_rate=("accepted", "mean")).reset_index()
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.4))
    for chain, group in am.groupby("chain"):
        sample = group.iloc[:: max(1, len(group) // 180)]
        axes[0].plot(sample["step"], sample["sse"], alpha=0.55, linewidth=0.9, label=f"链{chain}")
    axes[0].set_title("AM采样过程：SSE变化")
    axes[0].set_xlabel("步数")
    axes[0].set_ylabel("SSE，越低越好")
    axes[0].grid(alpha=0.25)
    axes[1].bar(chain_best["chain"].astype(str), chain_best["accept_rate"], color="#70AD47")
    axes[1].set_title("各链接受率")
    axes[1].set_xlabel("链编号")
    axes[1].set_ylabel("接受率")
    axes[1].set_ylim(0, 0.4)
    axes[1].grid(axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_network_png(nodes: pd.DataFrame, links: pd.DataFrame, weights: pd.Series, title: str, filename: str) -> Path:
    path = FIG_DIR / filename
    node_xy = nodes.dropna(subset=["x", "y"]).set_index("node")[["x", "y"]]
    fig, ax = plt.subplots(figsize=(13, 6.8))
    for _, link in links.iterrows():
        f, t = str(link["from_node"]), str(link["to_node"])
        if f in node_xy.index and t in node_xy.index:
            ax.plot([node_xy.loc[f, "x"], node_xy.loc[t, "x"]], [node_xy.loc[f, "y"], node_xy.loc[t, "y"]], color="#CBD5E1", linewidth=0.7, zorder=1)
    ax.scatter(node_xy["x"], node_xy["y"], s=9, color="#B7C9E2", alpha=0.75, zorder=2)
    plan = pd.read_csv(ANALYSIS_DIR / "0520_monitor_candidate_injection_plan.csv", encoding="utf-8-sig").set_index("node")
    for node in MONITOR_NODES:
        if node in node_xy.index:
            ax.scatter(node_xy.loc[node, "x"], node_xy.loc[node, "y"], s=60, color="#2F5597", zorder=4)
            ax.text(node_xy.loc[node, "x"], node_xy.loc[node, "y"], node, fontsize=8, color="#1F2937")
    max_weight = max(0.001, float(weights.max()))
    for node in CANDIDATE_NODES:
        if node in node_xy.index:
            value = float(weights.get(node, 0.0))
            color = "#D9A300" if node not in TRUTH_NODES else "#C00000"
            ax.scatter(node_xy.loc[node, "x"], node_xy.loc[node, "y"], s=80 + 900 * value / max_weight, color=color, alpha=0.65, edgecolor="#111827", linewidth=0.4, zorder=5)
            if value > 0.03 or node in TRUTH_NODES:
                ax.text(node_xy.loc[node, "x"], node_xy.loc[node, "y"], f"{node}\n{value*100:.1f}%", fontsize=8, color="#111827", zorder=6)
    legend_handles = [
        Line2D([0], [0], color="#CBD5E1", lw=1.4, label="管线"),
        Line2D([0], [0], marker="o", color="w", markerfacecolor="#2F5597", markersize=8, label="监测点"),
        Line2D([0], [0], marker="o", color="w", markerfacecolor="#C00000", markeredgecolor="#111827", markersize=9, label="真值注入点"),
        Line2D([0], [0], marker="o", color="w", markerfacecolor="#D9A300", markeredgecolor="#111827", markersize=9, label="候选/代偿点"),
    ]
    ax.legend(handles=legend_handles, loc="lower left", frameon=True, framealpha=0.92, fontsize=9)
    ax.set_aspect("equal", adjustable="box")
    ax.set_title(title)
    ax.axis("off")
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_network_html(nodes: pd.DataFrame, links: pd.DataFrame, shares: pd.DataFrame) -> Path:
    path = OUT_DIR / "0520中参数管网识别热力图.html"
    node_xy = nodes.dropna(subset=["x", "y"]).set_index("node")[["x", "y"]]
    edge_x, edge_y = [], []
    for _, link in links.iterrows():
        f, t = str(link["from_node"]), str(link["to_node"])
        if f in node_xy.index and t in node_xy.index:
            edge_x += [node_xy.loc[f, "x"], node_xy.loc[t, "x"], None]
            edge_y += [node_xy.loc[f, "y"], node_xy.loc[t, "y"], None]
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=edge_x, y=edge_y, mode="lines", line=dict(color="#CBD5E1", width=1), hoverinfo="skip", name="管线"))
    fig.add_trace(go.Scatter(
        x=node_xy["x"], y=node_xy["y"], mode="markers", marker=dict(size=4, color="#AABBD3"),
        text=node_xy.index, hovertemplate="节点 %{text}<extra></extra>", name="普通节点"
    ))
    for label, color, items in [
        ("监测点", "#2563EB", MONITOR_NODES),
        ("真值注入点", "#DC2626", TRUTH_NODES),
        ("候选点", "#F59E0B", CANDIDATE_NODES),
    ]:
        xs, ys, texts = [], [], []
        for node in items:
            if node in node_xy.index:
                xs.append(node_xy.loc[node, "x"])
                ys.append(node_xy.loc[node, "y"])
                texts.append(node)
        fig.add_trace(go.Scatter(
            x=xs, y=ys, mode="markers+text", text=texts, textposition="top center",
            marker=dict(size=10, color=color, line=dict(width=1, color="#111827")),
            hovertemplate=f"{label} %{{text}}<extra></extra>", name=label
        ))
    buttons = []
    solution_names = [("ga_best", "GA最佳"), ("posterior_best_map", "AM MAP"), ("posterior_median_summary", "后验中位数")]
    for sol_name, label in solution_names:
        s = solution_share_series(shares, sol_name)
        xs, ys, texts, sizes, colors = [], [], [], [], []
        max_w = max(0.001, float(s.max()))
        for node, value in s.items():
            if node in node_xy.index:
                xs.append(node_xy.loc[node, "x"])
                ys.append(node_xy.loc[node, "y"])
                texts.append(f"节点 {node}<br>比例 {value*100:.2f}%")
                sizes.append(8 + 42 * float(value) / max_w)
                colors.append(float(value))
        visible = sol_name == "posterior_best_map"
        fig.add_trace(go.Scatter(
            x=xs, y=ys, mode="markers", text=texts, visible=visible,
            marker=dict(size=sizes, color=colors, colorscale="YlOrRd", showscale=visible, colorbar=dict(title="比例")),
            hovertemplate="%{text}<extra>" + label + "</extra>",
            name=label + "热力"
        ))
    base_visible = [True, True, True, True, True, False, True, False]
    buttons = [
        dict(label="GA最佳", method="update", args=[{"visible": [True, True, True, True, True, True, False, False]}]),
        dict(label="AM MAP", method="update", args=[{"visible": [True, True, True, True, True, False, True, False]}]),
        dict(label="后验中位数", method="update", args=[{"visible": [True, True, True, True, True, False, False, True]}]),
    ]
    fig.update_layout(
        title="0520中参数管网识别热力图：蓝色为监测点，红色为真值注入点，黄色为候选点",
        updatemenus=[dict(buttons=buttons, direction="right", x=0.02, y=1.08)],
        template="plotly_white",
        font=dict(family="Microsoft YaHei, SimHei, Arial", size=13),
        height=780,
        legend=dict(orientation="h", yanchor="bottom", y=1.02),
        margin=dict(l=20, r=20, t=90, b=20),
        xaxis=dict(visible=False),
        yaxis=dict(visible=False, scaleanchor="x", scaleratio=1),
    )
    fig.write_html(path, include_plotlyjs="cdn", full_html=True)
    return path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, df: pd.DataFrame, widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        set_cell_shading(hdr[i], "D9EAF7")
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, float):
                text = f"{val:.4f}"
            else:
                text = str(val)
            cells[i].text = text
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(8.5)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(47, 84, 150)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)


def add_picture(doc: Document, path: Path, caption: str, width: float = 6.3) -> None:
    doc.add_picture(str(path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(89, 89, 89)


def build_report(inputs: dict[str, object], figures: dict[str, Path], html_path: Path, detection: pd.DataFrame, map_metrics: pd.DataFrame) -> Path:
    summary = inputs["summary"]
    data_summary = inputs["data_summary"]
    scores = inputs["scores"]
    shares = inputs["shares"]
    am = inputs["am"]
    ga_last = inputs["ga_last"]
    doc_path = OUT_DIR / "0520中参数版本结果分析汇报.docx"

    doc = Document()
    style_document(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("0520管网入流溯源中参数版本结果分析汇报")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)
    subtitle = doc.add_paragraph("基于9个监测点、20个候选点、5个真值注入点的GA-AM识别结果")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_heading(doc, "一、结论摘要", 1)
    add_paragraph(doc, "关键点1：本次中参数版本在曲线拟合上已经做到较好水平。AM MAP 的 mean NSE 为 0.936212，SSE 为 0.002591，是本轮曲线误差最小的结果。")
    add_paragraph(doc, "关键点2：本次没有稳定做到“5个真值点全部精确恢复”。按5%比例阈值统计，GA最佳解命中4个真值点，漏掉178；AM MAP命中3个真值点，漏掉103和304。")
    add_paragraph(doc, "关键点3：本轮主要问题是代偿。GA最佳解把一部分水量分给91和308；AM MAP把一部分水量分给64、308和241。也就是说，模型可以较好复现监测曲线，但部分相邻或同路径候选点仍会相互替代。")
    add_paragraph(doc, "关键点4：当前GA不是硬性找5个点。它是在20个候选点上分配总入流比例；因此结果需要同时看“拟合是否做到”和“定位是否被代偿”。")

    overview_df = pd.DataFrame([
        ["曲线拟合", "做到", "AM MAP mean NSE=0.936212，SSE=0.002591"],
        ["真值回放", "做到", "mean NSE=1.000000，说明注水链路和评分链路闭合"],
        ["全部真值点精确恢复", "未完全做到", "GA命中4/5，AM MAP命中3/5"],
        ["代偿控制", "仍存在", "主要代偿点为91、308、64、241"],
    ], columns=["判断项", "结论", "依据"])
    add_table(doc, overview_df, widths=[4.0, 3.0, 9.0])

    add_heading(doc, "二、本次实验条件", 1)
    add_paragraph(doc, f"本次模型采用36小时模拟，时间分辨率为10分钟，共{data_summary['rows']}个输出时刻。注入持续时间为0到24小时，注入过程使用INP文件中的[TIMESERIES]统一写入，没有使用generated_inflow运行时注入。")
    add_paragraph(doc, f"总入流量按“雨天排口总流出量减去旱天排口总流出量”计算：雨天排口积分为{data_summary['event_outfall_total_volume_m3']:.2f} m3，旱天排口积分为{data_summary['baseline_outfall_total_volume_m3']:.2f} m3，差值为{data_summary['outfall_delta_total_volume_m3']:.2f} m3。五个真值注入点为{ '、'.join(TRUTH_NODES) }，每个点折算体积约{list(data_summary['truth_scaled_volumes_m3'].values())[0]:.2f} m3。")
    add_paragraph(doc, f"候选点为20个，监测点为9个。真值回放检查mean NSE={data_summary['truth_replay_mean_nse']:.6f}，SSE={data_summary['truth_replay_sse']:.2e}，说明当前数据构造和[TIMESERIES]注入路径是一致的。")

    config = summary["config"]
    config_df = pd.DataFrame([
        ["GA种群数", config["ga_population_count"]],
        ["GA每种群个体数", config["ga_population_size"]],
        ["GA代数", config["ga_generations"]],
        ["GA目标", "9个监测点mean NSE最大"],
        ["AM链数", config["am_chain_count"]],
        ["AM每链步数", config["am_samples_per_chain"]],
        ["AM warmup", config["am_warmup"]],
        ["AM目标", "SSE最小，对应似然最大"],
        ["后验样本数", len(am) - config["am_chain_count"] * config["am_warmup"]],
    ], columns=["项目", "本次设置"])
    add_table(doc, config_df, widths=[5.0, 10.0])

    add_heading(doc, "三、总体结果", 1)
    add_paragraph(doc, "本次结果的核心特征是：AM MAP曲线拟合最好，但不是所有真值点都被单独分离出来；GA最佳解命中的真值点更多，但也带有明显代偿点。因此汇报时需要分清“曲线拟合程度”和“位置识别准确性”。")
    score_df = scores[scores["solution"].isin(["ga_best", "posterior_best_map"])].copy()
    score_df["mean_nse"] = score_df["mean_nse"].map(lambda x: f"{x:.6f}")
    score_df["sse"] = score_df["sse"].map(lambda x: f"{x:.6f}")
    score_df = score_df.rename(columns={"solution": "结果", "mean_nse": "mean NSE", "sse": "SSE", "primary_role": "含义"})
    add_table(doc, score_df, widths=[3.2, 2.5, 2.5, 8.0])
    add_picture(doc, figures["score_bar"], "图1  核心结果评分对比：GA看mean NSE，AM MAP看SSE。")

    add_heading(doc, "四、GA阶段结果", 1)
    add_paragraph(doc, f"GA阶段一共评估{len(inputs['ga_all'])}个候选解，最后一代保留{len(ga_last)}个唯一解。GA的目标是让9个监测点的平均NSE尽量高，本轮GA最佳mean NSE={summary['ga_last_score_stats']['max_mean_nse']:.6f}。")
    add_paragraph(doc, "需要特别说明：当前GA只在初始化时会生成1到5个激活点的稀疏解，后续交叉和变异后没有强制只保留5个点。因此本轮GA实际是在20个候选点上搜索比例组合，而不是严格找5个点。")
    ga_best = solution_share_series(shares, "ga_best")
    add_paragraph(doc, f"GA最佳解的主要比例为：10号{pct(ga_best['10'])}、42号{pct(ga_best['42'])}、91号{pct(ga_best['91'])}、308号{pct(ga_best['308'])}、103号{pct(ga_best['103'])}、304号{pct(ga_best['304'])}。其中10、42、103、304是真值点，91和308属于代偿点，178未被GA最佳解识别出来。")
    add_picture(doc, figures["ga_convergence"], "图2  GA逐代收敛过程：前几代提升明显，后期提升变慢。")

    add_heading(doc, "五、AM阶段结果", 1)
    add_paragraph(doc, f"AM阶段共8条链，每条链700步，前180步作为warmup，最终用于后验统计的样本为{len(inputs['am']) - 8 * 180}个。各链接受率在{min(summary['am_accept_rate_by_chain'].values()):.3f}到{max(summary['am_accept_rate_by_chain'].values()):.3f}之间，说明链条不是完全卡死，也不是无约束乱跳。")
    add_paragraph(doc, "AM MAP是后验样本中SSE最小的一组解。由于本轮AM没有把GA先验放入接受率，AM MAP本质上就是“在AM搜索过程中曲线误差最小的解”。")
    map_s = solution_share_series(shares, "posterior_best_map")
    add_paragraph(doc, f"AM MAP主要比例为：42号{pct(map_s['42'])}、178号{pct(map_s['178'])}、64号{pct(map_s['64'])}、10号{pct(map_s['10'])}、308号{pct(map_s['308'])}、241号{pct(map_s['241'])}。其中42、178、10是真值点，64、308、241为主要代偿点，103和304没有被单独稳定识别。")
    add_picture(doc, figures["am_trace"], "图3  AM多链搜索过程：SSE越低代表AM目标越好。")
    add_picture(doc, figures["monitor_fit"], "图4  AM MAP在9个监测点上的拟合曲线。")

    add_heading(doc, "六、代偿与定位分析", 1)
    add_paragraph(doc, "本轮代偿不是简单的程序错误现象，而是由管网结构、监测点响应和算法目标共同造成。只要两个候选点到监测点的水力响应很接近，算法就可能把一部分水量分给相邻或同一路径上的点，最后仍然得到较高拟合分数。")
    det = detection.copy()
    det["真值点总比例"] = det["真值点总比例"].map(lambda x: f"{x*100:.2f}%")
    det["非真值点总比例"] = det["非真值点总比例"].map(lambda x: f"{x*100:.2f}%")
    add_table(doc, det, widths=[3.0, 1.8, 2.2, 4.0, 1.8, 4.0, 4.0, 2.4, 2.4])
    add_paragraph(doc, "从5%阈值看，GA最佳解命中4个真值点，但保留了91和308两个高比例代偿点；AM MAP命中3个真值点，曲线误差更小，但把部分比例转移到64、308、241。也就是说，AM把曲线拟合进一步做优了，但定位解释性没有同步变强。")
    add_picture(doc, figures["share_compare"], "图5  真值与三类识别结果的比例对比。")
    add_picture(doc, figures["network_ga"], "图6  GA最佳解在管网结构上的热力分布：主要代偿为91和308。")
    add_picture(doc, figures["network_map"], "图7  AM MAP在管网结构上的热力分布：主要代偿为64、308和241。")

    add_heading(doc, "七、监测点响应分析", 1)
    add_paragraph(doc, "AM MAP的总体mean NSE达到0.936212，但各监测点贡献并不均匀。大多数监测点拟合较好，223号监测点是本轮较弱位置；189号响应接近零，对定位贡献有限。")
    metric_table = map_metrics.copy()
    metric_table = metric_table.sort_values("NSE", na_position="last")
    metric_table["NSE"] = metric_table["NSE"].map(lambda x: "近零响应" if pd.isna(x) else f"{x:.6f}")
    for col in ["SSE", "观测峰值", "模拟峰值", "峰值差"]:
        metric_table[col] = metric_table[col].map(lambda x: f"{x:.6g}")
    add_table(doc, metric_table, widths=[2.0, 2.3, 2.5, 2.5, 2.5, 2.5])

    add_heading(doc, "八、本次结论", 1)
    conclusions = [
        "本次中参数版本已经能够在曲线层面取得较高拟合，AM MAP的mean NSE为0.936212，SSE为0.002591。",
        "本次没有稳定恢复全部5个真值点。GA最佳解按5%阈值命中4个真值点，AM MAP按5%阈值命中3个真值点。",
        "代偿主要发生在同一路径或相邻响应区，例如304附近转移到64，103附近转移到308和241。这说明当前监测信息可以支持较好的水量拟合，但对部分近邻来源的区分仍不足。",
        "当前GA不是硬性找5个点，而是20个候选点比例搜索。若后续要让结果更接近“固定数量溯源”，需要在算法中明确加入点数约束或稀疏化规则；本报告只陈述本次结果，不把该项作为本轮已完成修改。",
    ]
    for item in conclusions:
        doc.add_paragraph(item, style=None)

    add_paragraph(doc, f"配套交互式热力图已生成：{html_path}")
    doc.save(doc_path)
    return doc_path


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    setup_matplotlib()
    inputs = load_inputs()
    scores = inputs["scores"]
    shares = inputs["shares"]
    observed = inputs["observed_delta"]
    map_delta = inputs["map_delta"]
    nodes = inputs["nodes"]
    links = inputs["links"]

    detection = compute_detection_table(shares)
    detection.to_csv(OUT_DIR / "0520中参数识别命中与代偿统计.csv", index=False, encoding="utf-8-sig")
    map_metrics = monitor_metrics(observed, map_delta)
    map_metrics.to_csv(OUT_DIR / "0520_AM_MAP监测点拟合指标.csv", index=False, encoding="utf-8-sig")

    figures = {
        "score_bar": save_score_bar(scores),
        "ga_convergence": save_ga_convergence(inputs["ga_history"]),
        "share_compare": save_share_comparison(shares),
        "active_count": save_active_count_distribution(inputs["ga_all"], inputs["ga_last"], inputs["initial_ppd"]),
        "monitor_fit": save_monitor_fit(observed, map_delta),
        "am_trace": save_am_trace(inputs["am"]),
    }
    figures["network_ga"] = save_network_png(nodes, links, solution_share_series(shares, "ga_best"), "GA最佳解在管网中的比例热力分布", "07_GA最佳解管网热力图.png")
    figures["network_map"] = save_network_png(nodes, links, solution_share_series(shares, "posterior_best_map"), "AM MAP解在管网中的比例热力分布", "08_AM_MAP管网热力图.png")
    figures["network_median"] = save_network_png(nodes, links, solution_share_series(shares, "posterior_median_summary"), "后验中位数在管网中的比例热力分布", "09_后验中位数管网热力图.png")
    html_path = save_network_html(nodes, links, shares)

    doc_path = build_report(inputs, figures, html_path, detection, map_metrics)
    manifest = {
        "report_docx": str(doc_path),
        "network_html": str(html_path),
        "figures": {key: str(path) for key, path in figures.items()},
        "tables": {
            "detection": str(OUT_DIR / "0520中参数识别命中与代偿统计.csv"),
            "monitor_metrics": str(OUT_DIR / "0520_AM_MAP监测点拟合指标.csv"),
        },
    }
    (OUT_DIR / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
