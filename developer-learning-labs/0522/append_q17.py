# -*- coding: utf-8 -*-
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading_styled(doc, text, level):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.keep_with_next = True
    
    # Set colors and fonts
    run = heading.runs[0]
    run.font.name = 'Microsoft YaHei'
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(31, 78, 121) # Dark Blue
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(46, 117, 182) # Medium Blue
    else:
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return heading

def add_paragraph_styled(doc, text="", space_after=6, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    p.alignment = align
    if text:
        run = p.add_run(text)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(10.5)
        run.font.bold = bold
        run.font.italic = italic
    return p

def add_bullet_styled(doc, label, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    # Bold part
    if label:
        run_l = p.add_run(label)
        run_l.font.name = 'Microsoft YaHei'
        run_l.font.size = Pt(10.5)
        run_l.font.bold = True
        
    run_t = p.add_run(text)
    run_t.font.name = 'Microsoft YaHei'
    run_t.font.size = Pt(10.5)
    return p

def main():
    filename = "物理综合题解析报告.docx"
    doc = docx.Document(filename)
    
    # Add a page break to start Q17 on a new page (standard formatting)
    doc.add_page_break()
    
    # ------------------ Q17 ------------------
    add_heading_styled(doc, "第 17 题：电路功率随电流变化曲线分析", level=1)
    
    add_heading_styled(doc, "题目分析：", level=2)
    add_paragraph_styled(doc, "本题探究的是定值电阻 R0 的发热功率 P0、滑动变阻器 R 消耗的电功率 P_R 和电源总功率 P_总 随电路电流 I 变化的关系。")
    add_bullet_styled(doc, "电路结构：", "电源 U、定值电阻 R0、滑动变阻器 R 串联接入电路。")
    add_bullet_styled(doc, "图像坐标：", "横轴表示电流 I，纵轴表示功率 P。")
    
    add_heading_styled(doc, "详细解答步骤：", level=2)
    add_bullet_styled(doc, "(1) 判定电源总功率图线：", "电源总功率计算公式为 P_总 = E × I，其中电源电动势 E 是定值。因此 P_总 随 I 的变化图线是一条过原点的倾斜直线。对应图乙中的图线 a。")
    add_paragraph_styled(doc, "根据图线 a 上的点 (3 A, 9 W) 可求得电源电动势为：E = P_总 / I = 9 W / 3 A = 3 V，这与图甲中电源标注的 3V 一致。", space_after=6)
    
    add_bullet_styled(doc, "(2) 确定定值电阻 R0：", "定值电阻 R0 消耗的电功率公式为 P0 = I² × R0。由于 R0 是定值，P0 是关于 I 的二次函数（开口向上的抛物线），对应图乙中的图线 b。")
    add_paragraph_styled(doc, "由图可知，当电路中电流最大为 I_max = 3 A（此时滑动变阻器接入电阻为 0）时，定值电阻消耗的功率为 9 W，由此可求得：R0 = P0 / I² = 9 W / (3 A)² = 1 Ω。", space_after=6)
    
    add_bullet_styled(doc, "(3) 计算滑动变阻器 R 消耗的最大功率：", "由串联电路特点，滑动变阻器消耗的电功率为 P_R = I² × R = I(E - I R0) = 3I - I²。其随电流 I 的变化图线为开口向下的抛物线，对应图乙中的图线 c。")
    
    add_heading_styled(doc, "求解最大功率的两种方法：", level=2)
    add_bullet_styled(doc, "方法一：配方法", "P_R = 3I - I² = -(I - 1.5)² + 2.25。显然，当电路中电流为 I = 1.5 A 时，滑动变阻器 R 消耗的电功率最大，最大值为 2.25 W。")
    add_bullet_styled(doc, "方法二：等效内阻法（物理常用推论）", "在串联电路中，当滑动变阻器的阻值 R 等于电路其余部分的等效电阻（此处即 R0）时，滑动变阻器消耗的功率最大。即 R = R0 = 1 Ω。此时电路总电阻 R_总 = R0 + R = 2 Ω，电路电流 I = E / R_总 = 3 V / 2 Ω = 1.5 A，滑动变阻器的最大功率为 P_Rmax = I² × R = 1.5² × 1 = 2.25 W。")
    
    # Save back to the same file
    doc.save(filename)
    print("Successfully appended Q17 to the end of the Word document!")

if __name__ == "__main__":
    main()
