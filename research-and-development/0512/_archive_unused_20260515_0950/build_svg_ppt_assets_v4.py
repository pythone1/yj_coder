from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "output"
SVG_DIR = OUT / "assets_svg_v4"
DOC_DIR = OUT / "doc"
SVG_DIR.mkdir(parents=True, exist_ok=True)
DOC_DIR.mkdir(parents=True, exist_ok=True)


W, H = 1600, 900


def svg_wrap(title, body):
    return f'''<svg xmlns="http://www.w3.org/2000/svg" width="{W}" height="{H}" viewBox="0 0 {W} {H}">
  <defs>
    <linearGradient id="water" x1="0" x2="1" y1="0" y2="1">
      <stop offset="0%" stop-color="#1e88e5"/>
      <stop offset="100%" stop-color="#00acc1"/>
    </linearGradient>
    <linearGradient id="pipe" x1="0" x2="1">
      <stop offset="0%" stop-color="#6aa6d9"/>
      <stop offset="50%" stop-color="#d9eefc"/>
      <stop offset="100%" stop-color="#5b94c8"/>
    </linearGradient>
    <linearGradient id="danger" x1="0" x2="1">
      <stop offset="0%" stop-color="#ff7043"/>
      <stop offset="100%" stop-color="#d84315"/>
    </linearGradient>
    <filter id="shadow" x="-20%" y="-20%" width="140%" height="140%">
      <feDropShadow dx="0" dy="10" stdDeviation="12" flood-color="#0f172a" flood-opacity="0.14"/>
    </filter>
    <filter id="soft" x="-20%" y="-20%" width="140%" height="140%">
      <feGaussianBlur stdDeviation="8"/>
    </filter>
    <style>
      .title {{ font-family: 'Microsoft YaHei','SimHei',sans-serif; font-size: 48px; font-weight: 700; fill:#172033; }}
      .subtitle {{ font-family: 'Microsoft YaHei','SimHei',sans-serif; font-size: 24px; fill:#5d6b82; }}
      .h {{ font-family: 'Microsoft YaHei','SimHei',sans-serif; font-size: 28px; font-weight: 700; fill:#172033; }}
      .t {{ font-family: 'Microsoft YaHei','SimHei',sans-serif; font-size: 20px; fill:#172033; }}
      .s {{ font-family: 'Microsoft YaHei','SimHei',sans-serif; font-size: 17px; fill:#5d6b82; }}
      .white {{ fill:#fff; }}
      .muted {{ fill:#5d6b82; }}
      .blue {{ fill:#2563a8; }}
      .cyan {{ fill:#0891a3; }}
      .green {{ fill:#2f8f5f; }}
      .orange {{ fill:#df8b21; }}
      .red {{ fill:#c94b41; }}
      .line {{ stroke:#d8e2ef; stroke-width:2; }}
      .pipe {{ stroke:url(#pipe); stroke-width:24; stroke-linecap:round; fill:none; }}
      .pipe2 {{ stroke:#2d6aa3; stroke-width:6; stroke-linecap:round; fill:none; opacity:.75; }}
      .dash {{ stroke-dasharray:8 10; }}
    </style>
  </defs>
  <rect width="{W}" height="{H}" fill="#ffffff"/>
  {body}
</svg>'''


def card(x, y, w, h, color, title, lines):
    text = f'<text x="{x+28}" y="{y+48}" class="h" fill="{color}">{title}</text>'
    yy = y + (84 if h < 170 else 92)
    line_class = "s" if h < 170 else "t"
    step = 30 if h < 170 else 40
    for line in lines:
        text += f'<text x="{x+32}" y="{yy}" class="{line_class}">• {line}</text>'
        yy += step
    return f'''<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="28" fill="#fff" stroke="{color}" stroke-width="4" filter="url(#shadow)"/>
    <circle cx="{x+48}" cy="{y+48}" r="16" fill="{color}" opacity=".15"/>
    {text}'''


def save(name, content):
    p = SVG_DIR / name
    p.write_text(content, encoding="utf-8")
    return p


def cover():
    body = '''
  <rect x="0" y="0" width="1600" height="900" fill="#f7fbff"/>
  <circle cx="1280" cy="190" r="260" fill="#dff5fb" opacity=".75"/>
  <circle cx="1380" cy="720" r="210" fill="#e8f3ff" opacity=".9"/>
  <text x="80" y="120" class="title">AI供水管网漏损检测</text>
  <text x="82" y="168" class="subtitle">DMA分区 · 动态基线 · 异常识别 · 工单闭环</text>
  <g transform="translate(70,250)">
    <path d="M60 350 C220 250,330 430,480 330 S760 220,910 345 S1170 485,1380 320" class="pipe"/>
    <path d="M60 350 C220 250,330 430,480 330 S760 220,910 345 S1170 485,1380 320" class="pipe2"/>
    <path d="M540 275 L620 230 L700 275 L700 420 L540 420 Z" fill="#fff" stroke="#2563a8" stroke-width="4" filter="url(#shadow)"/>
    <path d="M585 420 L585 330 L655 330 L655 420" fill="#e7f2ff" stroke="#2563a8" stroke-width="3"/>
    <rect x="575" y="290" width="35" height="28" rx="4" fill="#00acc1" opacity=".9"/>
    <rect x="632" y="290" width="35" height="28" rx="4" fill="#00acc1" opacity=".9"/>
    <circle cx="878" cy="352" r="52" fill="url(#danger)" filter="url(#shadow)"/>
    <path d="M865 323 C850 352,832 370,832 393 C832 421,853 444,880 444 C908 444,929 421,929 393 C929 370,907 350,895 323 C890 310,872 310,865 323 Z" fill="#fff" opacity=".92" transform="scale(.55) translate(720,300)"/>
    <g transform="translate(1040,95)" filter="url(#shadow)">
      <rect x="0" y="0" width="360" height="230" rx="28" fill="#172033"/>
      <path d="M50 155 L105 120 L160 135 L215 75 L295 105" fill="none" stroke="#00d4ff" stroke-width="8" stroke-linecap="round"/>
      <circle cx="215" cy="75" r="12" fill="#ff7043"/>
      <text x="34" y="54" fill="#fff" font-family="Microsoft YaHei,SimHei" font-size="24" font-weight="700">智能预警</text>
      <text x="34" y="198" fill="#b8c7dd" font-family="Microsoft YaHei,SimHei" font-size="17">残差持续扩大 · 建议复核</text>
    </g>
    <g transform="translate(260,100)">
      <rect x="0" y="0" width="180" height="86" rx="22" fill="#fff" stroke="#0891a3" stroke-width="4"/>
      <circle cx="45" cy="43" r="24" fill="#0891a3"/>
      <text x="86" y="52" class="t">流量计</text>
    </g>
    <g transform="translate(1120,360)">
      <rect x="0" y="0" width="190" height="86" rx="22" fill="#fff" stroke="#df8b21" stroke-width="4"/>
      <circle cx="45" cy="43" r="24" fill="#df8b21"/>
      <text x="86" y="52" class="t">压力点</text>
    </g>
  </g>
  <text x="82" y="820" class="subtitle">适用于对外教学演讲封面或章节开场</text>
'''
    return svg_wrap("cover", body)


def dma_scene():
    body = '''
  <text x="70" y="92" class="title">DMA现场感知与漏损定位场景</text>
  <text x="72" y="135" class="subtitle">入口计量判断是否异常，多点压力帮助收敛候选管段</text>
  <rect x="430" y="185" width="760" height="545" rx="62" fill="#edf6ff" stroke="#2563a8" stroke-width="5"/>
  <text x="740" y="235" class="h blue">DMA-A</text>
  <path d="M540 390 L730 330 L950 388 L1060 520 L850 605 L620 570 Z" class="pipe"/>
  <path d="M540 390 L850 605 M730 330 L850 605 M950 388 L850 605" class="pipe"/>
  <path d="M540 390 L730 330 L950 388 L1060 520 L850 605 L620 570 Z" class="pipe2"/>
  <path d="M540 390 L850 605 M730 330 L850 605 M950 388 L850 605" class="pipe2"/>
  <g>
    <circle cx="540" cy="390" r="23" fill="#fff" stroke="#2563a8" stroke-width="4"/><text x="528" y="398" class="s">J1</text>
    <circle cx="730" cy="330" r="23" fill="#fff" stroke="#2563a8" stroke-width="4"/><text x="718" y="338" class="s">J2</text>
    <circle cx="950" cy="388" r="23" fill="#fff" stroke="#2563a8" stroke-width="4"/><text x="938" y="396" class="s">J3</text>
    <circle cx="1060" cy="520" r="23" fill="#fff" stroke="#2563a8" stroke-width="4"/><text x="1048" y="528" class="s">J4</text>
    <circle cx="850" cy="605" r="23" fill="#fff" stroke="#2563a8" stroke-width="4"/><text x="838" y="613" class="s">J5</text>
    <circle cx="620" cy="570" r="23" fill="#fff" stroke="#2563a8" stroke-width="4"/><text x="608" y="578" class="s">J6</text>
  </g>
  <rect x="300" y="350" width="125" height="78" rx="18" fill="#0891a3"/><text x="329" y="382" class="white t">入口</text><text x="319" y="410" class="white t">流量计</text>
  <path d="M425 390 L520 390" stroke="#0891a3" stroke-width="8" stroke-linecap="round"/>
  <rect x="675" y="262" width="128" height="58" rx="16" fill="#df8b21"/><text x="697" y="299" class="white t">P1高点</text>
  <rect x="957" y="565" width="135" height="58" rx="16" fill="#df8b21"/><text x="974" y="602" class="white t">P2末梢</text>
  <rect x="1078" y="450" width="132" height="58" rx="16" fill="#2f8f5f"/><text x="1106" y="487" class="white t">边界阀</text>
  <circle cx="820" cy="500" r="58" fill="url(#danger)" filter="url(#shadow)"/><text x="792" y="493" class="white t">疑似</text><text x="792" y="523" class="white t">漏点</text>
  <g>
    ''' + card(80, 230, 270, 130, "#2563a8", "边界封闭", ["阀门状态可核验", "无隐性连通"]) + '''
    ''' + card(80, 420, 270, 130, "#0891a3", "计量闭合", ["入口/出口可计量", "低流量段可靠"]) + '''
    ''' + card(80, 610, 270, 130, "#df8b21", "压力可观测", ["高点/低点/末梢", "至少有代表点"]) + '''
    ''' + card(1260, 290, 270, 130, "#2f8f5f", "台账一致", ["GIS/SCADA/工单", "编码能关联"]) + '''
    ''' + card(1260, 520, 270, 130, "#c94b41", "反馈闭环", ["核查结果回填", "形成训练标签"]) + '''
  </g>
'''
    return svg_wrap("dma", body)


def mnf():
    body = '''
  <text x="70" y="92" class="title">MNF夜间最小流量拆分</text>
  <text x="72" y="135" class="subtitle">MNF是发现异常的入口，不等同于漏损量本身</text>
  <rect x="115" y="225" width="860" height="470" rx="12" fill="#fff" stroke="#d8e2ef" stroke-width="2"/>
  <g stroke="#d8e2ef" stroke-width="1">
    <line x1="115" y1="320" x2="975" y2="320"/><line x1="115" y1="415" x2="975" y2="415"/><line x1="115" y1="510" x2="975" y2="510"/><line x1="115" y1="605" x2="975" y2="605"/>
  </g>
  <rect x="210" y="225" width="95" height="470" fill="#ffd166" opacity=".35"/>
  <text x="218" y="270" class="s" fill="#df8b21">2:00-4:00</text><text x="220" y="296" class="s" fill="#df8b21">MNF窗口</text>
  <path d="M115 565 C160 590,190 555,210 580 C240 610,275 610,305 565 C360 530,420 510,505 470 C625 420,750 430,850 475 C910 502,945 530,975 555" fill="none" stroke="#2563a8" stroke-width="6"/>
  <line x1="115" y1="575" x2="975" y2="575" stroke="#c94b41" stroke-width="4"/>
  <text x="995" y="582" class="t red">MNF</text>
  <g transform="translate(1060,230)">
    <rect x="0" y="0" width="420" height="74" rx="18" fill="#e8f3ff" stroke="#2563a8" stroke-width="3"/><text x="150" y="46" class="t">入口MNF</text>
    <text x="-45" y="47" class="h muted">=</text>
    <rect x="0" y="92" width="420" height="74" rx="18" fill="#eefbf7" stroke="#0891a3" stroke-width="3"/><text x="130" y="138" class="t">合法夜间用水</text>
    <text x="-45" y="139" class="h muted">+</text>
    <rect x="0" y="184" width="420" height="74" rx="18" fill="#fff3e2" stroke="#df8b21" stroke-width="3"/><text x="150" y="230" class="t">突发漏损</text>
    <text x="-45" y="231" class="h muted">+</text>
    <rect x="0" y="276" width="420" height="74" rx="18" fill="#ffefee" stroke="#c94b41" stroke-width="3"/><text x="150" y="322" class="t">背景漏损</text>
    <text x="-45" y="323" class="h muted">±</text>
    <rect x="0" y="368" width="420" height="74" rx="18" fill="#f4f7fb" stroke="#5d6b82" stroke-width="3"/><text x="130" y="414" class="t">仪表/边界误差</text>
  </g>
  <rect x="230" y="770" width="1140" height="70" rx="22" fill="#f5f8fc" stroke="#d8e2ef" stroke-width="2"/>
  <text x="360" y="813" class="t">PPT讲法：先用MNF找到异常窗口，再用AI判断异常是否持续、是否符合漏损模式</text>
'''
    return svg_wrap("mnf", body)


def lstm():
    body = '''
  <text x="70" y="92" class="title">LSTM/GRU残差预警</text>
  <text x="72" y="135" class="subtitle">模型先学习正常运行曲线，再用实际值与预测值的偏差触发复核</text>
  <rect x="110" y="220" width="980" height="500" rx="12" fill="#fff" stroke="#d8e2ef" stroke-width="2"/>
  <g stroke="#d8e2ef" stroke-width="1">
    <line x1="110" y1="320" x2="1090" y2="320"/><line x1="110" y1="420" x2="1090" y2="420"/><line x1="110" y1="520" x2="1090" y2="520"/><line x1="110" y1="620" x2="1090" y2="620"/>
  </g>
  <rect x="735" y="220" width="355" height="500" fill="#ffefee" opacity=".55"/>
  <path d="M110 545 C170 492,230 575,290 520 S420 455,500 500 S640 540,735 455 S885 425,1090 380" fill="none" stroke="#2f8f5f" stroke-width="6"/>
  <path d="M110 550 C170 500,230 580,290 522 S420 458,500 505 S640 548,735 462 S835 365,910 340 S1020 320,1090 305" fill="none" stroke="#c94b41" stroke-width="6"/>
  <text x="770" y="265" class="t red">持续正残差区</text>
  <g transform="translate(1165,250)">
    <rect x="0" y="0" width="310" height="150" rx="24" fill="#fff" stroke="#d8e2ef" stroke-width="2" filter="url(#shadow)"/>
    <line x1="35" y1="50" x2="110" y2="50" stroke="#2f8f5f" stroke-width="6"/><text x="130" y="58" class="t">预测正常值</text>
    <line x1="35" y1="100" x2="110" y2="100" stroke="#c94b41" stroke-width="6"/><text x="130" y="108" class="t">实际监测值</text>
  </g>
  ''' + card(1165, 455, 310, 95, "#2563a8", "输入", ["历史流量/压力/日历"]) + '''
  ''' + card(1165, 585, 310, 95, "#2f8f5f", "输出", ["预测值/预测区间"]) + '''
  <rect x="300" y="770" width="1000" height="70" rx="22" fill="#f5f8fc" stroke="#d8e2ef" stroke-width="2"/>
  <text x="430" y="813" class="t">对外口径：LSTM负责动态基线，不单独承诺直接定位漏点</text>
'''
    return svg_wrap("lstm", body)


def algo_cards():
    body = '''
  <text x="70" y="92" class="title">漏损检测算法素材页</text>
  <text x="72" y="135" class="subtitle">按业务任务组合模型，而不是按“先进程度”堆算法</text>
  ''' + card(90, 215, 330, 220, "#2563a8", "时序预测", ["LSTM / GRU", "CNN-LSTM", "输出动态基线"]) + '''
  ''' + card(455, 215, 330, 220, "#0891a3", "少标签异常", ["孤立森林", "DBSCAN", "自编码器"]) + '''
  ''' + card(820, 215, 330, 220, "#2f8f5f", "管段风险", ["随机森林", "GBDT / HGB", "输出风险排序"]) + '''
  ''' + card(1185, 215, 330, 220, "#df8b21", "定位收敛", ["水力仿真", "GNN", "候选管段TopN"]) + '''
  <g transform="translate(140,515)">
    <rect x="0" y="0" width="1320" height="235" rx="30" fill="#172033" filter="url(#shadow)"/>
    <text x="48" y="60" fill="#fff" font-family="Microsoft YaHei,SimHei" font-size="28" font-weight="700">模型组合建议</text>
    <text x="50" y="115" fill="#cfe1f6" font-family="Microsoft YaHei,SimHei" font-size="22">第一阶段：MNF + LSTM残差 + 孤立森林，跑通异常发现和人工复核</text>
    <text x="50" y="165" fill="#cfe1f6" font-family="Microsoft YaHei,SimHei" font-size="22">第二阶段：RF/GBDT + 工单标签，形成管段风险排序</text>
    <text x="50" y="215" fill="#cfe1f6" font-family="Microsoft YaHei,SimHei" font-size="22">第三阶段：水力仿真 + GNN，在重点DMA做定位增强</text>
  </g>
'''
    return svg_wrap("algo", body)


def fusion_scene():
    body = '''
  <text x="70" y="92" class="title">水力机理模型 + AI模型双驱动</text>
  <text x="72" y="135" class="subtitle">机理模型给物理约束，AI模型给异常识别和候选排序</text>
  <g transform="translate(105,245)">
    <rect x="0" y="0" width="410" height="420" rx="32" fill="#fff" stroke="#2563a8" stroke-width="4" filter="url(#shadow)"/>
    <text x="36" y="58" class="h blue">水力机理模型</text>
    <path d="M70 170 C130 110,230 115,305 170 S345 290,250 320 S90 285,70 170" fill="#edf6ff" stroke="#2563a8" stroke-width="4"/>
    <text x="86" y="220" class="t">拓扑 · 管径 · 高程</text>
    <text x="86" y="258" class="t">泵阀工况 · 压力响应</text>
  </g>
  <g transform="translate(595,245)">
    <rect x="0" y="0" width="410" height="420" rx="32" fill="#fff" stroke="#2f8f5f" stroke-width="4" filter="url(#shadow)"/>
    <text x="36" y="58" class="h green">融合层</text>
    <rect x="75" y="135" width="260" height="70" rx="20" fill="#e9f8f1" stroke="#2f8f5f" stroke-width="3"/><text x="150" y="180" class="t">仿真样本库</text>
    <rect x="75" y="230" width="260" height="70" rx="20" fill="#e9f8f1" stroke="#2f8f5f" stroke-width="3"/><text x="160" y="275" class="t">特征工程</text>
    <rect x="75" y="325" width="260" height="70" rx="20" fill="#e9f8f1" stroke="#2f8f5f" stroke-width="3"/><text x="155" y="370" class="t">置信校验</text>
  </g>
  <g transform="translate(1085,245)">
    <rect x="0" y="0" width="410" height="420" rx="32" fill="#fff" stroke="#df8b21" stroke-width="4" filter="url(#shadow)"/>
    <text x="36" y="58" class="h orange">AI数据模型</text>
    <path d="M82 320 L145 225 L205 285 L270 150 L338 205" fill="none" stroke="#df8b21" stroke-width="9" stroke-linecap="round"/>
    <circle cx="270" cy="150" r="14" fill="#c94b41"/>
    <text x="84" y="105" class="t">LSTM残差 · 异常检测</text>
    <text x="84" y="365" class="t">RF/GBDT · GNN排序</text>
  </g>
  <rect x="360" y="740" width="880" height="72" rx="24" fill="#f5f8fc" stroke="#d8e2ef" stroke-width="2"/>
  <text x="510" y="785" class="t">输出：异常等级 + 候选管段TopN + 现场复核建议</text>
'''
    return svg_wrap("fusion", body)


def data_governance():
    body = '''
  <text x="70" y="92" class="title">AI漏损检测数据底座</text>
  <text x="72" y="135" class="subtitle">数据表能关联，模型结果才能进入派单和复盘</text>
  ''' + card(95, 220, 280, 175, "#2563a8", "设备表", ["设备编码", "量程精度", "在线状态"]) + '''
  ''' + card(425, 220, 280, 175, "#2f8f5f", "管网表", ["管段编号", "阀门状态", "DMA归属"]) + '''
  ''' + card(755, 220, 280, 175, "#0891a3", "时序表", ["采样时间", "流量压力", "缺失标记"]) + '''
  ''' + card(1085, 220, 280, 175, "#df8b21", "工单表", ["报警编号", "漏点位置", "处置结果"]) + '''
  <g transform="translate(240,495)">
    <rect x="0" y="0" width="1120" height="220" rx="34" fill="#172033" filter="url(#shadow)"/>
    <text x="52" y="62" fill="#fff" font-family="Microsoft YaHei,SimHei" font-size="30" font-weight="700">统一编码是关键</text>
    <text x="55" y="118" fill="#cfe1f6" font-family="Microsoft YaHei,SimHei" font-size="22">同一个流量计、压力点、管段、阀门，在SCADA、GIS、工单系统中必须能对应</text>
    <text x="55" y="170" fill="#cfe1f6" font-family="Microsoft YaHei,SimHei" font-size="22">否则AI只能发现“异常”，无法形成“可执行定位”和“可复盘标签”</text>
  </g>
'''
    return svg_wrap("data", body)


def roadmap():
    body = '''
  <text x="70" y="92" class="title">从试点到规模化实施路线</text>
  <text x="72" y="135" class="subtitle">先跑通一个可复盘DMA，再复制到全网</text>
  <g transform="translate(100,310)">
    <circle cx="60" cy="60" r="54" fill="#2563a8"/><text x="34" y="70" class="white h">01</text><text x="23" y="150" class="h blue">诊断</text><text x="-15" y="190" class="s">边界/仪表/GIS/工单</text>
    <circle cx="310" cy="60" r="54" fill="#0891a3"/><text x="284" y="70" class="white h">02</text><text x="273" y="150" class="h cyan">治理</text><text x="238" y="190" class="s">编码/对齐/补缺</text>
    <circle cx="560" cy="60" r="54" fill="#2f8f5f"/><text x="534" y="70" class="white h">03</text><text x="523" y="150" class="h green">建模</text><text x="488" y="190" class="s">MNF/LSTM/异常</text>
    <circle cx="810" cy="60" r="54" fill="#df8b21"/><text x="784" y="70" class="white h">04</text><text x="773" y="150" class="h orange">试点</text><text x="738" y="190" class="s">灰度报警/人工复核</text>
    <circle cx="1060" cy="60" r="54" fill="#c94b41"/><text x="1034" y="70" class="white h">05</text><text x="1023" y="150" class="h red">闭环</text><text x="988" y="190" class="s">派单/回填/再训练</text>
    <circle cx="1310" cy="60" r="54" fill="#695dbe"/><text x="1284" y="70" class="white h">06</text><text x="1273" y="150" class="h" fill="#695dbe">推广</text><text x="1238" y="190" class="s">跨DMA迁移/运营</text>
    <path d="M120 60 H250 M370 60 H500 M620 60 H750 M870 60 H1000 M1120 60 H1250" stroke="#d8e2ef" stroke-width="6" stroke-linecap="round"/>
  </g>
  <rect x="280" y="735" width="1040" height="76" rx="24" fill="#f5f8fc" stroke="#d8e2ef" stroke-width="2"/>
  <text x="430" y="782" class="t">试点选择：边界清晰、仪表稳定、历史事件较多、运维配合度高</text>
'''
    return svg_wrap("roadmap", body)


ASSETS = [
    ("01_封面_管网漏损AI.svg", cover(), "封面或章节开场"),
    ("02_DMA现场感知与漏损定位.svg", dma_scene(), "讲DMA现场架构、压力点和候选漏点"),
    ("03_MNF夜间最小流量拆分.svg", mnf(), "讲MNF与漏损量的关系"),
    ("04_LSTM残差预警.svg", lstm(), "讲LSTM/GRU动态基线"),
    ("05_算法组合素材页.svg", algo_cards(), "讲算法扩充和选型"),
    ("06_水力机理_AI双驱动.svg", fusion_scene(), "讲机理模型和AI融合"),
    ("07_数据治理底座.svg", data_governance(), "讲数据治理和标签体系"),
    ("08_试点到规模化路线.svg", roadmap(), "讲实施路径"),
]


def make_doc(paths):
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    p = doc.add_paragraph()
    r = p.add_run("AI供水管网漏损检测 SVG素材包 v4")
    r.font.name = "黑体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体"); r.font.size = Pt(20); r.font.bold = True
    p = doc.add_paragraph()
    r = p.add_run("说明：本包为真正SVG矢量文件，可直接插入PowerPoint。不是位图转SVG。")
    r.font.color.rgb = RGBColor(192,0,0)
    for i, (name, path, usage) in enumerate(paths, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i}. {name}")
        r.font.bold = True
        p = doc.add_paragraph(f"用途：{usage}")
        p = doc.add_paragraph(f"文件：{path}")
    out = DOC_DIR / "AI供水管网漏损检测_SVG素材包说明_v4.docx"
    doc.save(out)
    return out


def main():
    paths = []
    for name, content, usage in ASSETS:
        path = save(name, content)
        paths.append((name, path, usage))
        print(path)
    doc = make_doc(paths)
    print(doc)


if __name__ == "__main__":
    main()
