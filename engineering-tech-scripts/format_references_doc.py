# -*- coding: utf-8 -*-
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


BASE = r"D:\Users\xwechat_files\wxid_4668346683612_4126\msg\file\2026-04"
SRC_NAME = "参考文献.docx"
OUT_NAME = "参考文献_按年份梳理_格式统一版.docx"


def set_run_font(run, size=10.5, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text="", size=10.5, bold=False, align=None, hanging=False):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold)
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(20)
    if hanging:
        fmt.left_indent = Cm(0.74)
        fmt.first_line_indent = Cm(-0.74)
    return p


def clean_entry(entry):
    e = entry.strip()
    e = e.replace("[[M]", "[M]").replace("[s]", "[S]")
    e = e.replace("．", ".").replace("，", ",").replace("：", ":").replace("；", ";")
    e = e.replace("（", "(").replace("）", ")")
    e = re.sub(r"\s+", " ", e).strip()
    e = e.replace(" .", ".")
    return e


def normalize_key(entry):
    e = re.sub(r"\s+", "", entry)
    e = e.replace("．", ".").replace("：", ":").replace("，", ",").replace("；", ";")
    e = e.replace("“", "").replace("”", "").replace('"', "")
    e = e.replace("[[M]", "[M]").replace("[s]", "[S]")
    m = re.match(r"(.+?)\.(.+?)(\[[A-Za-z]\])", e)
    if m:
        return f"{m.group(1)}.{m.group(2)}{m.group(3)}".lower()
    return e.lower()


def type_of(entry):
    m = re.search(r"\[([A-Za-z])\]", entry)
    return m.group(1).upper() if m else ""


def year_of(entry):
    years = re.findall(r"(?:19|20)\d{2}", entry)
    return int(years[-1]) if years else 9999


def is_chinese(entry):
    letters = re.findall(r"[A-Za-z]", entry)
    chinese = re.findall(r"[\u4e00-\u9fff]", entry)
    if chinese and len(chinese) >= len(letters):
        return True
    if re.search(r"[\u4e00-\u9fff]", entry) and re.search(r"北京|上海|济南|郑州|中华书局|人民教育出版社|华东师范大学出版社", entry):
        return True
    return False


def category_group(entry):
    t = type_of(entry)
    zh = is_chinese(entry)
    if zh:
        if t == "M":
            return ("中文", "专著类")
        if t == "J":
            return ("中文", "期刊文章类")
        if t == "D":
            return ("中文", "学术论文类")
        if t == "S":
            return ("中文", "课程标准类")
        return ("中文", "期刊文章类")
    else:
        if t == "M":
            return ("外文", "专著类")
        return ("外文", "期刊论文类")


def read_entries():
    path = os.path.join(BASE, SRC_NAME)
    doc = Document(path)
    raw = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t or t == "参考文献":
            continue
        if "学位论文的撰写应本着严谨求实的科学态度" in t:
            break
        if t.startswith("[1] Sergio Manzettiab") or t.startswith("[2]***"):
            continue
        if any(t.startswith(prefix) for prefix in ["引用文献的作者不超过3位", "连续出版物:", "专 译 著:", "论 文 集:", "学位论文:", "专    利:", "技术标准:"]):
            continue
        raw.append(clean_entry(t))
    best = {}
    for e in raw:
        key = normalize_key(e)
        if key not in best or len(e) > len(best[key]):
            best[key] = e
    return list(best.values())


def build_doc():
    entries = read_entries()
    grouped = {
        ("中文", "专著类"): [],
        ("中文", "期刊文章类"): [],
        ("中文", "学术论文类"): [],
        ("中文", "课程标准类"): [],
        ("外文", "专著类"): [],
        ("外文", "期刊论文类"): [],
    }
    for e in entries:
        grouped[category_group(e)].append(e)
    for k in grouped:
        grouped[k] = sorted(grouped[k], key=lambda x: (year_of(x), x))

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    add_para(doc, "参考文献", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "（一）中文文献", size=10.5, bold=True)

    counter = 1
    for title, key in [
        ("1.专著类", ("中文", "专著类")),
        ("2.期刊文章类", ("中文", "期刊文章类")),
        ("3.学术论文类", ("中文", "学术论文类")),
        ("4.课程标准类", ("中文", "课程标准类")),
    ]:
        add_para(doc, title, size=10.5, bold=True)
        for e in grouped[key]:
            add_para(doc, f"[{counter}] {e}", hanging=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            counter += 1

    add_para(doc, "（二）外文文献", size=10.5, bold=True)
    for title, key in [
        ("1.专著类", ("外文", "专著类")),
        ("2.期刊论文类", ("外文", "期刊论文类")),
    ]:
        add_para(doc, title, size=10.5, bold=True)
        for e in grouped[key]:
            add_para(doc, f"[{counter}] {e}", hanging=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            counter += 1

    out_path = os.path.join(BASE, OUT_NAME)
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    print(build_doc())
