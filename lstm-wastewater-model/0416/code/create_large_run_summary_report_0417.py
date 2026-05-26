from __future__ import annotations

import json
import math
import sys
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from config_0416 import (  # noqa: E402
    CANDIDATE_NODES,
    MODEL_1D_INP,
    MONITOR_NODES,
    OUTFALL_NODE,
    DATA_SUMMARY_JSON,
    RESULT_DIR,
    TRUTH_INJECTION_NODES,
)


RUN_DIR = RESULT_DIR / "large_run"
ANALYSIS_DIR = RUN_DIR / "analysis_large_0417"
REPORT_DOCX = ANALYSIS_DIR / "0417大参数版本结果总结与布局优化依据.docx"
REPORT_MD = ANALYSIS_DIR / "0417大参数版本结果总结与布局优化依据.md"
NETWORK_FIG = ANALYSIS_DIR / "fig7_large_network_map_am_map.png"


def read_section(inp_path: Path, section_name: str) -> list[list[str]]:
    rows: list[list[str]] = []
    section = ""
    for raw in inp_path.read_text(encoding="gbk", errors="ignore").splitlines():
        line = raw.split(";", 1)[0].strip()
        if not line:
            continue
        if line.startswith("[") and line.endswith("]"):
            section = line[1:-1].upper()
            continue
        if section == section_name.upper():
            rows.append(line.split())
    return rows


def parse_network() -> tuple[dict[str, tuple[float, float]], list[tuple[str, str, str]], set[str], set[str]]:
    coords: dict[str, tuple[float, float]] = {}
    for row in read_section(MODEL_1D_INP, "COORDINATES"):
        if len(row) >= 3:
            coords[row[0]] = (float(row[1]), float(row[2]))
    conduits: list[tuple[str, str, str]] = []
    for row in read_section(MODEL_1D_INP, "CONDUITS"):
        if len(row) >= 3:
            conduits.append((row[0], row[1], row[2]))
    junctions = {row[0] for row in read_section(MODEL_1D_INP, "JUNCTIONS") if row}
    outfalls = {row[0] for row in read_section(MODEL_1D_INP, "OUTFALLS") if row}
    return coords, conduits, junctions, outfalls


def fmt(value: float, digits: int = 4) -> str:
    if value is None or (isinstance(value, float) and not math.isfinite(value)):
        return "-"
    return f"{value:.{digits}f}"


def pct(value: float, digits: int = 1) -> str:
    return f"{value * 100:.{digits}f}%"


def solution_top(row: pd.Series, limit: int = 6) -> str:
    values = [(node, float(row[node])) for node in CANDIDATE_NODES]
    values.sort(key=lambda kv: kv[1], reverse=True)
    return "；".join(f"{node}={value:.4f}" for node, value in values[:limit] if value > 1e-8)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        set_cell_shading(hdr[i], "DDEFEA")
        if widths:
            hdr[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Cm(widths[i])
    document.add_paragraph()


def add_picture_if_exists(document: Document, path: Path, caption: str, width_cm: float = 15.5) -> None:
    if not path.exists():
        return
    document.add_picture(str(path), width=Cm(width_cm))
    p = document.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(90, 90, 90)


def build_network_figure(solution_shares: pd.DataFrame) -> None:
    coords, conduits, _, outfalls = parse_network()
    shares = solution_shares.loc["posterior_best_map"]
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS"]
    plt.rcParams["axes.unicode_minus"] = False

    fig, ax = plt.subplots(figsize=(11.2, 7.4), dpi=180)
    for _, start, end in conduits:
        if start not in coords or end not in coords:
            continue
        x1, y1 = coords[start]
        x2, y2 = coords[end]
        ax.plot([x1, x2], [y1, y2], color="#7a837d", lw=1.1, alpha=0.45, zorder=1)

    non_special = [
        node for node in coords
        if node not in CANDIDATE_NODES and node not in MONITOR_NODES and node not in outfalls
    ]
    ax.scatter(
        [coords[n][0] for n in non_special],
        [coords[n][1] for n in non_special],
        s=10,
        color="#c9d2cc",
        alpha=0.8,
        label="普通节点",
        zorder=2,
    )

    values = [float(shares[node]) for node in CANDIDATE_NODES]
    sizes = [75 + 2600 * math.sqrt(max(v, 0)) for v in values]
    sc = ax.scatter(
        [coords[n][0] for n in CANDIDATE_NODES],
        [coords[n][1] for n in CANDIDATE_NODES],
        c=values,
        s=sizes,
        cmap="YlOrRd",
        vmin=0,
        vmax=max(max(values), 0.34),
        edgecolors="#26352f",
        linewidths=0.7,
        alpha=0.88,
        label="候选井 AM MAP 份额",
        zorder=4,
    )

    monitor_nodes = [n for n in MONITOR_NODES if n in coords]
    ax.scatter(
        [coords[n][0] for n in monitor_nodes],
        [coords[n][1] for n in monitor_nodes],
        marker="s",
        s=70,
        color="#1d4ed8",
        edgecolors="white",
        linewidths=0.8,
        label="监测点",
        zorder=5,
    )

    truth_nodes = [n for n in TRUTH_INJECTION_NODES if n in coords]
    ax.scatter(
        [coords[n][0] for n in truth_nodes],
        [coords[n][1] for n in truth_nodes],
        facecolors="none",
        edgecolors="#b91c1c",
        linewidths=2.1,
        s=700,
        label="真值注入点",
        zorder=6,
    )

    outfall_nodes = [n for n in outfalls if n in coords]
    ax.scatter(
        [coords[n][0] for n in outfall_nodes],
        [coords[n][1] for n in outfall_nodes],
        marker="D",
        s=92,
        color="#111827",
        label="排口",
        zorder=6,
    )

    label_nodes = set(TRUTH_INJECTION_NODES) | {"J1", "J2", "J21", "J84", "J72", "J86"} | set(MONITOR_NODES)
    for node in sorted(label_nodes):
        if node not in coords:
            continue
        x, y = coords[node]
        ax.text(x + 12, y + 8, node, fontsize=8.2, color="#1f2a24", weight="bold", zorder=7)

    ax.set_title("大参数 AM MAP 识别结果在管网结构上的分布", fontsize=14, weight="bold")
    ax.set_xlabel("模型 X 坐标")
    ax.set_ylabel("模型 Y 坐标")
    ax.grid(True, color="#e4e0d5", linewidth=0.5)
    ax.set_aspect("equal", adjustable="box")
    cbar = fig.colorbar(sc, ax=ax, fraction=0.026, pad=0.02)
    cbar.set_label("注入份额")
    ax.legend(loc="upper right", fontsize=8, frameon=True)
    fig.tight_layout()
    fig.savefig(NETWORK_FIG, bbox_inches="tight")
    plt.close(fig)


def make_markdown(summary: dict, data_summary: dict, analysis: dict, solution_shares: pd.DataFrame) -> str:
    large = analysis["large"]
    config = summary["config"]
    lines = [
        "# 0417 大参数版本结果总结与布局优化依据",
        "",
        "## 一、结论摘要",
        "",
        (
            f"本轮大参数版本在曲线拟合层面已经达到很高水平。AM 后验最优 MAP 的 mean NSE 为 "
            f"{summary['solution_scores']['posterior_best_map']['mean_nse']:.6f}，SSE 为 "
            f"{summary['solution_scores']['posterior_best_map']['sse']:.6g}，明显优于 GA 全局最优。"
        ),
        (
            "定位层面仍存在代偿：J11 基本稳定识别，J48 在 AM MAP 中恢复较好，"
            "J20 仍有一部分被 J2/J1 等邻近节点吸收。"
        ),
        (
            "因此下一轮不建议单纯继续放大参数规模，重点应放在监测点布局优化，"
            "尤其是 J20 邻近支路和 J48 邻近支路的区分性监测。"
        ),
        "",
        "## 二、本轮数据与参数",
        "",
        f"- 基线模型：{data_summary['baseline_inp']}",
        f"- 事件模型：{data_summary['truth_event_inp']}",
        f"- 时间尺度：{data_summary['rows']} 个 5 分钟样本，覆盖约 48 小时。",
        f"- 总入流量：排口事件总量减旱天总量 = {data_summary['outfall_delta_total_volume_m3']:.3f} m3。",
        f"- 真值注入：J20、J48、J11 各 {data_summary['truth_scaled_volumes_m3']['J20']:.3f} m3。",
        f"- 候选井：{len(CANDIDATE_NODES)} 个；监测点：{len(MONITOR_NODES)} 个。",
        f"- GA：{config['ga_population_count']} 个种群，每种群 {config['ga_population_size']} 个体，{config['ga_generations']} 代。",
        f"- AM：{config['am_chain_count']} 条链，每链 {config['am_samples_per_chain']} 步，warmup {config['am_warmup']} 步。",
        "",
        "## 三、阶段结果",
        "",
        (
            f"GA 共生成 {large['ga_rows']} 个样本，最终代 {large['ga_last_rows']} 个样本。"
            f"GA 最优 mean NSE={large['ga_best']:.6f}，最终代中位数={large['ga_last_median']:.6f}。"
        ),
        (
            f"initial PPD 保留 {large['initial_rows']} 个样本，中位 mean NSE={large['initial_median']:.6f}。"
            "这一阶段给 AM 提供起点和协方差信息。"
        ),
        (
            f"AM 共 {large['am_rows']} 个样本，warmup 后 PPD 为 {large['ppd_rows']} 个样本。"
            f"AM 平均接受率={large['am_accept_mean']:.3f}，说明链仍有足够移动能力。"
        ),
        "",
        "## 四、代偿分析",
        "",
        (
            f"GA 中 NSE>=0.99 的样本有 {large['high_ga']['0.99']['count']} 个，"
            f"这些高分样本的真值质量中位数只有 {large['high_ga']['0.99']['truth_mass_median']:.3f}，"
            f"非真值质量中位数为 {large['high_ga']['0.99']['outside_median']:.3f}。"
            "这说明高拟合分数可以由代偿解产生。"
        ),
        (
            "AM MAP 的真值质量提高到 0.903，非真值质量降到 0.097，"
            "但 J20 仍被 J2/J1 分走约 0.093，说明该区域监测信息仍不足。"
        ),
        "",
        "## 五、布局优化依据",
        "",
        "下一轮优化重点不是扩大 GA/AM 参数，而是提升监测点对相邻候选井的区分能力。",
        "重点区域包括 J20-J2-J1-J21 组和 J48-J84-J72-J86 组。",
        "监测点应优先布设在这些候选井下游但汇合前的位置，避免只在远端主干监测导致不同来源响应过于相似。",
    ]
    return "\n".join(lines) + "\n"


def setup_document_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = document.styles
    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_paragraph(document: Document, text: str, bold_prefix: str | None = None) -> None:
    p = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        rest = text[len(bold_prefix):]
        if rest:
            p.add_run(rest)
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(10.5)


def build_docx(summary: dict, data_summary: dict, analysis: dict, solution_shares: pd.DataFrame) -> None:
    large = analysis["large"]
    config = summary["config"]
    document = Document()
    setup_document_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("0417 大参数版本结果总结与布局优化依据")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("用于下一轮监测点布局优化测试")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(90, 90, 90)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    document.add_heading("一、结论摘要", level=1)
    add_paragraph(
        document,
        f"本轮大参数版本在曲线拟合层面已经达到很高水平。AM 后验最优 MAP 的 mean NSE 为 "
        f"{summary['solution_scores']['posterior_best_map']['mean_nse']:.6f}，SSE 为 "
        f"{summary['solution_scores']['posterior_best_map']['sse']:.6g}；GA 全局最优 mean NSE 为 "
        f"{summary['solution_scores']['ga_best']['mean_nse']:.6f}，SSE 为 "
        f"{summary['solution_scores']['ga_best']['sse']:.6g}。"
    )
    add_paragraph(
        document,
        "核心判断：算法已经能把观测曲线拟合到接近真值水平，但高分解并不必然等于真实点位。"
        "当前限制主要来自监测布局对近邻候选井的区分能力不足，而不是单纯参数规模不足。"
    )
    add_paragraph(
        document,
        "定位结果：J11 基本稳定识别，J48 在 AM MAP 中恢复较好；J20 仍被 J2/J1 等临近节点分走部分份额。"
        "下一轮应围绕 J20 邻近支路和 J48 邻近支路增加区分性监测，而不是继续单纯加大 GA/AM 规模。"
    )

    document.add_heading("二、实验数据与参数配置", level=1)
    add_table(
        document,
        ["项目", "本轮设置"],
        [
            ["基线模型", str(data_summary["baseline_inp"])],
            ["事件模型", str(data_summary["truth_event_inp"])],
            ["样本尺度", f"{data_summary['rows']} 个 5 分钟样本，首个样本为 00:05，覆盖约 48 小时"],
            ["总入流量计算", f"事件排口总量 - 旱天排口总量 = {data_summary['outfall_delta_total_volume_m3']:.3f} m3"],
            ["真值注入", f"J20、J48、J11 三点，各 {data_summary['truth_scaled_volumes_m3']['J20']:.3f} m3"],
            ["候选井与监测点", f"{len(CANDIDATE_NODES)} 个候选井；{len(MONITOR_NODES)} 个监测点"],
            ["真值回放", f"mean NSE={large['truth']['mean_nse']:.6f}，SSE={large['truth']['sse']:.3g}"],
        ],
        widths=[4.0, 12.5],
    )
    add_table(
        document,
        ["阶段", "参数规模", "说明"],
        [
            [
                "GA",
                f"{config['ga_population_count']} 个种群 x {config['ga_population_size']} 个体 x {config['ga_generations']} 代",
                f"共 {large['ga_rows']} 个样本；按 mean NSE 排序；种群内近似去重，保留多样性。",
            ],
            [
                "initial PPD",
                f"保留 {large['initial_rows']} 个样本",
                "作为 AM 起点池和协方差估计基础；权重为 rank-based，避免低质量样本污染。",
            ],
            [
                "AM",
                f"{config['am_chain_count']} 条链 x {config['am_samples_per_chain']} 步，warmup {config['am_warmup']} 步",
                f"AM prior 不进入接受率；平均接受率 {large['am_accept_mean']:.3f}。",
            ],
            [
                "后验验证",
                f"{summary['posterior_validation_sample_count']} 个后验预测样本",
                f"90% 覆盖率均值 {summary['posterior_coverage_mean']:.3f}。",
            ],
        ],
        widths=[3.2, 5.2, 8.0],
    )

    document.add_heading("三、识别结果分析", level=1)
    rows = []
    for item in large["solutions"]:
        row = solution_shares.loc[item["solution"]]
        rows.append(
            [
                item["solution"],
                fmt(item["mean_nse"], 6),
                f"{item['sse']:.6g}",
                pct(item["truth_mass"], 1),
                pct(item["outside_mass"], 1),
                solution_top(row),
            ]
        )
    add_table(
        document,
        ["方案", "mean NSE", "SSE", "真值质量", "非真值质量", "前六份额节点"],
        rows,
        widths=[3.6, 2.3, 2.4, 2.3, 2.3, 5.6],
    )
    add_paragraph(
        document,
        "GA 阶段：GA 全局最优曲线拟合已经较好，但份额中 J48 为 0，J84 和 J21 承担了大量代偿。"
        "这说明 GA 能快速找到高拟合区域，但在监测响应相似的区域会优先接受代偿解。"
    )
    add_paragraph(
        document,
        "AM 阶段：AM MAP 将真值质量提高到 90.3%，J11 和 J48 基本恢复，J20 仍偏低。"
        "AM 的作用是从 GA 高分区域继续局部探索，显著压低 SSE，但无法完全消除由监测布局造成的结构性不可辨识。"
    )
    add_paragraph(
        document,
        "后验中位数：后验中位数汇总解的 NSE 明显低于 AM MAP，且份额分散。"
        "原因是后验存在多峰代偿，逐节点取中位数会把多个模式混合到一起，因此只能作为不确定性描述，不宜作为唯一最终方案。"
    )

    add_picture_if_exists(document, ANALYSIS_DIR / "fig1_large_ga_convergence.png", "图 1  GA 收敛过程：高分解在前几代快速出现，后续提升有限。")
    add_picture_if_exists(document, ANALYSIS_DIR / "fig2_large_solution_shares.png", "图 2  关键解份额对比：AM MAP 明显优于 GA best 和后验中位数。")
    add_picture_if_exists(document, NETWORK_FIG, "图 3  AM MAP 在管网结构上的空间分布：J20 附近仍有 J2/J1 代偿。")

    document.add_heading("四、代偿现象与原因判断", level=1)
    add_table(
        document,
        ["GA 高分阈值", "样本数", "真值质量中位数", "非真值质量中位数", "说明"],
        [
            [
                "NSE >= 0.95",
                str(large["high_ga"]["0.95"]["count"]),
                fmt(large["high_ga"]["0.95"]["truth_mass_median"], 3),
                fmt(large["high_ga"]["0.95"]["outside_median"], 3),
                "大量高分样本仍包含明显非真值质量。",
            ],
            [
                "NSE >= 0.98",
                str(large["high_ga"]["0.98"]["count"]),
                fmt(large["high_ga"]["0.98"]["truth_mass_median"], 3),
                fmt(large["high_ga"]["0.98"]["outside_median"], 3),
                "拟合继续升高，但非真值质量仍占多数。",
            ],
            [
                "NSE >= 0.99",
                str(large["high_ga"]["0.99"]["count"]),
                fmt(large["high_ga"]["0.99"]["truth_mass_median"], 3),
                fmt(large["high_ga"]["0.99"]["outside_median"], 3),
                "高分并不能保证点位正确。",
            ],
            [
                "NSE >= 0.995",
                str(large["high_ga"]["0.995"]["count"]),
                fmt(large["high_ga"]["0.995"]["truth_mass_median"], 3),
                fmt(large["high_ga"]["0.995"]["outside_median"], 3),
                "极高分样本仍可能是代偿解。",
            ],
        ],
        widths=[3.0, 2.0, 3.0, 3.0, 5.2],
    )
    add_paragraph(
        document,
        "代偿原因 1：候选井之间存在水力响应相似性。若两个候选井位于同一支路或在监测点之前已经汇合，"
        "它们在下游监测点形成的流量增量曲线会非常接近，算法只能看到相似曲线，无法仅凭远端监测点稳定区分来源。"
    )
    add_paragraph(
        document,
        "代偿原因 2：当前总入流波形和总量固定。算法搜索的是 20 个候选井的份额分配，"
        "当监测点对局部位置不敏感时，把水量放在相邻候选井也能得到接近相同的监测曲线。"
    )
    add_paragraph(
        document,
        "代偿原因 3：主干远端监测点对总过程拟合很敏感，但对局部来源位置不一定敏感。"
        "因此 NSE 可以很高，但定位仍可能偏移。当前 J20 周边的 J2/J1/J21，以及 J48 周边的 J84/J72/J86 是主要代偿组。"
    )
    add_picture_if_exists(document, ANALYSIS_DIR / "fig3_large_posterior_weights.png", "图 4  后验权重：J11 稳定，J20/J48 周边仍呈多峰不确定性。")
    add_picture_if_exists(document, ANALYSIS_DIR / "fig4_large_per_monitor_nse.png", "图 5  各监测点 NSE：AM MAP 各点拟合都很高，但高拟合不能直接等价于唯一定位。")
    add_picture_if_exists(document, ANALYSIS_DIR / "fig5_large_coverage.png", "图 6  后验预测覆盖率：部分监测点覆盖率偏低，说明不确定性表达仍不均衡。")

    document.add_heading("五、监测布局优化依据", level=1)
    add_paragraph(
        document,
        "本轮结果给出的布局优化方向比较明确：下一轮要提高监测点对局部支路的区分能力，"
        "而不是继续依赖更大的参数规模。大参数 AM 已能把 MAP 推到接近真值，但仍不能完全消除 J20 附近代偿。"
    )
    add_table(
        document,
        ["重点区域", "本轮表现", "下一轮布局目标"],
        [
            [
                "J20 - J2 - J1 - J21",
                "AM MAP 中 J20=0.2435，J2+J1=0.0932，说明 J20 来源仍被上游或近邻节点分走。",
                "增加能区分 J20 与 J2/J1/J21 的下游但未完全汇合前监测点。",
            ],
            [
                "J48 - J84 - J72 - J86",
                "中参数版本 J48 被 J84 明显代偿；大参数 AM MAP 已恢复 J48=0.3293，但后验仍显示 J84 不确定性较高。",
                "保留或增强该区域关键监测，重点验证 J48 与 J84/J72/J86 的响应差异。",
            ],
            [
                "J11 周边",
                "J11 在 GA、AM、后验中都相对稳定，说明当前监测对该点识别较充分。",
                "不应把新增监测资源过多放在 J11 附近，除非用于冗余校核。",
            ],
            [
                "远端主干监测",
                "远端监测能提高整体拟合，但容易把不同来源看成相似总过程。",
                "保留少量总量约束点即可，避免过多远端点稀释局部敏感性。",
            ],
        ],
        widths=[3.3, 6.2, 6.5],
    )
    add_paragraph(
        document,
        "下一轮推荐原则：候选井数量仍保持 20 个时，监测点应优先覆盖“候选井下游、支路汇合前、响应差异最大”的位置。"
        "每个主要代偿组至少需要一个能分辨组内节点的局部监测点，同时保留少量主干点用于总量和整体过程约束。"
    )
    add_paragraph(
        document,
        "本轮可视化网页已经更新，可用于人工审查空间布局："
        f"{ANALYSIS_DIR / 'large_run_heatmap_dashboard.html'}"
    )
    add_picture_if_exists(document, ANALYSIS_DIR / "fig6_large_final_curves.png", "图 7  最终曲线对比：AM MAP 的过程曲线与观测增量高度一致。")

    document.add_heading("六、下一轮测试建议", level=1)
    add_paragraph(
        document,
        "建议下一轮只改监测布局，暂不改候选井数量和真值注入方案。这样可以把变量控制在“监测点位是否提升可辨识性”上，"
        "避免参数、候选集和数据同时变化导致结果无法归因。"
    )
    add_paragraph(
        document,
        "优先目标是验证两个问题：第一，J20 是否还能被 J2/J1/J21 高分代偿；第二，J48 是否还能被 J84/J72/J86 高分代偿。"
        "如果新布局下高 NSE 解的真值质量明显提高，说明当前问题主要来自监测布设；如果仍不提高，再考虑候选井分组、正则化或分段排查策略。"
    )

    document.save(REPORT_DOCX)


def main() -> None:
    ANALYSIS_DIR.mkdir(parents=True, exist_ok=True)
    summary = json.loads((RUN_DIR / "0417_summary.json").read_text(encoding="utf-8"))
    data_summary = json.loads(DATA_SUMMARY_JSON.read_text(encoding="utf-8"))
    analysis = json.loads((ANALYSIS_DIR / "large_analysis_summary.json").read_text(encoding="utf-8"))
    solution_shares = pd.read_csv(RUN_DIR / "0417_solution_shares.csv").set_index("solution")

    build_network_figure(solution_shares)
    REPORT_MD.write_text(make_markdown(summary, data_summary, analysis, solution_shares), encoding="utf-8")
    build_docx(summary, data_summary, analysis, solution_shares)
    print(REPORT_DOCX)
    print(REPORT_MD)
    print(NETWORK_FIG)


if __name__ == "__main__":
    main()
