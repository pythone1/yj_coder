# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "output" / "doc"
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "AI模型在供水管网DMA系统漏损检测中的应用_杨佳蓝色模型部分_汇报规范版.docx"

RED = RGBColor(192, 0, 0)
BLACK = RGBColor(0, 0, 0)
BLUE_FILL = "DDEBF7"


def set_font(run, *, color=RED, bold=False, size=10.5, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def setup_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    for style_name in ["Normal", "List Bullet"]:
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.color.rgb = BLACK


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI模型在供水管网DMA系统漏损检测中的应用")
    set_font(r, color=BLACK, bold=True, size=20, font="黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("杨佳蓝色模型部分专项扩写稿（汇报规范版，红色为补充写入内容）")
    set_font(r, color=RED, size=11, font="宋体")
    doc.add_paragraph()


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    size = 15 if level == 1 else 13 if level == 2 else 11.5
    for run in p.runs:
        set_font(run, color=BLACK, bold=True, size=size, font="黑体")


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, color=RED, size=10.5, font="宋体")


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r, color=RED, size=10.5, font="宋体")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, BLUE_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_font(r, color=BLACK, bold=True, size=9, font="黑体")

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(value)
            set_font(r, color=RED, size=8.5, font="宋体")
    doc.add_paragraph()


def build():
    doc = Document()
    setup_doc(doc)
    add_title(doc)

    add_para(
        doc,
        "本稿聚焦原提纲中由杨佳负责的模型相关内容，围绕供水管网DMA漏损检测的AI模型发展、技术选型、应用场景、建设路径和汇报表达进行扩充。文本采用正式汇报语体，可直接作为后续教学演讲稿和PPT内容底稿使用。",
    )

    add_heading(doc, "1.3 供水管网AI模型应用的发展（详细）", 1)
    add_heading(doc, "1.3.1 从经验阈值向动态基线演进", 2)
    add_para(
        doc,
        "供水管网漏损检测的智能化应用通常经历经验阈值、统计分析、机器学习、深度学习和智能协同等阶段。早期DMA漏损管控主要依靠固定阈值规则，例如入口流量超过设定上限、夜间最小流量高于经验值、压力低于控制值时触发报警。该类方法部署简单、逻辑清晰，适用于监测基础薄弱或业务启动阶段。",
    )
    add_para(
        doc,
        "随着DMA分区管理精细化和在线监测数据持续积累，固定阈值方法的适配能力逐步受到限制。不同DMA在用户结构、地形高差、压力制度、夜间用水习惯和管龄结构方面存在明显差异，统一阈值难以覆盖多类型分区。节假日、气温变化、二次供水补水、商业夜间用水等正常扰动，也会影响报警稳定性。",
    )
    add_para(
        doc,
        "AI模型的应用价值集中体现在分区动态基线构建。模型可基于每个DMA的历史流量、压力、日期、节假日、天气和用户结构，学习该分区在不同时间条件下的正常运行状态。当实际曲线持续偏离动态基线，并与漏损特征相吻合时，系统可形成复核建议、风险等级或派单依据。",
    )

    add_heading(doc, "1.3.2 从单点报警向多源联动识别演进", 2)
    add_para(
        doc,
        "传统报警多依赖单点流量或单点压力，容易受到局部设备波动和短时用水扰动影响。实际漏损事件通常表现为多源数据的组合变化，例如入口流量持续抬升、夜间底流增加、末梢压力下降、压力波动加剧，以及历史高风险管段附近重复出现异常。",
    )
    add_para(
        doc,
        "AI模型可将流量、压力、资产、工单、巡检和外部环境数据进行联合建模，提升异常识别的稳定性和业务解释能力。建议将模型输入分为五类：运行时序数据、时间上下文数据、空间资产数据、历史事件数据和业务状态数据。多源特征融合能够推动系统由单一超限判断转向综合风险判断。",
    )

    add_heading(doc, "1.3.3 从单一算法向模型组合演进", 2)
    add_para(
        doc,
        "供水管网漏损检测涉及预测、识别、定位、排序、解释和闭环等多类任务，单一算法难以覆盖完整业务链条。较为稳妥的技术路线，是根据不同任务构建模型组合。LSTM、GRU等时序模型用于预测正常运行曲线；孤立森林、DBSCAN、自编码器用于少标签异常识别；随机森林、GBDT等模型用于管段风险排序；水力仿真与机器学习融合用于候选管段收敛；知识图谱和大模型智能体用于结果解释、报告生成和业务协同。",
    )
    add_para(
        doc,
        "在该技术体系中，专业检测模型负责底层判断，数据平台负责跨系统数据组织，智能体负责汇总解释、报告生成和派单辅助。不同层级之间形成协同关系，共同支撑DMA漏损检测从数据监测向闭环治理升级。",
    )

    add_heading(doc, "1.3.4 供水漏损AI模型技术演进路线", 2)
    add_table(
        doc,
        ["阶段", "主要方法", "应用特点", "适用边界"],
        [
            ["经验阈值阶段", "固定上下限、MNF经验阈值、压力低限报警", "部署简单、逻辑透明、便于快速启用", "适用于基础监测阶段，误报率受分区差异影响较大"],
            ["统计分析阶段", "移动平均、同比环比、控制图、季节分解", "能够刻画历史偏离和周期变化", "对复杂非线性和多源联动识别能力有限"],
            ["机器学习阶段", "孤立森林、DBSCAN、随机森林、GBDT", "适合少标签异常识别和结构化风险排序", "依赖特征质量、样本覆盖和标签口径"],
            ["深度学习阶段", "LSTM、GRU、CNN-LSTM、自编码器、GNN", "适合高频时序、多变量耦合和拓扑关系学习", "需要较好的数据基础和持续验证机制"],
            ["智能协同阶段", "知识图谱、大模型智能体、工具调用", "适合解释、问答、报告和派单辅助", "需要保留专业模型、人工确认和权限控制"],
        ],
    )

    add_heading(doc, "2.1 智慧水务场景下的AI技术核心认知", 1)
    add_heading(doc, "2.1.1 AI模型的业务定位", 2)
    add_para(
        doc,
        "供水管网漏损检测具有较强的工程属性，AI模型的定位应聚焦辅助识别、辅助排序和辅助决策。模型通过分析高频监测数据和历史业务数据，提前发现异常、压缩排查范围、提升巡检效率，并为后续复盘提供数据基础。",
    )
    add_para(
        doc,
        "实际应用中，模型输出应面向业务处置，包含异常时段、异常证据、疑似原因、候选管段、复核方式和优先级等内容。相比单一异常分数，具备证据链和解释信息的模型结果更便于调度、巡检和管理人员采纳。",
    )

    add_heading(doc, "2.1.2 DMA作为AI建模基本单元", 2)
    add_para(
        doc,
        "DMA是AI模型建设的基本业务单元。不同DMA在用水结构、压力制度、管龄材质、地形高差和漏损历史方面存在差异，模型应充分体现分区特征。实际建设中可采用分区模型与全局模型结合的方式：分区模型用于学习本DMA的动态运行基线，全局模型用于识别跨DMA共性风险因素。",
    )
    add_para(
        doc,
        "居民区、商业区、工业区、学校医院片区和老旧管网片区的用水规律差异较大。模型特征和阈值策略应根据分区类型进行差异化设计，避免使用单一规则覆盖全部区域。",
    )

    add_heading(doc, "2.1.3 场景驱动的模型选型原则", 2)
    add_para(
        doc,
        "AI模型选型应先明确业务问题，再确定算法路线。以提前预警为目标时，应优先选择动态基线和异常检测模型；以缩小排查范围为目标时，应结合压力响应、水力模型和候选管段排序；以年度改造计划为目标时，应侧重管段风险评估；以压力优化为目标时，应引入需水预测和优化模型。",
    )
    add_para(
        doc,
        "模型选型还应结合数据基础。边界不清、计量不准、压力点不足、工单标签缺失时，应优先开展数据治理和基础模型建设。复杂模型应建立在可靠数据、稳定边界和可回填工单基础上。",
    )

    add_heading(doc, "2.1.4 数据质量与模型可信度", 2)
    add_para(
        doc,
        "数据质量直接决定模型效果上限。常见问题包括流量与压力时间戳不一致、设备编码在SCADA与GIS中不统一、工单记录缺少管段坐标、阀门状态长期未更新、DMA边界存在隐性连通等。上述问题会影响模型对异常原因和异常位置的判断。",
    )
    add_para(
        doc,
        "模型建设前应完成统一编码、时序对齐、缺失值处理、异常毛刺剔除、工单结构化和标签口径统一等工作。数据治理是模型可信运行的前置条件，也是后续教学汇报中需要重点强调的工程基础。",
    )

    add_heading(doc, "2.2 适配管网漏损检测的核心AI技术及业务价值", 1)
    add_heading(doc, "2.2.1 时序预测模型：LSTM、GRU与CNN-LSTM", 2)
    add_para(
        doc,
        "LSTM适合处理具有长期依赖关系的时间序列，可用于学习DMA流量和压力的日周期、周周期、季节性和短期扰动。GRU结构相对简化，参数量较少，适合算力有限或快速迭代场景。CNN-LSTM可通过一维卷积提取多传感器局部特征，再通过循环网络处理时间依赖，适用于多压力点与入口流量联合建模。",
    )
    add_para(
        doc,
        "在DMA漏损检测中，时序预测模型主要用于建立正常运行曲线。模型基于过去一段时间的流量、压力、天气、日期和节假日特征，输出未来短时窗口的正常预测值和预测区间。当实际值持续偏离预测区间，且偏离特征与漏损模式一致时，可形成疑似漏损预警信号。",
    )
    add_para(
        doc,
        "该类模型的业务价值在于提升预警灵敏度和分区适配能力。与固定阈值相比，动态预测模型能够更好适应不同DMA的运行规律，减少由分区差异造成的误报。",
    )

    add_heading(doc, "2.2.2 无监督异常检测：孤立森林、DBSCAN与自编码器", 2)
    add_para(
        doc,
        "供水漏损标签通常存在数量不足、时间不准和位置不精确等问题。无监督异常检测模型可在标签不足条件下识别候选异常，适合作为早期模型建设的重要组成部分。",
    )
    add_para(
        doc,
        "孤立森林适合对多维运行特征进行快速异常筛查；DBSCAN适合识别流量、压力特征空间中的低密度离群点和异常簇；自编码器适合学习正常状态下多变量之间的组合关系，并通过重构误差识别异常。",
    )
    add_para(
        doc,
        "无监督模型输出应作为人工复核清单使用。复核结果需要回填为真实漏损、合法用水、设备故障、边界异常、施工扰动等类别，为后续监督学习和模型迭代提供标签基础。",
    )

    add_heading(doc, "2.2.3 监督学习模型：随机森林、GBDT与HGB", 2)
    add_para(
        doc,
        "当资产台账和历史工单较为完整时，可采用监督学习模型开展管段风险排序和候选区域识别。随机森林、GBDT、HGB等树模型适合处理管龄、材质、口径、压力等级、历史维修、道路等级、投诉频次等结构化变量。",
    )
    add_para(
        doc,
        "监督学习模型可输出管段风险分、风险等级和主要影响因素，为检漏计划、巡检优先级和管网改造排序提供支撑。该类模型具有较好的业务解释空间，便于管理人员理解风险来源。",
    )

    add_heading(doc, "2.2.4 机理模型与AI模型融合", 2)
    add_para(
        doc,
        "水力模型与AI模型融合，是提升漏点定位能力的重要方向。典型做法包括将水力模型输出的压力、流量、压力敏感性和供水路径作为AI特征，或利用水力模型模拟不同漏点位置和漏量下的压力响应，构建仿真样本库。",
    )
    add_para(
        doc,
        "融合模型可在真实漏损样本不足时提供辅助训练数据，并将AI识别结果限定在水力学合理范围内。该方法的前提是水力模型经过校准，拓扑、阀门状态、粗糙系数和需水模式具备基本可信度。",
    )

    add_heading(doc, "2.2.5 图神经网络与知识图谱", 2)
    add_para(
        doc,
        "供水管网具有天然图结构，节点包括水源、泵站、阀门、压力点、用户和水池，边对应管段。图神经网络可将拓扑邻接关系和压力传播关系纳入模型，适用于多点压力数据较充分、拓扑质量较高的DMA。",
    )
    add_para(
        doc,
        "知识图谱可连接DMA、管段、设备、工单、维修、投诉、材质、管龄等对象，支撑异常原因解释、历史案例检索和智能问答。该技术更适合与业务平台结合，用于辅助分析和知识沉淀。",
    )

    add_heading(doc, "2.2.6 算法选型建议", 2)
    add_table(
        doc,
        ["业务任务", "推荐模型", "输入数据", "输出结果", "应用要点"],
        [
            ["动态基线预测", "LSTM、GRU、CNN-LSTM、移动基线", "入口流量、压力点、天气、日历", "预测值、预测区间、残差", "按DMA训练，采用时间切分验证"],
            ["少标签异常发现", "孤立森林、DBSCAN、自编码器", "MNF特征、残差、多维时序", "异常分数、异常时段", "作为复核清单，不直接确认漏损"],
            ["管段风险排序", "随机森林、GBDT、HGB", "管龄、材质、维修、压力、道路、投诉", "风险分、优先级", "依赖结构化工单和资产台账"],
            ["漏点定位收敛", "水力仿真+机器学习、GNN", "拓扑、多点压力、仿真样本", "候选片区、候选管段TopN", "依赖水力模型和压力点质量"],
            ["压力优化调度", "遗传算法、贝叶斯优化、强化学习", "泵阀状态、压力约束、需水预测", "调压策略、泵阀建议", "需满足供水安全和服务压力约束"],
            ["业务解释协同", "知识图谱、大模型智能体", "模型输出、GIS、工单、知识库", "问答、报告、派单建议", "保留人工确认和权限控制"],
        ],
    )

    add_heading(doc, "3. 核心应用：模型相关场景补充", 1)
    add_heading(doc, "3.1 事前防控：基于AI的DMA异常预警", 2)
    add_para(
        doc,
        "AI预警体系以动态基线和异常识别为核心。系统可为每个DMA建立日内基线、周周期基线和季节性基线，持续比较实际值与预测值之间的残差。当残差持续扩大，并在夜间低用水窗口或MNF时段表现明显，同时伴随压力点异常响应时，可提高漏损风险等级。",
    )
    add_para(
        doc,
        "预警结果应包含异常时间、持续时长、残差幅度、压力响应、相似历史事件、疑似原因和建议动作。该类输出能够帮助调度人员区分真实漏损、合法用水、设备故障和边界异常。",
    )

    add_heading(doc, "3.2 事中处置：漏点候选区域收敛", 2)
    add_para(
        doc,
        "事中处置阶段的模型目标，是将排查范围由整个DMA收敛至若干候选片区或候选管段。入口流量异常用于判断区域风险，多点压力响应和水力模型用于识别异常可能发生的位置范围。",
    )
    add_para(
        doc,
        "建议模型输出候选管段TopN清单，每个候选项包括管段编号、空间位置、置信度、主要证据、历史维修次数、影响用户数和建议复核方式。该输出形式更符合工程处置实际，也便于现场人员执行和反馈。",
    )

    add_heading(doc, "3.3 事后复盘：模型持续学习", 2)
    add_para(
        doc,
        "事后复盘是模型持续优化的关键环节。每一次报警、核查、维修和误报都应形成结构化记录，包括DMA编号、报警时间、异常特征、派单时间、现场结果、漏点位置、漏损类型、维修完成时间、估算漏量、误报原因和模型版本。",
    )
    add_para(
        doc,
        "复盘结果可用于更新特征、调整阈值、修正标签和优化模型。若某类误报频繁出现，例如二次供水补水、商业夜间用水或流量计漂移，应将其纳入特征或规则体系，降低后续误报。",
    )

    add_heading(doc, "3.4 调度类模型应用：压力优化与需水预测", 2)
    add_para(
        doc,
        "压力优化是模型应用的重要拓展方向。漏损流量通常与压力存在相关关系，高压会增加背景漏损和爆管风险。AI模型可结合需水预测、压力监测和水力模型，辅助形成分时压力控制、泵站启停和阀门调节建议。",
    )
    add_para(
        doc,
        "压力优化必须满足最不利点服务压力、消防保障、二次供水补水和用户体验等约束。模型应用目标应表述为在保障供水安全前提下，寻找降漏、节能和稳压之间的平衡。",
    )

    add_heading(doc, "3.5 DMA规划类模型应用：分区与监测点优化", 2)
    add_para(
        doc,
        "AI模型可辅助DMA分区方案和监测点布设优化。分区方案可根据拓扑结构、地形高差、用户数量、压力制度、入口数量、边界阀数量和施工成本进行多目标评价。",
    )
    add_para(
        doc,
        "监测点优化可通过水力仿真不同漏点情景，比较不同压力点组合对定位效果的贡献。该方法有助于在有限预算下提升监测点布设效率。",
    )

    add_heading(doc, "4. 模型建设与落地路径", 1)
    add_heading(doc, "4.1 数据治理：模型上线前的基础", 2)
    add_para(
        doc,
        "模型建设前应完成数据治理，至少形成设备表、管网表、时序表、工单表和标签表。设备表记录流量计、压力计、RTU和边缘网关；管网表记录管段、阀门、DMA归属和资产属性；时序表记录采样时间、流量压力、缺失标记和清洗版本；工单表记录报警、核查、维修和结果；标签表记录真实漏损、误报原因和模型版本。",
    )

    add_heading(doc, "4.2 模型训练：从基线模型到融合模型", 2)
    add_para(
        doc,
        "模型训练可按四步推进。第一步建立MNF和移动基线，形成基础预警能力。第二步训练LSTM或GRU动态基线模型，识别持续残差。第三步引入孤立森林、自编码器等异常检测模型，形成多模型交叉验证。第四步结合水力模型、资产数据和工单标签，建立候选管段定位和风险排序模型。",
    )

    add_heading(doc, "4.3 模型验证：算法指标与业务指标并重", 2)
    add_para(
        doc,
        "模型验证应同时关注算法指标和业务指标。算法指标包括MAE、RMSE、精确率、召回率、F1、AUC等；业务指标包括提前预警时间、误报工单比例、平均排查范围、TopN命中率、维修闭环时间、维修后MNF回落和复盘回填率。对外汇报中，业务指标通常更能体现应用价值。",
    )

    add_heading(doc, "4.4 模型运营：漂移监控与再训练", 2)
    add_para(
        doc,
        "DMA模型上线后需要持续运营。季节变化、用户结构变化、管网改造、阀门调整、传感器更换和压力制度改变都会影响数据分布。建议每月复盘误报警报，每季度检查模型漂移，每半年结合新增工单进行再训练，重大管网变更后重新评估模型基线。",
    )

    add_heading(doc, "5. PPT拆页建议（仅模型相关）", 1)
    add_bullet(doc, "第1页：供水漏损AI模型应用演进。建议配技术路线图，呈现经验阈值、统计分析、机器学习、深度学习和智能协同五个阶段。")
    add_bullet(doc, "第2页：DMA漏损检测模型总体架构。建议配数据层、模型层、业务层和闭环反馈层四层架构图。")
    add_bullet(doc, "第3页：动态基线与异常检测。建议配实际曲线、预测区间、残差和预警点示意图。")
    add_bullet(doc, "第4页：多模型组合。建议配LSTM、无监督异常检测、监督风险排序、水力仿真融合和知识图谱协同模块图。")
    add_bullet(doc, "第5页：漏点定位与候选管段收敛。建议配DMA拓扑、压力响应和候选管段TopN排序图。")
    add_bullet(doc, "第6页：模型落地闭环。建议配数据治理、训练验证、上线监控、工单回填和再训练闭环图。")

    add_heading(doc, "6. 汇报总结建议", 1)
    add_para(
        doc,
        "供水管网DMA漏损检测的AI应用，应以业务问题为牵引，以数据治理为基础，以动态基线、异常检测、风险排序和水力仿真融合为核心模型能力。模型建设的重点不只是提高算法精度，还包括提升预警稳定性、压缩排查范围、形成可解释证据链，并通过工单回填实现持续迭代。",
    )
    add_para(
        doc,
        "面向实际落地，建议按照先数据、再基线、再异常、再定位、再闭环的路径推进。前期以MNF、动态基线和无监督异常检测建立可用能力；中期引入监督学习和水力仿真提高定位精度；后期结合知识图谱和智能体，实现模型解释、报告生成和业务协同。",
    )

    add_heading(doc, "参考资料", 1)
    add_bullet(doc, "IWA水务漏损管理与DMA分区计量相关资料。")
    add_bullet(doc, "EPANET水力建模与管网仿真相关公开资料。")
    add_bullet(doc, "供水管网漏损检测、最小夜间流量、压力管理、机器学习异常检测、LSTM时序预测和图神经网络相关研究文献。")
    add_bullet(doc, "项目既有LSTM、HGB、Isolation Forest、DBSCAN、CNN-LSTM等模型经验及本地DMA漏损检测业务材料。")

    doc.save(DOCX)
    return DOCX


if __name__ == "__main__":
    print(build())
