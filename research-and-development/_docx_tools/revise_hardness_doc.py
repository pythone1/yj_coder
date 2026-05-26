from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement


SRC = Path(r"E:\PY\research\0511\LDHS2026F02613--一种深度水处理斜板沉淀池硬度调控方法及系统--发明--初稿.docx")
OUT = Path(r"E:\PY\research\0511\LDHS2026F02613--一种深度水处理斜板沉淀池硬度调控方法及系统--发明--批注修编_黄色标注.docx")


def insert_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph._parent.add_paragraph()
    p._p = new_p
    p.style = paragraph.style
    run = p.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p


def insert_formula_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph._parent.add_paragraph()
    p._p = new_p
    p.style = paragraph.style
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Cambria Math"
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p


def insert_block_after(paragraph, items):
    last = paragraph
    for kind, text in items:
        if kind == "formula":
            last = insert_formula_after(last, text)
        else:
            last = insert_after(last, text)
    return last


def replace_phrase_highlight(paragraph, old: str, new: str) -> bool:
    text = paragraph.text
    if old not in text:
        return False
    before, after = text.split(old, 1)
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    if before:
        paragraph.add_run(before)
    marked = paragraph.add_run(new)
    marked.font.highlight_color = WD_COLOR_INDEX.YELLOW
    if after:
        paragraph.add_run(after)
    return True


def main() -> None:
    doc = Document(SRC)

    ekf_anchor = (
        "将反应利用率和水力滞后时间作为状态变量，采用扩展卡尔曼滤波的标准预测步和更新步两阶段计算框架进行递归估计，"
        "预测步根据上一时刻状态和状态转移矩阵A，预测当前状态及其协方差；更新步获得当前时刻的实际出水硬度测量值，"
        "计算卡尔曼增益，然后利用测量残差来校正预测值，得到最优估计及其协方差。"
    )
    ekf_additions = [
        ("text", "具体估计时，设状态向量为："),
        ("formula", "xₜ=[ηₜ，τₜ]ᵀ"),
        ("text", "式中，ηₜ为t时刻的反应利用率，τₜ为t时刻的水力滞后时间；以进水硬度、碳酸钠投加量、pH、温度、流量、浊度和碱度构成输入向量uₜ，状态方程为："),
        ("formula", "xₜ=f(xₜ₋₁，uₜ₋₁)+wₜ，wₜ~N(0，Qₜ)"),
        ("text", "预测步中，先根据上一时刻后验状态计算当前先验状态，再对非线性状态方程求雅可比矩阵Fₜ，计算先验协方差："),
        ("formula", "xₜ|ₜ₋₁=f(xₜ₋₁|ₜ₋₁，uₜ₋₁)"),
        ("formula", "Pₜ|ₜ₋₁=FₜPₜ₋₁|ₜ₋₁Fₜᵀ+Qₜ"),
        ("text", "观测方程和更新步计算为："),
        ("formula", "zₜ=h(xₜ，dₜ)+vₜ，vₜ~N(0，Rₜ)"),
        ("formula", "rₜ=zₜ-h(xₜ|ₜ₋₁，dₜ)"),
        ("formula", "Sₜ=HₜPₜ|ₜ₋₁Hₜᵀ+Rₜ"),
        ("formula", "Kₜ=Pₜ|ₜ₋₁HₜᵀSₜ⁻¹"),
        ("formula", "xₜ|ₜ=xₜ|ₜ₋₁+Kₜrₜ"),
        ("formula", "Pₜ|ₜ=(I-KₜHₜ)Pₜ|ₜ₋₁"),
        ("text", "其中，zₜ为当前时刻的实际出水硬度测量值，dₜ为当前运行数据，Hₜ为观测方程雅可比矩阵；更新后将ηₜ限制在0至1之间，将τₜ限制在沉淀池允许的水力停留时间范围内。")
    ]

    stacking_anchor = (
        "TPE超参数优化，将每一组候选参数的验证误差作为观测值，将验证误差低于分位阈值的候选参数分布建模为优良参数密度函数，"
        "将验证误差不低于该分位阈值的候选参数分布建模为非优良参数密度函数；分位阈值由历史验证误差排序确定；"
        "在下一轮搜索时，通过最大化期望改进准则，确定下一轮搜索参数区域，优先选择能够提高期望改进的参数区域，"
        "参数包括学习率、树深度、叶节点数、正则系数、采样比例、时间窗口长度和预测步长。"
    )
    stacking_additions = [
        ("text", "数学建模时，将当前时刻及长度为L的历史窗口运行数据组成样本："),
        ("formula", "Xₜ=[Hᵢₙ，Hₒᵤₜ，Q，Mₛₒdₐ，pH，T，NTU，SS，ΔP，η，τ，派生指标]ₜ₋L:ₜ"),
        ("formula", "yₜ=[Hₒᵤₜ,ₜ₊ₖ，NTUₜ₊ₖ]"),
        ("text", "其中，k为预测步长；第m个基学习器输出和二级输入分别为："),
        ("formula", "ŷₜ⁽ᵐ⁾=gₘ(Xₜ；θₘ)"),
        ("formula", "Zₜ=[ŷₜ⁽¹⁾，…，ŷₜ⁽ᴹ⁾，ηₜ，τₜ]"),
        ("text", "以XGBoost为例，第q轮提升树的目标函数为："),
        ("formula", "Obj⁽q⁾=∑ᵢl(yᵢ，ŷᵢ⁽q−1⁾+f⁽q⁾(Xᵢ))+γT⁽q⁾+1/2λ∑ⱼwⱼ²"),
        ("text", "式中，l为硬度和浊度预测误差，T⁽q⁾为叶节点数量，wⱼ为叶节点权重，γ和λ为复杂度惩罚系数；LightGBM、CatBoost和极端随机树分别基于直方图叶生长、有序类别编码和随机特征/阈值分裂形成互补预测。"),
        ("text", "训练过程按时间顺序划分训练窗口、验证窗口和留出测试窗口，先分别训练各基学习器并生成验证窗口的折外预测，再用折外预测训练二级元学习器；二级元学习器采用TCN时以因果卷积和膨胀卷积提取跨周期依赖，采用Transformer编码器时以自注意力权重融合不同时间步和不同基学习器的输出，最终得到硬度和浊度的联合预测值。"),
        ("text", "参数优化过程为：初始化若干组候选参数θ并训练得到验证损失L(θ)；按分位阈值γ将历史样本划分为优良集合和非优良集合，分别估计密度："),
        ("formula", "l(θ)=p(θ|L<γ)"),
        ("formula", "g(θ)=p(θ|L≥γ)"),
        ("formula", "θ*=arg max l(θ)/g(θ)"),
        ("text", "训练后将新的(θ，L(θ))写入历史样本，循环至达到最大迭代次数或验证误差不再下降。")
    ]

    opt_anchor = "求解时采用基于帕累托前沿的进化算法或序列二次规划，输出最优的碳酸钠投加量、泵频率、搅拌强度和排泥周期。"
    opt_additions = [
        ("text", "权重确定采用安全优先与熵权修正结合的方法。先在滚动窗口内计算硬度偏差、药剂消耗、浊度偏差和污泥负荷四类指标的归一化序列rᵢₖ，再计算："),
        ("formula", "pᵢₖ=rᵢₖ/∑ᵢrᵢₖ"),
        ("formula", "Eₖ=−1/ln(n)∑ᵢpᵢₖln(pᵢₖ)"),
        ("formula", "dₖ=1−Eₖ"),
        ("formula", "w̄ₖ=dₖ/∑ₖdₖ"),
        ("text", "随后根据安全裕度修正基础权重：当预测出水硬度接近安全上限时提高硬度权重，当预测浊度或压差接近上限时提高澄清风险权重；修正系数记为αₖ，最终权重为："),
        ("formula", "wₖ=αₖw̄ₖ/∑ₖαₖw̄ₖ"),
        ("text", "具体求解时，在预测时域N内构造目标函数："),
        ("formula", "J(U)=∑ⱼ₌₁ᴺ[w₁((Ĥₜ₊ⱼ−Hₜₐᵣ)/(Hₘₐₓ−Hₜₐᵣ))²+w₂Mₛₒdₐ/Mₘₐₓ+w₃((NTÛₜ₊ⱼ−NTUₜₐᵣ)/(NTUₘₐₓ−NTUₜₐᵣ))²+w₄Sludge(U)/Sludgeₘₐₓ]+ρ‖ΔU‖²"),
        ("formula", "U=[Mₛₒdₐ，fₚᵤₘₚ，Iₘᵢₓ，Tₛludge]"),
        ("formula", "Uₘᵢₙ≤U≤Uₘₐₓ，|ΔU|≤ΔUₘₐₓ，Ĥ≤Hₛₐfe，NTÛ≤NTUₘₐₓ"),
        ("text", "采用帕累托前沿进化算法时，先围绕当前控制量和物理投加量生成候选控制序列，调用预测模型计算各候选序列的硬度、浊度和污泥负荷响应，再进行非支配排序和拥挤度筛选，最后从帕累托前沿中选择加权目标J(U)最小且满足安全约束的序列；采用序列二次规划时，以当前控制量为初值，对目标函数和约束进行局部二次近似，迭代求得可行最优解，并仅下发第一步控制量用于滚动执行。")
    ]

    inserted = {"ekf": False, "stacking": False, "opt": False, "fix": False}
    opt_seen = 0
    for p in list(doc.paragraphs):
        if p.text == ekf_anchor and not inserted["ekf"]:
            insert_block_after(p, ekf_additions)
            inserted["ekf"] = True
        elif p.text == stacking_anchor and not inserted["stacking"]:
            insert_block_after(p, stacking_additions)
            inserted["stacking"] = True
        elif p.text == opt_anchor and not inserted["opt"]:
            opt_seen += 1
            if opt_seen >= 3:
                insert_block_after(p, opt_additions)
                inserted["opt"] = True
        if "扩展卡尔曼滤波EKF会据此逐步提高反应利用率和水力滞后时间的估计值" in p.text:
            inserted["fix"] = replace_phrase_highlight(
                p,
                "扩展卡尔曼滤波EKF会据此逐步提高反应利用率和水力滞后时间的估计值",
                "扩展卡尔曼滤波EKF会据此逐步下调反应利用率估计值并上调水力滞后时间估计值",
            )

    missing = [key for key, ok in inserted.items() if not ok]
    if missing:
        raise RuntimeError(f"未完成插入/替换: {missing}")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
