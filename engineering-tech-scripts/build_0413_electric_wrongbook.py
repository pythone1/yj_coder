import os
from collections import OrderedDict, defaultdict

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE = r"E:\PY\tech"
PAGE_DIR = os.path.join(BASE, "tmp", "0413_allpages")
CROP_DIR = os.path.join(BASE, "tmp", "0413_electric_crops")
OUT_DIR = os.path.join(BASE, "output", "doc")
OUT_ORIG = os.path.join(OUT_DIR, "0413电学错题集_原卷版.docx")
OUT_ANS = os.path.join(OUT_DIR, "0413电学错题集_答案版.docx")


CATEGORIES = OrderedDict(
    [
        ("电路基础与串并联", "识别串并联、补全电路、开关状态和元件作用。"),
        ("电表示数与故障分析", "电流表、电压表读数，断路短路和故障定位。"),
        ("欧姆定律与动态电路", "I-U 图像、滑动变阻器、动态电路变化分析。"),
        ("电功率与电能综合", "额定值、电功、电功率、电能表和多挡位问题。"),
        ("传感器与自动控制", "气敏、热敏、报警和自动控制电路。"),
        ("电学实验与特殊测量", "探究规律、测电阻、测功率、等效替代法。"),
        ("电与磁", "奥斯特实验、电磁感应和相关应用。"),
    ]
)


def q(page_token, box):
    return {"page": page_token, "box": box}


ITEMS = [
    {
        "id": "A01",
        "title": "安全带提示系统电路设计",
        "category": "电路基础与串并联",
        "sources": ["23935120 第11题"],
        "segments": [q("23935120_p7", (35, 820, 800, 1135))],
        "answer": "参考答案：选能实现“未系安全带时指示灯亮、系上后熄灭”的连接方案，关键是把座椅开关和安全带开关按题意控制通断。",
    },
    {
        "id": "A02",
        "title": "并联电路干路电流实验改错",
        "category": "电学实验与特殊测量",
        "sources": ["23935120 第12题"],
        "segments": [q("23935120_p8", (35, 0, 800, 300))],
        "answer": "参考要点：电流表必须与被测支路串联，验证并联电路电流规律时应分别测干路和各支路电流。",
    },
    {
        "id": "A03",
        "title": "补电表符号并画实物图对应电路图",
        "category": "电路基础与串并联",
        "sources": ["23935120 第20题"],
        "segments": [q("23935120_p8", (35, 300, 800, 760))],
        "answer": "参考要点：先判断各元件是否串联或并联，再确定电流表串联、电压表并联。",
    },
    {
        "id": "A04",
        "title": "指纹解锁、反向充电、房卡取电、番茄电池",
        "category": "电路基础与串并联",
        "sources": ["23935123 第16题"],
        "segments": [
            q("23935123_p1", (35, 820, 800, 1170)),
            q("23935123_p2", (35, 0, 800, 350)),
        ],
        "answer": "参考答案要点：感应区相当于开关；反向充电时 A 手机相当于电源，能量由电能向其他形式再转化；房卡取电器相当于开关，应接在干路；番茄电池中电压表正接线柱所连金属片为正极。",
    },
    {
        "id": "A05",
        "title": "补全并联电路并改接电流表测支路电流",
        "category": "电学实验与特殊测量",
        "sources": ["23935123 第21题"],
        "segments": [q("23935123_p2", (35, 720, 800, 1145))],
        "answer": "参考要点：电源放在干路，电流表要与被测支路串联；若测 B 灯电流，就要把表改接到 B 所在支路。",
    },
    {
        "id": "B01",
        "title": "灯不亮但电压表有示数的故障判断",
        "category": "电表示数与故障分析",
        "sources": ["23935279 第8题"],
        "segments": [q("23935279_p1", (35, 330, 800, 760))],
        "answer": "参考答案：D。若灯短路或电阻断路都可能出现题述现象，单靠并联在灯两端的电压表示数不能可靠区分故障。",
    },
    {
        "id": "B02",
        "title": "双开关双电流表示数比较",
        "category": "电表示数与故障分析",
        "sources": ["23935279 第10题"],
        "segments": [q("23935279_p1", (35, 720, 800, 1135))],
        "answer": "参考要点：先根据指针位置和量程判断两表真实电流，再结合串并联关系比较各支路电流。",
    },
    {
        "id": "C01",
        "title": "两电阻 I-U 图像串并联综合",
        "category": "欧姆定律与动态电路",
        "sources": ["23935279 第20题"],
        "segments": [q("23935279_p2", (35, 780, 800, 1155))],
        "answer": "参考要点：图像斜率反映电阻大小；串联同流分压、并联同压分流，再用 P=UI 或 U=IR 计算。",
    },
    {
        "id": "F01",
        "title": "探究串联电压与并联电流规律",
        "category": "电学实验与特殊测量",
        "sources": ["23935279 第25题"],
        "segments": [
            q("23935279_p3", (35, 980, 800, 1170)),
            q("23935279_p4", (35, 0, 800, 900)),
        ],
        "answer": "参考结论：串联电路总电压等于各部分电压之和；并联电路干路电流等于各支路电流之和。实验中应通过改变电源或灯泡进行多次测量。",
    },
    {
        "id": "C02",
        "title": "滑动变阻器变化引起电压表示数变化",
        "category": "欧姆定律与动态电路",
        "sources": ["23935282 第6题"],
        "segments": [q("23935282_p1", (35, 40, 800, 350))],
        "answer": "参考答案：选 A。总电流增加 0.4A，对应总电阻减小 10Ω，可反推出电源电压，再据定值电阻电流变化求电压表示数增加量。",
    },
    {
        "id": "B03",
        "title": "两灯均不亮的故障分析",
        "category": "电表示数与故障分析",
        "sources": ["23935282 第8题"],
        "segments": [q("23935282_p1", (35, 310, 800, 675))],
        "answer": "参考要点：根据 M 接正极、N 依次触碰 A/B 都接近 3V，结合断路短路特征判断一灯短路、一灯断路的组合。",
    },
    {
        "id": "E01",
        "title": "天然气泄漏检测电路",
        "category": "传感器与自动控制",
        "sources": ["23935282 第10题"],
        "segments": [q("23935282_p1", (35, 620, 800, 915))],
        "answer": "参考要点：由传感器电压随浓度变化先判断传感器电阻变化，再由串联电路分压关系判断 U0 与 I 的图像。",
    },
    {
        "id": "C03",
        "title": "多开关动态电路图像判断",
        "category": "欧姆定律与动态电路",
        "sources": ["23935282 第11题"],
        "segments": [
            q("23935282_p1", (35, 900, 800, 1175)),
            q("23935282_p2", (35, 0, 800, 140)),
        ],
        "answer": "参考要点：先分状态分析电路结构，再结合图中两条线分别表示的电压或电流变化判断电阻值和电源电压。",
    },
    {
        "id": "D01",
        "title": "电阻额定参数与安全使用",
        "category": "电功率与电能综合",
        "sources": ["23935282 第14题"],
        "segments": [q("23935282_p2", (35, 120, 800, 380))],
        "answer": "参考要点：先由铭牌求允许最大电流或最大功率，再分别用串联、并联电路规律判断安全范围。",
    },
    {
        "id": "B04",
        "title": "辨析电表种类及示数比",
        "category": "电表示数与故障分析",
        "sources": ["23935282 第15题"],
        "segments": [q("23935282_p2", (35, 360, 800, 610))],
        "answer": "参考要点：先根据接法判断是电流表还是电压表，再结合串并联规律求示数之比。",
    },
    {
        "id": "C04",
        "title": "电子秤与滑动变阻器动态关系",
        "category": "欧姆定律与动态电路",
        "sources": ["23935282 第16题"],
        "segments": [q("23935282_p2", (35, 580, 800, 960))],
        "answer": "参考要点：由压力与弹簧形变量图确定滑片位置变化，再结合欧姆定律分析表的示数变化。",
    },
    {
        "id": "D02",
        "title": "4V 4W 灯泡与双挡动态功率综合",
        "category": "电功率与电能综合",
        "sources": ["23935282 第17题"],
        "segments": [
            q("23935282_p2", (35, 930, 800, 1175)),
            q("23935282_p3", (35, 0, 800, 330)),
        ],
        "answer": "参考要点：正常发光时 I=P/U，R=U²/P；分两种接法分别列总电流和总功率，再比较功率比并计算电能。",
    },
    {
        "id": "F02",
        "title": "用电阻箱和等效替代法测灯泡电阻",
        "category": "电学实验与特殊测量",
        "sources": ["23935282 第18题"],
        "segments": [q("23935282_p3", (35, 330, 800, 760))],
        "answer": "参考要点：先让小灯泡正常发光并记录电压，再用电阻箱替代小灯泡并调到电压表示数相同，此时电阻箱示数即近似等于灯泡该状态电阻。",
    },
    {
        "id": "F03",
        "title": "探究电流与电阻关系",
        "category": "电学实验与特殊测量",
        "sources": ["23935282 第19题"],
        "segments": [
            q("23935282_p3", (35, 760, 800, 1175)),
            q("23935282_p4", (35, 0, 800, 250)),
        ],
        "answer": "参考结论：电压一定时，电流与电阻成反比。实验中应通过滑动变阻器保持定值电阻两端电压不变。",
    },
    {
        "id": "B05",
        "title": "保护电表时滑动变阻器允许范围",
        "category": "电表示数与故障分析",
        "sources": ["23935282 第20题"],
        "segments": [q("23935282_p4", (35, 160, 800, 520))],
        "answer": "参考要点：分别用电流表和电压表的量程作约束，取同时满足两表安全的交集范围。",
    },
    {
        "id": "D03",
        "title": "相同灯泡总功率大小规律",
        "category": "电功率与电能综合",
        "sources": ["23935282 第6题（第4页）", "23935283 第6题"],
        "segments": [q("23935283_p1", (35, 40, 800, 350))],
        "answer": "参考答案：比较各电路等效电阻，电源电压相同则总功率 P=U²/R，总电阻越小，总功率越大。",
    },
    {
        "id": "D04",
        "title": "6V 3W 与 6V 9W 灯泡串联",
        "category": "电功率与电能综合",
        "sources": ["23935282 第7题（第4页）", "23935283 第7题"],
        "segments": [q("23935283_p1", (35, 330, 800, 610))],
        "answer": "参考答案：选 D。两灯串联电流相同，实际电压都偏离额定值，均不能正常发光。",
    },
    {
        "id": "C05",
        "title": "两电压表随电流变化图像判断",
        "category": "欧姆定律与动态电路",
        "sources": ["23935282 第11题（第4页）", "23935283 第11题"],
        "segments": [q("23935283_p1", (35, 600, 800, 930))],
        "answer": "参考要点：同一电流变化下，两条直线对应两个不同部分的电压，结合总电压不变和分压规律判断 A、B 所代表的物理量。",
    },
    {
        "id": "D05",
        "title": "3V 0.9W 灯泡安全调节范围",
        "category": "电功率与电能综合",
        "sources": ["23935283 第12题"],
        "segments": [q("23935283_p1", (35, 900, 800, 1175))],
        "answer": "参考要点：由额定值求灯泡电阻和额定电流，再结合电表量程和变阻器范围求安全电流、功率及电阻范围。",
    },
    {
        "id": "C06",
        "title": "两电阻 I-U 图像求串并联功率",
        "category": "欧姆定律与动态电路",
        "sources": ["23935283 第13题"],
        "segments": [q("23935283_p2", (35, 0, 800, 280))],
        "answer": "参考要点：图像斜率不同代表电阻不同，串联先求总电阻，并联先求总电流，再用 P=UI 或 P=U²/R。",
    },
    {
        "id": "D06",
        "title": "自制可调光台灯",
        "category": "电功率与电能综合",
        "sources": ["23935283 第14题"],
        "segments": [q("23935283_p2", (35, 250, 800, 520))],
        "answer": "参考要点：灯泡应与滑动变阻器串联接入；滑片移动导致接入电阻变化，从而改变灯泡实际功率；变阻器本身耗电增加并不等于更省电。",
    },
    {
        "id": "C07",
        "title": "闭合 S2 后电表和亮度变化",
        "category": "欧姆定律与动态电路",
        "sources": ["23935283 第15题"],
        "segments": [q("23935283_p2", (35, 500, 800, 760))],
        "answer": "参考要点：先判断闭合 S2 前后电路结构是否由串联变并联，再据此判断电流表示数、电压表示数和灯泡亮度变化。",
    },
    {
        "id": "D07",
        "title": "定值电阻、电阻箱与总功率综合",
        "category": "电功率与电能综合",
        "sources": ["23935283 第18题"],
        "segments": [q("23935283_p2", (35, 740, 800, 1170))],
        "answer": "参考要点：先用一组 U、I 求定值电阻，再由另一状态求电源电压、1min 电能，最后按不同支路求总功率。",
    },
    {
        "id": "F04",
        "title": "测小灯泡正常发光时的电阻",
        "category": "电学实验与特殊测量",
        "sources": ["23935283 第19题"],
        "segments": [q("23935283_p3", (35, 0, 800, 680))],
        "answer": "参考要点：先正确连接电路并使小灯泡正常发光，利用 U/I 求电阻；特殊法测量时要利用开关切换并保持某一物理量不变。",
    },
    {
        "id": "F05",
        "title": "双表格实验：电流与电压、电流与电阻",
        "category": "电学实验与特殊测量",
        "sources": ["23935283 第20题"],
        "segments": [
            q("23935283_p3", (35, 650, 800, 1175)),
            q("23935283_p4", (35, 0, 800, 340)),
        ],
        "answer": "参考结论：通过导体的电流与它两端电压成正比；电压一定时，电流与电阻成反比。第二组实验未能得出结论的根本原因是没有控制电阻两端电压不变。",
    },
    {
        "id": "D08",
        "title": "3V 0.3A 灯泡与多开关功率范围",
        "category": "电功率与电能综合",
        "sources": ["23935283 第21题"],
        "segments": [q("23935283_p4", (35, 300, 800, 780))],
        "answer": "参考要点：由灯泡 I-U 图像读出正常工作点，再分别分析不同开关闭合时的最小电流与总功率范围。",
    },
    {
        "id": "G01",
        "title": "动圈式话筒工作原理",
        "category": "电与磁",
        "sources": ["23935284 第4题"],
        "segments": [q("23935284_p1", (35, 40, 800, 240))],
        "answer": "参考答案：原理与发电机类似，属于电磁感应现象。",
    },
    {
        "id": "D09",
        "title": "电能表读数与电费计算",
        "category": "电功率与电能综合",
        "sources": ["23935284 第5题"],
        "segments": [q("23935284_p1", (35, 220, 800, 500))],
        "answer": "参考要点：先由两次示数求用电量，再算电费；由脉冲常数和闪烁次数可求消耗电能。",
    },
    {
        "id": "C08",
        "title": "恒流源电路中的量随滑片变化",
        "category": "欧姆定律与动态电路",
        "sources": ["23935284 第9题"],
        "segments": [q("23935284_p1", (35, 470, 800, 770))],
        "answer": "参考要点：恒流源条件下总电流恒定，重点分析支路电流、电压和总功率与电阻变化之间的关系。",
    },
    {
        "id": "B06",
        "title": "一灯熄灭时的故障与电表示数",
        "category": "电表示数与故障分析",
        "sources": ["23935284 第10题"],
        "segments": [q("23935284_p1", (35, 750, 800, 1045))],
        "answer": "参考要点：先判断原电路中两灯正常发光时的连接方式，再分别讨论一灯短路或断路时两表是否有示数。",
    },
    {
        "id": "G02",
        "title": "电学实验说法辨析",
        "category": "电与磁",
        "sources": ["23935284 第11题"],
        "segments": [q("23935284_p1", (35, 1030, 800, 1175))],
        "answer": "参考要点：逐项对应实验目的与现象，避免把控制变量、滑动变阻器作用和电表接法混淆。",
    },
    {
        "id": "E02",
        "title": "热敏电阻与滑动变阻器综合图像",
        "category": "传感器与自动控制",
        "sources": ["23935284 第12题"],
        "segments": [q("23935284_p2", (35, 120, 800, 760))],
        "answer": "参考要点：先根据已知工作状态求电源电压和 R 的取值范围，再由热敏电阻图像反推温度对应关系。",
    },
    {
        "id": "G03",
        "title": "奥斯特实验",
        "category": "电与磁",
        "sources": ["23935284 第13题"],
        "segments": [q("23935284_p2", (35, 760, 800, 980))],
        "answer": "参考答案：说明通电导体周围存在磁场，且磁场方向与电流方向有关。",
    },
    {
        "id": "C09",
        "title": "滑片左移时 V 与两电流之差变化",
        "category": "欧姆定律与动态电路",
        "sources": ["23935284 第16题"],
        "segments": [q("23935284_p2", (35, 955, 800, 1175))],
        "answer": "参考要点：先由并联电路求 V、A1、A2 的对应关系，再讨论滑片移动导致支路电阻变化后各物理量怎样变。",
    },
    {
        "id": "E03",
        "title": "温度报警电路设计",
        "category": "传感器与自动控制",
        "sources": ["23935284 第17题"],
        "segments": [q("23935284_p3", (35, 0, 800, 600))],
        "answer": "参考要点：设报警临界时电流为 20mA，由 U=IR 求总电阻，再结合热敏电阻阻值-温度图求报警温度范围与电源、变阻器参数。",
    },
    {
        "id": "F06",
        "title": "测小灯泡额定功率及电阻",
        "category": "电学实验与特殊测量",
        "sources": ["23935284 第18题"],
        "segments": [q("23935284_p3", (35, 580, 800, 1175))],
        "answer": "参考要点：先改正错线，再让滑片位于阻值最大处闭合开关；达到额定电压时读出电流求额定功率，特殊法部分利用两次等效测量求电阻。",
    },
    {
        "id": "D10",
        "title": "汽车坐垫高温低温两挡电路",
        "category": "电功率与电能综合",
        "sources": ["23935284 第19题"],
        "segments": [q("23935284_p4", (35, 0, 800, 520))],
        "answer": "参考要点：由高温和低温功率分别求总电阻，再结合开关位置判断 R1、R2 串并关系；吸热部分用 Q=cmΔt 计算。",
    },
    {
        "id": "A06",
        "title": "电学暗盒串并联判断",
        "category": "电路基础与串并联",
        "sources": ["24381232 第8题"],
        "segments": [q("24381232_p1", (35, 300, 800, 760))],
        "answer": "参考要点：依据“只闭合某开关时哪只灯亮、两只是否同时亮”判断暗盒内部连接方式。",
    },
    {
        "id": "C10",
        "title": "小灯泡与滑动变阻器动态图像题",
        "category": "欧姆定律与动态电路",
        "sources": ["24381232 第9题", "24498670 第9题"],
        "segments": [q("24381232_p1", (35, 740, 800, 1090))],
        "answer": "参考要点：由三只电压表和电流表变化图先确定电源电压和灯泡额定参数，再据滑片移动方向分析 R2 接入电阻变化。",
    },
    {
        "id": "D11",
        "title": "4V 1.6W 小灯泡多挡位功率综合",
        "category": "电功率与电能综合",
        "sources": ["24381232 第12题", "24498670 第12题", "24668125 第12题"],
        "segments": [q("24498670_p2", (35, 180, 800, 770))],
        "answer": "参考要点：由“只闭合 S、S2、S3 时 R2 功率为 1.8W”先求电源电压，再分别求不同开关状态下 R1 功率与灯泡实际功率比值、表值变化比和总功率范围。",
    },
    {
        "id": "E04",
        "title": "电热水器自动保温控制",
        "category": "传感器与自动控制",
        "sources": ["24668125 第14题"],
        "segments": [
            q("24668125_2__p1", (35, 860, 800, 1175)),
            q("24668125_2__p2", (35, 0, 800, 760)),
        ],
        "answer": "参考答案页可见：重新开始加热时水温为 40℃，保温 10s 产生热量 4400J。工作电路应结合继电器吸合/释放条件判断。",
    },
    {
        "id": "F07",
        "title": "测额定发光灯泡电阻与未知电阻",
        "category": "电学实验与特殊测量",
        "sources": ["24668125 第16题"],
        "segments": [
            q("24668125_2__p3", (35, 350, 800, 1175)),
            q("24668125_2__p4", (35, 0, 800, 540)),
        ],
        "answer": "参考答案页给出关键结果：灯泡正常发光电阻为 12.5Ω；特殊法部分核心是保持同一电流，利用电阻箱示数差求未知电阻。",
    },
]


def apply_font(run, size=11, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_text(doc, text="", size=11, bold=False, align=None, space_after=4):
    para = doc.add_paragraph()
    if text:
        run = para.add_run(text)
        apply_font(run, size=size, bold=bold)
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    return para


def find_page_image(token):
    matches = [n for n in os.listdir(PAGE_DIR) if token in n and n.lower().endswith(".png")]
    if not matches:
        raise FileNotFoundError(f"找不到页面图片：{token}")
    return os.path.join(PAGE_DIR, sorted(matches)[0])


def crop_segment(page_path, box, out_name):
    os.makedirs(CROP_DIR, exist_ok=True)
    out_path = os.path.join(CROP_DIR, out_name)
    img = Image.open(page_path)
    crop = img.crop(box)
    crop.save(out_path)
    return out_path


def build_crop_cache():
    cache = defaultdict(list)
    for item in ITEMS:
        for idx, seg in enumerate(item["segments"], start=1):
            page_path = find_page_image(seg["page"])
            out_name = f'{item["id"]}_{idx}.png'
            cache[item["id"]].append(crop_segment(page_path, seg["box"], out_name))
    return cache


def setup_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)
    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(1.9)
    sec.right_margin = Cm(1.9)
    return doc


def add_cover(doc, title, subtitle):
    add_text(doc, title, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_text(doc, subtitle, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_text(
        doc,
        "整理范围：0413 目录中的物理错题 PDF；按电学题型筛选，重复题合并并保留来源标注。",
        size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    add_text(doc, f"共整理 {len(ITEMS)} 组电学错题。", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def add_category_index(doc):
    add_text(doc, "题型分类索引", size=15, bold=True)
    grouped = defaultdict(list)
    for item in ITEMS:
        grouped[item["category"]].append(item)
    for cat, desc in CATEGORIES.items():
        if not grouped.get(cat):
            continue
        add_text(doc, cat, size=12, bold=True, space_after=2)
        add_text(doc, desc, size=10, space_after=2)
        for item in grouped[cat]:
            add_text(
                doc,
                f'[{item["id"]}] {item["title"]} - {"；".join(item["sources"])}',
                size=10,
                space_after=1,
            )
        add_text(doc, "", space_after=2)
    doc.add_page_break()


def add_questions(doc, crop_cache, with_answer=False):
    add_text(doc, "完整题目", size=15, bold=True)
    grouped = defaultdict(list)
    for item in ITEMS:
        grouped[item["category"]].append(item)
    first_cat = True
    for cat, _desc in CATEGORIES.items():
        if not grouped.get(cat):
            continue
        if not first_cat:
            doc.add_page_break()
        first_cat = False
        add_text(doc, cat, size=14, bold=True)
        for item in grouped[cat]:
            add_text(doc, f'[{item["id"]}] {item["title"]}', size=12, bold=True, space_after=2)
            add_text(doc, f'来源：{"；".join(item["sources"])}', size=10, space_after=2)
            for img_path in crop_cache[item["id"]]:
                doc.add_picture(img_path, width=Cm(15.7))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if with_answer:
                add_text(doc, item["answer"], size=10, space_after=4)
            else:
                add_text(doc, "分类备注：见前文题型索引。", size=9, space_after=4)


def save_docs(crop_cache):
    doc_orig = setup_doc()
    add_cover(doc_orig, "0413电学错题集 - 原卷版", "保留完整题图，用于刷题和二次打印")
    add_category_index(doc_orig)
    add_questions(doc_orig, crop_cache, with_answer=False)
    doc_orig.save(OUT_ORIG)

    doc_ans = setup_doc()
    add_cover(doc_ans, "0413电学错题集 - 答案版", "保留完整题图，并附参考答案与解题要点")
    add_category_index(doc_ans)
    add_questions(doc_ans, crop_cache, with_answer=True)
    doc_ans.save(OUT_ANS)


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    crop_cache = build_crop_cache()
    save_docs(crop_cache)
    print(OUT_ORIG)
    print(OUT_ANS)
    print(f"items={len(ITEMS)}")


if __name__ == "__main__":
    main()
