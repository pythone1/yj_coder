from __future__ import annotations

import json
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import plotly.graph_objects as go
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from matplotlib.lines import Line2D
from matplotlib.patches import Patch


ROOT = Path(__file__).resolve().parents[1]
RESULT_DIR = ROOT / "results_0520" / "medium_run"
DATA_DIR = ROOT / "data" / "generated_0520"
ANALYSIS_DIR = ROOT / "analysis_0520"
OUT_DIR = RESULT_DIR / "analysis_report_0520_medium"
FIG_DIR = OUT_DIR / "figures"

CANDIDATE_NODES = [
    "10", "62", "124", "42", "178", "63", "103", "241", "273", "215",
    "216", "60", "308", "310", "312", "118", "304", "64", "91", "85",
]
MONITOR_NODES = ["286", "223", "239", "267", "8", "251", "252", "189", "37"]
TRUTH_NODES = ["103", "304", "10", "178", "42"]
SOLUTION_LABELS = {
    "ga_best": "GA最佳解",
    "posterior_best_map": "AM最优解（MAP）",
}


def setup_matplotlib() -> None:
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "SimSun", "Arial Unicode MS"]
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.dpi"] = 150


def load_data() -> dict[str, object]:
    return {
        "summary": json.loads((RESULT_DIR / "0520_summary.json").read_text(encoding="utf-8")),
        "data_summary": json.loads((DATA_DIR / "0520_data_summary.json").read_text(encoding="utf-8")),
        "scores": pd.read_csv(RESULT_DIR / "0520_solution_scores.csv", encoding="utf-8-sig"),
        "shares": pd.read_csv(RESULT_DIR / "0520_solution_shares.csv", encoding="utf-8-sig"),
        "ga_history": pd.read_csv(RESULT_DIR / "0520_GA_best_by_generation.csv", encoding="utf-8-sig"),
        "ga_all": pd.read_csv(RESULT_DIR / "0520_GA_all.csv", encoding="utf-8-sig"),
        "ga_last": pd.read_csv(RESULT_DIR / "0520_GA_last_generation.csv", encoding="utf-8-sig"),
        "am": pd.read_csv(RESULT_DIR / "0520_AM_samples.csv", encoding="utf-8-sig"),
        "observed_delta": pd.read_csv(DATA_DIR / "0520_observed_delta_10min.csv", encoding="utf-8-sig"),
        "map_delta": pd.read_csv(RESULT_DIR / "0520_solution_posterior_best_map_delta.csv", encoding="utf-8-sig"),
        "nodes": pd.read_csv(ANALYSIS_DIR / "0520_nodes_classified.csv", encoding="utf-8-sig"),
        "links": pd.read_csv(ANALYSIS_DIR / "0520_links_classified.csv", encoding="utf-8-sig"),
    }


def pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def solution_shares(shares: pd.DataFrame, name: str) -> pd.Series:
    row = shares.loc[shares["solution"] == name].iloc[0]
    return row[CANDIDATE_NODES].astype(float)


def truth_shares() -> pd.Series:
    s = pd.Series(0.0, index=CANDIDATE_NODES)
    for node in TRUTH_NODES:
        s[node] = 1.0 / len(TRUTH_NODES)
    return s


def nse(obs: np.ndarray, sim: np.ndarray) -> float:
    denom = float(np.sum((obs - obs.mean()) ** 2))
    if denom <= 1e-12:
        return float("nan")
    return 1.0 - float(np.sum((sim - obs) ** 2)) / denom


def detection_table(shares: pd.DataFrame, threshold: float = 0.05) -> pd.DataFrame:
    rows = []
    truth = set(TRUTH_NODES)
    for name in ["ga_best", "posterior_best_map"]:
        s = solution_shares(shares, name)
        active = set(s[s > threshold].index)
        rows.append({
            "方案": SOLUTION_LABELS[name],
            "判断阈值": pct(threshold),
            "超过阈值点数": len(active),
            "命中真值点": "、".join([n for n in TRUTH_NODES if n in active]) or "无",
            "命中数量": f"{len(active & truth)}/5",
            "漏掉真值点": "、".join([n for n in TRUTH_NODES if n not in active]) or "无",
            "主要代偿点": "、".join(sorted(active - truth, key=lambda n: float(s[n]), reverse=True)) or "无",
            "真值点总比例": pct(float(s.loc[TRUTH_NODES].sum())),
            "非真值点总比例": pct(float(1.0 - s.loc[TRUTH_NODES].sum())),
        })
    return pd.DataFrame(rows)


def monitor_metric_table(observed: pd.DataFrame, simulated: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for node in MONITOR_NODES:
        obs = observed[node].to_numpy(dtype=float)
        sim = simulated[node].to_numpy(dtype=float)
        rows.append({
            "监测点": node,
            "单点NSE": nse(obs, sim),
            "单点SSE": float(np.sum((sim - obs) ** 2)),
            "观测峰值": float(np.max(obs)),
            "模拟峰值": float(np.max(sim)),
            "峰值差": float(np.max(sim) - np.max(obs)),
        })
    out = pd.DataFrame(rows).sort_values("单点NSE", na_position="last").reset_index(drop=True)
    return out


def save_score_figure(scores: pd.DataFrame) -> Path:
    path = FIG_DIR / "01_核心结果评分对比.png"
    selected = scores[scores["solution"].isin(["ga_best", "posterior_best_map"])].copy()
    labels = ["GA最佳解", "AM最优解"]
    mean_nse = selected["mean_nse"].to_numpy(float)
    sse = selected["sse"].to_numpy(float)
    fig, axes = plt.subplots(1, 2, figsize=(10, 4.2))
    x = np.arange(len(labels))

    bars = axes[0].bar(x, mean_nse, width=0.45, color="#2E75B6", label="平均NSE")
    axes[0].set_title("平均NSE：越高越好")
    axes[0].set_xticks(x)
    axes[0].set_xticklabels(labels)
    axes[0].set_ylim(0, 1.05)
    axes[0].set_ylabel("平均NSE")
    axes[0].grid(axis="y", alpha=0.25)
    for bar, val in zip(bars, mean_nse):
        axes[0].text(bar.get_x() + bar.get_width() / 2, val + 0.02, f"{val:.3f}", ha="center")

    bars = axes[1].bar(x, sse, width=0.45, color="#ED7D31", label="平方误差SSE")
    axes[1].set_title("平方误差SSE：越低越好")
    axes[1].set_xticks(x)
    axes[1].set_xticklabels(labels)
    axes[1].set_ylabel("平方误差SSE")
    axes[1].grid(axis="y", alpha=0.25)
    for bar, val in zip(bars, sse):
        axes[1].text(bar.get_x() + bar.get_width() / 2, val + max(sse) * 0.03, f"{val:.4f}", ha="center")

    fig.suptitle("本次中参数运行的核心结果评分")
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_ga_convergence(ga_history: pd.DataFrame) -> Path:
    path = FIG_DIR / "02_GA逐代收敛过程.png"
    best_by_gen = ga_history.groupby("generation")["best_mean_nse"].max()
    fig, ax = plt.subplots(figsize=(9, 4.5))
    ax.plot(best_by_gen.index, best_by_gen.values, marker="o", color="#2E75B6", label="每代最高平均NSE")
    ax.set_xlabel("GA代数")
    ax.set_ylabel("平均NSE")
    ax.set_title("GA阶段：逐代寻找监测曲线拟合更好的候选解")
    ax.grid(alpha=0.25)
    ax.legend()
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_share_figure(shares: pd.DataFrame) -> Path:
    path = FIG_DIR / "03_真值与识别比例对比.png"
    data = pd.DataFrame({
        "真值比例": truth_shares(),
        "GA最佳解": solution_shares(shares, "ga_best"),
        "AM最优解": solution_shares(shares, "posterior_best_map"),
    })
    x = np.arange(len(CANDIDATE_NODES))
    fig, ax = plt.subplots(figsize=(13, 5.2))
    width = 0.26
    colors = ["#A5A5A5", "#4472C4", "#ED7D31"]
    for i, col in enumerate(data.columns):
        ax.bar(x + (i - 1) * width, data[col].values, width=width, label=col, color=colors[i])
    ax.axhline(0.05, color="#C00000", linestyle="--", linewidth=1.0, label="5%判断阈值")
    ax.set_xticks(x)
    ax.set_xticklabels(CANDIDATE_NODES, rotation=45)
    ax.set_ylabel("注入比例")
    ax.set_title("20个候选点上的比例分配")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(ncol=4, loc="upper right")
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_am_trace(am: pd.DataFrame) -> Path:
    path = FIG_DIR / "04_AM多链搜索过程.png"
    chain_stats = am.groupby("chain").agg(min_sse=("sse", "min"), accept_rate=("accepted", "mean")).reset_index()
    best_chain = int(chain_stats.sort_values("min_sse").iloc[0]["chain"])
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.4))
    for chain, group in am.groupby("chain"):
        sample = group.iloc[:: max(1, len(group) // 180)]
        if int(chain) == best_chain:
            axes[0].plot(sample["step"], sample["sse"], color="#ED7D31", linewidth=1.4, label="误差最低链")
        else:
            axes[0].plot(sample["step"], sample["sse"], color="#7F7F7F", linewidth=0.8, alpha=0.35)
    axes[0].set_title("AM采样过程中的误差变化")
    axes[0].set_xlabel("采样步数")
    axes[0].set_ylabel("平方误差SSE")
    axes[0].grid(alpha=0.25)
    axes[0].legend()

    axes[1].bar(chain_stats["chain"].astype(str), chain_stats["accept_rate"], color="#70AD47", label="每条链接受率")
    axes[1].set_title("AM各链的接受率")
    axes[1].set_xlabel("链编号")
    axes[1].set_ylabel("接受率")
    axes[1].set_ylim(0, 0.4)
    axes[1].grid(axis="y", alpha=0.25)
    axes[1].legend()
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_monitor_fit(observed: pd.DataFrame, simulated: pd.DataFrame) -> Path:
    path = FIG_DIR / "05_AM最优解各监测点拟合曲线.png"
    hours = observed["relative_hour"].to_numpy(dtype=float)
    fig, axes = plt.subplots(3, 3, figsize=(13, 8.8), sharex=True)
    for ax, node in zip(axes.flatten(), MONITOR_NODES):
        obs = observed[node].to_numpy(dtype=float)
        sim = simulated[node].to_numpy(dtype=float)
        score = nse(obs, sim)
        label = "近零响应" if np.isnan(score) else f"单点NSE={score:.3f}"
        ax.plot(hours, obs, color="#1F4E79", linewidth=1.5, label="观测增量")
        ax.plot(hours, sim, color="#ED7D31", linewidth=1.2, linestyle="--", label="AM最优解模拟")
        ax.set_title(f"监测点{node}，{label}", fontsize=10)
        ax.grid(alpha=0.2)
    axes[-1, 0].set_xlabel("小时")
    axes[-1, 1].set_xlabel("小时")
    axes[-1, 2].set_xlabel("小时")
    axes[1, 0].set_ylabel("流量增量 m3/s")
    handles, labels = axes[0, 0].get_legend_handles_labels()
    fig.legend(handles, labels, loc="lower center", ncol=2, bbox_to_anchor=(0.5, -0.01))
    fig.suptitle("AM最优解在9个监测点上的拟合情况", y=0.995)
    fig.tight_layout(rect=[0, 0.035, 1, 0.965])
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def node_xy(nodes: pd.DataFrame) -> pd.DataFrame:
    return nodes.dropna(subset=["x", "y"]).assign(node=lambda d: d["node"].astype(str)).set_index("node")[["x", "y"]]


def save_network_figure(nodes: pd.DataFrame, links: pd.DataFrame, weights: pd.Series, title: str, filename: str) -> Path:
    path = FIG_DIR / filename
    xy = node_xy(nodes)
    fig, ax = plt.subplots(figsize=(13, 6.8))
    for _, link in links.iterrows():
        f, t = str(link["from_node"]), str(link["to_node"])
        if f in xy.index and t in xy.index:
            ax.plot([xy.loc[f, "x"], xy.loc[t, "x"]], [xy.loc[f, "y"], xy.loc[t, "y"]], color="#CBD5E1", linewidth=0.7, zorder=1)
    ax.scatter(xy["x"], xy["y"], s=9, color="#B7C9E2", alpha=0.75, zorder=2)
    for node in MONITOR_NODES:
        if node in xy.index:
            ax.scatter(xy.loc[node, "x"], xy.loc[node, "y"], s=55, color="#2F5597", zorder=4)
            ax.text(xy.loc[node, "x"], xy.loc[node, "y"], node, fontsize=8, color="#1F2937", zorder=6)
    max_weight = max(0.001, float(weights.max()))
    for node in CANDIDATE_NODES:
        if node not in xy.index:
            continue
        value = float(weights.get(node, 0.0))
        is_truth = node in TRUTH_NODES
        color = "#C00000" if is_truth else "#D9A300"
        size = 55 + 850 * value / max_weight
        ax.scatter(xy.loc[node, "x"], xy.loc[node, "y"], s=size, color=color, alpha=0.65, edgecolor="#111827", linewidth=0.4, zorder=5)
        if value >= 0.05 or is_truth:
            ax.text(
                xy.loc[node, "x"], xy.loc[node, "y"], f"{node}\n{value * 100:.1f}%",
                fontsize=8, color="#111827", zorder=6,
                bbox=dict(facecolor="white", edgecolor="none", alpha=0.6, pad=0.3),
            )
    legend_handles = [
        Line2D([0], [0], color="#CBD5E1", lw=1.4, label="管线"),
        Line2D([0], [0], marker="o", color="w", markerfacecolor="#2F5597", markersize=8, label="监测点"),
        Line2D([0], [0], marker="o", color="w", markerfacecolor="#C00000", markeredgecolor="#111827", markersize=9, label="真值注入点"),
        Line2D([0], [0], marker="o", color="w", markerfacecolor="#D9A300", markeredgecolor="#111827", markersize=9, label="非真值候选点"),
    ]
    ax.legend(handles=legend_handles, loc="lower left", frameon=True, framealpha=0.92, fontsize=9)
    ax.set_title(title)
    ax.set_aspect("equal", adjustable="box")
    ax.axis("off")
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_network_html(nodes: pd.DataFrame, links: pd.DataFrame, shares: pd.DataFrame) -> Path:
    path = OUT_DIR / "0520中参数管网识别热力图.html"
    xy = node_xy(nodes)
    edge_x, edge_y = [], []
    for _, link in links.iterrows():
        f, t = str(link["from_node"]), str(link["to_node"])
        if f in xy.index and t in xy.index:
            edge_x += [xy.loc[f, "x"], xy.loc[t, "x"], None]
            edge_y += [xy.loc[f, "y"], xy.loc[t, "y"], None]
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=edge_x, y=edge_y, mode="lines", line=dict(color="#CBD5E1", width=1), hoverinfo="skip", name="管线"))
    fig.add_trace(go.Scatter(x=xy["x"], y=xy["y"], mode="markers", marker=dict(size=4, color="#AABBD3"), text=xy.index, hovertemplate="节点 %{text}<extra></extra>", name="普通节点"))
    for label, color, items in [("监测点", "#2563EB", MONITOR_NODES), ("真值注入点", "#DC2626", TRUTH_NODES)]:
        xs, ys, texts = [], [], []
        for node in items:
            if node in xy.index:
                xs.append(xy.loc[node, "x"])
                ys.append(xy.loc[node, "y"])
                texts.append(node)
        fig.add_trace(go.Scatter(x=xs, y=ys, mode="markers+text", text=texts, textposition="top center", marker=dict(size=10, color=color, line=dict(width=1, color="#111827")), hovertemplate=f"{label} %{{text}}<extra></extra>", name=label))
    for sol_name, label, visible in [("ga_best", "GA最佳解", True), ("posterior_best_map", "AM最优解", False)]:
        s = solution_shares(shares, sol_name)
        xs, ys, texts, sizes, colors = [], [], [], [], []
        max_w = max(0.001, float(s.max()))
        for node, value in s.items():
            if node in xy.index:
                xs.append(xy.loc[node, "x"])
                ys.append(xy.loc[node, "y"])
                texts.append(f"节点 {node}<br>识别比例 {value * 100:.2f}%")
                sizes.append(8 + 42 * float(value) / max_w)
                colors.append(float(value))
        fig.add_trace(go.Scatter(x=xs, y=ys, mode="markers", text=texts, visible=visible, marker=dict(size=sizes, color=colors, colorscale="YlOrRd", showscale=visible, colorbar=dict(title="比例")), hovertemplate="%{text}<extra>" + label + "</extra>", name=label + "热力"))
    fig.update_layout(
        title="0520中参数管网识别热力图：蓝色为监测点，红色为真值点，热力圆越大表示识别比例越高",
        updatemenus=[dict(buttons=[
            dict(label="GA最佳解", method="update", args=[{"visible": [True, True, True, True, False]}]),
            dict(label="AM最优解", method="update", args=[{"visible": [True, True, True, False, True]}]),
        ], direction="right", x=0.02, y=1.08)],
        template="plotly_white",
        font=dict(family="Microsoft YaHei, SimHei, Arial", size=13),
        height=780,
        legend=dict(orientation="h", yanchor="bottom", y=1.02),
        margin=dict(l=20, r=20, t=90, b=20),
        xaxis=dict(visible=False),
        yaxis=dict(visible=False, scaleanchor="x", scaleratio=1),
    )
    fig.write_html(path, include_plotlyjs="cdn", full_html=True)
    return path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    for style_name in ["Normal", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 2"].font.bold = True


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10.5)


def add_table(doc: Document, df: pd.DataFrame) -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = str(row[col])
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(8.5)


def add_picture(doc: Document, path: Path, caption: str) -> None:
    doc.add_picture(str(path), width=Inches(6.35))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(89, 89, 89)


def build_doc(data: dict[str, object], figs: dict[str, Path], html_path: Path, det: pd.DataFrame, metrics: pd.DataFrame) -> Path:
    summary = data["summary"]
    data_summary = data["data_summary"]
    scores = data["scores"]
    shares = data["shares"]
    am = data["am"]
    doc_path = OUT_DIR / "0520中参数版本结果分析汇报_修正版.docx"
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("0520管网入流溯源中参数版本结果分析汇报")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)
    subtitle = doc.add_paragraph("重点说明：本轮哪些目标已经做到，哪些位置仍存在代偿")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("一、结论摘要", level=1)
    add_para(doc, "关键点1：本次中参数版本在曲线拟合上已经做到较好水平。AM最优解（MAP）的平均NSE为0.936212，平方误差SSE为0.002591，是本轮曲线误差最小的结果。")
    add_para(doc, "关键点2：本次没有稳定做到5个真值点全部精确恢复。按5%比例阈值统计，GA最佳解命中4个真值点，漏掉178；AM最优解命中3个真值点，漏掉103和304。")
    add_para(doc, "关键点3：本轮主要问题是代偿。GA最佳解把一部分水量分给91和308；AM最优解把一部分水量分给64、308和241。模型能较好复现监测曲线，但部分相邻或同路径候选点仍会相互替代。")
    add_table(doc, pd.DataFrame([
        ["曲线拟合", "做到", "AM最优解平均NSE=0.936212，平方误差SSE=0.002591"],
        ["真值回放", "做到", "平均NSE=1.000000，说明注水链路和评分链路闭合"],
        ["全部真值点精确恢复", "未完全做到", "GA命中4/5，AM最优解命中3/5"],
        ["代偿控制", "仍存在", "主要代偿点为91、308、64、241"],
    ], columns=["判断项", "结论", "依据"]))

    doc.add_heading("二、本次实验条件", level=1)
    add_para(doc, f"本次模型采用36小时模拟，时间分辨率为10分钟，共{data_summary['rows']}个输出时刻。注入持续时间为0到24小时，注入过程使用INP文件中的[TIMESERIES]统一写入，没有使用generated_inflow运行时注入。")
    add_para(doc, f"总入流量按雨天排口总流出量减去旱天排口总流出量计算：雨天排口积分为{data_summary['event_outfall_total_volume_m3']:.2f} m3，旱天排口积分为{data_summary['baseline_outfall_total_volume_m3']:.2f} m3，差值为{data_summary['outfall_delta_total_volume_m3']:.2f} m3。五个真值注入点为{'、'.join(TRUTH_NODES)}，每个点折算体积约{list(data_summary['truth_scaled_volumes_m3'].values())[0]:.2f} m3。")
    add_para(doc, f"候选点为20个，监测点为9个。真值回放检查平均NSE={data_summary['truth_replay_mean_nse']:.6f}，平方误差SSE={data_summary['truth_replay_sse']:.2e}，说明当前数据构造和[TIMESERIES]注入路径一致。")
    cfg = summary["config"]
    add_table(doc, pd.DataFrame([
        ["GA种群数", cfg["ga_population_count"]],
        ["GA每种群个体数", cfg["ga_population_size"]],
        ["GA代数", cfg["ga_generations"]],
        ["GA目标", "9个监测点平均NSE最大"],
        ["AM链数", cfg["am_chain_count"]],
        ["AM每链步数", cfg["am_samples_per_chain"]],
        ["AM预热步数", cfg["am_warmup"]],
        ["AM目标", "平方误差SSE最小"],
        ["后验有效样本数", len(am) - cfg["am_chain_count"] * cfg["am_warmup"]],
    ], columns=["项目", "本次设置"]))

    doc.add_heading("三、总体结果", level=1)
    add_para(doc, "本次结果要分两层看：第一层是曲线能不能拟合好，第二层是真值位置能不能被单独找出来。AM最优解的曲线误差最低，但定位上仍有代偿；GA最佳解命中的真值点更多，但拟合误差高于AM最优解。")
    score_table = scores[scores["solution"].isin(["ga_best", "posterior_best_map"])].copy()
    score_table["结果"] = score_table["solution"].map(SOLUTION_LABELS)
    score_table["平均NSE"] = score_table["mean_nse"].map(lambda x: f"{x:.6f}")
    score_table["平方误差SSE"] = score_table["sse"].map(lambda x: f"{x:.6f}")
    score_table["评分口径"] = ["GA阶段按平均NSE排序，数值越高越好", "AM阶段按平方误差SSE选择，数值越低越好"]
    add_table(doc, score_table[["结果", "平均NSE", "平方误差SSE", "评分口径"]])
    add_picture(doc, figs["score"], "图1 核心结果评分对比：左图看平均NSE，右图看平方误差SSE。")

    doc.add_heading("四、GA阶段结果", level=1)
    add_para(doc, f"GA阶段一共评估{len(data['ga_all'])}个候选解，最后一代保留{len(data['ga_last'])}个唯一解。GA目标是让9个监测点的平均NSE尽量高，本轮GA最佳平均NSE={summary['ga_last_score_stats']['max_mean_nse']:.6f}。")
    ga = solution_shares(shares, "ga_best")
    add_para(doc, f"GA最佳解主要比例为：10号{pct(ga['10'])}、42号{pct(ga['42'])}、91号{pct(ga['91'])}、308号{pct(ga['308'])}、103号{pct(ga['103'])}、304号{pct(ga['304'])}。其中10、42、103、304是真值点，91和308属于代偿点，178未被GA最佳解识别出来。")
    add_picture(doc, figs["ga"], "图2 GA逐代收敛过程：最高平均NSE逐步提高，但定位结果仍有代偿。")

    doc.add_heading("五、AM阶段结果", level=1)
    add_para(doc, f"AM阶段共8条链，每条链700步，前180步作为预热步，最终用于分析的有效样本为{len(am) - 8 * 180}个。各链接受率在{min(summary['am_accept_rate_by_chain'].values()):.3f}到{max(summary['am_accept_rate_by_chain'].values()):.3f}之间。")
    mp = solution_shares(shares, "posterior_best_map")
    add_para(doc, f"AM最优解主要比例为：42号{pct(mp['42'])}、178号{pct(mp['178'])}、64号{pct(mp['64'])}、10号{pct(mp['10'])}、308号{pct(mp['308'])}、241号{pct(mp['241'])}。其中42、178、10是真值点，64、308、241为主要代偿点，103和304没有被单独稳定识别。")
    add_picture(doc, figs["am"], "图3 AM多链搜索过程：平方误差SSE越低，表示AM阶段拟合越好。")
    add_picture(doc, figs["fit"], "图4 AM最优解在9个监测点上的拟合曲线。")

    doc.add_heading("六、代偿与定位分析", level=1)
    add_para(doc, "本轮代偿不是单纯的显示问题，而是结果本身反映出的定位不唯一。只要两个候选点到监测点的响应很接近，算法就可能把一部分水量分给相邻或同一路径上的点，同时仍然得到较高拟合分数。")
    add_table(doc, det)
    add_para(doc, "从5%阈值看，GA最佳解命中4个真值点，但保留了91和308两个高比例代偿点；AM最优解命中3个真值点，曲线误差更小，但把部分比例转移到64、308、241。")
    add_picture(doc, figs["share"], "图5 真值、GA最佳解与AM最优解的比例对比。")
    add_picture(doc, figs["net_ga"], "图6 GA最佳解在管网结构上的热力分布：主要代偿为91和308。")
    add_picture(doc, figs["net_map"], "图7 AM最优解在管网结构上的热力分布：主要代偿为64、308和241。")

    doc.add_heading("七、监测点响应分析", level=1)
    add_para(doc, "AM最优解整体拟合较好，但各监测点贡献不均匀。223号监测点拟合最弱，189号监测点响应接近零，对定位贡献有限。")
    metric_table = metrics.copy()
    metric_table["单点NSE"] = metric_table["单点NSE"].map(lambda x: "近零响应" if pd.isna(x) else f"{x:.6f}")
    for col in ["单点SSE", "观测峰值", "模拟峰值", "峰值差"]:
        metric_table[col] = metric_table[col].map(lambda x: f"{x:.6g}")
    add_table(doc, metric_table)

    doc.add_heading("八、本次结论", level=1)
    for item in [
        "本次中参数版本已经能够在曲线层面取得较高拟合，AM最优解的平均NSE为0.936212，平方误差SSE为0.002591。",
        "本次没有稳定恢复全部5个真值点。GA最佳解按5%阈值命中4个真值点，AM最优解按5%阈值命中3个真值点。",
        "代偿主要发生在同一路径或相邻响应区，例如304附近转移到64，103附近转移到308和241。",
        "当前GA不是硬性找5个点，而是在20个候选点上分配总入流比例，因此高拟合分数不能直接等同于精确定位。",
    ]:
        add_para(doc, item)
    add_para(doc, f"配套交互式热力图已生成：{html_path}")
    doc.save(doc_path)
    return doc_path


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    setup_matplotlib()
    data = load_data()
    shares = data["shares"]
    det = detection_table(shares)
    metrics = monitor_metric_table(data["observed_delta"], data["map_delta"])
    det.to_csv(OUT_DIR / "0520中参数识别命中与代偿统计.csv", index=False, encoding="utf-8-sig")
    metrics.to_csv(OUT_DIR / "0520_AM最优解监测点拟合指标.csv", index=False, encoding="utf-8-sig")

    figs = {
        "score": save_score_figure(data["scores"]),
        "ga": save_ga_convergence(data["ga_history"]),
        "share": save_share_figure(shares),
        "am": save_am_trace(data["am"]),
        "fit": save_monitor_fit(data["observed_delta"], data["map_delta"]),
        "net_ga": save_network_figure(data["nodes"], data["links"], solution_shares(shares, "ga_best"), "GA最佳解在管网中的比例热力分布", "06_GA最佳解管网热力图.png"),
        "net_map": save_network_figure(data["nodes"], data["links"], solution_shares(shares, "posterior_best_map"), "AM最优解在管网中的比例热力分布", "07_AM最优解管网热力图.png"),
    }
    html_path = save_network_html(data["nodes"], data["links"], shares)
    doc_path = build_doc(data, figs, html_path, det, metrics)
    manifest = {
        "report_docx": str(doc_path),
        "network_html": str(html_path),
        "figures": {k: str(v) for k, v in figs.items()},
        "tables": {
            "detection": str(OUT_DIR / "0520中参数识别命中与代偿统计.csv"),
            "monitor_metrics": str(OUT_DIR / "0520_AM最优解监测点拟合指标.csv"),
        },
    }
    (OUT_DIR / "manifest_clean.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
