# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = r"E:\PY\research\0519\output\doc\复杂管网智能漏损侦听与定位系统调研报告_一表版.docx"


def set_font(run, size=8.5, bold=False, color=None):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, size=8.5, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width
section.top_margin = Cm(1.2)
section.bottom_margin = Cm(1.2)
section.left_margin = Cm(1.2)
section.right_margin = Cm(1.2)

style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
style.font.size = Pt(8.5)

p = doc.add_paragraph()
r = p.add_run("复杂管网智能漏损侦听与定位系统调研报告 - 一表版")
set_font(r, size=16, bold=True, color="1F4E79")

p = doc.add_paragraph()
r = p.add_run("说明：每行只对应一个需求目标；算法列只写该目标直接使用的算法；参考列只放最贴切的一篇论文或一个案例。")
set_font(r, size=9, color="666666")

headers = [
    "需求目标",
    "实现路径",
    "采集数据",
    "设备/系统",
    "核心算法",
    "综合可行性",
    "最贴切参考论文/案例",
    "链接",
]

rows = [
    [
        "MEMS传感器阵列智能侦听节点",
        "在阀门井、管廊、外露管段或管件处安装低功耗侦听节点；节点本地完成采样、滤波、短时缓存和事件触发；上传声纹片段与基础特征。",
        "管壁振动、声压/水声、采样时间、节点位置、电池电量、通信强度、温湿度。",
        "MEMS三轴加速度计阵列；MEMS麦克风或MEMS水听器；低噪声前端；24-bit ADC；MCU/DSP；NB-IoT/4G通信；IP68外壳。",
        "STFT短时傅里叶变换；Welch功率谱估计；CUSUM变化检测。此处算法用于节点端事件检测，不承担最终AI识别。",
        "可实现工程样机和试点部署。关键在MEMS水听器封装、防水、安装耦合一致性和长期漂移校准。",
        "Micromachines 2023论文：AlN MEMS水听器用于管道泄漏检测，在31 m不锈钢管试验中完成漏/非漏识别和小泄漏定位。",
        "https://www.mdpi.com/2072-666X/14/3/654",
    ],
    [
        "背景噪声抑制能力≥30 dB",
        "节点端先做基础滤波，云端对泄漏目标频带做降噪增强；输出降噪前后SNR，作为噪声抑制结果。",
        "原始声纹、降噪后声纹、目标频带能量、背景噪声频谱、泵阀噪声频率、SNR。",
        "MEMS/压电声学节点；多阵元采样通道；云端降噪服务；频谱分析模块。",
        "VMD变分模态分解 + 小波软阈值降噪 + MVDR波束形成。三者分别用于模态分离、随机噪声抑制和阵列方向增强。",
        "可在目标频带内实现较大SNR提升；全频段环境噪声统一下降30 dB不作为工程口径。",
        "HWM PermaNET SU产品案例：固定噪声记录仪结合机器学习音频分析，用于长期在线漏损监测。",
        "https://www.hwmglobal.com/products/permanet-su/",
    ],
    [
        "工业管网多工况泄漏声纹识别，微小滴漏准确率≥95%",
        "将声纹切成固定长度片段，生成Mel谱图和包络谱；模型输出泄漏、非泄漏、疑似泄漏和噪声源类别；维修确认结果回流训练。",
        "原始音频、Mel谱图、STFT谱图、Hilbert包络谱、管径、管材、压力、流量、泄漏确认标签。",
        "侦听节点；声纹数据库；GPU训练服务器；云端AI推理服务。",
        "CNN-BiLSTM。CNN提取频谱局部纹理，BiLSTM提取泄漏声随时间变化，Softmax输出类别概率。",
        "单企业、已覆盖工况内可达到95%级别；跨企业泛化依赖样本覆盖和维修闭环标签质量。",
        "声学技术论文：供水管网泄漏声信号分类与识别，使用声学特征和分类模型识别泄漏声。",
        "https://www.sxjs.ac.cn/article/doi/10.16300/j.cnki.1000-3630.23122501",
    ],
    [
        "DN300以上管道多节点协同定位，误差≤0.25 m",
        "同一泄漏事件触发3个以上相邻节点同步采样；计算节点间到达时间差；结合管段长度、声速和GIS拓扑输出泄漏点坐标。",
        "多节点同步声纹、节点坐标、节点间管长、管径、管材、压力、互相关峰值、声速标定数据。",
        "多节点声学侦听器；高精度时钟同步模块；云端定位服务；GIS管网系统。",
        "GCC-PHAT时延估计 + TDOA定位 + 加权最小二乘。GCC-PHAT求时延，TDOA求位置，加权最小二乘融合多节点结果。",
        "试验管段和已标定管段可冲击0.25 m；复杂工业现场需依赖声速标定、节点同步和现场复核。",
        "Gutermann ZONESCAN HYDRO案例：水听器节点、NB-IoT上传、云端自动相关定位，公开强调高精度声学相关和≤1 m定位能力。",
        "https://en.gutermann-water.com/product/zonescan-hydro/",
    ],
    [
        "10万组以上泄漏声纹特征数据库",
        "所有节点每日上传固定时段样本和事件样本；音频切片、去重、异常剔除、标签生成、训练集划分后入库。",
        "原始音频、10秒切片、频谱图、节点ID、管段ID、工况、泄漏等级、维修确认、修复后复测结果。",
        "对象存储；关系数据库；时序数据库；数据标注后台；模型训练数据集管理模块。",
        "10秒滑窗切片；SHA-256完全去重；log-Mel余弦相似度近重复删除；Hampel异常剔除；GroupKFold按管段分组划分数据集。",
        "数据量可实现。200个节点每天5条样本，100天形成10万组；质量取决于泄漏维修确认和修复后复测。",
        "FIDO Thames Water案例：使用历史声学文件进行漏损分析，体现大规模声学数据用于AI诊断的行业做法。",
        "https://fido.tech/case-studies/thames-water-leak-team-gets-total-success-from-fido-led-work-orders/",
    ],
    [
        "云端AI漏损诊断服务系统",
        "节点数据进入云端后，依次完成声纹识别、异常告警、协同定位、GIS展示、派单、维修复测和标签回流。",
        "设备状态、声纹识别结果、定位结果、GIS管段、告警记录、工单状态、维修结果、复测声纹。",
        "设备接入服务；声纹数据湖；AI推理服务；定位服务；GIS管网库；告警中心；工单系统。",
        "EWMA时序异常检测 + DBSCAN告警聚类去重。EWMA发现连续异常，DBSCAN把同一空间和时间窗口内的重复告警合并。",
        "可支撑3家以上企业部署。核心工作是多租户数据隔离、企业管网GIS接入、告警规则配置和工单闭环。",
        "无锡市国资委案例：IPv6+噪声听漏物联网设备及AIOT云平台，用于供水噪声监测和漏损治理闭环。",
        "https://gzw.wuxi.gov.cn/doc/2024/10/21/4414908.shtml",
    ],
]

table = doc.add_table(rows=1, cols=len(headers))
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

for i, h in enumerate(headers):
    shade(table.rows[0].cells[i], "1F4E79")
    cell_text(table.rows[0].cells[i], h, size=8.5, bold=True, color="FFFFFF")

for row in rows:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cell_text(cells[i], value, size=8)

doc.save(OUT)
print(OUT)
