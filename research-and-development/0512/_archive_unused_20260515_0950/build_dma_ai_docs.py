from pathlib import Path
import math
import random

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DOC = ROOT / "output" / "doc"
OUT_ASSETS = ROOT / "output" / "assets"
OUT_DOC.mkdir(parents=True, exist_ok=True)
OUT_ASSETS.mkdir(parents=True, exist_ok=True)

FONT_CN = Path(r"C:\Windows\Fonts\simhei.ttf")
FONT = ImageFont.truetype(str(FONT_CN), 34)
FONT_BOLD = ImageFont.truetype(str(FONT_CN), 42)
FONT_SMALL = ImageFont.truetype(str(FONT_CN), 25)
FONT_TINY = ImageFont.truetype(str(FONT_CN), 21)

RED = RGBColor(192, 0, 0)
BLACK = RGBColor(0, 0, 0)
BLUE = (42, 104, 173)
TEAL = (23, 145, 132)
GREEN = (77, 157, 94)
AMBER = (231, 160, 59)
CORAL = (214, 94, 83)
INK = (31, 41, 55)
MUTED = (90, 101, 117)
PAPER = (248, 250, 252)
GRID = (220, 226, 235)


def wrap(draw, text, font, max_width):
    lines, line = [], ""
    for ch in text:
        test = line + ch
        if draw.textlength(test, font=font) <= max_width:
            line = test
        else:
            if line:
                lines.append(line)
            line = ch
    if line:
        lines.append(line)
    return lines


def center_text(draw, box, text, font, fill=INK, spacing=8):
    x1, y1, x2, y2 = box
    lines = []
    for raw in text.split("\n"):
        lines.extend(wrap(draw, raw, font, x2 - x1 - 24))
    h = sum(draw.textbbox((0, 0), line, font=font)[3] for line in lines) + spacing * (len(lines) - 1)
    y = y1 + (y2 - y1 - h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, font=font, fill=fill)
        y += (bbox[3] - bbox[1]) + spacing


def draw_arrow(draw, start, end, fill=INK, width=5):
    draw.line([start, end], fill=fill, width=width)
    ang = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 16
    pts = [
        end,
        (end[0] - size * math.cos(ang - math.pi / 6), end[1] - size * math.sin(ang - math.pi / 6)),
        (end[0] - size * math.cos(ang + math.pi / 6), end[1] - size * math.sin(ang + math.pi / 6)),
    ]
    draw.polygon(pts, fill=fill)


def canvas(title, subtitle=None):
    img = Image.new("RGB", (1600, 900), PAPER)
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, 1600, 96), fill=(18, 32, 53))
    d.text((56, 24), title, font=FONT_BOLD, fill=(255, 255, 255))
    if subtitle:
        d.text((56, 106), subtitle, font=FONT_SMALL, fill=MUTED)
    return img, d


def save_fig(img, name):
    path = OUT_ASSETS / name
    img.save(path, quality=95)
    return path


def fig_ai_history():
    img, d = canvas("人工智能与供水漏损模型演进", "从符号智能、深度学习到智能体，水务模型从规则阈值走向数据-机理协同")
    y = 455
    d.line((110, y, 1490, y), fill=GRID, width=8)
    items = [
        ("1950", "图灵测试\n机器能否表现出智能"),
        ("1956", "Dartmouth\nAI作为学科命名"),
        ("1980s", "专家系统\n规则库与推理机"),
        ("1997", "LSTM\n长序列时序建模"),
        ("2012", "深度学习\n感知能力跃迁"),
        ("2017", "Transformer\n注意力机制"),
        ("2022+", "生成式AI/智能体\n知识与工具协同"),
        ("现在", "DMA漏损AI\n预测、识别、定位、闭环"),
    ]
    colors = [BLUE, TEAL, GREEN, AMBER, CORAL, BLUE, TEAL, GREEN]
    xs = [130 + i * 190 for i in range(len(items))]
    for i, ((year, label), x) in enumerate(zip(items, xs)):
        color = colors[i]
        d.ellipse((x - 26, y - 26, x + 26, y + 26), fill=color)
        d.text((x - 42, y - 94), year, font=FONT_SMALL, fill=color)
        box = (x - 88, y + 48, x + 88, y + 210)
        d.rounded_rectangle(box, radius=18, fill=(255, 255, 255), outline=(205, 213, 224), width=2)
        center_text(d, box, label, FONT_TINY, INK)
    return save_fig(img, "fig01_ai_history_timeline.png")


def fig_dma_architecture():
    img, d = canvas("DMA漏损AI系统总体架构", "数据采集、模型训练、在线推理、工单闭环的工程化链路")
    cols = [
        ("感知层", ["流量计", "压力计", "远传水表", "阀门/泵站", "水质与声学"]),
        ("数据层", ["SCADA", "GIS/管网资产", "DMA边界", "营收抄表", "抢维修工单"]),
        ("模型层", ["LSTM/GRU预测", "孤立森林/DBSCAN", "自编码器", "GBDT/随机森林", "水力模型校准"]),
        ("应用层", ["异常预警", "漏损分区定位", "压力优化", "巡检派单", "复盘迭代"]),
    ]
    x0, gap, w, h = 80, 36, 345, 520
    for i, (title, rows) in enumerate(cols):
        x = x0 + i * (w + gap)
        color = [BLUE, TEAL, AMBER, GREEN][i]
        d.rounded_rectangle((x, 220, x + w, 760), radius=26, fill=(255, 255, 255), outline=color, width=4)
        d.rectangle((x, 220, x + w, 292), fill=color)
        center_text(d, (x, 220, x + w, 292), title, FONT, (255, 255, 255))
        for j, row in enumerate(rows):
            yy = 330 + j * 78
            d.rounded_rectangle((x + 34, yy, x + w - 34, yy + 48), radius=14, fill=(244, 247, 251), outline=(222, 229, 238))
            center_text(d, (x + 34, yy, x + w - 34, yy + 48), row, FONT_SMALL, INK)
        if i < len(cols) - 1:
            draw_arrow(d, (x + w + 6, 490), (x + w + gap - 10, 490), fill=(100, 116, 139), width=5)
    return save_fig(img, "fig02_dma_ai_architecture.png")


def fig_signal():
    img, d = canvas("LSTM漏损预警信号示意", "用预测残差区分正常用水波动和疑似漏损突变")
    left, top, right, bottom = 150, 190, 1460, 720
    d.rectangle((left, top, right, bottom), fill=(255, 255, 255), outline=GRID, width=2)
    for i in range(6):
        y = top + i * (bottom - top) / 5
        d.line((left, y, right, y), fill=(235, 240, 246), width=2)
    random.seed(7)
    points_actual, points_pred = [], []
    n = 120
    for i in range(n):
        x = left + i * (right - left) / (n - 1)
        base = 0.42 + 0.15 * math.sin(i / 8) + 0.05 * math.sin(i / 2.5)
        leak = 0 if i < 72 else 0.22 + min((i - 72) / 120, 0.12)
        noise = random.uniform(-0.025, 0.025)
        pred = base
        actual = base + leak + noise
        points_pred.append((x, bottom - pred * (bottom - top)))
        points_actual.append((x, bottom - actual * (bottom - top)))
    d.line(points_pred, fill=TEAL, width=5)
    d.line(points_actual, fill=CORAL, width=5)
    leak_x = left + 72 * (right - left) / (n - 1)
    d.line((leak_x, top + 15, leak_x, bottom - 15), fill=AMBER, width=4)
    d.text((leak_x + 16, top + 28), "疑似漏损开始", font=FONT_SMALL, fill=AMBER)
    d.rounded_rectangle((1020, 235, 1425, 335), radius=16, fill=(255, 255, 255), outline=GRID, width=2)
    d.line((1045, 270, 1120, 270), fill=TEAL, width=5)
    d.text((1135, 252), "模型预测正常流量", font=FONT_TINY, fill=INK)
    d.line((1045, 305, 1120, 305), fill=CORAL, width=5)
    d.text((1135, 287), "实际监测流量", font=FONT_TINY, fill=INK)
    d.text((left, bottom + 30), "时间", font=FONT_SMALL, fill=MUTED)
    d.text((45, top + 20), "流量/压力", font=FONT_SMALL, fill=MUTED)
    return save_fig(img, "fig03_lstm_signal_residual.png")


def fig_model_matrix():
    img, d = canvas("漏损检测算法选型矩阵", "按标签条件、数据形态和业务目标选择模型")
    headers = ["算法", "输入数据", "适用场景", "主要输出"]
    rows = [
        ["LSTM/GRU", "连续流量、压力", "趋势预测、提前预警", "未来值、残差、风险等级"],
        ["孤立森林", "多维运行特征", "少标签异常识别", "异常分数、异常时段"],
        ["DBSCAN", "夜间流量特征", "异常模式聚类", "异常簇、离群点"],
        ["自编码器", "高维时序片段", "非线性异常检测", "重构误差、异常概率"],
        ["随机森林/GBDT", "工单+管龄+材质+压力", "风险排序、管段评级", "风险分、优先级"],
        ["GNN/机理融合", "拓扑+传感器+仿真", "漏点定位、跨DMA泛化", "候选管段、定位置信度"],
    ]
    x, y, w, h = 80, 180, 1440, 92
    col_w = [260, 330, 370, 480]
    colors = [BLUE, TEAL, AMBER, GREEN]
    for i, head in enumerate(headers):
        x1 = x + sum(col_w[:i])
        d.rectangle((x1, y, x1 + col_w[i], y + h), fill=colors[i])
        center_text(d, (x1, y, x1 + col_w[i], y + h), head, FONT_SMALL, (255, 255, 255))
    for r, row in enumerate(rows):
        yy = y + h + r * 88
        fill = (255, 255, 255) if r % 2 == 0 else (244, 247, 251)
        for c, text in enumerate(row):
            x1 = x + sum(col_w[:c])
            d.rectangle((x1, yy, x1 + col_w[c], yy + 88), fill=fill, outline=GRID)
            center_text(d, (x1 + 8, yy, x1 + col_w[c] - 8, yy + 88), text, FONT_TINY, INK)
    return save_fig(img, "fig04_algorithm_selection_matrix.png")


def fig_hybrid_loop():
    img, d = canvas("水力机理模型 + AI模型双驱动闭环", "用机理解释边界，用AI吸收复杂扰动，用业务闭环校正模型")
    cx, cy, r = 800, 470, 230
    nodes = [
        ("数据治理", 800, 160, BLUE),
        ("水力仿真\nEPANET/WaterGEMS", 1180, 360, TEAL),
        ("AI预测与异常识别", 1030, 710, AMBER),
        ("巡检抢修\n工单反馈", 570, 710, GREEN),
        ("DMA运行策略\n压力/阀门/分区", 420, 360, CORAL),
    ]
    points = []
    for text, x, y, color in nodes:
        points.append((x, y))
        d.rounded_rectangle((x - 160, y - 64, x + 160, y + 64), radius=24, fill=(255, 255, 255), outline=color, width=4)
        center_text(d, (x - 150, y - 58, x + 150, y + 58), text, FONT_SMALL, INK)
    for i in range(len(points)):
        a = points[i]
        b = points[(i + 1) % len(points)]
        draw_arrow(d, a, b, fill=(100, 116, 139), width=4)
    d.ellipse((cx - r, cy - r, cx + r, cy + r), fill=(239, 246, 255), outline=(147, 197, 253), width=4)
    center_text(d, (cx - r + 18, cy - 80, cx + r - 18, cy + 80), "数字孪生底座\n持续学习\n人工复核", FONT, INK)
    return save_fig(img, "fig05_hydraulic_ai_closed_loop.png")


def fig_edge_cloud():
    img, d = canvas("端-边-云协同部署", "低时延预警在边缘侧完成，全局训练和知识沉淀在云端完成")
    boxes = [
        ("端侧传感", "流量计、压力计、RTU、远传水表\n分钟级/秒级数据采集", 100, 270, BLUE),
        ("边缘计算", "数据校验、缺失补齐、轻量异常判断\n断网缓存、本地告警", 510, 270, TEAL),
        ("云端平台", "模型训练、跨DMA对比、知识库\n报表、策略推荐、资产画像", 920, 270, AMBER),
    ]
    for title, desc, x, y, color in boxes:
        d.rounded_rectangle((x, y, x + 320, y + 270), radius=28, fill=(255, 255, 255), outline=color, width=5)
        d.rectangle((x, y, x + 320, y + 70), fill=color)
        center_text(d, (x, y, x + 320, y + 70), title, FONT, (255, 255, 255))
        center_text(d, (x + 18, y + 95, x + 302, y + 248), desc, FONT_SMALL, INK)
    draw_arrow(d, (420, 405), (500, 405), fill=MUTED, width=6)
    draw_arrow(d, (830, 405), (910, 405), fill=MUTED, width=6)
    draw_arrow(d, (1080, 555), (675, 690), fill=GREEN, width=5)
    draw_arrow(d, (675, 690), (260, 555), fill=GREEN, width=5)
    d.text((600, 705), "模型下发 / 参数更新 / 工单闭环", font=FONT_SMALL, fill=GREEN)
    return save_fig(img, "fig06_edge_cloud_deployment.png")


def fig_dma_schematic():
    img, d = canvas("DMA分区监测点位示意", "入口计量、关键压力点、边界阀和疑似漏点联动分析")
    # pipes
    nodes = {
        "A": (280, 300), "B": (520, 260), "C": (760, 330), "D": (1040, 260), "E": (1260, 360),
        "F": (390, 560), "G": (650, 610), "H": (890, 560), "I": (1160, 610)
    }
    edges = [("A", "B"), ("B", "C"), ("C", "D"), ("D", "E"), ("A", "F"), ("B", "F"), ("C", "G"), ("C", "H"), ("D", "H"), ("E", "I"), ("F", "G"), ("G", "H"), ("H", "I")]
    for a, b in edges:
        d.line((nodes[a], nodes[b]), fill=(94, 129, 172), width=18)
        d.line((nodes[a], nodes[b]), fill=(188, 214, 246), width=8)
    for name, (x, y) in nodes.items():
        d.ellipse((x - 24, y - 24, x + 24, y + 24), fill=(255, 255, 255), outline=BLUE, width=5)
        center_text(d, (x - 22, y - 20, x + 22, y + 20), name, FONT_TINY, INK)
    # devices
    d.rounded_rectangle((145, 258, 245, 340), radius=14, fill=TEAL)
    center_text(d, (145, 258, 245, 340), "入口\n流量计", FONT_TINY, (255, 255, 255))
    draw_arrow(d, (245, 300), (255, 300), fill=TEAL, width=5)
    for x, y, label in [(520, 205, "压力P1"), (890, 500, "压力P2"), (1160, 675, "压力P3")]:
        d.rounded_rectangle((x - 54, y - 32, x + 54, y + 32), radius=14, fill=AMBER)
        center_text(d, (x - 54, y - 32, x + 54, y + 32), label, FONT_TINY, (255, 255, 255))
    d.ellipse((805, 445, 885, 525), fill=CORAL, outline=(255, 255, 255), width=4)
    center_text(d, (805, 445, 885, 525), "疑似\n漏点", FONT_TINY, (255, 255, 255))
    d.rounded_rectangle((1120, 235, 1225, 295), radius=12, fill=GREEN)
    center_text(d, (1120, 235, 1225, 295), "边界阀", FONT_TINY, (255, 255, 255))
    d.rounded_rectangle((1020, 175, 1450, 235), radius=12, fill=(255, 255, 255), outline=GRID)
    center_text(d, (1020, 175, 1450, 235), "AI输出：疑似区域 H-C-G，建议声学复核", FONT_TINY, INK)
    return save_fig(img, "fig07_dma_sensor_layout.png")


def make_figures():
    return [
        ("图1 人工智能与供水漏损模型演进", fig_ai_history(), "适合放在1.3，讲AI发展史与水务AI的落点。"),
        ("图2 DMA漏损AI系统总体架构", fig_dma_architecture(), "适合放在2.1或实施方案总览页。"),
        ("图3 LSTM漏损预警信号示意", fig_signal(), "适合讲LSTM/GRU预测残差和预警逻辑。"),
        ("图4 漏损检测算法选型矩阵", fig_model_matrix(), "适合讲2.2算法选型，不同数据条件对应不同模型。"),
        ("图5 水力机理模型 + AI模型双驱动闭环", fig_hybrid_loop(), "适合讲AI与传统水力模型不是替代关系，而是融合关系。"),
        ("图6 端-边-云协同部署", fig_edge_cloud(), "适合讲4.4工程落地架构。"),
        ("图7 DMA分区监测点位示意", fig_dma_schematic(), "适合讲DMA系统数据端部署、定位逻辑和PPT案例页。"),
    ]


def set_doc_defaults(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "黑体"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def set_run(run, color=BLACK, bold=False, size=10.5, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run(r, BLACK, True, 20, "黑体")
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        set_run(r, RED, False, 11, "宋体")


def add_h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        set_run(r, BLACK, True, 16 if level == 1 else 13 if level == 2 else 11.5, "黑体")
    return p


def add_red(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run(r, RED, False, 10.5, "宋体")
    return p


def add_black(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run(r, BLACK, False, 10.5, "宋体")
    return p


def add_red_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run(r, RED, False, 10.5, "宋体")


def add_table(doc, headers, rows, red=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, BLACK, True, 9.5, "黑体")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            r = p.add_run(val)
            set_run(r, RED if red else BLACK, False, 9, "宋体")
    doc.add_paragraph()
    return table


def add_picture(doc, fig, caption, width=6.4, red_caption=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(fig), width=Inches(width))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    set_run(r, RED if red_caption else BLACK, False, 9, "宋体")


def make_main_doc(figs):
    doc = Document()
    set_doc_defaults(doc)
    add_title(doc, "AI模型在供水管网DMA系统漏损检测中的应用", "红色文字为本次补充写入内容；图片为原创PPT插图素材")
    add_red(doc, "写作边界：当前未获得具体历史项目数据、真实DMA点表和模型训练结果，因此本文按典型DMA漏损AI项目经验、LSTM时序预测经验和公开资料进行扩充；涉及具体城市、准确率、降漏率、投资收益等指标，建议后续用项目实测数据替换。")

    add_h(doc, "目录", 1)
    for item in [
        "1. 模型应用背景",
        "2. 技术基础",
        "3. 核心应用",
        "4. 实施路径",
        "5. 典型案例组织方式",
        "6. 总结与展望",
        "7. 参考资料",
    ]:
        add_black(doc, item)

    add_h(doc, "第一部分 模型应用背景", 1)
    add_h(doc, "1.1 供水管网漏损管控的国家战略与行业刚需（简略）", 2)
    add_black(doc, "• 介绍新供水条例的国家政策导向/供水管网漏损管控行业发展趋势")
    add_red(doc, "政策层面，公共供水管网漏损控制已经从企业内部降本事项，升级为城市韧性、水资源节约和基础设施精细化治理事项。住建部、国家发改委在建办城〔2022〕2号文中明确提出推进分区计量、压力调控和智能化建设，并提出到2025年全国城市公共供水管网漏损率力争控制在9%以内的目标[1]。")
    add_red(doc, "业务层面，漏损控制的难点不再只是“有没有检漏队”，而是能否持续回答四个问题：哪里异常、是否真实漏损、先查哪一段、处置后模型是否吸收经验。DMA系统为这些问题提供了边界和计量基础，AI模型则为高频数据解释、异常识别和优先级排序提供了工具。")

    add_h(doc, "1.2 供水管网机理模型的发展（详细）", 2)
    add_black(doc, "• 介绍模型算法演变：哈代克罗斯（平差算法）——牛顿拉夫逊（离散迭代算法）——特征线法（瞬变流常微分算法）")
    add_black(doc, "• 介绍模型软件及适用性：开源EPANET（供水）——开源SWMM（排水和给水支线）——国内商业模型华易、鸿业等——Bentley WaterGEMS/HAMMER——西门子FlowMaster——Autodesk Innovyze和DHI MIKE")
    add_red(doc, "机理模型的核心价值是把管网视为受连续性方程、能量方程、局部阻力和边界条件约束的水力系统。它擅长回答“在已知管网结构、管径、粗糙系数、阀泵状态和用水模式时，流量与压力应当如何分布”。因此，机理模型仍是方案校核、调度推演、压力分区、供水安全评估的基础。")
    add_red(doc, "EPANET是典型开源工具，可进行有压管网水力与水质的延时段仿真，输出管段流量、节点压力、水池水位、水龄和水质组分等结果[2]。商业软件在工程界进一步强化了GIS集成、模型校准、瞬变分析、成果制图和企业级协同能力。")
    add_red(doc, "但在DMA漏损场景中，机理模型的可靠性高度依赖基础数据。管网拓扑、阀门状态、管径材质、粗糙系数、用户用水模式、仪表同步性任一环节偏差，都会传导到漏损判断。AI模型并不是替代机理模型，而是补足其对高频扰动、非线性用水模式、噪声数据和历史工单经验吸收能力不足的问题。")

    add_h(doc, "1.3 供水管网AI模型应用的发展（详细）", 2)
    add_black(doc, "• 介绍AI模型和智能体发展")
    add_black(doc, "• 介绍供水管网漏损领域可适配的算法：如LSTM/GRU时序预测模型、孤立森林、DBSCAN聚类、自编码器异常识别、遗传算法、随机森林、梯度提升树等")
    add_picture(doc, figs[0][1], figs[0][0])
    add_red(doc, "人工智能的发展可以概括为四个阶段。第一阶段是符号智能阶段，重点是把专家规则、逻辑推理和知识表示写入计算机。第二阶段是统计学习阶段，重点是从数据中学习分类、回归和聚类规律。第三阶段是深度学习阶段，神经网络通过多层结构学习复杂特征，LSTM解决长序列依赖问题，CNN强化空间特征提取，Transformer通过注意力机制提升序列建模并支撑大模型发展[3][4][5]。第四阶段是智能体阶段，大模型不只是生成文本，而是连接知识库、业务系统和工具链，承担问答、诊断、报告生成、工单协同等任务。")
    add_red(doc, "供水管网AI模型的发展不是从“大模型”开始，而是从数据化和分区计量开始。早期主要依靠人工经验和阈值规则，例如夜间最小流量超过经验阈值即触发排查。随后进入统计模型阶段，开始使用滑动均值、同比环比、季节性分解和控制图识别异常。再往后进入机器学习阶段，利用流量、压力、管龄、材质、用户类型、历史维修记录等多源特征训练异常识别、风险分级和漏点排序模型。当前趋势是时序深度学习、图神经网络、机理-数据融合和智能体工作流的组合应用。")
    add_red(doc, "在DMA系统中，AI模型主要解决三类问题：第一，预测问题，即根据历史流量、压力、天气、节假日、工作日和用户结构预测正常运行曲线；第二，识别问题，即通过实际值与预测值的残差、异常模式和多点联动关系识别疑似漏损；第三，定位与决策问题，即结合管网拓扑、水力仿真、压力响应和历史工单，把异常收敛到候选片区、候选管段和巡检优先级。")
    add_red(doc, "LSTM/GRU适合处理DMA流量、压力的时间依赖关系。其典型做法是用过去若干小时或若干天的流量、压力、天气和日历特征预测未来短时窗口的正常值，再用“实际值-预测值”的残差判断异常。对漏损而言，模型不一定直接学习“漏点坐标”，更常见的是先学习正常运行模式，一旦出现持续性正残差、夜间底流抬升、压力异常下降或多点相关性改变，再触发漏损风险判断。")
    add_picture(doc, figs[2][1], figs[2][0])
    add_red(doc, "孤立森林、DBSCAN和自编码器适合标签不足的场景。很多水司没有完整、准确、结构化的漏损标签，直接训练监督模型容易过拟合。无监督模型可以先学习“多数正常样本”的分布，再识别离群点、异常簇或高重构误差时段。实际项目中，常用策略是“无监督发现候选异常 + 人工复核 + 工单回填 + 逐步沉淀监督标签”。")
    add_red(doc, "随机森林、梯度提升树和遗传算法适合风险评估与参数优化。树模型可处理管龄、材质、爆管历史、压力区间、道路荷载、土壤环境、维修频次等结构化变量，输出管段风险分和改造优先级。遗传算法、贝叶斯优化等方法可用于水力模型参数校准、压力控制策略寻优和DMA边界方案比选。")
    add_red(doc, "图神经网络和机理融合模型是近年研究重点。供水管网天然是图结构：节点是水表、阀门、用户或监测点，边是管段。GNN可以把拓扑邻接关系纳入模型，使异常识别不只看单点曲线，还看上下游传播关系。公开研究也开始探索将经典算法或水力知识注入图神经网络，以提升跨工况泛化能力[6]。")
    add_red(doc, "大模型和智能体在漏损检测中更适合作为“业务协同层”，而不是替代时序模型的核心检测器。它可以把模型输出翻译为管理语言，自动生成日报、周报、排查建议和复盘材料；也可以连接GIS、SCADA、工单系统和知识库，回答“为什么这个DMA今天异常”“过去相似事件如何处置”“先派哪条巡检路线”等问题。核心原则是：传感器时序判断用专业模型，跨系统解释和协同用智能体。")

    add_h(doc, "1.4 传统机理模型在DMA系统漏损管控应用的痛点分析（详细）", 2)
    add_black(doc, "• 介绍基础数据误差的蝴蝶效应、拓扑结构失真、用水模式时变复杂、压力-漏损耦合低、模型维护成本高")
    add_red(doc, "基础数据误差会沿水力计算链条放大：微小管径、粗糙系数、阀门开度或用户用水模式偏差，可能导致节点压力和管段流量偏移；这些偏移进入夜间最小流量或理论漏损量计算后，会表现为虚高、虚低或漏损定位跑偏。")
    add_red(doc, "拓扑结构失真是DMA应用中的常见问题。实际管网存在临时连通、阀门误开误关、历史竣工资料缺失、用户私接和二供系统边界不清等情况。机理模型若只保持“图纸正确”，而没有在线数据校验，就容易在异常发生时给出过度确定但实际错误的判断。")
    add_red(doc, "用水模式时变复杂。夜间最小流量不等于漏损量，夜间商业用水、二供水箱补水、消防或市政用水、节假日行为变化都会抬升底流。机理模型通常需要人工给定需求模式，而AI模型可以从长期时序中学习不同日期、天气和区域类型下的正常波动边界。")
    add_red(doc, "模型维护成本高。每年管网改造、换管、阀门调整、新增用户、传感器更换都会改变模型参数。纯机理模型需要持续更新拓扑和参数；AI模型则需要持续更新训练数据、标签和阈值。工程上应建立“机理模型变更同步 + AI模型定期再训练 + 工单闭环校验”的运维机制。")

    add_h(doc, "1.5 人工智能模型破解DMA系统漏损管控问题的核心价值（详细）", 2)
    add_black(doc, "数据集搭建——强抗误差、容错性高、对齐时序")
    add_black(doc, "自主学习——复杂规律、非线性耦合场景")
    add_black(doc, "自动决策——端到端自动训练、自动参数寻优、解放人工")
    add_black(doc, "迭代更新——动态学习，机理与AI双驱动")
    add_picture(doc, figs[4][1], figs[4][0])
    add_red(doc, "AI模型的第一项价值是建立可复用的数据资产。漏损项目不应只交付一次性分析结果，而应形成DMA级数据集：统一时间粒度、统一设备编码、统一缺失值和异常值处理规则、统一事件标签口径。数据集越标准，后续模型迁移、复盘和PPT教学展示越容易。")
    add_red(doc, "第二项价值是识别复杂非线性关系。漏损不是单变量问题，流量、压力、管龄、材质、阀泵运行、天气、节假日、用户结构都会共同影响结果。机器学习模型能在一定程度上学习这些变量之间的非线性耦合关系，减少单一阈值带来的误报。")
    add_red(doc, "第三项价值是把人工经验转化为可迭代规则。传统检漏依赖老师傅经验，AI系统可以把“异常曲线-派单-现场核查-维修结果-复盘标签”沉淀为训练样本，使经验从个人能力变成组织能力。")
    add_red(doc, "第四项价值是形成闭环决策。AI模型不是只给一个报警，而是给出异常等级、可能原因、候选区域、建议复核方式、巡检优先级和复盘动作。这样才能支撑从“发现异常”到“降低漏损”的完整业务闭环。")

    add_h(doc, "第二部分 技术基础", 1)
    add_h(doc, "2.1 智慧水务场景下的AI技术核心认知", 2)
    add_black(doc, "• AI技术在水务行业的应用边界与适配逻辑")
    add_black(doc, "• 管网漏损检测场景的AI技术选型核心原则")
    add_black(doc, "• AI模型在DMA系统中搭建的框架思路与要求")
    add_picture(doc, figs[1][1], figs[1][0])
    add_red(doc, "AI不是“自动替人判断一切”，而是把高频、复杂、重复的数据解释任务交给模型，把最终工程判断和资源调度保留给业务人员。漏损检测的核心仍然是水力常识、DMA边界、计量可靠性和现场复核，AI提供的是更早发现、更少漏报、更可追踪的辅助决策。")
    add_red(doc, "选型原则一：先看数据条件，再谈算法先进性。如果只有入口流量和少量压力点，优先使用统计阈值、LSTM预测残差、孤立森林等轻量方法；如果有完整拓扑、仿真模型和多点压力数据，可进一步引入水力-机器学习融合或GNN定位模型。")
    add_red(doc, "选型原则二：先做可解释闭环，再做复杂模型。水司最关心的是“为什么报、报了查哪里、查后如何复盘”。模型输出应包含异常时段、关键特征、对比基线、置信度和建议动作，避免只给一个黑箱分数。")
    add_red(doc, "选型原则三：以DMA为基本建模单元。每个DMA的用户结构、地形高差、管龄材质、压力制度和夜间用水模式不同，应建立分区模型或分区参数，而不是用一个全市统一阈值覆盖所有区域。")
    add_red(doc, "数据端部署建议包括：入口流量计保证准确计量，关键压力点覆盖高低压边界和末梢区域，边界阀状态可核验，SCADA与远传水表时间戳统一，工单系统能回填事件发生时间、位置、原因、处置结果和漏损量估算。")
    add_picture(doc, figs[6][1], figs[6][0])

    add_h(doc, "2.2 适配管网漏损检测的核心AI技术及业务价值", 2)
    add_h(doc, "2.2.1 时序数据分析与异常检测技术", 3)
    add_black(doc, "• 核心算法：LSTM/GRU时序预测模型、孤立森林、DBSCAN聚类、自编码器异常识别")
    add_black(doc, "• 业务价值：实现管网流量、压力数据的动态预测与异常识别，支撑漏损事前预警")
    add_red(doc, "时序模型的关键不是追求单点预测精度，而是稳定刻画“正常运行边界”。当实际曲线持续偏离预测区间，且偏离方向与漏损机理一致，例如夜间流量抬升、压力下降、相邻压力点响应异常，就应触发漏损风险。")
    add_red(doc, "异常检测模型可分为点异常、上下文异常和集体异常。单次毛刺通常是点异常，不一定代表漏损；节假日高用水可能是上下文异常，需要结合日历解释；持续数小时或数天的夜间底流抬升属于集体异常，更接近真实漏损信号。")
    add_h(doc, "2.2.2 机器学习与水力模型融合技术", 3)
    add_black(doc, "• 核心算法：遗传算法、随机森林、梯度提升树")
    add_black(doc, "• 业务价值：优化水力模型计算精度，实现漏点精准定位、DMA分区智能优化")
    add_red(doc, "融合技术的典型路径是：水力模型生成不同漏点、不同漏量、不同阀门状态下的仿真样本；AI模型学习压力响应模式和拓扑传播规律；在线数据进入系统后，与仿真样本库或AI定位模型进行匹配，输出候选漏点区域。")
    add_red(doc, "这种路径特别适合真实漏损标签不足的项目。通过仿真补齐样本，再用少量真实工单校准模型，可以降低直接依赖历史标签的风险。")
    add_h(doc, "2.2.3 大数据融合与智能决策技术", 3)
    add_black(doc, "• 核心技术：多源数据融合、知识图谱、边缘计算+云端AI协同")
    add_black(doc, "• 业务价值：打通全链条数据，实现漏损管控智能决策、巡检路径优化、改造优先级排序、全生命周期管理")
    add_picture(doc, figs[3][1], figs[3][0])
    add_red(doc, "漏损AI不是单一模型，而是一组模型组合。预测模型负责“正常应是多少”，异常模型负责“偏离是否显著”，分类模型负责“更像漏损还是用水扰动”，定位模型负责“在哪个区域”，优化模型负责“先查哪里、如何调压、如何安排改造”。")
    add_table(doc, ["模型类别", "推荐输入", "适用问题", "落地注意点"], [
        ["LSTM/GRU", "流量、压力、天气、日历", "短时预测、残差预警", "需要统一时间粒度，先按DMA分区训练"],
        ["孤立森林/DBSCAN", "夜间流量、压力波动、残差特征", "少标签异常识别", "误报需用工单回填迭代"],
        ["自编码器/LSTM-AE", "多维时序片段", "复杂异常模式识别", "需要足够正常样本，阈值需业务校准"],
        ["随机森林/GBDT", "管龄、材质、维修、压力、道路环境", "管段风险排序", "特征口径要稳定，避免历史工单偏差"],
        ["GNN/机理融合", "拓扑、传感器、仿真样本", "漏点定位、跨区泛化", "依赖拓扑质量和传感器布点"],
        ["智能体", "模型结果、知识库、GIS/工单接口", "报告生成、问答、派单建议", "必须保留人工确认和权限边界"],
    ])

    add_h(doc, "第三部分 核心应用", 1)
    add_h(doc, "3.1 事前防控：基于AI的管网漏损智能预警体系", 2)
    add_black(doc, "• DMA分区智能管控、管网运行状态动态预警、管网健康度智能评估")
    add_red(doc, "事前防控的目标是把漏损从“发现水量损失后再排查”前移到“异常刚形成时就识别”。系统可为每个DMA建立日内基线、周周期基线和季节性基线，结合LSTM预测值和残差阈值，形成蓝、黄、橙、红分级预警。")
    add_red(doc, "健康度评估可从管段维度进行：管龄越长、材质越脆弱、历史维修越频繁、压力波动越大、道路施工扰动越多，风险分越高。模型输出不应只是“高风险”，还应给出主要贡献因素，便于管理人员解释和安排改造。")
    add_h(doc, "3.2 事中处置：基于AI的漏点精准识别与定位", 2)
    add_black(doc, "• 水力模型+AI的漏点精准定位")
    add_red(doc, "事中处置强调缩小排查范围。流程上先由DMA模型发现异常，再由压力点联动和水力模型判断可能影响范围，最后结合工单、声学检漏和现场巡检确认。AI模型的价值是把“全区盲查”变成“候选片区优先查”。")
    add_red(doc, "定位输出建议采用候选清单，而不是单一坐标。例如输出Top 5候选管段、置信度、影响用户数、历史维修次数、推荐复核方式。这样符合工程实际，也能降低模型误差带来的决策风险。")
    add_h(doc, "3.3 事后管控：基于AI的漏损管控闭环优化", 2)
    add_black(doc, "• 漏损事件全生命周期管理：AI记录漏损事件处置全流程，形成案例库，持续优化预警与定位模型")
    add_red(doc, "事后闭环的关键是标签回填。每次异常都应记录：报警时间、DMA编号、异常特征、派单时间、现场核查结果、漏点位置、漏损类型、维修完成时间、估算漏量、误报原因。没有标签回填，AI系统会停留在展示层，无法持续变准。")
    add_h(doc, "3.4 调度类应用场景补充", 2)
    add_red(doc, "AI可用于压力优化调度。基于历史用水预测和压力敏感性分析，模型可识别高压区、低压区和压力波动时段，给出泵站启停、阀门调节、分区调蓄建议。目标不是简单降压，而是在满足服务压力的前提下减少高压漏损和爆管风险。")
    add_red(doc, "AI也可用于抢修资源调度。系统根据漏损风险等级、影响用户数、道路等级、历史漏损概率和队伍位置，生成巡检或抢修优先级。对教学演讲而言，可把它包装为“从发现异常到派单闭环”的典型应用。")
    add_h(doc, "3.5 DMA规划类应用场景补充", 2)
    add_red(doc, "DMA规划可引入优化模型，对分区边界、计量点、压力点和边界阀进行方案比选。评价指标包括分区规模、用户数量、入口数量、管网连通性、压力稳定性、施工成本、后期运维复杂度和漏损识别灵敏度。")
    add_red(doc, "在老城区，DMA边界往往受历史管网和道路条件约束，AI不应直接给出不可施工方案。更务实的做法是先生成若干可实施候选方案，再用模型评估各方案对漏损识别、压力控制和运维成本的影响。")

    add_h(doc, "第四部分 实施路径", 1)
    add_h(doc, "4.1 前期规划：明确建设目标与实施边界", 2)
    add_black(doc, "• 现状诊断、建设目标设定、实施路径规划")
    add_red(doc, "建议把目标拆成三层：管理目标是降低漏损率和无效巡检；技术目标是提升异常识别召回率、降低误报率、缩短发现时间；工程目标是形成可复制的DMA数据接入、模型训练、报警派单和复盘机制。")
    add_h(doc, "4.2 核心底座：多源数据融合与数据治理体系建设", 2)
    add_black(doc, "• 多源数据归集、数据标准化治理、数据中台建设")
    add_red(doc, "数据治理至少包括五项：设备编码统一、时间戳统一、采样粒度统一、异常值处理规则统一、工单标签统一。对LSTM模型而言，时间对齐尤其重要；如果流量是15分钟粒度、压力是5分钟粒度、工单只有日期没有时间，模型训练前必须先做时序重采样和标签修正。")
    add_h(doc, "4.3 模型建设：AI算法模型的选型、训练与优化", 2)
    add_black(doc, "• 业务场景拆解与模型选型、模型训练与验证、模型迭代机制")
    add_red(doc, "模型建设建议采用四步法：第一步建立基线模型，例如夜间最小流量阈值和移动均值；第二步建立时序预测模型，例如LSTM/GRU；第三步引入异常检测和分类模型，区分漏损、用水扰动和设备故障；第四步接入水力模型和GIS，形成定位与派单建议。")
    add_red(doc, "模型评价指标应同时覆盖算法和业务。算法指标包括MAE、RMSE、召回率、精确率、F1、AUC；业务指标包括提前预警时间、误报工单比例、平均定位范围、单次排查成本、维修闭环时间和漏损量下降贡献。")
    add_h(doc, "4.4 工程落地：系统集成与场景化落地", 2)
    add_black(doc, "• 硬件适配与对接、系统集成与开发、端边云协同架构、试点验证与优化")
    add_picture(doc, figs[5][1], figs[5][0])
    add_red(doc, "端侧负责稳定采集，边缘侧负责快速校验和本地告警，云端负责模型训练、跨DMA对比、知识沉淀和报表输出。对关键DMA，可在边缘侧部署轻量模型，保障网络中断时仍能完成基础异常判断。")
    add_red(doc, "试点建议选择数据质量较好、边界清晰、历史漏损事件较多、运维队伍配合度高的DMA。先用3到6个月跑通闭环，再扩展到更多分区。不要一开始追求全市覆盖，否则容易把主要精力消耗在数据接入和设备问题上。")
    add_h(doc, "4.5 长效运营：人员能力建设与运维保障体系", 2)
    add_black(doc, "• 人员能力建设、运维保障体系")
    add_red(doc, "人员培训应分层：管理层关注指标、投入产出和治理闭环；调度人员关注报警解释和策略执行；巡检人员关注候选区域、复核方法和反馈口径；算法人员关注数据质量、模型漂移和再训练。")
    add_red(doc, "长效运营需要设定模型巡检制度。建议每月复盘误报、漏报和高价值案例；每季度检查模型漂移和阈值有效性；每半年根据新增工单和管网改造情况进行再训练或参数更新。")

    add_h(doc, "第五部分 典型案例组织方式", 1)
    add_black(doc, "典型案例（原始素材和案例省规院提供，包装南大五维）")
    add_red(doc, "当前文档未提供真实案例素材，因此不建议虚构具体城市和效果数据。对外教学演讲可采用“教学型案例”组织：背景为某老城区DMA夜间最小流量持续抬升；数据包括入口流量、3个压力点、历史维修工单和GIS管网；模型采用LSTM预测残差+孤立森林异常识别+水力模型候选区定位；输出为异常等级、候选管段和巡检路线；复盘为维修结果回填并更新样本库。")
    add_red(doc, "案例PPT建议按五页展开：一页讲问题，一页讲数据，一页讲模型，一页讲处置闭环，一页讲价值指标。真实项目指标补齐后，可替换为“提前发现X小时、排查范围缩小X%、误报率下降X%、估算节水X吨”等量化表达。")

    add_h(doc, "第六部分 总结与展望", 1)
    add_h(doc, "6.1 总结", 2)
    add_black(doc, "• 核心认知：AI是管网漏损管控从“被动处置”向“主动防控”转型的核心驱动力")
    add_black(doc, "• 核心逻辑：以数据为基础、以算法为核心、人工智能与传统水力模型相结合、以业务场景为导向、解决实际问题")
    add_black(doc, "• 实施路径：先规划后建设、先试点后推广、兼顾技术先进性与业务实用性")
    add_red(doc, "本次扩充的核心观点是：DMA提供边界，传感器提供数据，机理模型提供水力约束，AI模型提供复杂模式识别，智能体提供跨系统协同。四者结合，才能把漏损管控从经验驱动升级为数据和模型驱动。")
    add_h(doc, "6.2 未来展望", 2)
    add_black(doc, "• 技术升级、架构升级、体系升级、行业升级")
    add_red(doc, "未来一段时间，供水漏损AI会向四个方向发展：一是更强的机理-数据融合，减少黑箱模型误判；二是更细颗粒度的边缘智能，实现近实时预警；三是更完整的知识库和智能体，让模型结果能自动转化为业务动作；四是更标准化的数据和评价体系，降低中小水司落地门槛。")

    add_h(doc, "第七部分 参考资料", 1)
    refs = [
        "[1] 住房和城乡建设部办公厅、国家发展改革委办公厅. 关于加强公共供水管网漏损控制的通知（建办城〔2022〕2号）. https://www.mohurd.gov.cn/gongkai/zc/wjk/art/2022/art_17339_764316.html",
        "[2] US EPA. EPANET: Application for Modeling Drinking Water Distribution Systems. https://www.epa.gov/water-research/epanet",
        "[3] Dartmouth. The birthplace of AI. Now shaping what comes next. https://home.dartmouth.edu/artificial-intelligence-dartmouth",
        "[4] Hochreiter S., Schmidhuber J. Long Short-Term Memory. Neural Computation, 1997. https://direct.mit.edu/neco/article/9/8/1735/6109/Long-Short-Term-Memory",
        "[5] Vaswani A. et al. Attention Is All You Need. arXiv:1706.03762, 2017. https://arxiv.org/abs/1706.03762",
        "[6] Zhang Z., Fink O. Algorithm-Informed Graph Neural Networks for Leakage Detection and Localization in Water Distribution Networks. arXiv:2408.02797, 2024. https://arxiv.org/abs/2408.02797",
        "[7] Adedeji K. B. et al. Leak detection in water distribution networks: an introductory overview. Smart Water, 2019. https://link.springer.com/article/10.1186/s40713-019-0017-x",
        "[8] Gómez-Coronel L. et al. Leak localization in an urban water distribution network using a LSTM deep neural network. IFAC-PapersOnLine, 2024. https://www.sciencedirect.com/science/article/pii/S2405896324002817",
    ]
    for ref in refs:
        add_red(doc, ref)

    path = OUT_DOC / "AI模型在供水管网DMA系统漏损检测中的应用_补充红色标注版.docx"
    doc.save(path)
    return path


def make_figure_doc(figs):
    doc = Document()
    set_doc_defaults(doc)
    add_title(doc, "AI模型在供水管网DMA系统漏损检测中的应用：插图素材文档", "所有插图均为本次原创生成，可直接用于后续PPT排版")
    add_black(doc, "素材使用建议：PPT中建议按“背景-架构-算法-案例-实施”顺序使用；如需要改色或导出透明底，可基于 output/assets 下PNG重绘。")
    for i, (caption, path, usage) in enumerate(figs, 1):
        add_h(doc, f"{i}. {caption}", 2)
        add_picture(doc, path, caption, width=6.6, red_caption=False)
        add_black(doc, f"文件路径：{path}")
        add_black(doc, f"PPT用途：{usage}")
    path = OUT_DOC / "AI模型在供水管网DMA系统漏损检测中的应用_插图素材文档.docx"
    doc.save(path)
    return path


def main():
    figs = make_figures()
    main_doc = make_main_doc(figs)
    fig_doc = make_figure_doc(figs)
    print(main_doc)
    print(fig_doc)
    for _, path, _ in figs:
        print(path)


if __name__ == "__main__":
    main()
