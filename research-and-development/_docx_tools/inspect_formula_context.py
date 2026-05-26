from pathlib import Path
from docx import Document


base = Path(r"E:\PY\research\0511")
draft = base / "LDHS2026F02613--一种深度水处理斜板沉淀池硬度调控方法及系统--发明--初稿.docx"
revised = base / "LDHS2026F02613--一种深度水处理斜板沉淀池硬度调控方法及系统--发明--批注修编_黄色标注.docx"

print("DRAFT", draft.exists(), draft)
d = Document(draft)
for i in range(24, 34):
    para = d.paragraphs[i - 1]
    print("\nPARA", i, repr(para.text))
    print(para._p.xml[:2200])

print("\nREVISED", revised.exists(), revised)
d = Document(revised)
for i in [162, 163, 164, 171, 172, 174, 182, 183, 184]:
    print(i, d.paragraphs[i - 1].text[:300])
