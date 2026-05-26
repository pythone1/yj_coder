# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = r"E:\PY\research\0519\output\doc\复杂管网智能漏损侦听与定位系统调研报告_中外案例版.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        for r in paragraph.runs:
            r.font.name = "微软雅黑"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            r.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(90, 90, 90)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        r.font.color.rgb = RGBColor(31, 78, 121)


def add_para(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(10.5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        r.font.size = Pt(10)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], "1F4E79")
        set_cell_text(table.rows[0].cells[i], h, bold=True, color="FFFFFF")
        if widths:
            table.rows[0].cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
            if widths:
                cells[i].width = Cm(widths[i])
    return table


sources = [
    ("C1", "无锡市国资委：IPv6+噪声听漏物联网设备及AIOT云平台案例", "https://gzw.wuxi.gov.cn/doc/2024/10/21/4414908.shtml"),
    ("C2", "无锡日报：无锡供水管网漏损率持续下降", "https://www.wxrb.com/doc/2023/05/15/276026.shtml"),
    ("C3", "济宁中山公用水务分区计量系统上线报道", "https://mhuanbao.in-en.com/html/huanbao-2397432.shtml"),
    ("C4", "济宁市政府：济宁中山公用水务智慧化建设基本信息", "https://www.jining.gov.cn/art/2024/5/27/art_79240_2800928.html"),
    ("C5", "济宁中山公用水务DMA系统降低漏损案例", "https://old.cuwa.org.cn/guoneizixun/18744.html"),
    ("C6", "新浪：济宁中山公用水务智慧水务平台与噪声监测", "https://kandian.sina.cn/article_1893761531_70e081fb02002y6cy.html?from=news&subch=onews"),
    ("C7", "济南水务漏损治理报道", "https://www.dxguanxian.org/index.php/news/showNews/nid/5797.shtml"),
    ("C8", "中国企业报：济南水务数智转型与管网漏控研讨会", "https://www.zqbao.com.cn/news/13733.html"),
    ("C9", "和达科技官网：渗漏预警、水务云和漏损控制方案", "https://www.hddznet.com/"),
    ("C10", "长江网：和达科技智能检漏设备报道", "https://news.cjn.cn/csqpd/wh_20004/202303/t4500049.htm"),
    ("C11", "威派格：NB远传型噪声监测仪", "https://www.shwpg.com/intelligent-hardware/pro9/pro39"),
    ("C12", "威派格：智慧水务解决方案", "https://www.shwpg.com/solution/smart-water"),
    ("C13", "威派格：数据驱动漏损决策模型", "https://www.shwpg.com/about-weipaig/news/70.html"),
    ("C14", "新天科技：DMA漏损管控分析", "https://wap.suntront.com/product/product24.html"),
    ("C15", "新天科技：DMA漏控咨询服务", "https://www.suntront.com/WisdomWater/html/ConsultingPlanning/DMA.html"),
    ("C16", "嘉定自来水DMA小区漏损智慧管控案例", "https://m.thepaper.cn/newsDetail_forward_28877298"),
    ("C17", "住房城乡建设部/发改委：加强公共供水管网漏损控制通知", "https://www.mdx.gov.cn/__local/B/F4/C3/674943FF141A3D09868CE8F6FBE_01E624FA_2570F.pdf"),
    ("C18", "国务院：城市公共供水管网漏损治理可复制政策机制清单", "https://www.gov.cn/zhengce/zhengceku/202402/P020240223404578546694.pdf"),
    ("C19", "标准资料：CJJ 92-2016 城镇供水管网漏损控制及评定标准", "https://www.biaozhun.org/hangye/85446.html"),
    ("C20", "北京建筑大学/北京自来水等：微泄漏声学检测与定位论文", "https://pubs.rsc.org/doi/d3ew00686g"),
    ("C21", "声学技术：供水管网泄漏声信号分类与识别", "https://www.sxjs.ac.cn/article/doi/10.16300/j.cnki.1000-3630.23122501"),
    ("C22", "无锡水务/威派格：使用噪声记录仪控制漏损的方法探究", "https://twsds.org.tw/wp-content/uploads/2023/11/%E3%80%8A%E5%BB%BA%E7%AD%91%E7%BB%99%E6%B0%B4%E6%8E%92%E6%B0%B4%E3%80%8B2023%E5%B9%B4%E7%AC%AC1%E6%9C%9F.pdf"),
    ("C23", "MDPI：AlN MEMS水听器管道泄漏监测", "https://www.mdpi.com/2072-666X/14/3/654"),
    ("C24", "MDPI：AlN MEMS水听器IoT泄漏检测系统", "https://www.mdpi.com/2073-4441/12/11/2966"),
    ("C25", "MDPI：MEMS麦克风阵列管道泄漏检测集成系统", "https://www.mdpi.com/2072-666X/17/1/140"),
    ("G1", "Gutermann ZONESCAN HYDRO", "https://en.gutermann-water.com/product/zonescan-hydro/"),
    ("G2", "HWM PermaNET SU", "https://www.hwmglobal.com/products/permanet-su/"),
    ("G3", "Ovarro Enigma", "https://ovarro.com/en/global/solutions/monitoring--control-devices/data-loggers--leak-noise-loggers/leak-noise-loggers--correlators/3/enigma/2/"),
    ("G4", "Mueller Systems EchoShore-DX", "https://muellersystems.com/leak-detection-echoshoredx/"),
    ("G5", "FIDO DC Water传输干管案例", "https://fido.tech/case-studies/dc-water-detects-unseen-water-leak-on-transmission-main-using-fido-ai/"),
    ("G6", "FIDO AI", "https://fido.tech/fido-ai/"),
    ("G7", "AP Sensing管线监测", "https://www.apsensing.com/en/application/process-automation-and-pipeline-monitoring/pipeline-monitoring"),
    ("G8", "Badger Meter PipeMinder ONE", "https://www.badgermeter.com/products/pressure-and-leak-monitoring/transient-management/pipeminder-one-standard/"),
]


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

styles = doc.styles
styles["Normal"].font.name = "微软雅黑"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
styles["Normal"].font.size = Pt(10.5)

add_title(
    doc,
    "复杂管网智能漏损侦听与定位系统调研报告",
    "中外行业方案、本土案例、MEMS声学阵列、AI声纹识别与DN300以上定位可行性评估 | 2026-05-19",
)

add_heading(doc, "一、图标化结论", 1)
add_para(doc, "图例：● 低风险/可行；▲ 有条件可行；◆ 中高风险；■ 高风险。")
summary_rows = [
    ("● 云端AI诊断与3家部署", "可行", "国内外智慧水务平台已普遍覆盖设备接入、GIS、DMA、告警、工单和SaaS化服务。"),
    ("● 10万组样本库", "可行", "数据量本身不难；难点是维修闭环标签和跨企业数据标准。"),
    ("▲ MEMS阵列侦听节点", "有条件可行", "MEMS水听器、加速度计、麦克风阵列已有论文原型；工程量产仍需封装、防水、温漂和现场耦合验证。"),
    ("▲ 微小滴漏识别≥95%", "有条件可行", "本域、已标定管段可达；跨管材、跨噪声、跨企业泛化需要主动学习和企业域微调。"),
    ("◆ 背景噪声抑制≥30 dB", "中高风险", "建议定义为目标频带内SNR提升或等效检出率提升，不建议定义为全频带绝对降噪。"),
    ("■ DN300以上误差≤0.25 m", "高风险", "固定在线系统公开指标更接近1 m级；0.25 m应作为现场复核或标定试验段指标。"),
]
add_table(doc, ["指标", "判断", "依据"], summary_rows, [4.2, 3.0, 9.5])

add_heading(doc, "二、中国本土方案与案例", 1)
add_heading(doc, "2.1 本土行业路线", 2)
add_para(
    doc,
    "中国本土漏损治理目前以DMA分区、夜间最小流量、压力/流量监测、GIS、工单闭环为底座；声学噪声记录仪和AIOT平台正在从辅助巡检升级为主动预警。与海外声学固定相关产品相比，本土公开案例更强调平台集成、分区管理、政策达标和运营闭环，公开披露的声学定位精度、微漏准确率和训练数据规模较少。",
)
add_bullets(
    doc,
    [
        "政策牵引：2022年住建部、发改委要求加强公共供水管网漏损控制；2024年国务院发布可复制政策机制清单，强调老旧管网改造、分区计量、周期检测和精细化管控。[C17][C18]",
        "标准底座：CJJ 92-2016适用于城镇供水管网漏损分析、控制及评定；其条文说明强调管网漏点监测设备和数据分析系统对主动监测、诊断评估的价值。[C19]",
        "工程形态：水司更倾向采购“硬件+平台+运维服务”，而不是单点声学算法；典型组合为智能水表/流量计、压力计、噪声记录仪、GIS、DMA、移动检漏和工单闭环。",
    ],
)

local_rows = [
    ("无锡市水务集团", "IPv6+噪声听漏物联网设备及AIOT云平台", "融合物联网噪声监测和AI智能分析，构筑供水噪声AIOT智慧平台，提供SaaS化服务，实现漏损监控治理闭环。[C1]", "本项目可参考其“设备+AIOT云平台+闭环服务”模式。公开资料未披露30 dB降噪或0.25 m定位指标。"),
    ("济宁中山公用水务", "DMA分区计量+流量压力+噪声监测", "以DMA为最小管理单元，结合夜间最小流量、总分表分析、流量压力和管道噪声监测，支撑快速定位和精准判别。[C3][C6]", "适合作为国内水司落地样板：先做区域定位，再用声学设备现场复核。"),
    ("济南水务集团", "噪声控漏体系+AI和数据模型", "公开报道提到结合AI和数据模型打造噪声控漏体系，2024年减少漏水量近300万立方米；研讨会中分享AI人工智能体系建设与落地成效。[C7][C8]", "证明AI+噪声控漏在国内大水司有应用牵引，但细分算法指标需现场验证。"),
    ("上海嘉定自来水", "多级DMA小区漏损智慧管控", "构建公司、管理站、一级分区和四级分区组成的分区计量体系，结合总分表分析掌握最小单元格水量变化。[C16]", "本项目在工业园区应先建分区和边界计量，否则声学告警难以形成治理闭环。"),
    ("和达科技", "渗漏预警云平台+智能检漏设备", "官网显示覆盖水务云、分区计量、GIS、水力模型、渗漏预警系列；报道提到噪声水音监测仪报警后，现场设备以高灵敏振动检测、无线音频传输和相关定位算法指导施工。[C9][C10]", "本土“平台+噪声水音+现场相关定位”路线清晰，适合对标产品化。"),
    ("威派格", "NB远传噪声监测仪+漏损决策模型", "噪声监测仪为NB远传型，用于大面积供水网络监测；智慧水务方案强调传感器采集、水平衡分析和漏损控制能力；公开内容提到数据驱动漏损决策模型。[C11][C12][C13]", "可参考其NB-IoT低功耗硬件和数据驱动漏损决策方案。"),
    ("新天科技", "DMA漏损管控分析+咨询服务", "系统按IWA管网漏损控制策略和DMA方法研发，提供漏损数据分析、决策支持和咨询服务。[C14][C15]", "偏计量和管理决策路线，可与声学AI模块互补。"),
]
add_table(doc, ["主体", "方案", "公开资料要点", "对本项目启示"], local_rows, [2.8, 3.4, 6.1, 5.0])

add_heading(doc, "2.2 国内本土结论", 2)
add_bullets(
    doc,
    [
        "本土方案的成熟度在“DMA+流量压力+平台闭环”最高，在“声纹AI精度公开验证”和“MEMS阵列量产”较弱。",
        "本项目若定位为工业供水管网，应避免只做市政水司DMA复制，应强化工业泵站、电机、阀门、长距离输送和大口径管段的噪声库。",
        "本土客户通常关心漏损率下降、少开挖、工单闭环和监管报表；验收指标应把技术指标转化为“告警命中率、误报率、平均定位误差、修复闭环时长、节水量”。",
    ],
)

add_heading(doc, "三、海外头部产品对标", 1)
global_rows = [
    ("Gutermann ZONESCAN HYDRO", "水听器+NB-IoT+云端每日相关", "公开强调高精度相关、≤1 ms同步、≤1 m定位能力。[G1]", "可作为本项目多节点TDOA和水听器增强模式对标。"),
    ("HWM PermaNET SU", "固定噪声记录仪+蜂窝通信+云端", "公开强调永久在线、机器学习音频分析和AI相关。[G2]", "说明ML已成为商业声学漏损监测标配。"),
    ("Ovarro Enigma", "多点噪声记录/水听器+云端相关", "面向漏点噪声记录和相关定位。[G3]", "与本项目“多节点协同定位”技术路径一致。"),
    ("Mueller/Echologics EchoShore", "固定声学监测+云端分析", "覆盖小口径和大口径漏损监测产品线。[G4]", "大口径管线可采用专用水听器/声学节点，不宜只依赖外贴传感。"),
    ("FIDO AI", "传感器采样+AI云服务", "DC Water案例覆盖24英寸、12英寸和8英寸传输管，发现非地表漏。[G5][G6]", "证明AI声纹在复杂真实管网有商业落地，但公开资料不等同于可复现验收指标。"),
    ("AP Sensing DAS", "分布式光纤声学/温度/应变", "用于长距离管线泄漏和第三方入侵监测。[G7]", "工业长输干线若已有光缆，应评估DAS作为增强层。"),
    ("Badger Meter/Syrinix PipeMinder", "高分辨压力瞬变监测", "面向瞬变、爆管、压力事件和管网平静化。[G8]", "作为声学AI的辅助输入，可降低误报。"),
]
add_table(doc, ["产品", "架构", "公开能力", "对本项目启示"], global_rows, [3.4, 4.2, 5.8, 4.2])

add_heading(doc, "四、MEMS阵列与30 dB噪声抑制", 1)
add_para(
    doc,
    "学术研究已经覆盖MEMS水听器、MEMS加速度计、MEMS麦克风阵列和PMUT/CMUT类器件。2023年AlN MEMS水听器论文在31 m不锈钢管试验中报告了约97.69%的漏/非漏分类准确率，并给出小泄漏定位相对误差；2025年MEMS麦克风阵列集成系统论文表明MEMS-MASI路线仍处在快速发展阶段。[C23][C25]",
)
add_bullets(
    doc,
    [
        "硬件可行点：MEMS低功耗、小尺寸、易阵列化，适合常设节点；水听器对大口径和弱信号更有利。",
        "工程短板：封装、防水、长期漂移、安装耦合、现场一致性和强工业噪声下SNR仍是主要风险。",
        "30 dB判断：16阵元理论阵列增益约12 dB，64阵元约18 dB；若再叠加窄带增强、自适应滤波、深度降噪、多夜累积和多节点相关，特定频带可接近30 dB等效提升。",
        "验收建议：将指标定义为“泄漏目标频带内SNR提升≥30 dB”或“等效检出率提升”，不要定义为全频段环境噪声下降30 dB。",
    ],
)

add_heading(doc, "五、AI声纹识别与10万组数据库", 1)
add_para(
    doc,
    "国内论文已出现CNN、残差块、注意力机制、VMD/小波去噪等用于供水管网泄漏声信号分类识别的研究；北京建筑大学、北京自来水等团队2024年发表的微泄漏声学检测与定位研究覆盖低压、小孔径和三通场景，给出相对误差约1%的定位结果。[C20][C21]",
)
add_bullets(
    doc,
    [
        "95%准确率在实验室和单企业本域数据集有可行性；跨企业、跨管材、跨噪声场景需用未见站点测试，不能只用随机切分。",
        "建议指标从单一准确率改为：微漏召回率、误报率、F1、AUC、平均提前发现时间和维修闭环命中率。",
        "10万组样本库可通过200-500个节点在1-4个月内积累原始音频，但高置信标签需要6-12个月运维闭环。",
        "数据治理需保存原始波形、谱图、设备状态、压力/流量、管材管径、安装位置、天气/工况、人工复核和修复后复测结果。",
    ],
)

add_heading(doc, "六、DN300以上多节点协同定位", 1)
add_para(
    doc,
    "两节点TDOA定位可用 x=(L+vΔt)/2 表示。若声速取1000-1400 m/s，要把时间误差贡献控制在0.25 m内，时间差估计误差需小于约0.36-0.50 ms；同时还要求GIS管长、节点同步、有效波速和传感器耦合误差都很小。海外固定声学产品公开精度更常见为1 m级，0.25 m对DN300以上工业管道属于高风险目标。",
)
add_bullets(
    doc,
    [
        "理论瓶颈：有效波速受管材、内衬、压力、埋设条件和支路影响；强噪声会使互相关峰值展宽或偏移。",
        "工程瓶颈：DN300以上泄漏声衰减、多径、泵阀噪声、并行管道串扰、管网GIS误差都会放大定位误差。",
        "推荐验收拆分：云端自动初定位≤1 m；现场临时高密度节点/水听器/相关仪复核≤0.25 m；标定试验管段可挑战自动≤0.25 m。",
    ],
)

add_heading(doc, "七、推荐综合方案", 1)
arch_rows = [
    ("① 感知层", "MEMS/压电振动、水听器、压力瞬变、流量计、智能表、DAS可选", "覆盖低成本广域和关键干线增强。"),
    ("② 边缘层", "夜间采样、带通滤波、事件触发、低功耗唤醒、特征压缩", "减少通信和云存储成本。"),
    ("③ 网络层", "NB-IoT/LTE-M/4G/LoRaWAN/有线光纤", "按井下信号、功耗和企业网络选型。"),
    ("④ 云平台", "MQTT接入、时序库、对象存储、GIS、DMA、水力模型、AI诊断", "形成设备-管段-告警-工单一体化。"),
    ("⑤ AI算法", "CNN/CRNN/Transformer、GCC-PHAT、多节点RANSAC、波速校准", "检测、分级、定位和复核联动。"),
    ("⑥ 闭环处置", "派单、现场复核、维修、修后复测、标签回流、模型更新", "把AI系统变成持续变好的运营系统。"),
]
add_table(doc, ["模块", "配置", "作用"], arch_rows, [3.0, 7.2, 5.8])

add_heading(doc, "八、2027目标可行性与里程碑", 1)
milestone_rows = [
    ("0-6个月", "实验室+试验管段", "完成MEMS/压电/水听器对比，定义微漏流量范围，建立DN100-DN500测试台。"),
    ("6-12个月", "工业现场试点", "选择2-3家企业，每家50-150个节点，采集真实噪声与漏损闭环标签。"),
    ("12-18个月", "规模化优化", "完成10万组样本库、企业域模型微调、DN300以上水听器/临时相关增强。"),
    ("18-24个月", "验收部署", "完成3家外部企业部署，形成第三方测试、SaaS运维、API和标准交付包。"),
]
add_table(doc, ["阶段", "任务", "验收重点"], milestone_rows, [3.0, 4.5, 8.5])

add_heading(doc, "九、最终建议", 1)
add_bullets(
    doc,
    [
        "项目定位应从“单一听漏设备”升级为“工业管网漏损AI运营系统”。",
        "本土交付应先打通DMA/GIS/工单/维修闭环，再叠加声纹AI和MEMS阵列，否则样本标签和ROI难成立。",
        "DN300以上大口径目标建议采用“双层定位”：常设节点做广域预警，临时高精度水听器/相关仪/DAS做0.25 m复核。",
        "技术指标要有测试协议：管径、管材、压力、泄漏孔径、噪声工况、节点间距、采样率、同步方式、未见站点测试集。",
        "2027年前可实现有竞争力的产品化版本，但0.25 m自动定位和30 dB通用降噪需要作为高风险攻关项管理。",
    ],
)

add_heading(doc, "附录：参考链接", 1)
for sid, name, url in sources:
    p = doc.add_paragraph()
    r = p.add_run(f"[{sid}] {name}: ")
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(9)
    add_hyperlink(p, url, url)

doc.save(OUT)
print(OUT)
