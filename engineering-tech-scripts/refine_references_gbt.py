# -*- coding: utf-8 -*-
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE = r"D:\Users\xwechat_files\wxid_4668346683612_4126\msg\file\2026-04"
SRC_NAME = "参考文献.docx"
OUT_NAME = "参考文献_按年份梳理_细修版.docx"


def set_run_font(run, size=10.5, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text="", size=10.5, bold=False, align=None, hanging=False):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold)
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(20)
    if hanging:
        fmt.left_indent = Cm(0.74)
        fmt.first_line_indent = Cm(-0.74)
    return p


def read_entries():
    src = os.path.join(BASE, SRC_NAME)
    doc = Document(src)
    entries = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t or t == "参考文献":
            continue
        if "学位论文的撰写应本着严谨求实的科学态度" in t:
            break
        if t.startswith("[1] Sergio Manzettiab") or t.startswith("[2]***"):
            continue
        if any(
            t.startswith(prefix)
            for prefix in [
                "引用文献的作者不超过3位",
                "连续出版物",
                "专著",
                "论文集",
                "学位论文",
                "专利",
                "技术标准",
            ]
        ):
            continue
        entries.append(t)
    return entries


def clean_spaces(text):
    text = re.sub(r"\s+", " ", text).strip()
    text = text.replace(" .", ".")
    text = text.replace(" ,", ",")
    text = text.replace(" :", ":")
    text = text.replace(" ;", ";")
    return text


def normalize_basic(text):
    text = text.strip()
    text = text.replace("．", ".").replace("，", ",").replace("：", ":").replace("；", ";")
    text = text.replace("（", "(").replace("）", ")")
    text = text.replace("【", "[").replace("】", "]")
    text = text.replace("[[M]", "[M]").replace("[[J]", "[J]").replace("[[D]", "[D]").replace("[[S]", "[S]")
    text = text.replace("[s]", "[S]").replace("[m]", "[M]").replace("[j]", "[J]").replace("[d]", "[D]")
    text = text.replace(" andDaiyo", " and Daiyo")
    text = text.replace("P ,", "P,")
    text = text.replace(" H ,", " H,")
    text = clean_spaces(text)
    return text


def normalize_key(entry):
    e = normalize_basic(entry)
    e = re.sub(r":\d+(?:-\d+)?\.?$", "", e)
    e = re.sub(r"\s+", "", e)
    e = e.replace('"', "")
    e = e.replace("'", "")
    return e.lower()


def dedupe(entries):
    best = {}
    for entry in entries:
        c = normalize_basic(entry)
        key = normalize_key(c)
        if key not in best or len(c) > len(best[key]):
            best[key] = c
    return list(best.values())


def year_of(entry):
    years = re.findall(r"(?:19|20)\d{2}", entry)
    return int(years[-1]) if years else 9999


def type_of(entry):
    m = re.search(r"\[([A-Za-z])\]", entry)
    return m.group(1).upper() if m else ""


def is_chinese(entry):
    chinese = re.findall(r"[\u4e00-\u9fff]", entry)
    latin = re.findall(r"[A-Za-z]", entry)
    if chinese and len(chinese) >= len(latin):
        return True
    if re.search(r"[\u4e00-\u9fff]", entry) and re.search(
        r"北京|上海|济南|郑州|中国人民大学出版社|华东师范大学出版社|辽宁教育出版社",
        entry,
    ):
        return True
    return False


def category_group(entry):
    typ = type_of(entry)
    zh = is_chinese(entry)
    if zh:
        if typ == "M":
            return ("中文", "专著类")
        if typ == "J":
            return ("中文", "期刊文章类")
        if typ == "D":
            return ("中文", "学术论文类")
        if typ == "S":
            return ("中文", "课程标准类")
        return ("中文", "期刊文章类")
    if typ == "M":
        return ("外文", "专著类")
    return ("外文", "期刊论文类")


def ensure_terminal_period(text):
    text = text.rstrip(" ;,.")
    return text + "."


def normalize_author_title_split(text):
    text = re.sub(r",\[(?=[MJDS]\])", "[", text)
    text = re.sub(r"\.(\[[MJDS]\])", r"\1", text)
    text = re.sub(r"(\[[MJDS]\])(?=[A-Za-z\u4e00-\u9fff])", r"\1.", text)
    text = re.sub(r"\.\.", ".", text)
    return text


def refine_book(text):
    text = normalize_basic(text)
    text = normalize_author_title_split(text)
    text = re.sub(r"\[M\](?=[^\.\s])", "[M].", text)
    return ensure_terminal_period(text)


def refine_journal(text):
    text = normalize_basic(text)
    text = normalize_author_title_split(text)
    text = re.sub(r"\[J\](?=[^\.\s])", "[J].", text)
    text = re.sub(r",\s*\((\d+)\)", r",(\1)", text)
    text = re.sub(r"(\d{4})\s*,\s*(\d+\(\d+\))", r"\1,\2", text)
    text = re.sub(r"(\d{4})\((\d+)\)", r"\1,(\2)", text)
    text = clean_spaces(text)
    return ensure_terminal_period(text)


def refine_thesis(text):
    text = normalize_basic(text)
    text = normalize_author_title_split(text)
    text = re.sub(r"\[D\](?=[^\.\s])", "[D].", text)
    return ensure_terminal_period(text)


def refine_standard(text):
    text = normalize_basic(text)
    text = normalize_author_title_split(text)
    text = re.sub(r"\[S\](?=[^\.\s])", "[S].", text)
    return ensure_terminal_period(text)


def refine_entry(text):
    typ = type_of(text)
    if typ == "M":
        return refine_book(text)
    if typ == "J":
        return refine_journal(text)
    if typ == "D":
        return refine_thesis(text)
    if typ == "S":
        return refine_standard(text)
    return ensure_terminal_period(normalize_basic(text))


def build_doc():
    entries = [refine_entry(e) for e in dedupe(read_entries())]
    grouped = {
        ("中文", "专著类"): [],
        ("中文", "期刊文章类"): [],
        ("中文", "学术论文类"): [],
        ("中文", "课程标准类"): [],
        ("外文", "专著类"): [],
        ("外文", "期刊论文类"): [],
    }
    for e in entries:
        grouped[category_group(e)].append(e)
    for key in grouped:
        grouped[key] = sorted(grouped[key], key=lambda x: (year_of(x), x))

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)

    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    add_para(doc, "参考文献", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "（一）中文文献", bold=True)

    counter = 1
    order = [
        ("1.专著类", ("中文", "专著类")),
        ("2.期刊文章类", ("中文", "期刊文章类")),
        ("3.学术论文类", ("中文", "学术论文类")),
        ("4.课程标准类", ("中文", "课程标准类")),
        ("（二）外文文献", None),
        ("1.专著类", ("外文", "专著类")),
        ("2.期刊论文类", ("外文", "期刊论文类")),
    ]

    for title, key in order:
        add_para(doc, title, bold=True)
        if key is None:
            continue
        for e in grouped[key]:
            add_para(doc, f"[{counter}] {e}", hanging=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            counter += 1

    out = os.path.join(BASE, OUT_NAME)
    doc.save(out)
    print(out)


if __name__ == "__main__":
    build_doc()
