# -*- coding: utf-8 -*-
import os

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


BASE = r"D:\Users\Desktop"
SRC = os.path.join(BASE, "Z2321008吴雨.docx")
OUT = os.path.join(BASE, "Z2321008吴雨_审校标注版.docx")


PARA_NOTES = {
    59: "【审校意见】“课堂观察发等研究方法”中“发”为笔误。建议改为“课堂观察法等研究方法”。",
    69: "【审校意见】本句数量前后矛盾：“仅检索到相关文献17篇”与后文“学位论文20篇，期刊论文3篇”无法对应。建议核对检索结果后统一总数与分项数据。",
    73: "【审校意见】本段中“他进一步将这一概念界定为……”与后文同义表述存在重复，建议删除其中一处重复界定，避免表述冗余。",
    82: "【审校意见】“学界通常对从学生、教师、学科教学与社会发展四个层面来讨论这一问题”语法不通。建议改为“学界通常从学生、教师、学科教学与社会发展四个层面讨论这一问题”。",
    98: "【审校意见】本段有两处需核改：1.“理解，评价，创造”建议统一为“理解、评价、创造”；2.“话读”若非原文术语，建议核对来源文献并改为准确表述。",
    101: "【审校意见】“古代论说文因为因为逻辑结构严密”重复用词。建议改为“古代论说文因为逻辑结构严密”或“古代论说文因逻辑结构严密”。",
    122: "【审校意见】“课堂观察法是作为获取现场印证的关键方法”表述不通。建议改为“课堂观察法是获取现场印证的关键方法”。",
    198: "【审校意见】表号写法不统一：前文为“表2-1”，此处写为“表2.1”。建议统一为同一套编号格式，如“表2-1、表2-2”。",
    201: "【审校意见】“续表2-2 教师问卷维度划分情况”当前使用四级标题样式，不符合一般论文表题格式。建议改为与前表一致的表题/正文样式，而非标题样式。",
    220: "【审校意见】“表2-8教师量表旋转后的成分矩阵a”表题后缺少空格。建议改为“表2-8 教师量表旋转后的成分矩阵a”。",
    221: "【审校意见】本行与上一行“表2-8教师量表旋转后的成分矩阵a”重复出现。建议删除其中一行重复表题。",
    230: "【审校意见】“针对50份学生试测问卷进行效度分析如表2-11所示”中“分析”后缺少停顿。建议改为“针对50份学生试测问卷进行效度分析，结果如表2-11所示”。",
    234: "【审校意见】“Q1-Q24均只在单个维度上的载荷均大于0.5”中“均”重复。建议改为“Q1-Q24均只在单个维度上的载荷大于0.5”或“Q1-Q24在单个维度上的载荷均大于0.5”。",
    237: "【审校意见】本表题写为“学生问卷基本信息（修改后）”，与正文“修改后的学生问卷题项与维度划分见表2-14”不一致。建议统一为与正文一致的表题表述。",
    239: "【审校意见】“0.925大于0.8”之间缺少连接性表达。建议改为“0.925，大于0.8”或“0.925，说明其高于0.8”。",
    249: "【审校意见】前文写“随机选取5位高中语文教师和2位教研组长”，合计应为7位；此处写“8位访谈对象”，前后矛盾。建议核对后统一人数。",
    257: "【审校意见】表题格式不统一，“表2-18第一次访谈维度”表号后缺少空格。建议改为“表2-18 第一次访谈维度”。",
    258: "【审校意见】表题格式不统一，“表2-19第二次访谈维度”表号后缺少空格。建议改为“表2-19 第二次访谈维度”。",
    289: "【审校意见】本段中“整体维度均值3.81”与上一段“维度均值为3.68”前后不一致。建议核对原始统计结果后统一均值。",
    324: "【审校意见】本段有三处问题：1.“分析结果由表3-13所示”应改为“分析结果如表3-13所示”；2.此处讨论的是“教学实施”维度，不应写成“教学设计方面”；3.“不同地区之前表现出显著差异”应改为“不同地区之间表现出显著差异”。",
    327: "【审校意见】本段所处部分讨论“教学实施”维度，但句中写为“在古代论说文教学设计上表现出显著差异”。建议将“教学设计”改为“教学实施”。",
    331: "【审校意见】“教学实施能力表现出碎职称晋升而递增的态势”中“碎”为笔误。建议改为“随职称晋升而递增的态势”。",
    339: "【审校意见】“均为呈现出显著差异”中“为”误用。建议改为“均未呈现出显著差异”。",
    341: "【审校意见】本段讨论“教学效果与评价”维度，但句中写为“在古代论说文教学设计上表现显著差异”。建议将“教学设计”改为“教学效果与评价”。",
    344: "【审校意见】“总的来说，越大的教师……”主语缺失且搭配错误。建议改为“总的来说，教龄越大的教师在古代论说文思辨性阅读教学的效果与评价方面表现越好”。",
    378: "【审校意见】本段前半句写“不同性别学生之间”，后半句实际讨论的是“不同地区的学生”。建议将前半句改为“不同地区学生之间”。",
    391: "【审校意见】本段前半句写“不同性别学生之间”，后半句实际讨论的是“不同地区的学生”。建议将前半句改为“不同地区学生之间”。",
    415: "【审校意见】本句前半写“p值<0.001”，后半却写“存在显著差异（p>0.05）”，前后矛盾。若按前文数据，应将括号内改为“p<0.05”。",
    640: "【审校意见】本段引述单元导语后引号未闭合。建议在“提高学习能力”后补上右引号，再接后文“教师据此可将教学目标细化为……”。",
    643: "【审校意见】“传统的语文古代论说文的阅读教学”成分赘余。建议改为“传统的古代论说文阅读教学”或“传统语文课堂中的古代论说文阅读教学”。",
    644: "【审校意见】“仁义不施而攻守之势异也的结论”引号缺失。建议改为“‘仁义不施而攻守之势异也’的结论”。",
    719: "【审校意见】参考文献著录中页码写法总体较统一，但正文建议继续核查全表，确保所有页码连接号统一为半角“-”。本条当前无误，可作为统一格式参照。",
    784: "【审校意见】外文专著作者著录疑有缺漏：“Richard Pirozzi. Gretchen Stark-Martin Julie Dziewisz.”作者之间缺少明确分隔。建议核对原书版权页后按作者全名逐一分隔著录。",
    819: "【审校意见】“二级（. ）”中出现错误符号。建议改为“二级（  ）”。",
    843: "【审校意见】“您在备课时，您在设计……”重复使用“您在”。建议删去其中一个，改为“您在备课时，设计古代论说文教学目标时，如何平衡……”。",
    845: "【审校意见】本句语序不顺，“在您平常的古代论说教学的课堂中，是采用什么样的教学方法……”建议改为“在您平常的古代论说文教学课堂中，通常采用什么样的教学方法、设计什么活动来促进学生进行思辨性阅读？”",
    846: "【审校意见】“您认为或者遇到的……”表达不通。建议改为“您认为高中学段实施思辨性阅读的困难有哪些？或者，您在教学中实际遇到过哪些困难？”",
    876: "【审校意见】“学校或教研组是否学校或教研组是否组织过……”重复。建议删去一处重复，改为“学校或教研组是否组织过古代论说文思辨性阅读的专题培训、优秀案例分析？”。",
}


TABLE_NOTES = [
    (212, "【审校意见】表2-5“组统计量（平均值±标准差）”列中，多处数据缺少“±”符号，格式不规范；其中Q10高分组数据写成“4.00 (0.68”，括号也不完整。建议统一改为“平均值±标准差”格式，如“2.86±1.46”“4.00±0.68”。"),
    (352, "【审校意见】表3-22表下注释“a. 使用了值1对二分组进行制表。”在表格最后一行被重复填入多个单元格，属于格式错误。建议保留为表下注或单独说明，不要重复铺满整行单元格。"),
]


def insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


def set_red(run, size=10.5, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


def color_para_red(paragraph):
    for run in paragraph.runs:
        size = run.font.size.pt if run.font.size else 10.5
        set_red(run, size=size, bold=bool(run.bold))


def add_note_after(paragraph, text):
    note = insert_paragraph_after(paragraph)
    note.style = paragraph.style
    run = note.add_run(text)
    set_red(run)
    return note


def color_table_cell_red(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            size = r.font.size.pt if r.font.size else 10.5
            set_red(r, size=size, bold=bool(r.bold))


def main():
    doc = Document(SRC)
    original_paras = list(doc.paragraphs)

    guide = add_note_after(
        original_paras[38],
        "【说明】本稿红色标注仅针对需要修改的问题，不涉及一般性润色。每一处均附有明确修改建议，便于逐项核改。",
    )
    guide.style = original_paras[38].style

    for idx, note in sorted(PARA_NOTES.items(), reverse=True):
        if 0 <= idx < len(original_paras):
            color_para_red(original_paras[idx])
            add_note_after(original_paras[idx], note)

    for idx, note in sorted(TABLE_NOTES, key=lambda x: x[0], reverse=True):
        if 0 <= idx < len(original_paras):
            color_para_red(original_paras[idx])
            add_note_after(original_paras[idx], note)

    # Mark representative problematic table cells in red
    # Table 14: malformed statistic formatting
    try:
        color_table_cell_red(doc.tables[14].rows[11].cells[2])  # "4.00 (0.68"
        color_table_cell_red(doc.tables[14].rows[2].cells[1])   # representative missing ±
        color_table_cell_red(doc.tables[14].rows[2].cells[2])
    except Exception:
        pass

    # Table 48: repeated footnote row
    try:
        for cell in doc.tables[48].rows[9].cells:
            color_table_cell_red(cell)
    except Exception:
        pass

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
