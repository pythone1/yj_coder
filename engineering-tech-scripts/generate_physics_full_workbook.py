import os
import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE = r"E:\PY\tech"
OUT_DIR = os.path.join(BASE, "output", "doc")
IMG_DIR = os.path.join(BASE, "tmp_question_images")
OUT_PATH = os.path.join(OUT_DIR, "王睿熙物理错题与同类题精练_完整版.docx")


def set_run_font(run, size=11, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_p(doc, text="", size=11, bold=False, align=None):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, size=10)


def save_clip(pdf_path, page_index, rect, out_name, zoom=2.2):
    pdf = fitz.open(pdf_path)
    page = pdf[page_index]
    clip = fitz.Rect(rect)
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), clip=clip, alpha=False)
    out_path = os.path.join(IMG_DIR, out_name)
    pix.save(out_path)
    pdf.close()
    return out_path


def add_image(doc, path, width_cm=15.8):
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(width_cm))


def add_question_text(doc, title, source, stem, answer, analysis):
    add_p(doc, title, size=12, bold=True)
    add_p(doc, f"来源：{source}", size=10)
    add_p(doc, "题目：", size=11, bold=True)
    add_p(doc, stem, size=11)
    add_p(doc, f"答案：{answer}", size=11, bold=True)
    add_p(doc, "解析：", size=11, bold=True)
    for para in analysis:
        add_p(doc, para, size=11)


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    os.makedirs(IMG_DIR, exist_ok=True)

    wrong_pdf = [os.path.join(BASE, n) for n in os.listdir(BASE) if n.lower().endswith(".pdf")][0]

    # Wrong-question screenshots from original PDF
    img_paths = {
        "wq10": save_clip(wrong_pdf, 0, (35, 85, 560, 268), "wq10.png"),
        "wq11": save_clip(wrong_pdf, 0, (35, 260, 560, 497), "wq11.png"),
        "wq12": save_clip(wrong_pdf, 0, (35, 485, 560, 600), "wq12.png"),
        "wq14a": save_clip(wrong_pdf, 0, (35, 595, 560, 770), "wq14a.png"),
        "wq14b": save_clip(wrong_pdf, 1, (35, 35, 560, 170), "wq14b.png"),
        "wq15a": save_clip(wrong_pdf, 1, (35, 490, 560, 770), "wq15a.png"),
        "wq15b": save_clip(wrong_pdf, 2, (35, 35, 560, 255), "wq15b.png"),
        "wq16a": save_clip(wrong_pdf, 2, (35, 250, 560, 770), "wq16a.png"),
        "wq16b": save_clip(wrong_pdf, 3, (35, 35, 560, 372), "wq16b.png"),
        "wq17": save_clip(wrong_pdf, 3, (35, 375, 560, 680), "wq17.png"),
    }

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)

    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)

    add_p(doc, "王睿熙物理错题与同类题精练", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "完整版：保留原错题截图，并补充同类完整题目与详细解析", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "说明：这次文档不再只做摘要。每个类型先放原错题截图，再放2-3道同类题。原错题以“重新审题、纠错思路”为主，同类题给出完整题干、答案与解析。", size=10)

    # Type 1
    doc.add_page_break()
    add_p(doc, "类型一：简单机械与机械效率", size=15, bold=True)
    add_p(doc, "原错题1（PDF第10题）", size=12, bold=True)
    add_image(doc, img_paths["wq10"])
    add_p(doc, "纠错思路：先写机械效率公式 η=W有/W总=Gh/Fs。若题图弹簧测力计示数读作 5N，则 G=ηFs/h=0.8×5×5/2=10N；摩擦力 f=F-Gh/s=5-10×2/5=1N；斜面能省力，不能省功。", size=11)

    add_question_text(
        doc,
        "同类题1",
        "《精品解析：2025年江苏省南京市秦淮区中考一模物理试题（解析版）》第17题",
        "如图所示，利用动滑轮匀速提升物体，已知物重400N，动滑轮重60N，在100s内将物体提升10m，此过程中动滑轮机械效率为80%，有用功为_______J，拉力的功率为_______W，克服动滑轮重所做的功占总功的_______%。",
        "4000J；50W；12%。",
        [
            "先算有用功：W有=Gh=400N×10m=4000J。",
            "再由机械效率 η=W有/W总=80% 得总功 W总=4000/0.8=5000J。",
            "拉力功率 P=W总/t=5000J/100s=50W。",
            "克服动滑轮重所做的额外功 W额=G动h=60N×10m=600J，占总功的 600/5000=12%。",
            "这类题最稳的主线就是“有用功→总功→功率/额外功”。"
        ]
    )
    add_question_text(
        doc,
        "同类题2",
        "《2026年江苏省南京市鼓楼区名校联合一模物理试题》 第5题",
        "晓敏在观看杂技表演，有一项跳板节目引起了她的注意。如图所示，两名质量均为60kg的男演员同时从3m高处自由下落到跳板的一端，站立于跳板另一端的一名质量为45kg的女演员被向上弹起4m高。晓敏利用所学的知识粗略地计算出跳板的机械效率为（ ）\nA．90%  B．75%  C．50%  D．45%",
        "C。",
        [
            "输入能量来自两名男演员减小的重力势能：W总=2×60×10×3=3600J。",
            "有用能量是女演员获得的重力势能：W有=45×10×4=1800J。",
            "机械效率 η=W有/W总=1800/3600=50%。",
            "这类题要先看“谁提供能量”，再看“谁真正得到要利用的能量”。"
        ]
    )

    # Type 2
    doc.add_page_break()
    add_p(doc, "类型二：受力、惯性与摩擦力", size=15, bold=True)
    add_p(doc, "原错题1（PDF第11题）", size=12, bold=True)
    add_image(doc, img_paths["wq11"])
    add_p(doc, "纠错思路：第一问先抓“匀速直线运动”，说明合力为0；第二问先抓“突然停止”，说明A、B因惯性仍保持原速度，若都不受外力且初速度相同，则二者间距保持不变。", size=11)

    add_question_text(
        doc,
        "同类题1",
        "《精品解析：2025年江苏省南京市秦淮区中考一模物理试题（解析版）》第10题",
        "关于滑雪的物理知识，下列说法正确的是（ ）\nA. 滑雪板的重力和运动员对滑雪板的压力是一对平衡力\nB. 运动员能一直下滑是由于其受到的重力大于惯性\nC. 若运动员在滑行过程中受到的所有外力突然消失，其滑行速度将越来越大\nD. 用杆撑地加速滑行，说明力可以改变物体的运动状态",
        "D。",
        [
            "A错：平衡力必须作用在同一物体上，这里不是。",
            "B错：惯性不是力，不能与重力比较大小。",
            "C错：若外力全消失，物体会保持原来的速度做匀速直线运动。",
            "D对：撑地后速度变化，说明力可以改变物体的运动状态。"
        ]
    )
    add_question_text(
        doc,
        "同类题2",
        "《精品解析：2025年江苏省南京市秦淮区中考一模物理试题（解析版）》第22题",
        "在探究“影响滑动摩擦力大小的因素”实验中：\n（1）图甲中，将一个木块放在水平长木板上，用弹簧测力计沿水平方向拉动，使其做_______运动，根据_______知识可知木块所受滑动摩擦力大小等于拉力大小。\n（2）图乙中，在木块上放一重物，重复上面的操作。比较甲、乙两图中弹簧测力计的示数，可探究滑动摩擦力大小与_______的关系；以下事例中，用到了该探究结论的是_______。\nA．鞋底的槽纹  B．机器转轴处加润滑油  C．体操运动员在手上涂镁粉  D．自行车刹车时用力捏闸",
        "（1）匀速直线；二力平衡。（2）压力大小；D。",
        [
            "只有木块做匀速直线运动时，水平方向合力为0，拉力才等于滑动摩擦力。",
            "甲乙相比，接触面粗糙程度不变、压力改变，所以探究的是摩擦力与压力大小的关系。",
            "自行车刹车时用力捏闸，本质上是在接触面粗糙程度一定时增大压力，从而增大摩擦力。"
        ]
    )

    # Type 3
    doc.add_page_break()
    add_p(doc, "类型三：电学综合计算与自动控制", size=15, bold=True)
    add_p(doc, "原错题1（PDF第12题）", size=12, bold=True)
    add_image(doc, img_paths["wq12"])
    add_p(doc, "订正：0.6度=0.6kW·h，转盘转数 n=0.6×3000=1800转。10min=1/6h，现有总功率 P=0.6÷(1/6)=3.6kW；电能表最大允许功率 P最大=220×20=4.4kW，所以还能增加 0.8kW=800W。", size=11)
    add_p(doc, "原错题2（PDF第14题）", size=12, bold=True)
    add_image(doc, img_paths["wq14a"])
    add_image(doc, img_paths["wq14b"])
    add_p(doc, "纠错提醒：继电器/热敏电阻题一定先找“临界电流”对应的总电阻，再反推热敏电阻阻值和温度，不要一上来就代功率比。", size=11)

    add_question_text(
        doc,
        "同类题1",
        "《2026年江苏省南京市鼓楼区名校联合一模物理试题》 第17题",
        "某电能表上标有“220V 20A 3000r/（kW•h）”字样，当家用电器全部工作时，该表10min消耗了0.6度电，那么电能表的铝盘转了_______转；若该家全部用电器工作时的总电流不超过电能表的额定电流，还可以再增加_______W以下的用电器。",
        "1800转；800W。",
        [
            "消耗0.6kW·h时，铝盘转数 n=0.6×3000=1800转。",
            "总功率 P=W/t=0.6÷(1/6)=3.6kW。",
            "额定最大功率为 UI=220×20=4400W，因此还能增加 4400-3600=800W。"
        ]
    )
    add_question_text(
        doc,
        "同类题2",
        "《精品解析：2025年江苏省南京市秦淮区中考一模物理试题（解析版）》第19题",
        "小红用图甲的手持挂烫机帮家人熨烫衣物，挂烫机内的水被加热汽化后喷出。图乙为挂烫机的等效电路，R1为加热电阻，R2为限流电阻，S为自动温控开关。将挂烫机接入电路，S在位置“1”，前端电热板被加热，电路中的电流为5.0A；当温度升到130℃时，S切换到位置“2”，挂烫机处于待机状态，电路中的电流为0.02A，当温度降到100℃时，S又回到位置“1”。\n（1）求R1的阻值；\n（2）求挂烫机待机状态的功率；\n（3）熨烫过程中，电路的电流随时间变化如图丙，求挂烫机0∼150s消耗的电能。",
        "（1）44Ω；（2）4.4W；（3）110220J。",
        [
            "加热状态下，R1单独工作，R1=U/I=220/5.0=44Ω。",
            "待机状态电流为0.02A，待机功率 P=UI=220×0.02=4.4W。",
            "总电能要分加热段和待机段分别算，再相加。题目给了电流-时间图，本质上就是分段用 W=UIt。"
        ]
    )

    # Type 4
    doc.add_page_break()
    add_p(doc, "类型四：实验探究与误差分析", size=15, bold=True)
    add_p(doc, "原错题1（PDF第15题）", size=12, bold=True)
    add_image(doc, img_paths["wq15a"])
    add_image(doc, img_paths["wq15b"])
    add_p(doc, "纠错提醒：密度实验最容易错在“是否擦干”“是否回到同一液面标记”“天平零点偏差影响方向”。先把每一步测量的物理量写清，再列式。", size=11)
    add_p(doc, "原错题2（PDF第16题）", size=12, bold=True)
    add_image(doc, img_paths["wq16a"])
    add_image(doc, img_paths["wq16b"])
    add_p(doc, "纠错提醒：电学实验要按“接线是否正确→是否调零→故障判断→控制变量→误差来源”这一顺序检查。", size=11)

    add_question_text(
        doc,
        "同类题1",
        "《精品解析：2025年江苏省南京市秦淮区中考一模物理试题（解析版）》第26题",
        "图甲是小明探究“电流与电阻关系”的实验，选用的实验器材有电源（3V），电流表（0~0.6A），电压表（0~3V），滑动变阻器（20Ω 2A），阻值分别为5Ω、10Ω、20Ω的三个定值电阻，开关及导线若干。\n（1）将图甲实物电路连接完整（要求滑片向左移动时电流表示数变大）。\n（2）电路连接正确后，闭合开关，发现电压表有示数但电流表无示数，此时出现的故障可能是_______。\n（3）排除故障后，接入5Ω电阻，调节滑动变阻器，使电压表示数为U；仅将电阻R由5Ω换成10Ω，应将滑片向_______移动，使电压表示数仍为U。当电阻R换成20Ω时，发现无论怎样移动滑片，电压表示数都无法达到U，为了完成实验，U的取值不能小于_______V。\n（4）多次实验，可得出结论：_______。",
        "（2）电阻R断路；（3）向右；1.5V；（4）导体两端电压一定时，通过导体的电流与导体电阻成反比。",
        [
            "电压表有示数但电流表无示数，说明主电路断路，而电压表并联支路还能和电源连通，所以最可能是定值电阻断路。",
            "把5Ω换成10Ω后，定值电阻分压会变大。要让它恢复到原来的U，就要让滑动变阻器分到更多电压，也就是增大滑动变阻器接入电阻，所以滑片要向右。",
            "当滑动变阻器调到最大20Ω时，20Ω定值电阻与20Ω滑动变阻器串联，分压相等，所以定值电阻最小电压是3V的一半，即1.5V。",
            "最后结论一定带前提：导体两端电压一定。"
        ]
    )
    add_question_text(
        doc,
        "同类题2",
        "《2025-2026学年九年级上学期（中考一模）物理调研试题》 第15题",
        "小明利用图甲所示的电路探究“通过导体的电流与电阻的关系”，根据实验的数据绘出了I-R图像，如图乙所示。分析图像可知，当电流分别为0.25A和0.5A时，接入电路的导体的电阻之比为_______，实验过程中，小明控制导体两端的电压为_______。",
        "电阻之比为2∶1；导体两端电压保持不变。",
        [
            "在电压一定时，I与R成反比。电流从0.25A变到0.5A，变为2倍，则电阻应变为原来的一半，所以电阻之比为2∶1。",
            "这题虽然短，但核心是在提醒你：研究I-R关系时，必须控制电压不变。"
        ]
    )

    # Type 5
    doc.add_page_break()
    add_p(doc, "类型五：压强、功率与能量转化综合计算", size=15, bold=True)
    add_p(doc, "原错题1（PDF第17题）", size=12, bold=True)
    add_image(doc, img_paths["wq17"])
    add_p(doc, "订正：倒车雷达利用超声波传递信息。静止时压力 F=G=mg=1.5×10^3×10=1.5×10^4N，受力面积 S=2000cm²=0.2m²，压强 p=F/S=7.5×10^4Pa。匀速行驶 3.6km 用时 360s，有用功 W=Pt=1.5×10^4×360=5.4×10^6J；效率30%，则燃料释放总能量 Q=5.4×10^6/0.3=1.8×10^7J；汽油质量 m=Q/q=1.8×10^7/(4.5×10^7)=0.4kg。", size=11)

    add_question_text(
        doc,
        "同类题1",
        "《精品解析：2025年江苏省南京市秦淮区中考一模物理试题（解析版）》第27题",
        "如图所示是餐厅中的送餐机器人，已知托着饮料的机器人总质量为52kg，与地面的接触面积为0.01m²。\n（1）机器人眼睛的主要元件是发光二极管，其核心材料是_______（选填“导体”“半导体”或“超导体”）。\n（2）若机器人沿水平地面匀速运动了20m，用时20s。求机器人的速度。\n（3）机器人对地面压强是多少Pa？",
        "（1）半导体；（2）1m/s；（3）52000Pa。",
        [
            "速度 v=s/t=20/20=1m/s。",
            "机器人对地压力等于重力：F=mg=52×10=520N。",
            "压强 p=F/S=520/0.01=52000Pa。",
            "综合题先统一单位，再列公式，最后代值。"
        ]
    )
    add_question_text(
        doc,
        "同类题2",
        "《2025-2026学年九年级上学期（中考一模）物理调研试题》 第26题",
        "如图是小明家购置的新型能源轿车。假期全家驾车游玩，相关数据如下：轿车质量1050kg，车轮与地面接触总面积0.15m²，轿车在高速公路上匀速行驶90km所用时间为1h，匀速行驶时牵引力2000N，行车20km消耗天然气1m³，天然气热值为8.0×10^7J/m³。求：\n（1）轿车停放在水平地面时对地面的压强；\n（2）轿车在1h内牵引力做功的功率；\n（3）轿车发动机的效率。",
        "（1）7.0×10^4Pa；（2）5.0×10^4W；（3）50%。",
        [
            "压强：F=mg=1050×10=10500N，p=10500/0.15=7.0×10^4Pa。",
            "功率：速度 v=90km/h=25m/s，所以 P=Fv=2000×25=5.0×10^4W。",
            "90km耗气 4.5m³，燃料总能量 Q=4.5×8.0×10^7=3.6×10^8J。",
            "有用功 W=Fs=2000×90000=1.8×10^8J。",
            "效率 η=W/Q=1.8×10^8 / 3.6×10^8=50%。"
        ]
    )

    doc.add_page_break()
    add_p(doc, "最后建议", size=15, bold=True)
    add_bullets(doc, [
        "接下来最值得做的不是再刷大量新题，而是把这份文档里的每道同类题独立再做一遍。",
        "原错题请你先遮住我的订正，自己重做；做完再对照“纠错思路”，这样提分最快。",
        "如果你需要，我下一步可以把这份完整版再拆成两份：一份“学生空白训练版”，一份“老师解析版”。"
    ])

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
