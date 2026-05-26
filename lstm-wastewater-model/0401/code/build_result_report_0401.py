from __future__ import annotations

import json
import math
import re
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from config_0401 import CANDIDATE_NODES, MONITOR_NODES, TRUTH_INJECTION_NODES


ROOT = Path(r"E:\PY\LSTM")
PROJECT_DIR = ROOT / "0401"
MODEL_DIR = PROJECT_DIR / "models" / "current_confirmed_models"
RESULTS_DIR = PROJECT_DIR / "results"
OUTPUT_DIR = PROJECT_DIR / "output" / "doc"
ASSET_DIR = OUTPUT_DIR / "0401_结果分析汇报材料_assets"


plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS"]
plt.rcParams["axes.unicode_minus"] = False


def ensure_dirs() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def find_csv_by_columns(run_dir: Path, required_cols: set[str]) -> tuple[Path, pd.DataFrame]:
    for path in run_dir.glob("*.csv"):
        try:
            df = pd.read_csv(path, encoding="utf-8-sig")
        except Exception:
            continue
        if required_cols.issubset(set(df.columns)):
            return path, df
    raise FileNotFoundError(f"No CSV in {run_dir} with columns {required_cols}")


def parse_rpt_metrics(path: Path) -> dict:
    text = path.read_text(encoding="utf-8", errors="ignore")
    lines = text.splitlines()

    def extract_last(label: str) -> float | None:
        values: list[float] = []
        for idx, line in enumerate(lines):
            if label in line:
                nums = re.findall(r"[-+]?\d+\.\d+", line)
                if nums:
                    values.append(float(nums[-1]))
        return values[-1] if values else None

    return {
        "path": str(path),
        "dry_weather_inflow": extract_last("Dry Weather Inflow"),
        "wet_weather_inflow": extract_last("Wet Weather Inflow"),
        "groundwater_inflow": extract_last("Groundwater Inflow"),
        "rdii_inflow": extract_last("RDII Inflow"),
        "external_inflow": extract_last("External Inflow"),
        "external_outflow": extract_last("External Outflow"),
        "flooding_loss": extract_last("Flooding Loss"),
        "evaporation_loss": extract_last("Evaporation Loss"),
        "exfiltration_loss": extract_last("Exfiltration Loss"),
        "initial_stored_volume": extract_last("Initial Stored Volume"),
        "final_stored_volume": extract_last("Final Stored Volume"),
        "continuity_error_pct": extract_last("Continuity Error (%)"),
        "no_nodes_flooded": "No nodes were flooded." in text,
    }


def load_model_metrics() -> tuple[dict, dict]:
    metrics = [parse_rpt_metrics(path) for path in MODEL_DIR.glob("*.rpt")]
    baseline = min(metrics, key=lambda item: abs(item.get("external_inflow") or 0.0))
    event = max(metrics, key=lambda item: item.get("external_inflow") or 0.0)
    return baseline, event


def total_inflow_million_liters(metrics: dict) -> float:
    keys = [
        "dry_weather_inflow",
        "wet_weather_inflow",
        "groundwater_inflow",
        "rdii_inflow",
        "external_inflow",
    ]
    return sum(float(metrics.get(key) or 0.0) for key in keys)


def balance_residual_million_liters(metrics: dict) -> float:
    incoming = total_inflow_million_liters(metrics) + float(metrics.get("initial_stored_volume") or 0.0)
    accounted = (
        float(metrics.get("external_outflow") or 0.0)
        + float(metrics.get("flooding_loss") or 0.0)
        + float(metrics.get("evaporation_loss") or 0.0)
        + float(metrics.get("exfiltration_loss") or 0.0)
        + float(metrics.get("final_stored_volume") or 0.0)
    )
    return incoming - accounted


def load_layout_tables() -> tuple[pd.DataFrame, pd.DataFrame]:
    node_csv = None
    link_csv = None
    for path in MODEL_DIR.glob("*.csv"):
        header = path.read_text(encoding="utf-8-sig", errors="ignore").splitlines()[0]
        if header.startswith("节点名称"):
            node_csv = path
        elif header.startswith("连接名称"):
            link_csv = path
    if node_csv is None or link_csv is None:
        raise FileNotFoundError("Layout CSV files not found")
    return (
        pd.read_csv(node_csv, encoding="utf-8-sig"),
        pd.read_csv(link_csv, encoding="utf-8-sig"),
    )


def load_run_bundle(run_name: str) -> dict:
    run_dir = RESULTS_DIR / run_name
    summary = read_json(run_dir / "0401_结果汇总.json")
    _, ga_history = find_csv_by_columns(run_dir, {"generation", "population", "best_mean_nse", "best_sse"})
    _, weights = find_csv_by_columns(run_dir, {"node", "posterior_mean", "posterior_median", "p05", "p95"})
    _, coverage = find_csv_by_columns(run_dir, {"monitor", "coverage_90"})
    return {
        "summary": summary,
        "ga_history": ga_history,
        "weights": weights,
        "coverage": coverage,
    }


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def save_current_layout_figure(nodes: pd.DataFrame, links: pd.DataFrame) -> Path:
    role_colors = {
        "真值注水点": "#d73027",
        "监测点": "#1f78b4",
        "候选布设点": "#fdae61",
        "泵站链路节点": "#984ea3",
        "末端关键节点": "#33a02c",
        "结构排口": "#238b45",
        "敏感节点": "#fb6a4a",
        "关联节点": "#bdbdbd",
    }

    node_map = nodes.set_index("节点名称")[["X", "Y", "节点角色"]].to_dict("index")
    fig, ax = plt.subplots(figsize=(11, 8))

    for _, row in links.iterrows():
        start = row["起点"]
        end = row["终点"]
        if start in node_map and end in node_map:
            ax.plot(
                [node_map[start]["X"], node_map[end]["X"]],
                [node_map[start]["Y"], node_map[end]["Y"]],
                color="#d9d9d9",
                linewidth=0.6,
                zorder=1,
            )

    role_order = [
        "关联节点",
        "候选布设点",
        "监测点",
        "真值注水点",
        "末端关键节点",
        "结构排口",
        "泵站链路节点",
        "敏感节点",
    ]
    for role in role_order:
        subset = nodes[nodes["节点角色"] == role]
        if subset.empty:
            continue
        size = 10 if role == "关联节点" else 38
        ax.scatter(
            subset["X"],
            subset["Y"],
            s=size,
            c=role_colors.get(role, "#666666"),
            label=role,
            zorder=3,
            edgecolors="none",
            alpha=0.95 if role != "关联节点" else 0.6,
        )

    key_labels = {"J76", "J124", "J140", "J141", "J139", "J145", "J231", "J132", "J74", "J78", "J123", "J126"}
    for _, row in nodes.iterrows():
        if row["节点名称"] in key_labels:
            ax.text(row["X"] + 12, row["Y"] + 12, row["节点名称"], fontsize=8)

    ax.set_title("当前正式布设方案与关键节点分布")
    ax.set_xlabel("X")
    ax.set_ylabel("Y")
    ax.legend(loc="upper left", fontsize=8, ncols=2)
    ax.grid(alpha=0.15)
    path = ASSET_DIR / "figure_01_current_layout.png"
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return path


def save_local_cluster_figure(nodes: pd.DataFrame, links: pd.DataFrame) -> Path:
    focus_nodes = {
        "J123",
        "J124",
        "J125",
        "J126",
        "J137",
        "J138",
        "J139",
        "J140",
        "J141",
        "J142",
        "J145",
        "J146",
        "J67",
        "J231",
        "J132",
    }
    node_subset = nodes[nodes["节点名称"].isin(focus_nodes)].copy()
    node_map = node_subset.set_index("节点名称")[["X", "Y", "节点角色"]].to_dict("index")
    fig, ax = plt.subplots(figsize=(9, 7))

    for _, row in links.iterrows():
        start = row["起点"]
        end = row["终点"]
        if start in node_map and end in node_map:
            ax.plot(
                [node_map[start]["X"], node_map[end]["X"]],
                [node_map[start]["Y"], node_map[end]["Y"]],
                color="#9e9ac8",
                linewidth=1.5,
                zorder=1,
            )

    for _, row in node_subset.iterrows():
        role = row["节点角色"]
        if row["节点名称"] in {"J140"}:
            color = "#d73027"
        elif row["节点名称"] in {"J145"}:
            color = "#1f78b4"
        elif row["节点名称"] in {"J141", "J139", "J231"}:
            color = "#33a02c"
        else:
            color = "#969696"
        ax.scatter(row["X"], row["Y"], s=70, c=color, zorder=3)
        ax.text(row["X"] + 10, row["Y"] + 10, row["节点名称"], fontsize=9)

    ax.set_title("J140/J145 局部支路与关键监测点")
    ax.set_xlabel("X")
    ax.set_ylabel("Y")
    ax.grid(alpha=0.2)
    path = ASSET_DIR / "figure_02_local_cluster.png"
    fig.tight_layout()
    fig.savefig(path, dpi=240, bbox_inches="tight")
    plt.close(fig)
    return path


def save_truth_replay_figure(truth_replay: dict) -> Path:
    per_node = truth_replay["per_node_nse"]
    fig, ax = plt.subplots(figsize=(10, 4.2))
    ax.bar(list(per_node.keys()), list(per_node.values()), color="#3182bd")
    ax.set_ylim(0.95, 1.01)
    ax.set_ylabel("NSE")
    ax.set_title(f"Truth replay: {len(per_node)} monitors, node-wise NSE")
    ax.grid(axis="y", alpha=0.2)
    path = ASSET_DIR / "figure_03_truth_replay.png"
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return path


def save_summary_comparison_figure(medium: dict, large: dict) -> Path:
    labels = ["GA最佳", "后验中位数", "覆盖率均值"]
    medium_vals = [
        medium["summary"]["ga_best_mean_nse"],
        medium["summary"]["posterior_median_nse"],
        medium["summary"]["posterior_coverage_mean"],
    ]
    large_vals = [
        large["summary"]["ga_best_mean_nse"],
        large["summary"]["posterior_median_nse"],
        large["summary"]["posterior_coverage_mean"],
    ]

    fig, ax = plt.subplots(figsize=(8.5, 5))
    x = range(len(labels))
    width = 0.34
    ax.bar([v - width / 2 for v in x], medium_vals, width=width, label="Medium", color="#4daf4a")
    ax.bar([v + width / 2 for v in x], large_vals, width=width, label="Large", color="#377eb8")
    for idx, val in enumerate(medium_vals):
        ax.text(idx - width / 2, val + 0.015, f"{val:.3f}", ha="center", fontsize=8)
    for idx, val in enumerate(large_vals):
        ax.text(idx + width / 2, val + 0.015, f"{val:.3f}", ha="center", fontsize=8)
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels)
    ax.set_ylim(0, 1.0)
    ax.set_ylabel("数值")
    ax.set_title("Medium 与 Large 结果总览")
    ax.legend()
    ax.grid(axis="y", alpha=0.2)
    path = ASSET_DIR / "figure_04_summary_comparison.png"
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return path


def save_ga_convergence_figure(medium: dict, large: dict) -> Path:
    med = medium["ga_history"].groupby("generation", as_index=False)["best_mean_nse"].max()
    lar = large["ga_history"].groupby("generation", as_index=False)["best_mean_nse"].max()
    fig, ax = plt.subplots(figsize=(8.5, 5))
    ax.plot(med["generation"], med["best_mean_nse"], marker="o", linewidth=2, label="Medium", color="#4daf4a")
    ax.plot(lar["generation"], lar["best_mean_nse"], marker="o", linewidth=2, label="Large", color="#377eb8")
    ax.set_xlabel("GA 代数")
    ax.set_ylabel("当代最佳 mean NSE")
    ax.set_title("GA 收敛轨迹比较")
    ax.legend()
    ax.grid(alpha=0.25)
    path = ASSET_DIR / "figure_05_ga_convergence.png"
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return path


def save_posterior_compare_figure(medium: dict, large: dict) -> Path:
    focus_nodes = ["J76", "J124", "J140", "J145", "J125", "J78"]
    truth = {"J76": 0.2368421053, "J124": 0.3421052632, "J140": 0.4210526316}
    med = medium["weights"].set_index("node")
    lar = large["weights"].set_index("node")
    truth_vals = [truth.get(node, 0.0) for node in focus_nodes]
    med_vals = [float(med.loc[node, "posterior_median"]) if node in med.index else 0.0 for node in focus_nodes]
    lar_vals = [float(lar.loc[node, "posterior_median"]) if node in lar.index else 0.0 for node in focus_nodes]

    fig, ax = plt.subplots(figsize=(10, 5))
    x = range(len(focus_nodes))
    width = 0.24
    ax.bar([v - width for v in x], truth_vals, width=width, label="真值份额", color="#d73027")
    ax.bar(list(x), med_vals, width=width, label="Medium 后验中位数", color="#4daf4a")
    ax.bar([v + width for v in x], lar_vals, width=width, label="Large 后验中位数", color="#377eb8")
    ax.set_xticks(list(x))
    ax.set_xticklabels(focus_nodes)
    ax.set_ylabel("份额")
    ax.set_title("真值份额与后验中位数对比")
    ax.legend(fontsize=8)
    ax.grid(axis="y", alpha=0.2)
    path = ASSET_DIR / "figure_06_posterior_compare.png"
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return path


def save_compensation_figure(comp: dict) -> Path:
    labels = ["真值", "J140→J145", "J140→J125", "J140一半→J145"]
    vals = [
        comp["truth"]["mean_nse"],
        comp["J140_to_J145"]["mean_nse"],
        comp["J140_to_J125"]["mean_nse"],
        comp["half_J140_to_J145"]["mean_nse"],
    ]
    fig, ax = plt.subplots(figsize=(9, 5))
    colors = ["#d73027", "#1f78b4", "#fdae61", "#756bb1"]
    ax.bar(labels, vals, color=colors)
    for idx, val in enumerate(vals):
        ax.text(idx, val + 0.015, f"{val:.3f}", ha="center", fontsize=9)
    ax.set_ylim(0, 1.05)
    ax.set_ylabel("Mean NSE")
    ax.set_title("J140 代偿试算：拟合能力比较")
    ax.grid(axis="y", alpha=0.2)
    path = ASSET_DIR / "figure_07_compensation.png"
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return path


def save_coverage_compare_figure(medium: dict, large: dict) -> Path:
    med = medium["coverage"].rename(columns={"coverage_90": "medium"})
    lar = large["coverage"].rename(columns={"coverage_90": "large"})
    merged = med.merge(lar, on="monitor", how="outer").fillna(0.0)
    fig, ax = plt.subplots(figsize=(10.5, 5.2))
    x = range(len(merged))
    width = 0.36
    ax.bar([v - width / 2 for v in x], merged["medium"], width=width, label="Medium", color="#4daf4a")
    ax.bar([v + width / 2 for v in x], merged["large"], width=width, label="Large", color="#377eb8")
    ax.set_xticks(list(x))
    ax.set_xticklabels(merged["monitor"])
    ax.set_ylim(0, 1.0)
    ax.set_ylabel("90% 区间覆盖率")
    ax.set_title("监测点 posterior predictive coverage 比较")
    ax.legend()
    ax.grid(axis="y", alpha=0.2)
    path = ASSET_DIR / "figure_08_coverage_compare.png"
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return path


def paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(4)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def set_doc_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        st = doc.styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def build_markdown(
    out_path: Path,
    baseline_metrics: dict,
    event_metrics: dict,
    truth_replay: dict,
    medium: dict,
    large: dict,
    comp: dict,
    network_stats: dict,
) -> None:
    baseline_total = total_inflow_million_liters(baseline_metrics)
    event_total = total_inflow_million_liters(event_metrics)
    baseline_residual = balance_residual_million_liters(baseline_metrics)
    event_residual = balance_residual_million_liters(event_metrics)
    md = []
    md.append("# 0401 当前结果与识别稳定性分析汇报\n")
    md.append("导语：当前正式实验口径如下。\n")
    md.append("- 基线模型：`0327_由真值模型去注水重建_旱天基线模型_10分钟_泵站0.5开0.2关.inp`\n")
    md.append("- 真值事件模型：`0327_由旱天基线重建_三点注水模型_0.3倍.inp`\n")
    md.append(f"- 网络规模：{network_stats['node_count']} 个节点，{network_stats['link_count']} 条连接\n")
    md.append(f"- 候选井数：{network_stats['candidate_count']}\n")
    md.append(f"- 监测点数：{network_stats['monitor_count']}\n")
    md.append(f"- 真值注入点：{'、'.join(TRUTH_INJECTION_NODES)}\n")
    md.append(f"- 当前监测点：{'、'.join(MONITOR_NODES)}\n")
    md.append("\n## 第一章 结果分析\n")
    md.append("### 1.1 真值回灌与实验闭合性\n")
    md.append(f"- 真值回灌 mean NSE = {truth_replay['mean_nse']:.3f}\n")
    md.append(f"- 真值回灌 SSE = {truth_replay['sse']:.3e}\n")
    md.append("- 当前正式基线模型与正式真值事件模型之间的注水与评分链已经闭合，因此后续误差不再来自模板错配。\n")
    md.append("### 1.2 Medium 结果\n")
    md.append(f"- 参数规模：{medium['summary']['config']['ga_population_count']}×{medium['summary']['config']['ga_population_size']}×{medium['summary']['config']['ga_generations']} + {medium['summary']['config']['am_chain_count']}×{medium['summary']['config']['am_samples_per_chain']}\n")
    md.append(f"- GA 最优 mean NSE = {medium['summary']['ga_best_mean_nse']:.4f}\n")
    md.append(f"- AM posterior median NSE = {medium['summary']['posterior_median_nse']:.4f}\n")
    md.append(f"- posterior predictive coverage 均值 = {medium['summary']['posterior_coverage_mean']:.4f}\n")
    md.append(f"- 预测 top3 = {', '.join(medium['summary']['predicted_top3'])}\n")
    md.append("- Medium 更像高分代偿解：分数高，但 top3 仍被 J145 / J125 / J124 主导。\n")
    md.append("### 1.3 Large 结果\n")
    md.append(f"- 参数规模：{large['summary']['config']['ga_population_count']}×{large['summary']['config']['ga_population_size']}×{large['summary']['config']['ga_generations']} + {large['summary']['config']['am_chain_count']}×{large['summary']['config']['am_samples_per_chain']}\n")
    md.append(f"- GA 最优 mean NSE = {large['summary']['ga_best_mean_nse']:.4f}\n")
    md.append(f"- AM posterior median NSE = {large['summary']['posterior_median_nse']:.4f}\n")
    md.append(f"- posterior predictive coverage 均值 = {large['summary']['posterior_coverage_mean']:.4f}\n")
    md.append(f"- 预测 top3 = {', '.join(large['summary']['predicted_top3'])}\n")
    md.append("- Large 的点解拟合值下降，但后验结构更接近真值，J140 被抬起，J125 被压低。\n")
    md.append("### 1.4 结果中的代偿事实\n")
    md.append(f"- `J140 -> J145` 代偿试算 mean NSE = {comp['J140_to_J145']['mean_nse']:.4f}\n")
    md.append(f"- `J140 -> J125` 代偿试算 mean NSE = {comp['J140_to_J125']['mean_nse']:.4f}\n")
    md.append("- `J140 -> J145` 的代偿能力明显强于 `J140 -> J125`，说明当前不稳定性的核心是 J140/J145 的局部等效性。\n")
    md.append("\n## 第二章 数据分析\n")
    md.append("### 2.1 连续性误差是什么\n")
    md.append("连续性误差不是识别误差，而是 SWMM 在整个仿真时段上的**系统级水量平衡残差**。\n")
    md.append("它衡量的是：进入系统的总水量，是否能与排出系统的水量、损失掉的水量以及最终留存在系统中的水量闭合。\n")
    md.append("概念式：`Continuity Error ≈ (总入流 + 初始存储 - 外排 - 损失 - 末时刻存储) / 总入流 × 100%`\n")
    md.append("### 2.2 连续性误差不是漏算了存储\n")
    md.append("不是。这里已经把 `Initial Stored Volume` 和 `Final Stored Volume` 都算进去了，而且它们来自 SWMM 报表的系统级存储汇总，不是只统计单独某一个储池点。\n")
    md.append(f"- 基线：总入流 `{baseline_total:.3f}`、末时刻存储 `{baseline_metrics['final_stored_volume']:.3f}`、平衡残差 `{baseline_residual:.3f}`（单位均为 `10^6 L`），连续性误差 `{baseline_metrics['continuity_error_pct']:.3f}%`\n")
    md.append(f"- 真值事件：总入流 `{event_total:.3f}`、末时刻存储 `{event_metrics['final_stored_volume']:.3f}`、平衡残差 `{event_residual:.3f}`（单位均为 `10^6 L`），连续性误差 `{event_metrics['continuity_error_pct']:.3f}%`\n")
    md.append("即便把存储量算进去，真值事件工况下仍然存在很大的系统平衡残差，因此这里反映的是数值水量平衡没有完全闭合，而不是汇报时漏算了存储。\n")
    md.append("### 2.3 当前数据是否有问题\n")
    md.append("- 从模板一致性看：truth replay 严格成立，因此模板错配和注水接口错误已经被排除。\n")
    md.append(f"- 从数值质量看：基线连续性误差为 `{baseline_metrics['continuity_error_pct']:.3f}%`，真值事件连续性误差为 `{event_metrics['continuity_error_pct']:.3f}%`，说明事件工况的数值平衡明显更粗糙。\n")
    md.append("- 从辨识结构看：J140/J145 的强代偿是数据中客观存在的事实，不是算法后处理造成的假象。\n")
    md.append(f"- 从参数与布设看：当前问题是 {network_stats['candidate_count']} 候选井、{network_stats['monitor_count']} 监测点、20 维 simplex 份额反演；监测布设已把问题推进到局部簇识别阶段，但尚未唯一切开 J140/J145。\n")
    md.append("\n## 第三章 论文适用性分析\n")
    md.append("### 3.1 中文论文的适用性\n")
    md.append("- 中文论文更偏工程应用，监测点更多，并且有实地监测/实测数据参与。\n")
    md.append("- 它更适合解释：为什么监测点更充分、工程信息更完整时，定位结果会更稳。\n")
    md.append("### 3.2 英文论文的适用性\n")
    md.append("- 英文论文更偏方法可行性与不确定性分析，典型案例监测点更少，重点是 GA → initial PPD → AM。\n")
    md.append("- 它更适合解释：为什么在少量监测点条件下，后验会是多峰的、会出现代偿和不确定性。\n")
    md.append("### 3.3 当前项目与两篇论文的差距\n")
    md.append(f"- 当前 0401 项目是一个 {network_stats['candidate_count']} 候选井、{network_stats['monitor_count']} 监测点、20 维 simplex 份额反演问题，网络规模为 {network_stats['node_count']} 个节点、{network_stats['link_count']} 条连接。\n")
    md.append("- 相比英文论文的少监测点示例，当前网络更复杂、局部支路更多，代偿更容易出现。\n")
    md.append("- 相比中文论文的工程监测条件，当前监测信息仍然不够充分，因此还不能等价看作工程上的唯一稳定定位。\n")
    md.append("- 当前实验最准确的定位是：方法链已经可行，但适用性仍受监测布设、局部支路等效性和事件数值平衡粗糙度的共同限制。\n")
    out_path.write_text("".join(md), encoding="utf-8")


def build_docx(
    out_path: Path,
    baseline_inp: Path,
    event_inp: Path,
    baseline_metrics: dict,
    event_metrics: dict,
    truth_replay: dict,
    medium: dict,
    large: dict,
    comp: dict,
    figures: dict[str, Path],
    monitor_images: list[Path],
    network_stats: dict,
) -> None:
    baseline_total = total_inflow_million_liters(baseline_metrics)
    event_total = total_inflow_million_liters(event_metrics)
    baseline_residual = balance_residual_million_liters(baseline_metrics)
    event_residual = balance_residual_million_liters(event_metrics)

    doc = Document()
    set_doc_style(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("0401 当前结果与识别稳定性分析汇报")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("基于当前正式测试文件、Medium/Large 结果与中英文参考论文的客观分析").font.size = Pt(10.5)

    paragraph(doc, "说明：本报告只呈现当前事实、数据与原因分析，不包含调优建议。")

    doc.add_heading("第一章 结果分析", level=1)
    doc.add_heading("1.1 真值回灌与实验闭合性", level=2)
    paragraph(doc, f"正式基线模型：{baseline_inp.name}")
    paragraph(doc, f"正式真值事件模型：{event_inp.name}")
    paragraph(doc, f"真值注入点为 {'、'.join(TRUTH_INJECTION_NODES)}；当前监测点为 {'、'.join(MONITOR_NODES)}。")
    paragraph(doc, f"当前正式网络范围共有 {network_stats['node_count']} 个节点、{network_stats['link_count']} 条连接；反演对象为 {network_stats['candidate_count']} 个候选井上的 20 维份额向量。")
    paragraph(doc, "当前两份正式测试文件已经完全闭合：从真值事件模型去掉 3 个注入点即可得到正式基线模型；再把真值时序按同样机制注回基线，可 1:1 复现真值事件。")
    paragraph(doc, f"真值回灌汇总：mean NSE = {truth_replay['mean_nse']:.3f}，SSE = {truth_replay['sse']:.3e}。这说明当前实验口径下，数据链、注水机制、监测提取和评分链是一致的。")
    doc.add_picture(str(figures["layout"]), width=Inches(6.6))
    add_caption(doc, "图 1 当前正式布设方案与关键节点分布")
    doc.add_picture(str(figures["truth_replay"]), width=Inches(6.5))
    add_caption(doc, "图 2 真值回灌校验：10 个监测点逐点 NSE 均为 1.0")

    doc.add_heading("1.2 GA 结果分析", level=2)
    add_table(
        doc,
        ["项目", "Medium", "Large"],
        [
            ["GA 最优 mean NSE", f"{medium['summary']['ga_best_mean_nse']:.4f}", f"{large['summary']['ga_best_mean_nse']:.4f}"],
            [
                "GA 参数规模",
                f"{medium['summary']['config']['ga_population_count']}×{medium['summary']['config']['ga_population_size']}×{medium['summary']['config']['ga_generations']}",
                f"{large['summary']['config']['ga_population_count']}×{large['summary']['config']['ga_population_size']}×{large['summary']['config']['ga_generations']}",
            ],
        ],
    )
    paragraph(doc, "Medium 的 GA 最优值高于 Large；但这并不代表 Large 的搜索方向更差，而是说明当前问题存在高分代偿盆地。Large 的 GA 搜索更广，但当前代数仍低于英文论文示例中的 25 代。")
    doc.add_picture(str(figures["ga"]), width=Inches(6.4))
    add_caption(doc, "图 3 GA 收敛轨迹对比")

    doc.add_heading("1.3 AM 与后验证结果分析", level=2)
    add_table(
        doc,
        ["项目", "Medium", "Large"],
        [
            ["AM posterior median NSE", f"{medium['summary']['posterior_median_nse']:.4f}", f"{large['summary']['posterior_median_nse']:.4f}"],
            ["Posterior best NSE", f"{medium['summary']['posterior_best_nse']:.4f}", f"{large['summary']['posterior_best_nse']:.4f}"],
            ["Posterior coverage mean", f"{medium['summary']['posterior_coverage_mean']:.4f}", f"{large['summary']['posterior_coverage_mean']:.4f}"],
            ["预测 top3", " / ".join(medium['summary']['predicted_top3']), " / ".join(large['summary']['predicted_top3'])],
            [
                "AM 参数规模",
                f"{medium['summary']['config']['am_chain_count']}×{medium['summary']['config']['am_samples_per_chain']}, warmup={medium['summary']['config']['am_warmup']}",
                f"{large['summary']['config']['am_chain_count']}×{large['summary']['config']['am_samples_per_chain']}, warmup={large['summary']['config']['am_warmup']}",
            ],
        ],
    )
    paragraph(doc, "Medium 的 posterior median 分数更高，但它更像高分代偿解；Large 的 posterior median 分数更低，却把后验结构推向了更接近真值的位置。Large 的 top3 已经从 J145 / J125 / J124 变成 J124 / J140 / J145。")
    doc.add_picture(str(figures["summary"]), width=Inches(6.4))
    add_caption(doc, "图 4 Medium 与 Large 关键指标对比")
    doc.add_picture(str(figures["posterior"]), width=Inches(6.5))
    add_caption(doc, "图 5 真值份额与 Medium/Large 后验中位数对比")
    doc.add_picture(str(figures["coverage"]), width=Inches(6.5))
    add_caption(doc, "图 6 监测点 posterior predictive coverage 对比")

    doc.add_heading("1.4 代偿结果分析", level=2)
    paragraph(doc, f"把真值中的 J140 全部替换为 J145，单独试算仍有 mean NSE = {comp['J140_to_J145']['mean_nse']:.4f}；而 Medium 的 posterior median NSE = {medium['summary']['posterior_median_nse']:.4f}，两者几乎相同。这说明 Medium 当前主要落在一个高分代偿解附近。")
    paragraph(doc, f"相比之下，把 J140 全部替换为 J125 后，mean NSE 降到 {comp['J140_to_J125']['mean_nse']:.4f}。这表明 J145 是当前真正强势的代偿点，J125 属于次级竞争点。")
    doc.add_picture(str(figures["cluster"]), width=Inches(5.8))
    add_caption(doc, "图 7 J140/J145 局部支路与关键监测点")
    doc.add_picture(str(figures["compensation"]), width=Inches(6.2))
    add_caption(doc, "图 8 J140 代偿试算拟合能力比较")

    if monitor_images:
        doc.add_heading("1.5 关键监测点时序示例", level=2)
        paragraph(doc, "以下时序图用于说明新增监测点 J139 之后，J140 片区响应已经被更清楚地截获，但 J140/J145 仍然存在可竞争解释。")
        for img in monitor_images:
            doc.add_picture(str(img), width=Inches(6.1))
            add_caption(doc, f"图 {monitor_images.index(img) + 9} {img.stem}")

    doc.add_heading("第二章 数据分析", level=1)
    doc.add_heading("2.1 连续性误差是什么", level=2)
    paragraph(
        doc,
        "连续性误差不是识别误差，也不是 GA/AM 的拟合误差。它是 SWMM 在整个仿真时段上的水量平衡残差，用来描述“进入系统的总水量”与“排出系统、损失掉、以及最终留存在系统中的水量”之间是否闭合。",
    )
    paragraph(
        doc,
        "概念上可以理解为：Continuity Error ≈ (总入流 + 初始存储 - 外排 - 溢流损失 - 蒸发/渗漏损失 - 末时刻存储) / 总入流 × 100%。因此，连续性误差高，表示系统的整体水量平衡没有很好闭合；连续性误差低，表示整体水量平衡更可信。",
    )
    add_table(
        doc,
        ["项目", "基线模型", "真值事件模型"],
        [
            ["总入流 (10^6 L)", f"{baseline_total:.3f}", f"{event_total:.3f}"],
            ["初始存储 (10^6 L)", f"{baseline_metrics['initial_stored_volume']:.3f}", f"{event_metrics['initial_stored_volume']:.3f}"],
            ["末时刻存储 (10^6 L)", f"{baseline_metrics['final_stored_volume']:.3f}", f"{event_metrics['final_stored_volume']:.3f}"],
            ["外部出流 (10^6 L)", f"{baseline_metrics['external_outflow']:.3f}", f"{event_metrics['external_outflow']:.3f}"],
            ["Flooding Loss (10^6 L)", f"{baseline_metrics['flooding_loss']:.3f}", f"{event_metrics['flooding_loss']:.3f}"],
            ["平衡残差 (10^6 L)", f"{baseline_residual:.3f}", f"{event_residual:.3f}"],
            ["Continuity Error (%)", f"{baseline_metrics['continuity_error_pct']:.3f}", f"{event_metrics['continuity_error_pct']:.3f}"],
        ],
    )
    paragraph(
        doc,
        f"基线模型总入流约 {baseline_total:.3f} ×10^6 L，末时刻存储约 {baseline_metrics['final_stored_volume']:.3f} ×10^6 L，平衡残差约 {baseline_residual:.3f} ×10^6 L，因此连续性误差仅 {baseline_metrics['continuity_error_pct']:.3f}%。",
    )
    paragraph(
        doc,
        f"真值事件模型总入流约 {event_total:.3f} ×10^6 L，末时刻存储约 {event_metrics['final_stored_volume']:.3f} ×10^6 L，平衡残差约 {event_residual:.3f} ×10^6 L，因此连续性误差达到 {event_metrics['continuity_error_pct']:.3f}%。这说明当前事件工况虽然无溢流，但整体水量平衡比基线工况更紧、更粗糙。",
    )
    doc.add_heading("2.2 这次连续性误差不是因为漏算了存储", level=2)
    paragraph(doc, "不是。这里已经把 Initial Stored Volume 和 Final Stored Volume 都算进去了，而且它们来自 SWMM 报表中的系统级存储汇总，不是只统计单独某一个储池点。")
    paragraph(doc, f"基线工况中，总入流约 {baseline_total:.3f} ×10^6 L，末时刻存储约 {baseline_metrics['final_stored_volume']:.3f} ×10^6 L，平衡残差约 {baseline_residual:.3f} ×10^6 L，因此连续性误差仅 {baseline_metrics['continuity_error_pct']:.3f}%。")
    paragraph(doc, f"真值事件工况中，总入流约 {event_total:.3f} ×10^6 L，末时刻存储约 {event_metrics['final_stored_volume']:.3f} ×10^6 L，平衡残差约 {event_residual:.3f} ×10^6 L，因此连续性误差达到 {event_metrics['continuity_error_pct']:.3f}%。也就是说，即便把存储量算进去，事件工况依然存在很大的系统平衡残差。")

    doc.add_heading("2.3 当前数据本身是否存在问题", level=2)
    paragraph(doc, "从模板一致性看，当前数据链是闭合的：正式真值事件模型去掉三处注入即可得到正式基线模型，再把真值时序注回去可以严格达到 truth replay = 1.0。")
    add_table(
        doc,
        ["项目", "基线模型", "真值事件模型"],
        [
            ["外部入流 (10^6 L)", f"{baseline_metrics['external_inflow']:.3f}", f"{event_metrics['external_inflow']:.3f}"],
            ["Flooding Loss (10^6 L)", f"{baseline_metrics['flooding_loss']:.3f}", f"{event_metrics['flooding_loss']:.3f}"],
            ["节点溢流", "无", "无"],
            ["Continuity Error (%)", f"{baseline_metrics['continuity_error_pct']:.3f}", f"{event_metrics['continuity_error_pct']:.3f}"],
        ],
    )
    paragraph(doc, f"从数值质量看，基线连续性误差仅为 {baseline_metrics['continuity_error_pct']:.3f}%，而真值事件连续性误差达到 {event_metrics['continuity_error_pct']:.3f}%。因此，当前事件工况虽然无节点溢流，但系统级水量平衡比基线更粗糙。")
    paragraph(doc, f"从辨识结构看，把真值中的 J140 全部换成 J145 后，拟合仍然可达到 {comp['J140_to_J145']['mean_nse']:.4f}；把 J140 全部换成 J125 后也还有 {comp['J140_to_J125']['mean_nse']:.4f}。这说明强代偿关系本身就存在于当前数据之中，而不是算法后处理造成的幻象。")

    doc.add_heading("2.4 参数规模、参数维度与监测方案对结果的影响", level=2)
    add_table(
        doc,
        ["指标", "Medium", "Large"],
        [
            ["GA 最优 mean NSE", f"{medium['summary']['ga_best_mean_nse']:.4f}", f"{large['summary']['ga_best_mean_nse']:.4f}"],
            ["Posterior median NSE", f"{medium['summary']['posterior_median_nse']:.4f}", f"{large['summary']['posterior_median_nse']:.4f}"],
            ["Posterior best NSE", f"{medium['summary']['posterior_best_nse']:.4f}", f"{large['summary']['posterior_best_nse']:.4f}"],
            ["Posterior coverage mean", f"{medium['summary']['posterior_coverage_mean']:.4f}", f"{large['summary']['posterior_coverage_mean']:.4f}"],
            ["预测 top3", " / ".join(medium['summary']['predicted_top3']), " / ".join(large['summary']['predicted_top3'])],
            [
                "参数规模",
                f"{medium['summary']['config']['ga_population_count']}×{medium['summary']['config']['ga_population_size']}×{medium['summary']['config']['ga_generations']} + "
                f"{medium['summary']['config']['am_chain_count']}×{medium['summary']['config']['am_samples_per_chain']}",
                f"{large['summary']['config']['ga_population_count']}×{large['summary']['config']['ga_population_size']}×{large['summary']['config']['ga_generations']} + "
                f"{large['summary']['config']['am_chain_count']}×{large['summary']['config']['am_samples_per_chain']}",
            ],
        ],
    )
    paragraph(
        doc,
        "从参数规模上看，Large 已经显著加大了 GA 和 AM 的计算量；但结果并没有单调变得更“漂亮”。Medium 的 posterior median 分数更高，Large 的 posterior median 分数更低，却把后验结构推向了更接近真值的位置。这说明当前问题的主要限制不是“算得不够多”，而是参数空间中存在多个可竞争解释。",
    )
    doc.add_picture(str(figures["summary"]), width=Inches(6.4))
    add_caption(doc, "图 12 Medium 与 Large 关键指标对比")
    doc.add_picture(str(figures["ga"]), width=Inches(6.4))
    add_caption(doc, "图 13 GA 收敛轨迹：Large 搜索更广，但当前代数仍低于英文论文的 25 代")
    doc.add_picture(str(figures["posterior"]), width=Inches(6.5))
    add_caption(doc, "图 14 真值份额与 Medium/Large 后验中位数对比")
    doc.add_picture(str(figures["coverage"]), width=Inches(6.5))
    add_caption(doc, "图 15 监测点 posterior predictive coverage 比较")
    paragraph(
        doc,
        "当前参数维度一共 20 维，约束是非负且总和为 1。与一般无约束参数不同，这种份额反演会天然带来“一个点份额抬升、其余 19 个点份额同步被迫调整”的耦合效应，因此更容易出现多个局部代偿结构。",
    )
    paragraph(
        doc,
        f"当前监测方案已调整为 {len(MONITOR_NODES)} 个点，并用 J141 + J139 + J145 + J231 这组局部监测来强化 J140/J145 支路分离。这个方案是否真的优于上一版，需要靠新一轮运行结果验证；当前最客观的判断仍然是：问题已推进到局部簇识别阶段，但尚未证明能够唯一识别。",
    )

    doc.add_heading("第三章 论文适用性分析", level=1)
    doc.add_heading("3.1 当前实验与两篇论文的对应关系", level=2)
    add_table(
        doc,
        ["维度", "中文论文", "英文论文", "当前 0401 实验"],
        [
            ["研究场景", "工程定位导向，带实地监测/实测信息", "受控合成场景，验证方法与不确定性分析", "受控真值验证场景，用终极基线与真值事件闭合检验"],
            ["监测点特征", "监测点更多，布设更工程化", "典型示例监测点较少，可少至 3 个", f"当前 {len(MONITOR_NODES)} 个监测点，属于两者之间的中间形态"],
            ["核心约束", "总入流量约束明确", "GA→initial PPD→AM 的不确定性框架", "同时采用总过程约束与 GA-AM 框架"],
            ["反演对象", "错接节点入流量/定位", "缺陷诊断与后验概率分布", "20 候选井上的 20 维份额向量"],
            ["目标函数", "更强调水位/误差指标", "监测流量 mean NSE", "监测点流量增量的 SSE/mean NSE"],
            ["算法规模", "偏工程优化，不直接对应 GA-AM", "文中 GA 标准量级为 5×80×25，AM 最大 2000", "本轮 Large 实际为 5×80×10 + 4×2000"],
            ["结果解读", "强调工程定位与监测充分性", "强调后验分布与不确定性", "当前已经能稳定收缩到真值簇，但仍存在局部代偿"],
        ],
    )
    paragraph(doc, "中文论文更偏工程应用，监测点更多，而且有实地监测/实测数据支撑，因此更适合解释“为什么在监测信息更充分时，定位结果可以更稳定”。")
    paragraph(doc, "英文论文更偏方法可行性与不确定性分析，典型案例监测点较少，因此更适合解释“为什么在监测点较少、局部支路复杂时，后验会呈现多峰和代偿”。")
    paragraph(doc, "当前 0401 实验位于两者之间：它已经充分支持方法可行性，但并不等价于中文论文那类更充分监测条件下的工程唯一定位。")

    doc.add_heading("3.2 当前结果与中文论文的适用性边界", level=2)
    paragraph(doc, "因此，当前结果可以支持“方法在受控场景下可行”，但还不能直接外推为“在工程实测条件下已经达到中文论文那种稳定定位精度”。")

    doc.add_heading("3.3 当前结果与英文论文的适用性边界", level=2)
    paragraph(doc, "与英文论文相比，当前实验在方法链上是高度一致的：都采用 GA → initial PPD → AM → posterior predictive validation。Large 结果已经表现出典型的不确定性特征：高分代偿解与更接近真值的结构解并存。")
    paragraph(doc, "但当前问题的参数几何更难：我们反演的是 20 候选井上的 simplex 份额向量，而不是更简单的缺陷参数；当前网络也比英文论文的小规模示例更复杂，因此多峰后验和代偿更容易出现。")

    doc.add_heading("3.4 当前实验最准确的定位", level=2)
    paragraph(doc, "当前实验已经客观证明：在正式测试文件完全闭合、真值回灌严格为 1 的前提下，GA-AM 框架能够稳定识别出真值结构中的主要部分，并且把第三个真值源收缩到一个局部竞争簇。")
    paragraph(doc, "但当前实验同样客观表明：在 20 候选井、10 监测点的当前布设下，结果并不支持“唯一且稳定地逐点恢复 3 个真值源”这一更强命题。更准确的结论应当是：当前方法已经能够完成高概率簇的识别与排序，但在局部支路簇上仍然存在不可忽视的代偿与多峰后验。")

    doc.save(out_path)


def main() -> None:
    ensure_dirs()

    inp_files = sorted(MODEL_DIR.glob("*.inp"), key=lambda p: p.stat().st_size)
    if len(inp_files) < 2:
        raise FileNotFoundError("Expected two formal INP files in current_confirmed_models")
    baseline_inp = inp_files[0]
    event_inp = inp_files[-1]

    baseline_metrics, event_metrics = load_model_metrics()
    truth_replay = read_json(RESULTS_DIR / "truth_replay_check_new_monitor.json")
    compensation = read_json(RESULTS_DIR / "compensation_point_fit_check.json")
    medium = load_run_bundle("medium_run")
    large = load_run_bundle("large_run")
    nodes, links = load_layout_tables()
    network_stats = {
        "node_count": int(len(nodes)),
        "link_count": int(len(links)),
        "candidate_count": int(len(CANDIDATE_NODES)),
        "monitor_count": int(len(MONITOR_NODES)),
    }

    figures = {
        "layout": save_current_layout_figure(nodes, links),
        "cluster": save_local_cluster_figure(nodes, links),
        "truth_replay": save_truth_replay_figure(truth_replay),
        "summary": save_summary_comparison_figure(medium, large),
        "ga": save_ga_convergence_figure(medium, large),
        "posterior": save_posterior_compare_figure(medium, large),
        "compensation": save_compensation_figure(compensation),
        "coverage": save_coverage_compare_figure(medium, large),
    }

    curve_dir = RESULTS_DIR / "monitor_curves"
    monitor_images = []
    for prefix in ["J139", "J145", "J231"]:
        matches = sorted(curve_dir.glob(f"{prefix}*.png"))
        if matches:
            monitor_images.append(matches[0])

    md_path = OUTPUT_DIR / "0401_结果分析汇报材料.md"
    docx_path = OUTPUT_DIR / "0401_结果分析汇报材料.docx"
    build_markdown(md_path, baseline_metrics, event_metrics, truth_replay, medium, large, compensation, network_stats)
    build_docx(
        docx_path,
        baseline_inp,
        event_inp,
        baseline_metrics,
        event_metrics,
        truth_replay,
        medium,
        large,
        compensation,
        figures,
        monitor_images,
        network_stats,
    )
    print(json.dumps({"docx": str(docx_path), "markdown": str(md_path), "asset_dir": str(ASSET_DIR)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
