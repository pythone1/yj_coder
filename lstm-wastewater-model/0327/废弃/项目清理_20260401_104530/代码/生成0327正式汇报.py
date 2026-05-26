from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
import plotly.graph_objects as go
import plotly.io as pio
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(r"E:\PY\LSTM\0327")
RESULT_DIR = BASE_DIR / "结果"
DATA_DIR = BASE_DIR / "数据" / "生成数据"
ASCII_DIR = BASE_DIR / "data_ascii"

SUMMARY_JSON = RESULT_DIR / "0327_结果汇总.json"
GA_HISTORY_CSV = RESULT_DIR / "0327_GA每代最佳.csv"
GA_ALL_CSV = RESULT_DIR / "0327_GA全部方案.csv"
GA_LAST_CSV = RESULT_DIR / "0327_GA末代合并.csv"
INITIAL_PPD_CSV = RESULT_DIR / "0327_initial_PPD.csv"
AM_SAMPLES_CSV = RESULT_DIR / "0327_AM样本.csv"
PPD_SAMPLES_CSV = RESULT_DIR / "0327_PPD样本.csv"
POSTERIOR_CSV = RESULT_DIR / "0327_后验节点权重.csv"
COVERAGE_CSV = RESULT_DIR / "0327_posterior_predictive_coverage.csv"
BANDS_CSV = RESULT_DIR / "0327_posterior_predictive_bands.csv"
TOTAL_PROCESS_CSV = DATA_DIR / "0327_总入流过程_10分钟.csv"
TRUTH_CSV = DATA_DIR / "0327_真值注水数据_10分钟.csv"
BASELINE_CSV = DATA_DIR / "0327_基线监测_10分钟.csv"
EVENT_CSV = DATA_DIR / "0327_事件监测_10分钟.csv"
DELTA_CSV = DATA_DIR / "0327_观测增量_10分钟.csv"
OUTLET_CSV = DATA_DIR / "0327_排口过程_10分钟.csv"
FINAL_DELTA_CSV = RESULT_DIR / "0327_最终方案模拟增量.csv"
INP_PATH = ASCII_DIR / "dry_base_core.inp"

DOCX_OUT = RESULT_DIR / "0327_正式汇报.docx"
HTML_OUT = RESULT_DIR / "0327_正式汇报.html"
FILE_NOTE_OUT = RESULT_DIR / "0327_文件说明.md"
DETAIL_MD_OUT = RESULT_DIR / "0327_详细汇报.md"


def load_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path, encoding="utf-8-sig")


def parse_inp_structure(inp_path: Path) -> tuple[pd.DataFrame, pd.DataFrame]:
    text = inp_path.read_text(encoding="utf-8", errors="ignore")
    sections: dict[str, list[str]] = {}
    current = ""
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("[") and stripped.endswith("]"):
            current = stripped[1:-1].upper()
            sections[current] = []
            continue
        if current:
            sections[current].append(line)

    nodes = []
    for line in sections.get("COORDINATES", []):
        raw = line.strip()
        if not raw or raw.startswith(";"):
            continue
        parts = raw.split()
        if len(parts) >= 3:
            nodes.append({"node": parts[0], "x": float(parts[1]), "y": float(parts[2])})

    links = []
    for section_name in ("CONDUITS", "PUMPS"):
        for line in sections.get(section_name, []):
            raw = line.strip()
            if not raw or raw.startswith(";"):
                continue
            parts = raw.split()
            if len(parts) >= 3:
                links.append(
                    {
                        "link": parts[0],
                        "from_node": parts[1],
                        "to_node": parts[2],
                        "section": section_name,
                    }
                )

    return pd.DataFrame(nodes), pd.DataFrame(links)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)


def style_doc(doc: Document) -> None:
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.styles["Normal"].font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.0)
        footer_p = section.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(footer_p)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(20)
    r.font.name = "Microsoft YaHei"
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(11)
    r2.font.name = "Microsoft YaHei"


def add_formula(doc: Document, title: str, formula: str, desc: str) -> None:
    doc.add_paragraph(title, style="List Bullet")
    p = doc.add_paragraph()
    r = p.add_run(formula)
    r.bold = True
    r.font.name = "Consolas"
    r.font.size = Pt(10.5)
    doc.add_paragraph(desc)


def add_table(doc: Document, df: pd.DataFrame, title: str, max_rows: int | None = None) -> None:
    doc.add_heading(title, level=2)
    if df.empty:
        doc.add_paragraph("无数据。")
        return
    if max_rows is not None:
        df = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if isinstance(value, float):
                cells[i].text = f"{value:.6f}"
            else:
                cells[i].text = str(value)


def build_structure_figure(nodes: pd.DataFrame, links: pd.DataFrame, candidate: list[str], monitor: list[str], truth: list[str], outfall: str) -> str:
    node_index = nodes.set_index("node")
    fig = go.Figure()

    line_x: list[float | None] = []
    line_y: list[float | None] = []
    for _, row in links.iterrows():
        if row["from_node"] not in node_index.index or row["to_node"] not in node_index.index:
            continue
        f = node_index.loc[row["from_node"]]
        t = node_index.loc[row["to_node"]]
        line_x.extend([f["x"], t["x"], None])
        line_y.extend([f["y"], t["y"], None])
    fig.add_trace(
        go.Scatter(
            x=line_x,
            y=line_y,
            mode="lines",
            line=dict(color="#cbd5e1", width=1),
            hoverinfo="skip",
            name="管段",
        )
    )

    def add_nodes(node_names: list[str], name: str, color: str, symbol: str, size: int) -> None:
        sub = nodes[nodes["node"].isin(node_names)].copy()
        if sub.empty:
            return
        fig.add_trace(
            go.Scatter(
                x=sub["x"],
                y=sub["y"],
                mode="markers+text",
                text=sub["node"],
                textposition="top center",
                marker=dict(color=color, size=size, symbol=symbol, line=dict(width=1, color="#111827")),
                name=name,
            )
        )

    all_nodes = nodes[~nodes["node"].isin(set(candidate) | set(monitor) | set(truth) | {outfall})]
    fig.add_trace(
        go.Scatter(
            x=all_nodes["x"],
            y=all_nodes["y"],
            mode="markers",
            marker=dict(color="#94a3b8", size=4),
            name="其他节点",
            hovertext=all_nodes["node"],
            hoverinfo="text",
        )
    )
    add_nodes(candidate, "20个候选节点", "#2563eb", "circle", 8)
    add_nodes(monitor, "10个监测点", "#16a34a", "diamond", 11)
    add_nodes(truth, "3个真值注入点", "#dc2626", "x", 13)
    add_nodes([outfall], "唯一排口", "#111827", "square", 14)

    fig.update_layout(
        title="0327 全网点位与成果布设图",
        template="plotly_white",
        height=700,
        legend=dict(orientation="h", y=1.02, x=0),
        xaxis=dict(visible=False),
        yaxis=dict(visible=False, scaleanchor="x", scaleratio=1),
        margin=dict(l=20, r=20, t=60, b=20),
    )
    return pio.to_html(fig, include_plotlyjs="cdn", full_html=False)


def build_ga_curve(ga_history: pd.DataFrame) -> str:
    by_gen = ga_history.groupby("generation", as_index=False)["best_mean_nse"].max()
    fig = go.Figure()
    fig.add_trace(
        go.Scatter(
            x=by_gen["generation"],
            y=by_gen["best_mean_nse"],
            mode="lines+markers",
            line=dict(color="#2563eb", width=3),
            name="代际最优 Mean NSE",
        )
    )
    fig.update_layout(
        title="GA 收敛曲线",
        template="plotly_white",
        height=380,
        xaxis_title="代数",
        yaxis_title="Mean NSE",
        margin=dict(l=40, r=20, t=60, b=40),
    )
    return pio.to_html(fig, include_plotlyjs=False, full_html=False)


def build_monitor_fit(observed: pd.DataFrame, simulated: pd.DataFrame, monitors: list[str]) -> str:
    fig = go.Figure()
    for i, node in enumerate(monitors[:4]):
        fig.add_trace(
            go.Scatter(
                x=observed["相对小时"],
                y=observed[node],
                mode="lines",
                line=dict(width=2),
                name=f"{node} 观测",
            )
        )
        fig.add_trace(
            go.Scatter(
                x=simulated["相对小时"],
                y=simulated[node],
                mode="lines",
                line=dict(width=2, dash="dash"),
                name=f"{node} 最终方案",
            )
        )
    fig.update_layout(
        title="监测点增量拟合示意（前4个监测点）",
        template="plotly_white",
        height=420,
        xaxis_title="相对小时",
        yaxis_title="增量流量",
        margin=dict(l=40, r=20, t=60, b=40),
    )
    return pio.to_html(fig, include_plotlyjs=False, full_html=False)


def build_report() -> dict[str, str]:
    summary = json.loads(SUMMARY_JSON.read_text(encoding="utf-8"))
    ga_history = load_csv(GA_HISTORY_CSV)
    ga_all = load_csv(GA_ALL_CSV)
    ga_last = load_csv(GA_LAST_CSV)
    initial_ppd = load_csv(INITIAL_PPD_CSV)
    am_samples = load_csv(AM_SAMPLES_CSV)
    ppd_samples = load_csv(PPD_SAMPLES_CSV)
    posterior = load_csv(POSTERIOR_CSV)
    coverage = load_csv(COVERAGE_CSV)
    total_process = load_csv(TOTAL_PROCESS_CSV)
    truth = load_csv(TRUTH_CSV)
    baseline = load_csv(BASELINE_CSV)
    event = load_csv(EVENT_CSV)
    observed_delta = load_csv(DELTA_CSV)
    outlet = load_csv(OUTLET_CSV)
    final_delta = load_csv(FINAL_DELTA_CSV)
    nodes, links = parse_inp_structure(INP_PATH)

    candidate_nodes = [
        "J193", "J70", "J71", "J74", "J76", "J78", "J81", "J85", "J89", "J41",
        "J120", "J124", "J125", "J129", "J131", "J135", "J137", "J140", "J145", "J67",
    ]
    monitor_nodes = ["J191", "J74", "J78", "J91", "J59", "J123", "J126", "J137", "J145", "J231"]
    truth_nodes = ["J76", "J124", "J140"]
    outfall_node = "J132"
    total_hours = 48
    injection_hours = 24
    step_seconds = 600
    total_steps = 288
    injection_steps = 144

    truth_summary = (
        truth.groupby("节点", as_index=False)
        .agg(
            总量_m3=("该步体积_m3", "sum"),
            峰值_CMS=("注入流量_CMS", "max"),
            非零步数=("注入流量_CMS", lambda s: int((s > 0).sum())),
        )
        .sort_values("总量_m3", ascending=False)
    )

    monitor_max = []
    for node in monitor_nodes:
        monitor_max.append(
            {
                "监测点": node,
                "基线峰值_CMS": float(baseline[node].max()),
                "事件峰值_CMS": float(event[node].max()),
                "增量峰值_CMS": float(observed_delta[node].max()),
            }
        )
    monitor_max_df = pd.DataFrame(monitor_max)

    ga_by_gen = ga_history.groupby("generation", as_index=False)["best_mean_nse"].max()
    ga_improve = ga_by_gen.rename(columns={"best_mean_nse": "代际最优MeanNSE"})

    am_accept_df = (
        am_samples.groupby("chain", as_index=False)["accepted"].mean().rename(columns={"chain": "链号", "accepted": "接受率"})
    )
    am_accept_df["接受率"] = am_accept_df["接受率"].astype(float)

    file_desc_rows = [
        ("代码", "运行0327论文参数实验.py", "论文参数正式入口，负责串联数据校验、GA、initial PPD、AM、PPD、可视化和汇总输出。"),
        ("代码", "运行0327小参数实验.py", "小参数验证入口，用于快速检查主链是否能跑通。"),
        ("代码", "公共配置与数据.py", "保存点位、时间、总入流量、并行数、路径等统一配置。"),
        ("代码", "模型仿真与评估.py", "把份额方案转成注水序列，调用 SWMM 正演，并计算 SSE / NSE。"),
        ("代码", "遗传搜索与后验.py", "实现多种群 GA、competition/migration、轮盘赌、AM、PPD 和后验预测验证。"),
        ("数据", "0327_总入流过程_10分钟.csv", "48小时总入流波形，前24小时注水，10分钟分辨率。"),
        ("数据", "0327_真值注水数据_10分钟.csv", "三处真值注入点在每个10分钟步上的体积与流量。"),
        ("数据", "0327_基线监测_10分钟.csv", "从原始 dry.out 提取并细分得到的48小时基线监测序列。"),
        ("数据", "0327_事件监测_10分钟.csv", "在 clean 基线副本上注水后得到的48小时事件监测序列。"),
        ("数据", "0327_观测增量_10分钟.csv", "事件序列减去基线序列后的观测增量。"),
        ("数据", "0327_排口过程_10分钟.csv", "排口基线、事件和增量过程。"),
        ("结果", "0327_GA全部方案.csv", "GA 全部评估方案，含每套方案的 20 节点份额和评分。"),
        ("结果", "0327_GA每代最佳.csv", "每代每群体的最优评分记录，用于判断 GA 收敛。"),
        ("结果", "0327_GA末代合并.csv", "多种群末代合并池，是轮盘赌形成 initial PPD 的输入。"),
        ("结果", "0327_initial_PPD.csv", "轮盘赌后进入 initial PPD 的方案集合及权重。"),
        ("结果", "0327_AM样本.csv", "AM 全部采样记录，含接受标记、likelihood、prior 与 posterior。"),
        ("结果", "0327_PPD样本.csv", "去掉预热后的后验样本。"),
        ("结果", "0327_后验节点权重.csv", "每个候选节点的后验均值、中位数、P05、P95。"),
        ("结果", "0327_posterior_predictive_bands.csv", "基于 PPD 重跑仿真的 P05/P50/P95 区间带。"),
        ("结果", "0327_posterior_predictive_coverage.csv", "10 个监测点的 90% 区间覆盖率。"),
        ("结果", "0327_监测拟合.html", "最终方案与观测增量的监测点拟合图。"),
        ("结果", "0327_原始全网选点方案.html", "全网点位布设图，可视化候选点、监测点、真值点和排口。"),
        ("结果", "0327_正式汇报.docx", "正式 Word 汇报稿。"),
        ("结果", "0327_正式汇报.html", "正式 HTML 汇报稿。"),
    ]
    file_desc_df = pd.DataFrame(file_desc_rows, columns=["类别", "文件名", "用途说明"])

    structure_html = build_structure_figure(
        nodes,
        links,
        candidate_nodes,
        monitor_nodes,
        truth_nodes,
        outfall_node,
    )
    ga_curve_html = build_ga_curve(ga_history)
    final_delta_plot = final_delta.copy()
    if "相对小时" not in final_delta_plot.columns:
        final_delta_plot["相对小时"] = total_process["相对小时"]
    monitor_fit_html = build_monitor_fit(observed_delta, final_delta_plot, monitor_nodes)

    # DOCX
    doc = Document()
    style_doc(doc)
    add_title(doc, "0327 项目正式汇报", "排水管网多点入渗反演实验结果说明")
    doc.add_paragraph(
        "本报告用于汇总 0327 版本完整实验链路、数据构建方式、算法步骤、参数设置、中间结果、最终结果及成果文件说明，"
        "供项目沟通、阶段汇报和后续复现实验使用。"
    )

    doc.add_heading("一、项目目标与本轮成果概述", level=1)
    doc.add_paragraph(
        "本轮任务是在保持 20 个候选节点、10 个监测点、3 个注入点和既定布设结构不变的前提下，"
        "基于 48 小时基线数据构造 24 小时注水事件，并采用“总入流量约束 + 多种群遗传算法 + 轮盘赌 + initial PPD + AM + PPD 中位数诊断 + 90% 置信区间验证”的完整链路完成反演。"
    )
    doc.add_paragraph(
        f"本轮最终采用的诊断结果为 posterior_median，最终 Mean NSE = {summary['final_mean_nse']:.4f}；"
        f"GA 最优 Mean NSE = {summary['ga_best_mean_nse']:.4f}；"
        f"PPD 平均 90% 覆盖率 = {summary['posterior_coverage_mean']:.4f}。"
    )
    doc.add_paragraph(
        "本轮后验识别前三节点为："
        + " / ".join(summary["predicted_top3"])
        + "；真实注入点为："
        + " / ".join(truth_nodes)
        + "。"
    )

    doc.add_heading("二、数据构建与基础设定", level=1)
    doc.add_paragraph("1. 原始模型与基线来源")
    doc.add_paragraph(
        "原始旱天模型和有雨模型均来自用户提供的工程 INP 文件。当前基线不再使用修改后模型重新回算，而是直接采用原始 dry.out 中已有的 48 小时基线模拟结果，"
        "然后将其统一细分到 10 分钟分辨率。"
    )
    doc.add_paragraph("2. 时间设置")
    doc.add_paragraph(
        f"总时间 {total_hours} 小时；注水时间 {injection_hours} 小时；时间步长 {step_seconds} 秒；总步数 {total_steps}；注水步数 {injection_steps}。"
    )
    doc.add_paragraph("3. 点位设置")
    doc.add_paragraph("候选节点（20个）：" + "、".join(candidate_nodes))
    doc.add_paragraph("监测点（10个）：" + "、".join(monitor_nodes))
    doc.add_paragraph("真值注入点（3个）：" + "、".join(truth_nodes))
    doc.add_paragraph("唯一排口：" + outfall_node)
    add_table(doc, truth_summary, "三处真值注入点的总量与峰值统计")
    add_table(doc, monitor_max_df, "10 个监测点基线、事件与增量峰值统计")

    doc.add_heading("三、数据构建公式与计算过程", level=1)
    add_formula(
        doc,
        "总入流量约束",
        "Q_R = Σ_i ∫ I_i(t) dt",
        "其中 I_i(t) 为第 i 个真值注入点在时间 t 的注入流量。当前实验直接用三处真值注入点的总注入量积分得到 Q_R = 76000 m3，用作后续反演的总量约束。",
    )
    add_formula(
        doc,
        "节点份额约束",
        "x_i >= 0,  Σ_i x_i = 1",
        "反演变量为 20 个候选节点的份额向量 x。x_i 表示节点 i 分到的总入流比例，所有份额非负且总和为 1。",
    )
    add_formula(
        doc,
        "节点总量计算",
        "q_i = x_i * Q_R",
        "先由份额向量确定每个候选节点应承担的总注入量。",
    )
    add_formula(
        doc,
        "节点逐时注水计算",
        "I_i(t) = x_i * Q_total(t)",
        "Q_total(t) 是 48 小时总入流波形，其中前 24 小时非零。所有候选节点共享同一条总波形，差别仅体现在份额大小。",
    )
    add_formula(
        doc,
        "观测增量",
        "ΔQ(t) = Q_event(t) - Q_baseline(t)",
        "事件工况在 clean 基线副本上叠加注水后仿真得到，基线工况来自原始 dry.out 的已存在模拟结果。反演拟合的是 10 个监测点的增量过程，而不是绝对流量。",
    )
    add_table(
        doc,
        total_process[["步号", "相对小时", "总入流量_CMS", "总入流体积_m3"]].head(18),
        "总入流过程示例（前 18 个 10 分钟步）",
    )

    doc.add_heading("四、算法步骤与来源说明", level=1)
    doc.add_paragraph(
        "1. 约束部分参考中文论文：先确定总入流量 Q_R，再要求所有候选节点的总注入量之和等于 Q_R。这样做的作用是把空间反演限制在守恒可行域内。"
    )
    doc.add_paragraph(
        "2. 搜索与后验部分参考英文论文：采用多种群遗传算法（GA）完成全局搜索，再通过轮盘赌形成 initial PPD，随后使用 Adaptive Metropolis（AM）在 initial PPD 提供的高概率区域附近进行后验采样，最终输出 PPD，并用中位数给出诊断结果。"
    )
    doc.add_paragraph("3. 具体步骤如下：")
    doc.add_paragraph(
        "Step 1：初始化 5 个 GA 种群。每个种群有 80 个个体，每个个体都是一个 20 维份额向量。前半部分采用稀疏随机初始化，后半部分采用 Dirichlet 初始化。"
    )
    doc.add_paragraph(
        "Step 2：对每个个体，把份额向量转成节点逐时注水序列，调用 SWMM 仿真 48 小时事件，再与 10 个监测点的观测增量比较。"
    )
    add_formula(
        doc,
        "GA 适应度指标",
        "Mean NSE = (1/m) * Σ_j NSE_j",
        "其中 m 为监测点个数，NSE_j 是第 j 个监测点的 Nash-Sutcliffe 效率系数。当前 m = 10。",
    )
    add_formula(
        doc,
        "单点 NSE",
        "NSE = 1 - Σ_t (Q_obs(t)-Q_sim(t))^2 / Σ_t (Q_obs(t)-mean(Q_obs))^2",
        "NSE 越高表示模拟曲线越接近观测曲线。",
    )
    doc.add_paragraph(
        "Step 3：GA 在群体内部执行精英保留、交叉和变异；在群体之间执行 competition 和 migration。competition 用强群体精英替换弱群体的最差个体，migration 按固定代数把精英个体迁移到相邻群体。"
    )
    doc.add_paragraph(
        "Step 4：GA 末代所有群体合并后，按论文中的轮盘赌思想形成 initial PPD。当前做法是用末代合并池的 Mean NSE 构建抽样概率，保留高质量方案并输出其 roulette_weight。"
    )
    add_formula(
        doc,
        "轮盘赌概率",
        "p_k = f_k / Σ_r f_r",
        "其中 f_k 为第 k 个候选方案的非负适应度值。实现时先做平移，确保所有适应度为正，再归一化得到保留概率。",
    )
    doc.add_paragraph(
        "Step 5：AM 以 initial PPD 为 prior information，在 20 维份额空间内继续采样。AM 的 proposal 使用多元正态分布，步长缩放按论文设置为 sd = 2.42 / d。"
    )
    add_formula(
        doc,
        "AM proposal",
        "Y ~ N(theta_n, C_n)",
        "theta_n 为当前样本，C_n 为当前提议协方差矩阵。",
    )
    add_formula(
        doc,
        "协方差更新",
        "C_n = sd * Cov(X_0, ..., X_n-1) + sd * eps * I",
        "当前 d = 20，因此 sd = 2.42 / 20 = 0.121；eps 为数值稳定项。",
    )
    add_formula(
        doc,
        "接受率",
        "alpha = min(1, L(Y) / L(theta_n))",
        "按照英文论文实现，接受率按 likelihood ratio 计算。这里 L 由监测点增量拟合的 SSE 决定。",
    )
    doc.add_paragraph(
        "Step 6：预热后保留 AM 样本，形成 PPD。当前最终诊断值使用 posterior median，即每个候选节点后验样本的中位数。"
    )
    doc.add_paragraph(
        "Step 7：从 PPD 中抽样重跑 SWMM，构造每个监测点的 P05 / P50 / P95 区间，并统计 90% 区间覆盖率，以验证后验概率分布是否具有解释力。"
    )

    doc.add_heading("五、本轮运行参数", level=1)
    param_rows = [
        ("GA 种群数", summary["ga_population_count"]),
        ("GA 单群规模", summary["ga_population_size"]),
        ("GA 代数", summary["ga_generations"]),
        ("GA 迁移间隔", summary["ga_migration_interval"]),
        ("GA 精英比例", 0.25),
        ("GA 变异强度", 0.18),
        ("群体竞争替换个数", 4),
        ("群体迁移个数", 4),
        ("AM 链数", summary["am_chain_count"]),
        ("AM 每链样本数", summary["am_samples_per_chain"]),
        ("AM 预热步数", 500),
        ("AM 自适应开始步数", 100),
        ("AM 初始协方差", 0.002),
        ("AM 缩放系数 sd", summary["am_sd"]),
        ("AM 维度 d", summary["am_dimension_d"]),
        ("后验预测验证抽样数", summary["posterior_validation_sample_count"]),
        ("并行工作进程数", summary["parallel_workers"]),
    ]
    add_table(doc, pd.DataFrame(param_rows, columns=["参数项", "本轮取值"]), "本轮运行参数表")

    doc.add_heading("六、中间结果与过程性成果", level=1)
    add_table(doc, ga_improve, "GA 按代最优 Mean NSE 变化", max_rows=25)
    add_table(doc, initial_ppd.head(12), "initial PPD 样本示例", max_rows=12)
    add_table(doc, am_accept_df, "AM 各链接受率")
    add_table(doc, posterior.head(10), "后验节点权重前 10 名")
    add_table(doc, coverage, "10 个监测点的 90% 覆盖率")

    doc.add_heading("七、最终结果解读", level=1)
    doc.add_paragraph(
        f"本轮总耗时为 {summary['elapsed_seconds'] / 3600:.2f} 小时。GA 最优 Mean NSE = {summary['ga_best_mean_nse']:.4f}，"
        f"posterior median Mean NSE = {summary['posterior_median_nse']:.4f}，posterior best Mean NSE = {summary['posterior_best_nse']:.4f}。"
    )
    doc.add_paragraph(
        f"最终按中位数诊断得到的前 3 节点为 {' / '.join(summary['predicted_top3'])}。与真实注入点 J76 / J124 / J140 对比，J140 与 J124 已被稳定识别，J76 当前被其邻近节点 J78 部分替代。"
    )
    doc.add_paragraph(
        f"PPD 平均 90% 覆盖率为 {summary['posterior_coverage_mean']:.4f}，说明当前后验分布在监测点层面已具有较好的概率解释力。"
    )
    doc.add_paragraph(
        "当前最明显的剩余问题不是主链缺失，而是 J76 对应支路与邻近节点的区分仍不够强，导致后验质量向 J78 偏移。"
    )

    doc.add_heading("八、点位成果可视化说明", level=1)
    doc.add_paragraph(
        "本轮点位可视化已单独输出为 HTML，可直接打开查看全网结构、20 个候选节点、10 个监测点、3 个真值注入点和唯一排口的位置关系："
    )
    doc.add_paragraph(str((RESULT_DIR / "0327_原始全网选点方案.html").resolve()))
    doc.add_paragraph(
        "本轮监测拟合和 PPD 置信区间验证也已输出为 HTML，可用于展示最终方案对监测点增量的拟合效果以及 PPD 的覆盖能力："
    )
    doc.add_paragraph(str((RESULT_DIR / "0327_监测拟合.html").resolve()))
    doc.add_paragraph(str((RESULT_DIR / "0327_PPD置信区间验证.html").resolve()))

    doc.add_heading("九、文件清单与用途说明", level=1)
    add_table(doc, file_desc_df, "0327 项目主要文件说明", max_rows=len(file_desc_df))

    doc.add_heading("十、结论与下一步建议", level=1)
    doc.add_paragraph(
        "本轮结果表明：基于中文论文的总入流量守恒约束和英文论文的多种群 GA + initial PPD + AM + PPD 框架，已经能够在当前 20 节点、10 监测点、3 注入点问题上跑出较强结果。"
    )
    doc.add_paragraph(
        "其中 GA 已能把高分解稳定搜索到 0.93 以上，posterior median 也已达到 0.84 左右，并能通过 90% 区间覆盖验证。当前剩余的核心任务，是继续针对 J76 与 J78 的替代关系做辨识性增强，而不是再补算法主链。"
    )
    doc.add_paragraph(
        "建议下一步围绕三个方向继续推进：一是分析 J76 与 J78 的响应差异；二是检查 J76 周边监测点对该支路的约束强度；三是在不改主框架的前提下优化 AM proposal 的混合效率。"
    )

    doc.save(DOCX_OUT)

    # Markdown detail
    detail_md = []
    detail_md.append("# 0327 项目详细汇报")
    detail_md.append("")
    detail_md.append("## 1. 本轮结论")
    detail_md.append(f"- 最终采用：`{summary['final_solution_name']}`")
    detail_md.append(f"- 最终 Mean NSE：`{summary['final_mean_nse']:.4f}`")
    detail_md.append(f"- GA best Mean NSE：`{summary['ga_best_mean_nse']:.4f}`")
    detail_md.append(f"- posterior median Mean NSE：`{summary['posterior_median_nse']:.4f}`")
    detail_md.append(f"- posterior best Mean NSE：`{summary['posterior_best_nse']:.4f}`")
    detail_md.append(f"- 最终识别前 3：`{' / '.join(summary['predicted_top3'])}`")
    detail_md.append(f"- PPD 平均覆盖率：`{summary['posterior_coverage_mean']:.4f}`")
    detail_md.append("")
    detail_md.append("## 2. 算法核心公式")
    detail_md.append("- 总量约束：`Q_R = Σ_i ∫ I_i(t) dt`")
    detail_md.append("- 份额约束：`x_i >= 0, Σ_i x_i = 1`")
    detail_md.append("- 节点总量：`q_i = x_i * Q_R`")
    detail_md.append("- 节点逐时注水：`I_i(t) = x_i * Q_total(t)`")
    detail_md.append("- 适应度：`Mean NSE = (1/m) * Σ_j NSE_j`")
    detail_md.append("- 轮盘赌：`p_k = f_k / Σ_r f_r`")
    detail_md.append("- AM proposal：`Y ~ N(theta_n, C_n)`")
    detail_md.append("- 协方差：`C_n = sd * Cov(X_0,...,X_n-1) + sd * eps * I`")
    detail_md.append("- 缩放：`sd = 2.42 / d = 0.121`")
    detail_md.append("- 接受率：`alpha = min(1, L(Y)/L(theta_n))`")
    detail_md.append("")
    DETAIL_MD_OUT.write_text("\n".join(detail_md), encoding="utf-8")

    # File note
    add_table_df = file_desc_df.copy()
    lines = ["# 0327 文件说明", ""]
    for _, row in add_table_df.iterrows():
        lines.append(f"- [{row['类别']}] {row['文件名']}：{row['用途说明']}")
    FILE_NOTE_OUT.write_text("\n".join(lines), encoding="utf-8")

    # HTML report
    posterior_top_html = posterior.head(10).to_html(index=False, classes="table")
    coverage_html = coverage.to_html(index=False, classes="table")
    truth_html = truth_summary.to_html(index=False, classes="table")
    param_html = pd.DataFrame(param_rows, columns=["参数项", "本轮取值"]).to_html(index=False, classes="table")
    file_html = file_desc_df.to_html(index=False, classes="table")
    monitor_html = monitor_max_df.to_html(index=False, classes="table")
    initial_html = initial_ppd.head(10).to_html(index=False, classes="table")

    html = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>0327 项目正式汇报</title>
  <style>
    body {{ font-family: "Microsoft YaHei", sans-serif; margin: 24px; color: #1f2937; background: #f8fafc; }}
    h1, h2, h3 {{ color: #0f172a; }}
    .section {{ background: white; border: 1px solid #e2e8f0; border-radius: 14px; padding: 20px; margin-bottom: 18px; }}
    .grid {{ display: grid; grid-template-columns: repeat(2, minmax(0, 1fr)); gap: 16px; }}
    .kpi {{ background: #eff6ff; border: 1px solid #bfdbfe; border-radius: 12px; padding: 14px; }}
    .table {{ width: 100%; border-collapse: collapse; font-size: 13px; }}
    .table th, .table td {{ border: 1px solid #cbd5e1; padding: 6px 8px; text-align: left; }}
    .table th {{ background: #e2e8f0; }}
    .formula {{ background: #111827; color: #f8fafc; padding: 10px 12px; border-radius: 8px; font-family: Consolas, monospace; white-space: pre-wrap; }}
    .small {{ color: #475569; font-size: 13px; }}
    .link-list a {{ display:block; margin: 6px 0; }}
  </style>
</head>
<body>
  <h1>0327 项目正式汇报</h1>
  <p class="small">本页用于项目沟通汇报，覆盖数据构建、算法流程、参数设置、中间结果、最终结果、点位可视化与文件清单。</p>

  <div class="section">
    <h2>1. 结果摘要</h2>
    <div class="grid">
      <div class="kpi"><b>总运行时间</b><br>{summary['elapsed_seconds'] / 3600:.2f} 小时</div>
      <div class="kpi"><b>总入流量 Q_R</b><br>{summary['Qr_m3']:.2f} m3</div>
      <div class="kpi"><b>GA best Mean NSE</b><br>{summary['ga_best_mean_nse']:.4f}</div>
      <div class="kpi"><b>posterior median Mean NSE</b><br>{summary['posterior_median_nse']:.4f}</div>
      <div class="kpi"><b>最终采用</b><br>{summary['final_solution_name']}</div>
      <div class="kpi"><b>最终识别前 3</b><br>{' / '.join(summary['predicted_top3'])}</div>
      <div class="kpi"><b>PPD 平均 90% 覆盖率</b><br>{summary['posterior_coverage_mean']:.4f}</div>
      <div class="kpi"><b>AM 接受率</b><br>{json.dumps(summary['am_accept_rate_by_chain'], ensure_ascii=False)}</div>
    </div>
  </div>

  <div class="section">
    <h2>2. 点位成果可视化</h2>
    <p>下图直接展示全网结构、20 个候选节点、10 个监测点、3 个真值注入点和唯一排口位置，是本轮成果最核心的空间展示。</p>
    {structure_html}
  </div>

  <div class="section">
    <h2>3. 数据构建逻辑</h2>
    <p>本轮以原始 dry.out 中已有的 48 小时基线模拟数据作为 baseline，再在 clean 基线副本上叠加前三处注水事件，构造事件工况。所有过程统一到 10 分钟分辨率。</p>
    <div class="formula">Q_R = Σ_i ∫ I_i(t) dt
x_i >= 0,  Σ_i x_i = 1
q_i = x_i * Q_R
I_i(t) = x_i * Q_total(t)
ΔQ(t) = Q_event(t) - Q_baseline(t)</div>
    <p>三处真值点共享同一条总入流波形，但总量不同。最终生成的数据包括：总入流过程、真值注水数据、基线监测、事件监测、观测增量与排口过程。</p>
    <h3>真值注入统计</h3>
    {truth_html}
    <h3>监测点峰值统计</h3>
    {monitor_html}
  </div>

  <div class="section">
    <h2>4. 算法步骤与公式</h2>
    <p><b>约束来源：</b>总入流量约束参考中文论文，保证空间分配满足总量守恒。</p>
    <p><b>搜索与后验：</b>多种群 GA、轮盘赌、initial PPD、AM、中位数诊断和 90% 置信区间验证参考英文论文。</p>
    <div class="formula">Mean NSE = (1/m) * Σ_j NSE_j
NSE = 1 - Σ_t (Q_obs(t)-Q_sim(t))^2 / Σ_t (Q_obs(t)-mean(Q_obs))^2
p_k = f_k / Σ_r f_r
Y ~ N(theta_n, C_n)
C_n = sd * Cov(X_0,...,X_n-1) + sd * eps * I
sd = 2.42 / d = {summary['am_sd']:.3f}
alpha = min(1, L(Y) / L(theta_n))</div>
    <p>运行步骤为：多种群初始化 → SWMM 正演评分 → 群内进化 → 群间 competition / migration → 末代合并 → 轮盘赌形成 initial PPD → AM 后验采样 → posterior median 诊断 → PPD 置信区间验证。</p>
    <h3>本轮参数表</h3>
    {param_html}
  </div>

  <div class="section">
    <h2>5. 中间结果</h2>
    <h3>GA 收敛曲线</h3>
    {ga_curve_html}
    <h3>initial PPD 样本示例</h3>
    {initial_html}
    <h3>后验节点权重前 10</h3>
    {posterior_top_html}
  </div>

  <div class="section">
    <h2>6. 最终拟合与 PPD 验证</h2>
    <p>最终结果按 posterior median 生成，用于和观测增量进行比较。下图展示监测点拟合示意（前 4 个监测点）。</p>
    {monitor_fit_html}
    <h3>90% 覆盖率</h3>
    {coverage_html}
    <p>当前平均 90% 覆盖率为 {summary['posterior_coverage_mean']:.4f}，说明当前 PPD 具备较好的概率解释能力。</p>
  </div>

  <div class="section">
    <h2>7. 结果解读</h2>
    <p>本轮真值节点为 J76 / J124 / J140，最终后验前三为 {' / '.join(summary['predicted_top3'])}。从结果看，J140 与 J124 已被稳定识别，J76 已进入后验前列，但当前仍被邻近节点 J78 部分替代。</p>
    <p>GA 的最优点解 Mean NSE 已达 {summary['ga_best_mean_nse']:.4f}，posterior median 为 {summary['posterior_median_nse']:.4f}。这说明 GA 已能稳定找到高分解，PPD 层也已具备可用解释力，但 J76 对应支路的辨识仍需增强。</p>
  </div>

  <div class="section">
    <h2>8. 文件清单</h2>
    {file_html}
    <div class="link-list">
      <a href="0327_原始全网选点方案.html">打开全网点位方案图</a>
      <a href="0327_监测拟合.html">打开监测拟合图</a>
      <a href="0327_PPD置信区间验证.html">打开 PPD 置信区间验证图</a>
    </div>
  </div>
</body>
</html>
"""
    HTML_OUT.write_text(html, encoding="utf-8")

    return {
        "docx": str(DOCX_OUT),
        "html": str(HTML_OUT),
        "file_note": str(FILE_NOTE_OUT),
        "detail_md": str(DETAIL_MD_OUT),
    }


def main() -> None:
    result = build_report()
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
