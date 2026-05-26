import os
import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE = r"E:\PY\tech"
PDF_DIR = os.path.join(BASE, "tmp_pdf_from_docx")
OUT_DIR = os.path.join(BASE, "output", "doc")
IMG_DIR = os.path.join(BASE, "tmp_question_images_v2")
OUT_PATH = os.path.join(OUT_DIR, "王睿熙物理错题与同类题精练_含题图完整版.docx")


def font(run, size=11, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def p(doc, text="", size=11, bold=False, align=None):
    para = doc.add_paragraph()
    if text:
        r = para.add_run(text)
        font(r, size=size, bold=bold)
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = Pt(4)
    return para


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        r = para.add_run(item)
        font(r, size=10)


def save_clip(pdf_path, page_index, rect, out_name, zoom=2.2):
    pdf = fitz.open(pdf_path)
    page = pdf[page_index]
    clip = fitz.Rect(rect)
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), clip=clip, alpha=False)
    out_path = os.path.join(IMG_DIR, out_name)
    pix.save(out_path)
    pdf.close()
    return out_path


def add_img(doc, path, width=15.7):
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(width))


def add_similar(doc, title, source, img_path, answer, analysis):
    p(doc, title, size=12, bold=True)
    p(doc, f"来源：{source}", size=10)
    add_img(doc, img_path)
    p(doc, f"答案：{answer}", size=11, bold=True)
    p(doc, "解析：", size=11, bold=True)
    for item in analysis:
        p(doc, item, size=11)


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    os.makedirs(IMG_DIR, exist_ok=True)

    wrong_pdf = [os.path.join(BASE, n) for n in os.listdir(BASE) if n.lower().endswith(".pdf")][0]
    qinhuai_pdf = os.path.join(PDF_DIR, "精品解析：2025年江苏省南京市秦淮区中考一模物理试题（解析版）.pdf")
    gulou_pdf = os.path.join(PDF_DIR, "[56318892]2026年江苏省南京市鼓楼区名校联合一模物理试题[t].pdf")

    # Original wrong questions
    original = {
        "w10": save_clip(wrong_pdf, 0, (35, 85, 560, 268), "orig10.png"),
        "w11": save_clip(wrong_pdf, 0, (35, 260, 560, 497), "orig11.png"),
        "w12": save_clip(wrong_pdf, 0, (35, 485, 560, 600), "orig12.png"),
        "w14a": save_clip(wrong_pdf, 0, (35, 595, 560, 770), "orig14a.png"),
        "w14b": save_clip(wrong_pdf, 1, (35, 35, 560, 170), "orig14b.png"),
        "w15a": save_clip(wrong_pdf, 1, (35, 490, 560, 770), "orig15a.png"),
        "w15b": save_clip(wrong_pdf, 2, (35, 35, 560, 255), "orig15b.png"),
        "w16a": save_clip(wrong_pdf, 2, (35, 250, 560, 770), "orig16a.png"),
        "w16b": save_clip(wrong_pdf, 3, (35, 35, 560, 372), "orig16b.png"),
        "w17": save_clip(wrong_pdf, 3, (35, 375, 560, 680), "orig17.png"),
    }

    # Similar questions with full screenshots
    sim = {
        "q17": save_clip(qinhuai_pdf, 10, (35, 485, 560, 760), "sim_q17.png"),
        "q10": save_clip(qinhuai_pdf, 5, (35, 60, 560, 395), "sim_q10.png"),
        "q22": save_clip(qinhuai_pdf, 16, (35, 225, 560, 760), "sim_q22.png"),
        "q19": save_clip(qinhuai_pdf, 12, (35, 390, 560, 760), "sim_q19.png"),
        "q28": save_clip(qinhuai_pdf, 24, (35, 250, 560, 760), "sim_q28.png"),
        "q23": save_clip(qinhuai_pdf, 17, (35, 155, 560, 760), "sim_q23.png"),
        "q26": save_clip(qinhuai_pdf, 21, (35, 155, 560, 760), "sim_q26.png"),
        "q27": save_clip(qinhuai_pdf, 23, (35, 450, 560, 760), "sim_q27.png"),
        "g5": save_clip(gulou_pdf, 0, (35, 150, 560, 365), "sim_g5.png"),
    }

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)

    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)

    p(doc, "王睿熙物理错题与同类题精练", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "题图增强版：原错题和同类题都尽量保留整题截图", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "这次重点修正：同类题不再只给文字，而是直接嵌入题库里的整题截图。", size=10)

    # Type 1
    doc.add_page_break()
    p(doc, "类型一：简单机械与机械效率", size=15, bold=True)
    p(doc, "原错题（PDF第10题）", size=12, bold=True)
    add_img(doc, original["w10"])
    p(doc, "纠错：机械效率题先列 η=W有/W总=Gh/Fs，再求额外功或摩擦力；斜面省力不省功。", size=11)
    add_similar(
        doc, "同类题1",
        "《精品解析：2025年江苏省南京市秦淮区中考一模物理试题（解析版）》第17题",
        sim["q17"],
        "4000J；50W；12%。",
        [
            "W有=Gh=400×10=4000J。",
            "W总=W有/η=4000/0.8=5000J。",
            "P=W总/t=5000/100=50W。",
            "克服动滑轮重的功为 60×10=600J，占总功 12%。"
        ],
    )
    add_similar(
        doc, "同类题2",
        "《2026年江苏省南京市鼓楼区名校联合一模物理试题》 第5题",
        sim["g5"],
        "50%。",
        [
            "输入能量看男演员损失的重力势能，有用能量看女演员增加的重力势能。",
            "η=W有/W总=（45×10×4）/[2×60×10×3]=50%。"
        ],
    )

    # Type 2
    doc.add_page_break()
    p(doc, "类型二：受力、惯性与摩擦力", size=15, bold=True)
    p(doc, "原错题（PDF第11题）", size=12, bold=True)
    add_img(doc, original["w11"])
    p(doc, "纠错：匀速先判合力为0；突然停止先想到惯性。", size=11)
    add_similar(
        doc, "同类题1",
        "《精品解析：2025年江苏省南京市秦淮区中考一模物理试题（解析版）》第10题",
        sim["q10"],
        "D。",
        [
            "惯性不是力，不能和重力比较。",
            "若外力全部消失，物体应保持原速度做匀速直线运动。",
            "速度变化说明力改变了物体的运动状态。"
        ],
    )
    add_similar(
        doc, "同类题2",
        "《精品解析：2025年江苏省南京市秦淮区中考一模物理试题（解析版）》第22题",
        sim["q22"],
        "（1）匀速直线；二力平衡。（2）压力大小；D。",
        [
            "只有匀速直线拉动时，测力计示数才等于滑动摩擦力。",
            "甲乙若只改变木块上方重物，就是在研究摩擦力与压力大小的关系。"
        ],
    )

    # Type 3
    doc.add_page_break()
    p(doc, "类型三：电学综合与自动控制", size=15, bold=True)
    p(doc, "原错题（PDF第12题、第14题）", size=12, bold=True)
    add_img(doc, original["w12"])
    add_img(doc, original["w14a"])
    add_img(doc, original["w14b"])
    p(doc, "纠错：电能表题先算现有总功率；继电器题先找吸合/释放临界电流。", size=11)
    add_similar(
        doc, "同类题1",
        "《精品解析：2025年江苏省南京市秦淮区中考一模物理试题（解析版）》第19题",
        sim["q19"],
        "（1）44Ω；（2）4.4W；（3）110220J。",
        [
            "加热状态直接用 R=U/I。",
            "待机功率直接用 P=UI。",
            "总电能按电流-时间图分段求和。"
        ],
    )
    add_similar(
        doc, "同类题2",
        "《精品解析：2025年江苏省南京市秦淮区中考一模物理试题（解析版）》第28题",
        sim["q28"],
        "（1）a、b；（2）2Lx；（3）15Ω，45J；（4）更长。",
        [
            "临界吸合时 I=0.1A，所以控制电路总电阻先定出来。",
            "再由总电阻减去已知电阻，反推光敏电阻和电阻箱。",
            "最后用 W=I²Rt 计算电阻箱耗能。"
        ],
    )

    # Type 4
    doc.add_page_break()
    p(doc, "类型四：实验探究与误差分析", size=15, bold=True)
    p(doc, "原错题（PDF第15题、第16题）", size=12, bold=True)
    add_img(doc, original["w15a"])
    add_img(doc, original["w15b"])
    add_img(doc, original["w16a"])
    add_img(doc, original["w16b"])
    p(doc, "纠错：实验题先看调零、接线、故障、控制变量，再看误差方向。", size=11)
    add_similar(
        doc, "同类题1",
        "《精品解析：2025年江苏省南京市秦淮区中考一模物理试题（解析版）》第23题",
        sim["q23"],
        "（1）C；（2）高度差；（3）乙、丙；上窄下宽。",
        [
            "U形管压强计靠液面高度差反映压强大小。",
            "深度越大，液体压强越大，所以大坝要做成上窄下宽。"
        ],
    )
    add_similar(
        doc, "同类题2",
        "《精品解析：2025年江苏省南京市秦淮区中考一模物理试题（解析版）》第26题",
        sim["q26"],
        "（2）电阻R断路；（3）向右；1.5V；（4）导体两端电压一定时，电流与电阻成反比。",
        [
            "电压表有示数但电流表无示数，优先判断与电压表并联的定值电阻断路。",
            "研究电流与电阻关系时一定保持电压不变。"
        ],
    )

    # Type 5
    doc.add_page_break()
    p(doc, "类型五：压强、功率与能量转化综合", size=15, bold=True)
    p(doc, "原错题（PDF第17题）", size=12, bold=True)
    add_img(doc, original["w17"])
    p(doc, "纠错：综合计算题建议固定顺序“单位换算→压强→时间/速度→功率→效率→燃料质量”。", size=11)
    add_similar(
        doc, "同类题1",
        "《精品解析：2025年江苏省南京市秦淮区中考一模物理试题（解析版）》第27题",
        sim["q27"],
        "（1）半导体；（2）1m/s；（3）52000Pa。",
        [
            "v=s/t=20/20=1m/s。",
            "压力等于重力，F=mg=52×10=520N。",
            "压强 p=F/S=520/0.01=52000Pa。"
        ],
    )
    add_similar(
        doc, "同类题2",
        "《精品解析：2025年江苏省南京市秦淮区中考一模物理试题（解析版）》第17题",
        sim["q17"],
        "4000J；50W；12%。",
        [
            "虽然这题主线是机械效率，但同时也训练功率与能量分配，和你的第17题综合计算能力是相通的。",
            "做法仍然是先分清有用功、总功，再求功率。"
        ],
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
