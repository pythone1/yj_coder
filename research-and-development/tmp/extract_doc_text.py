from docx import Document
from pathlib import Path
p=Path(r'E:\PY\research\output\doc\今世缘酒业生产模块AI工艺分析与实施路径_详细版.docx')
doc=Document(str(p))
paras=[x.text.strip() for x in doc.paragraphs if x.text.strip()]
print('paras', len(paras), 'tables', len(doc.tables))
for i,t in enumerate(paras[:160],1):
    print(f'{i:03d} {t}')
