# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = r"E:\PY\research\0519\output\doc\复杂管网智能漏损侦听与定位系统调研报告_需求对应版_算法细化.docx"


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        shade(table.rows[0].cells[i], "1F4E79")
        cell_text(table.rows[0].cells[i], h, bold=True, color="FFFFFF", size=9)
        if widths:
            table.rows[0].cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell_text(cells[i], text, size=9)
            if widths:
                cells[i].width = Cm(widths[i])
    return table


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        set_font(r, size=16 if level == 1 else 13, bold=True, color="1F4E79")


def add_para(doc, text=""):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        set_font(r)
    return p


def add_label_para(doc, label, text):
    p = doc.add_paragraph()
    r = p.add_run(label)
    set_font(r, bold=True)
    r = p.add_run(text)
    set_font(r)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r, size=10)


def add_refs(doc, refs):
    p = doc.add_paragraph()
    r = p.add_run("参考案例/论文：")
    set_font(r, bold=True)
    for idx, (name, url) in enumerate(refs, start=1):
        p = doc.add_paragraph()
        r = p.add_run(f"{idx}. {name}：")
        set_font(r, size=9)
        add_hyperlink(p, url, url)


refs = {
    "wuxi": ("无锡市国资委，IPv6+噪声听漏物联网设备及AIOT云平台", "https://gzw.wuxi.gov.cn/doc/2024/10/21/4414908.shtml"),
    "weipai": ("威派格，NB远传型噪声监测仪", "https://www.shwpg.com/intelligent-hardware/pro9/pro39"),
    "heda": ("和达科技，渗漏预警、水务云、分区计量与GIS方案", "https://www.hddznet.com/"),
    "jining": ("济宁中山公用水务，DMA分区计量与噪声监测", "https://mhuanbao.in-en.com/html/huanbao-2397432.shtml"),
    "xintian": ("新天科技，DMA漏损管控分析", "https://wap.suntront.com/product/product24.html"),
    "mems1": ("Micromachines 2023，AlN MEMS水听器用于管道泄漏检测", "https://www.mdpi.com/2072-666X/14/3/654"),
    "mems2": ("Water 2020，AlN MEMS水听器IoT泄漏检测系统", "https://www.mdpi.com/2073-4441/12/11/2966"),
    "chinese_ai": ("声学技术，供水管网泄漏声信号分类与识别", "https://www.sxjs.ac.cn/article/doi/10.16300/j.cnki.1000-3630.23122501"),
    "micro_leak": ("RSC，微泄漏声学检测与定位研究", "https://pubs.rsc.org/doi/d3ew00686g"),
    "gutermann": ("Gutermann ZONESCAN HYDRO，多点声学相关定位", "https://en.gutermann-water.com/product/zonescan-hydro/"),
    "hwm": ("HWM PermaNET SU，固定噪声记录与机器学习分析", "https://www.hwmglobal.com/products/permanet-su/"),
    "ovarro": ("Ovarro Enigma，多点噪声记录与相关定位", "https://ovarro.com/en/global/solutions/monitoring--control-devices/data-loggers--leak-noise-loggers/leak-noise-loggers--correlators/3/enigma/2/"),
    "fido": ("FIDO AI，漏损声纹AI识别平台", "https://fido.tech/fido-ai/"),
    "fido_dc": ("FIDO DC Water，24英寸传输干管漏损识别案例", "https://fido.tech/case-studies/dc-water-detects-unseen-water-leak-on-transmission-main-using-fido-ai/"),
    "fido_tw": ("FIDO Thames Water，历史声学文件分析案例", "https://fido.tech/case-studies/thames-water-leak-team-gets-total-success-from-fido-led-work-orders/"),
    "xylem": ("Xylem/Pure SmartBall和Sahara，大口径管道检测", "https://www.xylem.com/en-us/resources/blog-posts/asset-management-begins-with-leak-detection/"),
    "ap": ("AP Sensing，长距离管线DAS监测", "https://www.apsensing.com/en/application/process-automation-and-pipeline-monitoring/pipeline-monitoring"),
}


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
style.font.size = Pt(10.5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("复杂管网智能漏损侦听与定位系统调研报告")
set_font(r, size=22, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("需求对应版：实现方案、具体算法、案例与论文链接 | 2026-05-20")
set_font(r, size=11, color="666666")

add_heading(doc, "一、需求与实现对应表", 1)
rows = [
    ("MEMS智能侦听节点", "MEMS三轴加速度计阵列 + MEMS麦克风/水听器 + 24-bit ADC + MCU边缘处理 + NB-IoT/4G通信", "阵列采集、边缘滤波、事件触发、特征提取"),
    ("背景噪声抑制≥30 dB", "频带选择 + 陷波 + VMD/CEEMDAN分解 + 波束形成 + 深度降噪", "目标泄漏频带内SNR提升"),
    ("微小滴漏识别≥95%", "声纹数据库 + CNN-LSTM/CRNN/Transformer分类模型", "泄漏/非泄漏/疑似/噪声源分类"),
    ("DN300以上定位误差≤0.25 m", "多节点同步采样 + TDOA + GCC-PHAT + 声速校准 + GIS管段约束", "泄漏点坐标和管段位置"),
    ("10万组样本数据库", "原始音频、频谱、管网参数、工况、维修确认结果统一入库", "训练集、验证集、现场复测数据"),
    ("云端AI诊断系统", "设备接入、声纹分析、定位计算、GIS展示、告警派单、维修复测", "在线监测、自动告警、处置记录"),
]
add_table(doc, ["需求目标", "实现内容", "输出结果"], rows, [4.2, 8.0, 4.8])

add_heading(doc, "二、MEMS传感器阵列智能侦听节点", 1)
add_label_para(doc, "实现方案：", "在阀门井、管廊支架、外露管段或管件处安装低功耗侦听节点。节点由MEMS三轴加速度计阵列采集管壁振动，MEMS麦克风或水听器采集声压信号，24-bit ADC完成高动态范围采样，MCU完成边缘滤波和事件触发，通信模块上传声纹片段和特征。")
add_label_para(doc, "硬件配置：", "4-16路MEMS阵元，8-48 kHz采样，24-bit ADC，低噪声前端放大，温湿度和电池状态采集，NB-IoT/4G通信，IP68外壳。")
add_label_para(doc, "特征计算：", "短时能量、均方根RMS、峰值因子、峭度、谱熵、功率谱密度PSD、Hilbert包络谱。")
add_label_para(doc, "使用算法：", "Butterworth带通滤波、IIR陷波滤波、Welch功率谱估计、STFT短时傅里叶变换、CUSUM变化检测。")
add_label_para(doc, "节点输出：", "原始音频片段、STFT谱图、Mel谱图、泄漏频带能量、设备状态、电量、信号强度。")
add_refs(doc, [refs["mems1"], refs["mems2"], refs["weipai"], refs["wuxi"]])

add_heading(doc, "三、背景噪声抑制≥30 dB", 1)
add_label_para(doc, "实现方案：", "节点先做基础滤波，云端再做多级降噪。处理对象不是全部声音，而是泄漏声所在的目标频带。每个节点保存降噪前后SNR，作为验收数据。")
add_label_para(doc, "处理流程：", "原始声纹 -> 带通滤波 -> 工频/泵频陷波 -> VMD或CEEMDAN分解 -> 小波降噪 -> MVDR波束形成 -> U-Net时频谱降噪 -> 泄漏频带增强。")
add_label_para(doc, "使用算法：", "Butterworth带通滤波、IIR陷波滤波、LMS自适应滤波、RLS自适应滤波、Wiener滤波、谱减法、VMD变分模态分解、CEEMDAN、软阈值小波降噪、Delay-and-Sum波束形成、MVDR波束形成、U-Net时频掩膜降噪。")
add_label_para(doc, "输出指标：", "目标泄漏频带内SNR提升值、降噪后泄漏概率、降噪前后谱图、噪声源类型。")
add_refs(doc, [refs["hwm"], refs["gutermann"], refs["ovarro"], refs["ap"]])

add_heading(doc, "四、多工况泄漏声纹AI识别", 1)
add_label_para(doc, "实现方案：", "把每段声纹转换为STFT谱图、Mel谱图和包络谱，输入深度学习模型。模型输出泄漏、非泄漏、疑似泄漏和噪声源类别，同时输出微小滴漏概率。")
add_label_para(doc, "特征输入：", "STFT幅值谱、Mel谱、Hilbert包络谱、Welch PSD、频带能量比。")
add_label_para(doc, "模型结构：", "声纹输入 -> 频谱特征 -> CNN提取局部频谱特征 -> BiLSTM提取时间变化 -> Attention加权 -> Softmax分类。")
add_label_para(doc, "使用算法：", "Random Forest、XGBoost、CNN、CNN-BiLSTM、CRNN、Transformer Encoder、Conformer。训练损失采用Cross Entropy或Focal Loss。")
add_label_para(doc, "输出结果：", "泄漏概率、微漏等级、噪声类型、置信度、触发定位标志。")
add_refs(doc, [refs["fido"], refs["fido_dc"], refs["chinese_ai"], refs["micro_leak"]])

add_heading(doc, "五、DN300以上多节点协同定位", 1)
add_label_para(doc, "实现方案：", "同一泄漏事件由3个以上相邻节点同步采样。云端计算节点间到达时间差，把TDOA结果投影到GIS管段上，再用多节点结果合并得到泄漏点位置。")
add_label_para(doc, "定位公式：", "x=(L+v*dt)/2。x为泄漏点距节点A距离，L为两节点间管长，v为管段声速，dt为两个节点接收泄漏声的时间差。")
add_label_para(doc, "使用算法：", "互相关时延估计、GCC-PHAT、SCOT加权互相关、TDOA定位、加权最小二乘定位、RANSAC异常节点剔除、卡尔曼滤波位置平滑。")
add_label_para(doc, "输出结果：", "泄漏管段、泄漏点坐标、距上下游节点距离、定位置信度、参与定位节点编号、互相关峰值。")
add_refs(doc, [refs["gutermann"], refs["ovarro"], refs["xylem"], refs["fido_dc"]])

add_heading(doc, "六、10万组泄漏声纹数据库", 1)
add_label_para(doc, "实现方案：", "所有节点按日上传固定时段声纹和事件声纹。数据库同时保存音频、频谱、管网信息、运行工况和维修确认结果。")
add_label_para(doc, "数据结构：", "样本ID、节点ID、采集时间、原始音频、STFT谱图、Mel谱图、管径、管材、压力、流量、节点位置、背景噪声类型、泄漏等级、维修确认结果、修复后复测结果。")
add_label_para(doc, "数据量计算：", "200个节点每天每节点5条样本，100天形成100000组；500个节点每天每节点5条样本，40天形成100000组。")
add_label_para(doc, "样本分段算法：", "原始音频统一重采样到16 kHz或48 kHz；按10秒窗口切片，5秒步长滑动；每段先做去直流和幅值归一化；用RMS能量门限保留有效片段，门限取同节点近7天RMS中位数+3倍MAD。")
add_label_para(doc, "重复样本处理算法：", "完全重复用SHA-256哈希删除；近重复先生成64维log-Mel均值向量，再计算余弦相似度，相似度>0.98且采集时间间隔<10分钟的样本只保留SNR最高的一条。")
add_label_para(doc, "异常样本剔除算法：", "对RMS、谱质心、谱带宽、零交叉率、峭度5个特征做Hampel滤波；超过中位数±3倍MAD的样本标为异常。设备离线、削顶失真、全零音频、饱和占比>1%的样本剔除。")
add_label_para(doc, "标签生成算法：", "维修确认泄漏且修复后目标频带能量下降≥6 dB，标为泄漏样本；维修复测无泄漏且后续30天无同点位告警，标为非泄漏样本；只有模型报警但无维修确认的样本标为疑似样本，不进入最终测试集。")
add_label_para(doc, "数据集划分算法：", "按企业ID、管段ID、节点ID分组做GroupKFold，训练集/验证集/测试集按70%/15%/15%划分；同一管段、同一维修事件、同一天相邻切片不能同时进入训练集和测试集。")
add_label_para(doc, "样本抽样算法：", "先用CNN倒数第二层生成声纹向量，再用K-means按噪声类型聚成20类；每类按泄漏/非泄漏/疑似比例分层抽样；人工复核优先选择softmax最大概率<0.6的低置信样本。")
add_refs(doc, [refs["fido_tw"], refs["jining"], refs["heda"], refs["xintian"]])

add_heading(doc, "七、云端AI漏损诊断服务系统", 1)
add_label_para(doc, "实现方案：", "云端接收节点数据，完成声纹识别、降噪、定位、GIS展示、告警派单和维修复测记录。维修结果回写数据库，作为下一轮模型训练标签。")
add_label_para(doc, "系统组成：", "设备接入服务、声纹数据湖、时序数据库、GIS管网库、AI识别服务、TDOA定位服务、告警中心、工单系统、模型训练服务。")
add_label_para(doc, "使用算法：", "EWMA时序异常检测、CUSUM突变检测、CNN-BiLSTM声纹分类、GCC-PHAT时延估计、TDOA定位、DBSCAN告警聚类去重、PSI模型漂移检测。")
add_label_para(doc, "部署结果：", "支持3家以上企业接入，每家企业独立设备台账、管网GIS、模型配置、告警规则和处置记录。")
add_refs(doc, [refs["wuxi"], refs["heda"], refs["jining"], refs["weipai"], refs["xintian"]])

add_heading(doc, "八、系统最终组成", 1)
final_rows = [
    ("前端节点", "MEMS三轴加速度计阵列、MEMS麦克风/水听器、24-bit ADC、MCU、NB-IoT/4G"),
    ("边缘处理算法", "Butterworth带通、IIR陷波、Welch PSD、STFT、CUSUM"),
    ("AI识别算法", "Random Forest、XGBoost、CNN、CNN-BiLSTM、CRNN、Transformer Encoder、Conformer"),
    ("降噪算法", "LMS、RLS、Wiener滤波、谱减法、VMD、CEEMDAN、小波软阈值、MVDR、U-Net"),
    ("定位算法", "互相关、GCC-PHAT、SCOT、TDOA、加权最小二乘、RANSAC、卡尔曼滤波"),
    ("数据库处理算法", "10秒滑窗切片、RMS中位数+3MAD门限、SHA-256去重、log-Mel余弦相似度近重复删除、Hampel异常剔除、修复前后6 dB能量差标签生成、GroupKFold分组划分、K-means声纹聚类抽样"),
    ("云端诊断算法", "EWMA、CUSUM、CNN-BiLSTM、GCC-PHAT、TDOA、DBSCAN、PSI漂移检测"),
]
add_table(doc, ["模块", "具体内容"], final_rows, [4.0, 12.5])

doc.save(OUT)
print(OUT)
