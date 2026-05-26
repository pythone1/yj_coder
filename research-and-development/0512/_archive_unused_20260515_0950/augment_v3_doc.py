from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SRC = ROOT / "output" / "doc" / "AI模型在供水管网DMA系统漏损检测中的应用_完善扩充红色标注版_v3.docx"
OUT = ROOT / "output" / "doc" / "AI模型在供水管网DMA系统漏损检测中的应用_完善扩充红色标注版_v3_增强.docx"
RED = RGBColor(192, 0, 0)
BLACK = RGBColor(0, 0, 0)


def style_run(run, red=True, bold=False, size=10.5, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RED if red else BLACK


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        style_run(r, False, True, 16 if level == 1 else 13, "黑体")


def p(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(21)
    para.paragraph_format.line_spacing = 1.25
    r = para.add_run(text)
    style_run(r, True)


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.line_spacing = 1.15
    r = para.add_run(text)
    style_run(r, True)


doc = Document(str(SRC))

h(doc, "附录A 算法扩展讲稿素材", 1)
p(doc, "本附录用于后续PPT备注页或演讲展开，不建议全部放进PPT正文。PPT正文应保留图、关键词和结论，具体算法解释放在讲者备注中。")

h(doc, "A.1 LSTM/GRU在DMA时序预测中的讲法", 2)
p(doc, "LSTM/GRU的核心不是“知道哪里漏了”，而是“知道正常情况下应该是什么样”。供水DMA流量通常具有明显的日周期、周周期和季节性，同时又受到天气、节假日、学校开学、商业营业、二供补水等因素影响。传统阈值只能判断是否超过固定上限，LSTM/GRU可以学习每个DMA自己的正常曲线。")
p(doc, "实际项目中，输入特征可包括过去24小时、48小时或7天的入口流量、关键压力点、时段、星期、节假日、温度、降雨、用户类型比例等。输出可以是未来1小时、3小时或24小时的正常流量预测值，也可以是预测区间。若实际值连续高于预测区间，且主要发生在夜间低用水窗口，就形成疑似漏损信号。")
p(doc, "讲解时要强调残差而非神经网络结构。听众更容易理解“模型预测正常值，实际值持续偏高，所以怀疑漏损”，而不是理解输入门、遗忘门、输出门。门控结构可放在技术附录，不必作为主线。")

h(doc, "A.2 孤立森林、DBSCAN、自编码器的工程定位", 2)
p(doc, "孤立森林适合做快速异常筛查。它的直观解释是：正常样本彼此相似，需要多次切分才能被隔离；异常样本少且不同，较容易被隔离。因此它适合对多个DMA、多维特征进行日常巡检，输出异常分数和异常清单。")
p(doc, "DBSCAN适合识别运行模式中的离群簇。比如某些夜间流量和压力组合长期处于正常密度区域，一旦出现新的低密度点簇，就可能代表边界异常、设备故障或真实漏损。DBSCAN的缺点是参数敏感，不同DMA密度差异大时，需要分区调参。")
p(doc, "自编码器适合多变量联合异常。它先学习正常状态下流量、压力、阀门、泵站等变量之间的组合关系，再用重构误差衡量新样本是否异常。若仅单点流量看起来正常，但多变量组合不符合历史规律，自编码器可能比单阈值更早发现问题。")

h(doc, "A.3 RF/GBDT在管段风险排序中的讲法", 2)
p(doc, "随机森林和梯度提升树更适合做管段风险排序，而不是实时曲线预测。它们可以把管龄、材质、口径、压力等级、道路等级、历史维修、爆管记录、周边施工、用户投诉等变量组合起来，输出每条管段的风险分。")
p(doc, "树模型的优势是可解释性相对较好。对管理人员而言，知道“这个管段高风险，是因为管龄长、维修频次高、压力波动大、附近道路施工多”比单纯看到一个黑箱分数更有价值。风险排序可用于年度检漏计划、换管计划和改造优先级。")

h(doc, "A.4 GNN和机理融合的讲法", 2)
p(doc, "供水管网天然是图结构，节点是水源、泵站、阀门、压力点、用户和水池，边是管段。GNN的价值是把这种拓扑关系纳入模型，而不是把每个传感器当作孤立点。发生漏损时，压力响应会沿拓扑传播，不同位置的漏点会产生不同的多点响应模式。")
p(doc, "但GNN不应作为第一阶段落地重点。它依赖较高质量的拓扑、足够压力点和可靠样本。对大多数项目，建议先用MNF、LSTM残差、无监督异常和树模型跑通闭环，再在重点DMA上尝试GNN或图信号处理。")
p(doc, "机理融合是更稳妥的中间路线。先用水力模型生成不同漏点和漏量的压力响应样本，再用机器学习模型学习这些响应模式。上线后，用真实事件持续校准仿真样本库。")

h(doc, "A.5 大模型智能体的边界", 2)
p(doc, "大模型智能体不应直接替代漏损检测模型。它更适合作为协同层：读取模型输出、GIS、工单、历史案例和管理制度，自动生成日报、异常解释、巡检建议和复盘报告。")
p(doc, "例如调度人员可以问：“昨晚哪些DMA异常最值得优先处理？”智能体应调用异常检测结果、MNF变化、压力点响应和历史工单，返回排序清单和理由。但最终派单、阀门操作、停水安排仍需要人工确认和权限控制。")

h(doc, "附录B PPT拆页建议", 1)
p(doc, "建议最终PPT控制在22到28页。不要把Word长段落直接搬到PPT，PPT只保留关键句和图。")
for item in [
    "封面：AI模型在供水管网DMA系统漏损检测中的应用。",
    "背景1：政策目标与行业刚需，讲漏损控制从运维问题升级为城市治理问题。",
    "背景2：传统漏损管控痛点，讲人工巡检、固定阈值和机理模型维护成本。",
    "DMA基础1：DMA是什么，讲边界、计量、压力、台账、工单。",
    "DMA基础2：DMA分区建设蓝图，用拓扑素材图讲入口流量、压力点和边界阀。",
    "MNF页：讲夜间最小流量不是漏损量，需要扣除合法夜间用水和边界误差。",
    "AI总览页：讲DMA+AI总体架构，强调AI不是单个算法。",
    "LSTM页：讲预测正常曲线和残差预警。",
    "异常检测页：讲孤立森林、DBSCAN、自编码器适合少标签场景。",
    "风险排序页：讲RF/GBDT如何做管段风险和改造优先级。",
    "机理融合页：讲水力模型提供物理约束，AI提供识别和排序。",
    "数据治理页：讲设备表、管网表、时序表、工单表、标签表。",
    "实施路径页：讲诊断、治理、建模、试点、闭环、推广。",
    "案例1：异常曲线和MNF抬升。",
    "案例2：LSTM残差和异常分数。",
    "案例3：候选管段TopN和现场复核。",
    "案例4：维修后回落和模型再训练。",
    "总结页：DMA提供边界，模型提供识别，工单提供闭环。"
]:
    bullet(doc, item)

doc.save(str(OUT))
print(OUT)
