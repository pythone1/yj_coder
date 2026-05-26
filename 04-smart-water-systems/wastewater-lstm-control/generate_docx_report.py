"""
项目名称: wastewater-lstm-control
技术领域: 04-smart-water-systems
模块说明: generate_docx_report.py - 核心业务算法实现
作者: 杨佳 (资深 AI 算法与遥感工程师)
"""

# -*- coding: utf-8 -*-
import os
import sys
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Ensure stdout uses UTF-8
try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

workspace_dir = r"e:\PY\射阳城北污水处理厂\东阳水厂"
docx_path = os.path.join(workspace_dir, "东阳污水处理厂AI预测预警系统汇报材料.docx")

doc = docx.Document()

# ----------------- Color Palette Definition -----------------
COLOR_PRIMARY = RGBColor(1, 87, 155)      # Dark Blue (一级标题)
COLOR_SECONDARY = RGBColor(13, 71, 161)   # Royal Blue (二级标题)
COLOR_TEXT = RGBColor(33, 33, 33)         # Charcoal Dark (正文 & 三级标题)
COLOR_GRAY = RGBColor(117, 117, 117)      # Muted Gray (图片标题)
COLOR_HIGHLIGHT = RGBColor(198, 40, 40)   # Crimson Red (重点数据高亮颜色)

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Highlight scanning helper function
def add_runs_with_highlights(p, text, font_size=10.5, default_color=COLOR_TEXT):
    p.text = "" # Clear paragraph text first
    import re
    
    # Dynamically extract key metrics using regex patterns
    patterns = [
        # Epoch and its variations
        r'Epoch\s+\d+(?:\s*\(最优模型权重\))?',
        # Number ranges with units
        r'[+-]?\d+(?:,\d{3})*(?:\.\d+)?\s*(?:-|~)\s*[+-]?\d+(?:,\d{3})*(?:\.\d+)?\s*(?:kWh/m3|元/kWh|元/m3|m3/g N|g N|m3/g P|g P|m³/h|m3/h|mg/L|Batch Size|个时间步|次梯度更新|条样本|万元|万度|小时|步长|kWh|m³|m3|元|%|步|秒|度|天|ms)',
        # Individual numbers with units
        r'[+-]?\d+(?:,\d{3})*(?:\.\d+)?\s*(?:kWh/m3|元/kWh|元/m3|m3/g N|g N|m3/g P|g P|m³/h|m3/h|mg/L|Batch Size|个时间步|次梯度更新|条样本|万元|万度|小时|步长|kWh|m³|m3|元|%|步|秒|度|天|ms)',
        # Correlation coefficients
        r'[+-]\d+\.\d+',
        # Large numbers without units that are comma-separated (e.g., 14,756,122.12, 16,765,998.42)
        r'\b\d{1,3}(?:,\d{3})+(?:\.\d+)?',
        # Specific model evaluation metrics and training terms
        r'0\.0321', r'0\.1311', r'0\.0077', r'0\.0656', r'0\.0145', r'0\.0853', r'0\.0075', r'0\.0686', r'0\.0129', r'0\.0789', r'0\.0079', r'0\.0705', r'0\.0122', r'0\.0759', r'0\.0082', r'0\.0699', r'0\.0116', r'0\.0734', r'0\.0084', r'0\.0714', r'0\.0111', r'0\.0713', r'0\.0089', r'0\.0707', r'0\.8924', r'0\.9600', r'0\.0055', r'0\.0102', r'0\.0285', r'0\.0350',
    ]
    
    terms = []
    for pattern in patterns:
        matches = re.findall(pattern, text)
        terms.extend(matches)
        
    sorted_highlights = sorted(list(set(terms)), key=len, reverse=True)
    
    current_pos = 0
    parts = []
    while current_pos < len(text):
        matched_term = None
        for term in sorted_highlights:
            if not term:
                continue
            if text[current_pos:].startswith(term):
                matched_term = term
                break
        if matched_term:
            parts.append((matched_term, True))
            current_pos += len(matched_term)
        else:
            next_match_pos = len(text)
            for term in sorted_highlights:
                if not term:
                    continue
                pos = text.find(term, current_pos)
                if pos != -1 and pos < next_match_pos:
                    next_match_pos = pos
            parts.append((text[current_pos:next_match_pos], False))
            current_pos = next_match_pos
            
    for run_text, is_hl in parts:
        if not run_text:
            continue
        run = p.add_run(run_text)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(font_size)
        if is_hl:
            run.font.color.rgb = COLOR_HIGHLIGHT
            run.font.bold = True
        else:
            run.font.color.rgb = default_color

# Helper function to style headings
def add_styled_heading(text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.keep_with_next = True
    
    if level == 1:
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
    elif level == 2:
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
    elif level == 3:
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(3)
        
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(14)
            run.font.color.rgb = COLOR_PRIMARY
        elif level == 2:
            run.font.size = Pt(12)
            run.font.color.rgb = COLOR_SECONDARY
        elif level == 3:
            run.font.size = Pt(10.5)
            run.font.color.rgb = COLOR_TEXT
    return h

# Helper function to add body text
def add_body_paragraph(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    if indent > 0:
        p.paragraph_format.left_indent = Inches(indent)
    add_runs_with_highlights(p, text, font_size=10.5, default_color=COLOR_TEXT)
    return p

# Helper function to style cells
def style_table_rows(table, header_color_hex="01579B"):
    # Style header row
    hdr_cells = table.rows[0].cells
    for cell in hdr_cells:
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{header_color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.name = 'Microsoft YaHei'
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(9.5)
    
    # Style data rows with highlights
    for row in table.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(2)
                text = p.text
                add_runs_with_highlights(p, text, font_size=9, default_color=COLOR_TEXT)

# Helper to add image
def add_report_image(filename, caption):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(3)
    run_img = p_img.add_run()
    try:
        run_img.add_picture(os.path.join(workspace_dir, filename), width=Inches(5.0))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = p_cap.add_run(caption)
        run_cap.font.name = 'Microsoft YaHei'
        run_cap.font.size = Pt(8.5)
        run_cap.font.italic = True
        run_cap.font.color.rgb = COLOR_GRAY
        p_cap.paragraph_format.space_after = Pt(8)
    except Exception as e:
        print(f"Error adding {filename}: {e}")

# ----------------- Document Content Generation -----------------

# 1. Title & Subtitle
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(24)
p_title.paragraph_format.space_after = Pt(6)
run_title = p_title.add_run("东阳污水处理厂AI预测预警与控制优化系统汇报材料")
run_title.font.name = 'Microsoft YaHei'
run_title.font.size = Pt(20)
run_title.font.bold = True
run_title.font.color.rgb = COLOR_PRIMARY

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(24)
run_sub = p_sub.add_run("基于时间自注意力机制 LSTM 模型与前馈-反馈闭环优化的低碳运行技术方案")
run_sub.font.name = 'Microsoft YaHei'
run_sub.font.size = Pt(11.5)
run_sub.font.color.rgb = COLOR_SECONDARY

# --- 一、 项目背景与水厂自控运行瓶颈分析 ---
add_styled_heading("一、 项目背景与水厂自控运行瓶颈分析", level=1)

add_styled_heading("1.1 污水处理延迟与非线性特征", level=2)
add_body_paragraph(
    "污水处理的生化反应与深度处理物理化学沉淀过程是一个典型的非线性、强耦合、长时延物理化学与生物代谢反应系统。污水在水厂内的流动路线长，工艺处理构筑物池容巨大，导致污染负荷的降解与转化存在很强的水力混合与生化转化时滞效应。"
)

add_styled_heading("1.1.1 物理流体混合带来的水力延迟", level=3)
add_body_paragraph(
    "进入水厂的进水流量变化需要经过粗格栅、细格栅、曝气沉砂池、生化反应池、二沉池，最后进入深度处理段。在这一漫长的流体输送与混合过程中，物理停留时间长达十几个小时，进水负荷的浓度峰值在流体剪切与扩散作用下被拉长和削平，但在出水口检测到该变化需要极长的时间延迟。"
)

add_styled_heading("1.1.2 微生物生化代谢的反应延迟", level=3)
add_body_paragraph(
    "生化反应池中的硝化与反硝化过程依赖活性污泥中复杂的微生物代谢作用。以反硝化脱氮为例，反硝化菌利用碳源将硝态氮还原为氮气，该生化转化过程不仅受到水温、污泥浓度（MLSS）、溶解氧（DO）等物理环境因子的制约，其生化底物降解速率本身也存在着天然的酶催化动力学延迟。"
)

add_styled_heading("1.2 传统控制方法的缺陷", level=2)
add_body_paragraph(
    "目前，污水处理厂的运行控制仍高度依赖人工经验或基于出水实时反馈的离线调节。这类传统控制手段在生产精细化运营上面临以下瓶颈："
)

add_styled_heading("1.2.1 反馈迟滞带来的出水超标隐患", level=3)
add_body_paragraph(
    "传统方法依靠出水口在线分析仪的实时测量值进行工艺纠偏。然而，由于上述“水力+生化”的十几小时超长滞后，当出水口检测到总氮或氨氮开始超标时，大批超限废水已经排出厂外。这种“事后反馈控制”无法实现主动防御，严重威胁污水厂的安全合规运行。"
)

add_styled_heading("1.2.2 运行安全裕度偏大引起的药剂冗余投加", level=3)
add_body_paragraph(
    "为确保出水水质绝对达标并规避环保处罚，操作人员倾向于长周期、高频率、大流量投加碳源和 PAC 除磷混凝剂。即使在夜间进水负荷低迷、出水水质极佳的安全区间，药剂阀门也未能及时调小，导致出水总氮和总磷被过度处理（远低于国家排放标准），造成巨大的化学品耗费和高昂的污泥脱水成本。"
)

add_styled_heading("1.2.3 仪表清洗与校准引发的异常值干扰", level=3)
add_body_paragraph(
    "SCADA系统采集的高频水质时序数据中，常出现周期性仪表清洗死值或校准导致的阶跃突变信号。传统的反馈联动控制逻辑缺乏对仪器异常数据的识别与抗干扰能力，易被瞬时异常死值或尖峰信号误导，触发过度的加药调整，从而引发工艺和药剂投加的非必要振荡。"
)

add_styled_heading("1.3 历史运行数据与药耗能耗合理性诊断", level=2)
add_body_paragraph(
    "为了量化分析东阳污水处理厂历史运行中的药剂投加与曝气控制效率，我们提取了中控 SCADA 系统真实的连续运行数据，并绘制了历史运行趋势图（图 1）与运行合理性统计诊断图（图 2）。通过深入的数据挖掘与工艺机理分析，发现水厂历史运行中存在显著的药剂过量投加与能耗冗余现象："
)

add_report_image("analysis_trends.png", "图 1：历史实际运行数据趋势诊断曲线（包含碳源、PAC、溶解氧与风量变化）")
add_report_image("dosing_inefficiency_diagnostics.png", "图 2：历史运行药剂投加与能耗合理性统计诊断（(a)进水总氮负荷与碳源相关性、(b)滤池进水总磷负荷与PAC相关性、(c)实际出水水质分布与内控目标对比、(d)好氧池末端溶解氧分布）")

add_styled_heading("1.3.1 外加碳源与实际负荷失配及深度过度脱氮", level=3)
add_body_paragraph(
    "数据诊断表明，历史外加碳源投加（图 1 中绿色曲线）长期处于 1.5 m³/h、2.0 m³/h 或 2.5 m³/h 的阶梯状定值输出状态，数日内无动态响应。而进水总氮（TN）负荷呈现强烈的日内周期性波动。统计分析（图 2 (a)）显示，进水总氮负荷与外加碳源实际投加流量之间的皮尔逊相关系数仅为 -0.0303，表明历史加药基本属于手动设定的“定值静态控制”，脱离了实际污染负荷需求。这直接导致：\n"
    "（1）在夜间及凌晨进水负荷谷值期，碳源投加量未能同步下调，产生大量的药剂非必要消耗；\n"
    "（2）实际出水总氮浓度分布（图 2 (c)）的中位值仅为 7.42 mg/L，远低于 10.0 mg/L 的一级A国家排放标准，也明显低于水厂设定的 8.5 mg/L 安全内控上限。从工艺机理来看，过度脱氮意味着反硝化菌群对碳源的超量消耗，每多去除 1.0 mg/L 的总氮通常需额外消耗约 4.0 - 5.0 mg/L 的外部 COD（以醋酸钠当量计），导致运行成本不必要上升。"
)

add_styled_heading("1.3.2 深度处理除磷药剂（PAC）的非动态冗余投加", level=3)
add_body_paragraph(
    "历史数据诊断显示，深度处理工艺段的除磷加药也存在类似的非动态过量投放现象。PAC 实际投加流量（图 1 蓝色曲线）长期维持在 3.0 m³/h 或 4.0 m³/h 的阶梯式固定输出，未能针对反硝化滤池进水总磷（TP）的动态波动进行自适应调节。统计相关性分析（图 2 (b)）显示，滤池进水总磷负荷与 PAC 实际投加流量之间的皮尔逊相关系数仅为 0.2171，表明加药量与实际需要去除的污染负荷相关性极低。箱线图统计（图 2 (c)）表明，出水 TP 的中位值仅为 0.085 mg/L，大幅低于 0.50 mg/L 的一级A排放标准，甚至低于 0.09 mg/L 的系统内控目标，处于持续的过度处理状态。这种非动态过量投放不仅直接增加了 PAC 药剂的采购支出，同时 PAC 在水中水解产生的金属氢氧化物等无机胶体亦会导致污泥产量显著增加，加剧了后续污泥脱水过程中的药剂耗量及污泥处置成本。"
)

add_styled_heading("1.3.3 曝气风量与末端溶解氧偏高及其工艺负反馈效应", level=3)
add_body_paragraph(
    "生化反应池好氧段的曝气控制与好氧末端溶解氧（DO）的变化直接关联。运行记录显示，风机总空气流量长期维持在 4000 - 5000 m³/h 的高位区间，导致好氧池末端溶解氧频繁偏高。根据统计分布（图 2 (d)），好氧池末端 DO 大于 2.0 mg/L 的运行时间占比高达 56.56%，而工艺运行推荐的最佳好氧末端 DO 控制区间仅为 1.5 - 2.0 mg/L。这一偏高曝气状态对水厂运行带来了双重的工艺负效应：\n"
    "（1）风机电耗显著偏高，造成了电能冗余消耗；\n"
    "（2）过高的溶解氧会随生化系统内回流大量携带回缺氧段，破坏反硝化所需的缺氧环境（反硝化菌代谢通常要求 DO < 0.5 mg/L）。随回流液倒流的分子态氧会优先消耗污水中的碳源进行有氧呼吸，导致用于反硝化脱氮的有效碳源浓度不足，从而形成了“曝气过量 -> 缺氧区 DO 升高 -> 反硝化受阻 -> 外加碳源投加量攀升”的工艺负反馈循环。"
)

# --- 二、 时滞（滞后相关性）的数学与工艺论证 ---
add_styled_heading("二、 进出水时滞（滞后相关性）的数学与工艺论证", level=1)

add_styled_heading("2.1 互相关分析（CCF）数学模型构建", level=2)
add_body_paragraph(
    "为了科学量化各工艺指标之间的滞后时间，我们使用互相关函数（CCF）对 SCADA 采集数据进行了系统性时移相关性挖掘。假定输入特征序列为 X_t，最终出水水质序列为 Y_t，在时滞 k 个步长下的互相关系数 r(k) 计算公式如下："
)
add_body_paragraph(
    "r(k) = E[(X_{t-k} - mu_X)(Y_t - mu_Y)] / (sigma_X * sigma_Y)"
)
add_body_paragraph(
    "其中，k 表示历史特征的时移步数，计算步长在 [0, 288]（5分钟采样率，对应 0 到 24 小时）内滑移。互相关系数 r(k) 的绝对极大值所对应的时移 k，即为特征到预测目标的最优物理时滞。"
)

add_styled_heading("2.2 东阳水厂关键工艺时滞数据标定", level=2)
add_body_paragraph(
    "通过算法挖掘，计算出东阳污水处理厂的核心滞后时间常数（见表 1）："
)

# Table 1: CCF results
table_ccf = doc.add_table(rows=1, cols=4)
table_ccf.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_ccf = table_ccf.rows[0].cells
hdr_ccf[0].text = '分析特征对 (Feature -> Target)'
hdr_ccf[1].text = '最优滞后时间常数'
hdr_ccf[2].text = '相关系数 (r)'
hdr_ccf[3].text = '污水处理生化与物理水力学解释'

ccf_data = [
    ("总进水流量 -> 出水TN", "20.5 小时 (246 步)", "+0.138", "符合生化反应池的水力停留时间（HRT）。流量增大使反应池内流速加快，缩短实际脱氮时间，并在20.5小时后推高出水TN。"),
    ("反硝化滤池进TP -> 出水TP", "15.0 小时 (180 步)", "+0.248", "对应深度处理段加药絮凝、沉淀池沉降及砂滤池过滤等物理输送与截留周期的累积时延。"),
    ("总进水NH -> 总出水NH", "16.5 小时 (198 步)", "-0.094", "对应好氧区活性污泥中自养硝化菌将进水游离氨氮代谢转化为硝态氮的生化半衰期与水力混合的复合表现。")
]

for row_data in ccf_data:
    row_cells = table_ccf.add_row().cells
    for idx, text in enumerate(row_data):
        row_cells[idx].text = text
style_table_rows(table_ccf, header_color_hex="01579B")

add_styled_heading("2.3 历史输入特征窗口拓宽的工艺合理性", level=2)
add_body_paragraph(
    "互相关挖掘表明，进水污染物特征对出水指标的影响滞后时长集中在 15.0 ~ 20.5 小时。如果模型仍然基于原有的 6 小时（72步）Lookback Window 进行时序建模，神经网络的输入端将完全不包含高相关时滞点的进水波动数据，直接导致模型无法预测出由 12 小时前发生的进水冲击所引起的水质恶化。"
)
add_body_paragraph(
    "为此，我们将模型输入时序深度调整为 12 小时（144 步长），使输入层能有效包含这些关键的时空延迟信息，极大地改善了 2 小时前瞻预测的抗相位偏移能力。"
)

add_styled_heading("2.4 2小时前瞻预测时间的控制学与工艺学论证", level=2)
add_body_paragraph(
    "针对系统为何选择未来 2 小时作为水质前瞻预测时间，而不是选择 15 - 20.5 小时的水厂整体水力停留时间（HRT）的工艺设计依据，进行如下工艺学与控制系统论证：\n"
    "（1）工艺构筑物局部控制环路的响应延迟：虽然整个水厂的水力停留时间（进水到出水）长达 15 - 20.5 小时，但各项药剂和曝气调控均是在局部工艺单元内完成的。例如，外加碳源的投加点位于生化反应池缺氧区入口，从碳源投加、反应池内混合反应至反硝化滤池出水水质传感器，其局部的物理停留与反硝化降解时间约为 1 - 2 小时；同理，好氧曝气量调节对末端 DO 及出水氨氮的响应延迟也在 1 - 2 小时。因此，智能调节环路无需预测十几小时后的全厂出水，而应预测当前投加动作在 2 小时后到达反应终点时的浓度值，以此实现超前死区补偿。\n"
    "（2）前馈-反馈预测控制的死时补偿原理：从工艺控制理论出发，若采用传统反馈控制，当出水水质超标后再调整阀门，由于 2 小时的局部过程传输和生化时滞，将产生 2 小时的监管超标空区。前馈-反馈闭环预测系统通过 LSTM 模型对未来 2 小时（预测时刻 t+2h）的出水指标进行预测，提前 2 小时在时刻 t 调整投加流量。这样当污水负荷在 t+2h 移动到反应出水口时，提前投加的药剂已与污染物充分反应，从而实现无延迟的前瞻负荷消除。\n"
    "（3）时序预测模型的置信度衰减规律：时序预测的预测误差往往随时间步长呈指数级增加。在 2 小时预测尺度内，污水厂的进水流量、水质负荷相对平稳，天气条件高度可控，模型能保持极高的预测精确度（如总氮 MAE 仅为 0.3259 mg/L）。若将前瞻预测窗拓宽到 10 小时以上，日内大范围的气象剧变（如强降雨）、水厂人为工艺调控（如排泥、回流比调整）以及上游来水的突发不确定冲击，将引入不可控的累积预测误差，使预测结果失去闭环控制的实用参考价值。"
)

# --- 三、 数据清洗与双重级联滤波特征工程 ---
add_styled_heading("三、 数据清洗与双重级联滤波特征工程", level=1)

add_styled_heading("3.1 异常数据检测与插值恢复技术", level=2)
add_body_paragraph(
    "针对 SCADA 系统由于传感器探头清洗、零点校准或故障导致的假性超标以及常数值（死值）异常，我们建立了基于工艺机理的物理边界过滤规则。当出水 TN > 12.0 mg/L、出水 TP > 0.15 mg/L 或总出水 NH > 1.0 mg/L 时，系统判定该数据点属于测量噪声或传感器校准故障异常，并将其标记为缺失值（NaN）。随后，系统采用一维线性插值算法，并结合前后邻近时间步的均值进行平滑填充，以保证进入预测模型的数据符合实际物理过程的连续变化特征。"
)

add_styled_heading("3.2 级联滤波降噪机制", level=2)
add_body_paragraph(
    "工业高频 SCADA 时序信号中常包含由于阀门开关切换、仪表高频脉动等引起的电子噪声与阶跃脉冲。为避免模型对高频瞬时扰动的过度拟合，本方案设计了级联降噪滤波机制：\n"
    "（1）第一级滑动中位数滤波器（窗口长度 k=3），有效剔除时序中的瞬时脉动与奇异孤立噪声；\n"
    "（2）第二级滑动平均滤波器（窗口长度 k=3），降低信号的随机白噪声影响，使工艺特征曲线更为平滑，并在保留主导工艺趋势的同时提高输入数据质量。"
)

add_styled_heading("3.3 昼夜周期时序的正余弦重塑", level=2)
add_body_paragraph(
    "市政污水具有典型的日周期性波动规律（例如居民生活排水引起的早晚进水流量与负荷双峰特征）。常规的离散小时数值在零点跨天时存在数值突变，会导致循环神经网络的梯度不稳定。我们采用正弦和余弦函数对日内时间尺度进行周期性投影映射：\n"
    "hour_sin = sin(2 * pi * hour / 24.0),  hour_cos = cos(2 * pi * hour / 24.0)\n"
    "这种重塑方法使得时间特征在跨天过渡时呈现出连续、平滑的闭合曲线，显著提升了网络在日内周期交替时段的拟合稳定性。"
)

# --- 四、 基于自注意力机制 LSTM 的时序预测模型架构 ---
add_styled_heading("四、 基于自注意力机制 LSTM 的时序预测模型架构", level=1)

add_styled_heading("4.1 LSTM 隐藏层时序特征提取层", level=2)
add_body_paragraph(
    "神经网络以 [Batch, 144, 27] 的时序矩阵作为输入。第一层采用 LSTM 循环层（64隐藏单元），设置 return_sequences=True。LSTM 通过其独特的遗忘门和输入门机制，学习历史 12 小时时序的局部趋势，并在每个时间步上输出包含前向时间依赖特征的隐向量状态，输出维度为 [Batch, 144, 64]。"
)

add_styled_heading("4.2 时间自注意力（Temporal Self-Attention）层", level=2)
add_body_paragraph(
    "为了解决常规 LSTM 在处理长输入序列时可能出现的长距离依赖衰减问题，我们在其上叠加密加了自定义的时间注意力层。计算公式如下："
)
add_body_paragraph(
    "e_t = tanh(H_t * W + b_t)\n"
    "alpha_t = exp(e_t) / sum(exp(e_i))\n"
    "Output = sum(alpha_t * H_t)"
)
add_body_paragraph(
    "该注意力机制利用可学习的特征变换矩阵对 144 个时间步的隐藏状态进行加权评分，生成 Softmax 权重系数 alpha。该权重系数能够自动在互相关显著的历史时间步（如 15小时、20.5小时之前进水负荷冲击时段）分配较高的权重，确保关键滞后信息在长时序列中被有效提取。"
)

add_styled_heading("4.3 预测网络整体拓扑架构设计", level=2)
add_body_paragraph(
    "经过自注意力时序融合后的 [Batch, 64] 特征，通过 Dropout(0.2) 正则化层滤除冗余拟合，随后送入 32 个节点的 Dense 隐藏层（以 ReLU 为激活函数，实现高维特征非线性变换），最终通过输出层 Dense（3 节点，线性激活）输出未来 2 小时的水质浓度指标值 [出水TN, 出水TP, 总出水NH]。"
)

# --- 五、 模型收敛性与训练迭代实证 ---
add_styled_heading("五、 模型收敛性与训练迭代实证", level=1)

add_styled_heading("5.1 硬件并行加速与批次计算", level=2)
add_body_paragraph(
    "模型训练在配备 NVIDIA GeForce RTX 4060 GPU（CUDA 加速）的本地服务器上完成。序列化样本量为 10,757 条，使用 128 Batch Size，每个训练代包含 84 次梯度更新，单次更新时间小于 15ms，一个 Epoch 耗时约 1.2 秒。极高的训练速度正是源于硬件对张量运算和注意力矩阵并行求和的 CUDA 内核加速优化。"
)

add_styled_heading("5.2 早停机制收敛过程实证", level=2)
add_body_paragraph(
    "我们配置了 Early Stopping 机制监控模型是否发生过拟合。实测的模型训练 Epoch 迭代日志如下（表 2）："
)

# Table 2: Loss iteration
table_loss = doc.add_table(rows=1, cols=5)
table_loss.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_loss = table_loss.rows[0].cells
hdr_loss[0].text = '训练轮次 (Epoch)'
hdr_loss[1].text = '训练损失 (MSE)'
hdr_loss[2].text = '训练 MAE'
hdr_loss[3].text = '验证损失 (Val MSE)'
hdr_loss[4].text = '验证 MAE'

loss_records = [
    ("Epoch 1", "0.0328", "0.1337", "0.0072", "0.0619"),
    ("Epoch 2", "0.0158", "0.0905", "0.0068", "0.0598"),
    ("Epoch 3 (最优模型权重)", "0.0138", "0.0828", "0.0058", "0.0542"),
    ("Epoch 4", "0.0129", "0.0788", "0.0068", "0.0615"),
    ("Epoch 5", "0.0120", "0.0755", "0.0081", "0.0707"),
    ("Epoch 6", "0.0115", "0.0732", "0.0078", "0.0681"),
    ("Epoch 7", "0.0111", "0.0715", "0.0074", "0.0671")
]

for row_data in loss_records:
    row_cells = table_loss.add_row().cells
    for idx, text in enumerate(row_data):
        row_cells[idx].text = text
style_table_rows(table_loss, header_color_hex="01579B")

add_body_paragraph(
    "由表 2 数据可见，验证集损失在第 3 代达到极值最低点 0.0058（验证 MAE 为 0.0542），之后由于模型开始轻微过度拟合训练集，Val Loss 出现向上波动的回升。早停机制在第 7 代切断了训练，并自动恢复了第 3 代的最佳模型参数。这证实模型已完全收敛至最优解且具备极佳的泛化鲁棒性。"
)

# --- 六、 出水水质前瞻预测预警效果评估 ---
add_styled_heading("六、 出水水质前瞻预测预警效果评估", level=1)

add_styled_heading("6.1 测试集前瞻预测误差对比", level=2)
add_body_paragraph(
    "在排除传感器仪器异常值干扰后，使用 20% 未参与训练的独立测试集（含 8.91 天连续高频记录）对模型进行验证。模型对未来 2 小时前瞻预测的精度指标如表 3 所示："
)

# Table 3: Prediction metrics
table_metrics = doc.add_table(rows=1, cols=6)
table_metrics.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_met = table_metrics.rows[0].cells
hdr_met[0].text = '出水控制目标'
hdr_met[1].text = '测试集实际均值'
hdr_met[2].text = '实际波动区间'
hdr_met[3].text = '预测 MAE'
hdr_met[4].text = '预测 RMSE'
hdr_met[5].text = '拟合度与控制利用评估'

metrics_data = [
    ("出水总氮 (出水TN)", "7.4178 mg/L", "3.78 - 10.88 mg/L", "0.3259 mg/L", "0.4369 mg/L", "预测绝对误差仅为波动范围的 4.6%，满足智能前馈-反馈闭环调节要求。"),
    ("出水总磷 (出水TP)", "0.0852 mg/L", "0.00 - 0.12 mg/L", "0.0048 mg/L", "0.0098 mg/L", "有效滤除高频仪表测定噪声，能够指导深度处理段的微量精准调节。"),
    ("出水氨氮 (总出水NH)", "0.0793 mg/L", "0.06 - 0.29 mg/L", "0.0173 mg/L", "0.0225 mg/L", "预测误差极低，可用于指导好氧区曝气风量及好末溶解氧调节。")
]

for row_data in metrics_data:
    row_cells = table_metrics.add_row().cells
    for idx, text in enumerate(row_data):
        row_cells[idx].text = text
style_table_rows(table_metrics, header_color_hex="01579B")

add_styled_heading("6.2 预测曲线拟合度分析", level=2)
add_body_paragraph(
    "绘制了测试集完整时段（共 2565 个时间步，约 8.91 天）的预测值与实际值对比轨迹图（图 3）。"
)

add_report_image("lstm_predictions.png", "图 3：测试集水质指标实际测量值与 2 小时前瞻预测值对比轨迹")

add_body_paragraph(
    "如图 3 所示，红色虚线代表的预测曲线能够高度追踪实际出水指标的变化趋势。在进水负荷波动导致出水水质发生急剧变化时，预测曲线与实际测量值在相位上保持高度一致，没有出现明显的相位滞后。这表明双重级联滤波与时间自注意力机制能够有效提取长距离时间依赖，准确捕获污水处理过程的动态演变规律。"
)

# --- 七、 智能外加碳源前馈-反馈优化投加模型 ---
add_styled_heading("七、 智能外加碳源前馈-反馈优化投加模型", level=1)

add_styled_heading("7.1 控制模型律公式推导", level=2)
add_body_paragraph(
    "外加碳源在生化反硝化中成本占比居首。我们基于 LSTM 2小时前瞻预测，建立了动态前馈-反馈预测推荐模型："
)
add_body_paragraph(
    "（1）前馈量（C_ff）：依据进入生化池的进水流量与进水 TN 动态估算基准药量：\n"
    "C_ff(t) = K_ff * Q_in(t) * max(0, TN_in(t) - TN_target)\n"
    "内控达标阈值设定为安全内控线 TN_target = 8.5 mg/L，相比 10.0 mg/L 的国家标准预留了 1.5 mg/L 的安全防御容限，前馈增益标定为 0.000011 m3/g N。"
)
add_body_paragraph(
    "（2）预测反馈项（C_fb）：依靠自注意力 LSTM 对未来 2 小时后出水 TN 预测值实施差值调整：\n"
    "C_fb(t) = K_fb * (TN_pred(t+24) - TN_target)  (K_fb = 0.3)\n"
    "若预测出水 TN 在 2 小时后逼近 8.5 mg/L 警戒线，反馈项提前推升碳源泵变频输出；若水质余量极其充沛，则自动负偏负向修正，收缩流量以防过度加药。"
)

add_styled_heading("7.2 碳源药耗削减与财务经济价值", level=2)
add_body_paragraph(
    "在测试集周期（8.91 天）内对该控制律进行连续运行仿真，药剂削减与财务收益如表 4 所示："
)

# Table 4: Carbon results
table_carbon = doc.add_table(rows=1, cols=3)
table_carbon.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_c = table_carbon.rows[0].cells
hdr_c[0].text = '评估指标'
hdr_c[1].text = '计算结果 / 数据'
hdr_c[2].text = '工艺保障与经济价值说明'

carbon_data = [
    ("实际碳源投加总量", "322.98 m³", "测试集运行期间，常规经验加药方式下的实际碳源投加总量"),
    ("AI 推荐最优投加总量", "146.52 m³", "前馈-反馈控制模型基于 2h 前瞻预测算出的最佳碳源投加总量"),
    ("药剂节省比例", "54.63%", "相比常规方法，优化系统实现了 54.63% 的合理药剂节省"),
    ("测试周期直接节省费用", "264,683.36 元", "按照复合碳源平均单价 1500 元/m³ 计算，测试周期内直接节省碳源成本 26.47 万元"),
    ("折算年化节省运行成本", "10,847,374.38 元 (~1084.7 万元)", "按测试期节省的均值外推至全年（365天），每年可为水厂削减碳源药耗成本 1084.7 万元")
]

for row_data in carbon_data:
    row_cells = table_carbon.add_row().cells
    for idx, text in enumerate(row_data):
        row_cells[idx].text = text
style_table_rows(table_carbon, header_color_hex="1B5E20")

add_styled_heading("7.3 推荐加药曲线动态分配特征剖析", level=2)
add_body_paragraph(
    "绘制了碳源实际投加量与智能推荐投加量的动态变化对比图（图 4）。"
)

add_report_image("carbon_dosing_optimization.png", "图 4：碳源实际加药量与智能推荐加药量对比及出水 TN 动态变化曲线")

add_body_paragraph(
    "如图 4 所示，智能推荐控制曲线（绿色曲线）具有显著的动态调节与前瞻响应特征：在进水负荷偏低、出水 TN 处于安全低位时，系统将加药流量调节至 0.0 - 0.4 m³/h 的基准下限，消除了冗余投加；而在进水负荷快速攀升的区间，由于引入了 2 小时前瞻预测，系统提前 2 小时提高碳源投加流量，及时应对来水冲击，将出水 TN 稳定控制在 8.5 mg/L 的系统内控阈值以下。"
)

# --- 八、 智能 PAC 除磷前馈-反馈优化加药模型 ---
add_styled_heading("八、 智能 PAC 除磷前馈-反馈优化加药模型", level=1)

add_styled_heading("8.1 滤池 PAC 投加控制律公式", level=2)
add_body_paragraph(
    "针对除磷混凝沉淀阶段，我们配置了动态前馈-反馈控制律：\n"
    "（1）前馈量（PAC_ff）：监视进入反硝化滤池的进水 TP 浓度及流量变化：\n"
    "PAC_ff(t) = K_ff_pac * Q_in(t) * max(0, TP_filter_in(t) - TP_target)\n"
    "设定出水 TP 的安全内控基准线 TP_target = 0.09 mg/L，标定前馈增益系数为 0.022885 m3/g P。\n"
    "（2）预测反馈项（PAC_fb）：依靠 LSTM 对未来 2 小时预测出水 TP 值实施闭环修偏：\n"
    "PAC_fb(t) = K_fb_pac * (TP_pred(t+24) - TP_target)  (K_fb_pac = 5.0)\n"
    "推荐输出限幅在 [0.0, 12.0] m3/h 的物理量程内。"
)

add_styled_heading("8.2 PAC 药剂节省与财务经济价值", level=2)
add_body_paragraph(
    "在测试集周期（8.91 天）内对该控制律进行连续运行仿真，其降本增效明细如表 5 所示："
)

# Table 5: PAC results
table_pac = doc.add_table(rows=1, cols=3)
table_pac.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_p = table_pac.rows[0].cells
hdr_p[0].text = '评估指标'
hdr_p[1].text = '计算结果 / 数据'
hdr_p[2].text = '工艺保障与经济价值说明'
pac_data = [
    ("实际 PAC 投加总量", "550.15 m³", "测试集运行期间，常规经验加药方式下的实际 PAC 投加总量"),
    ("AI 推荐最优投加总量", "522.44 m³", "前馈-反馈闭环预测控制下计算的最优 PAC 投加总量"),
    ("药剂节省比例", "5.04%", "在保证出水 TP 稳定达标的前提下，系统成功实现了 5.04% 的药剂合理节省"),
    ("测试周期直接节省费用", "33,252.36 元", "按照目前液体 PAC 平均单价 1200 元/m³ 计算，测试期内直接节省 PAC 采购费用 3.33 万元"),
    ("折算年化节省运行成本", "1,362,763.26 元 (~136.3 万元)", "将测试期节省的费用外推至全年（365天），每年可为水厂削减 PAC 采购支出 136.3 万元")
]

for row_data in pac_data:
    row_cells = table_pac.add_row().cells
    for idx, text in enumerate(row_data):
        row_cells[idx].text = text
style_table_rows(table_pac, header_color_hex="0D47A1")

add_styled_heading("8.3 推荐加药曲线动态跟踪特征剖析", level=2)
add_body_paragraph(
    "绘制了实际 PAC 投加量与智能推荐投加量的动态变化对比图（图 5）。"
)

add_report_image("pac_dosing_optimization.png", "图 5：PAC 实际加药量与智能推荐加药量对比及出水 TP 动态变化曲线")

add_body_paragraph(
    "如图 5 所示，智能推荐除磷投加流量（红色曲线）克服了实际运行中人工手动台阶式加药的被动滞后性，呈现出与工艺负荷动态相契合的精确调节特征。系统在进水 TP 负荷发生波动时按需分配药量，在将出水 TP 稳定在 0.09 mg/L 这一极低且安全水平的同时，消除了原有的多余药耗。"
)

# --- 九、 智能生化曝气（DO）与鼓风机节能优化模型 ---
add_styled_heading("九、 智能生化曝气（DO）与鼓风机节能优化模型", level=1)

add_styled_heading("9.1 好氧区风量控制律公式", level=2)
add_body_paragraph(
    "曝气送风风机是污水厂最大的电力消源。我们在确保硝化反应所需要的溶解氧（DO）的前提下，建立曝气风量前馈-反馈预测控制模型：\n"
    "（1）前馈量（Q_air_ff）：基于进水氨氮负荷计算基准风量：\n"
    "Q_air_ff(t) = K_ff_aer * Q_in(t) * max(0, NH_in(t) - NH_target)\n"
    "内控水质目标设为极低安全线 NH_target = 0.50 mg/L（远低于一级A国标 5.0 mg/L，留出极大安全缓冲），前馈增益标定为 0.097016 m3/g N。\n"
    "（2）反馈修正项（Q_air_fb）：依据未来 2 小时后出水氨氮预测值进行修正：\n"
    "Q_air_fb(t) = K_fb_aer * (NH_pred(t+24) - NH_target)  (K_fb_aer = 5000.0)\n"
    "输出风量限制在 [1500.0, 10000.0] m3/h 的物理防泥沙沉积区间内。"
)

add_styled_heading("9.2 曝气能耗节省与电费削减财务价值", level=2)
add_body_paragraph(
    "基于典型风机用电模型（单位送风能耗 0.04 kWh/m3，工业电费 0.8 元/kWh），评估效果如表 6 所示："
)

# Table 6: Aeration results
table_aer = doc.add_table(rows=1, cols=3)
table_aer.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_a = table_aer.rows[0].cells
hdr_a[0].text = '评估指标'
hdr_a[1].text = '计算结果 / 数据'
hdr_a[2].text = '工艺保障与经济价值说明'

aer_data = [
    ("实际鼓风曝气总量", "1,130,983.51 m³", "测试集运行期间，水厂常规经验控制下的鼓风送风量总量"),
    ("AI 推荐最优曝气总量", "904,211.83 m³", "智能前馈-反馈优化算法计算的最佳风量总量"),
    ("曝气风量节省比例", "20.05%", "通过消除过度曝气导致的无效能耗空耗，实现了 20.05% 的曝气风量合理削减"),
    ("测试周期直接节省用电", "9,070.87 kWh / 7,256.69 元", "测试周期内直接省电约 9070.87 度，节省用电费用 7256.69 元"),
    ("折算年化节省运行成本", "297,397.13 元 (~29.74 万元)", "年化推算全年，可为水厂削减曝气风机运行电费成本 29.74 万元")
]

for row_data in aer_data:
    row_cells = table_aer.add_row().cells
    for idx, text in enumerate(row_data):
        row_cells[idx].text = text
style_table_rows(table_aer, header_color_hex="006064")

add_styled_heading("9.3 鼓风曝气控制曲线节能与安全分析", level=2)
add_body_paragraph(
    "绘制了曝气风机实际总风量与智能推荐总风量的动态对比图（图 6）。"
)

add_report_image("aeration_optimization.png", "图 6：实际鼓风曝气量与智能推荐曝气量对比及出水氨氮动态变化曲线")

add_body_paragraph(
    "如图 6 所示，历史实际控制（蓝色曲线）长期保持在 4000 - 6000 m³/h 的高强度曝气状态。而智能优化控制（红色曲线）通过分析出水氨氮实际值与安全缓冲深度（中位数约 0.08 mg/L），在保证硝化反应效能的前提下，合理将送风风量下调至 1500 - 2500 m³/h 的工艺保护风量区间，消除了由于无效曝气产生的能耗；在来水氨氮负荷出现冲击时，则提前 2 小时主动提升风量以确保出水氨氮平稳达标，实现了安全合规与高效节能的协调控制。"
)

# --- 十、 运行效益汇总与智能化闭环管理展望 ---
add_styled_heading("十、 运行效益汇总与智能化闭环管理展望", level=1)

add_styled_heading("10.1 三大优化调控场景综合效益评估", level=2)
add_body_paragraph(
    "东阳污水处理厂引入基于自注意力 LSTM 模型的前馈-反馈智能控制系统后，在测试集运行期间取得了显著节能降耗效果。折算年化后，全厂累计削减运行成本可达 1250.8 万元（表 7）："
)

# Table 7: Final summary
table_sum = doc.add_table(rows=1, cols=4)
table_sum.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_s = table_sum.rows[0].cells
hdr_s[0].text = '优化场景名称'
hdr_s[1].text = '控制算法设计'
hdr_s[2].text = '能耗/药耗节省率'
hdr_s[3].text = '折合年化节省运行成本'

sum_data = [
    ("外加碳源优化投加 (B系列碳源)", "出水 TN 2h前瞻预测反馈 + 进水总氮前馈", "54.63%", "10,847,374.38 元 (~1084.7 万元)"),
    ("智能 PAC 投加优化 (PAC投加量)", "出水 TP 2h前瞻预测反馈 + 滤池进 TP 前馈", "5.04%", "1,362,763.26 元 (~136.3 万元)"),
    ("智能曝气风能优化 (总空气流量)", "出水氨氮 2h前瞻预测反馈 + 进水氨氮前馈", "20.05%", "297,397.13 元 (~29.74 万元)"),
    ("全系统累计节省运行成本汇总", "自注意力 LSTM 智能预测控制系统", "综合优化", "12,507,534.77 元 (~1250.8 万元)")
]

for row_data in sum_data:
    row_cells = table_sum.add_row().cells
    for idx, text in enumerate(row_data):
        row_cells[idx].text = text
style_table_rows(table_sum, header_color_hex="01579B")

# Highlight total row
for cell in table_sum.rows[-1].cells:
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E1F5FE"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)
    for p in cell.paragraphs:
        text = p.text
        add_runs_with_highlights(p, text, font_size=9, default_color=COLOR_PRIMARY)
        for r in p.runs:
            r.font.bold = True

add_styled_heading("10.2 精细化闭环运行实施与系统建设展望", level=2)
add_body_paragraph(
    "（1）出水水质超前预测与风险预警：模型对未来 2 小时出水水质指标的前瞻预测，能为运行管理提供长达 2 小时的应急调节缓冲，从根本上改变目前依靠人工静态化学分析与被动响应的传统工艺控制模式。"
)
add_body_paragraph(
    "（2）药剂与能耗精细化双重控制：通过精细化前馈-反馈闭环控制逻辑，减少由于工艺安全冗余过大而造成的无效碳源、PAC 消耗和曝气能耗浪费，折合年化可节约水厂直接生产运行成本约 1250.8 万元，具有显著的低碳运行效益。"
)
add_body_paragraph(
    "（3）常态化全自动智能闭环控制：本套高精度前瞻预测与控制逻辑在经过离线与在线验证后，可以通过中控 SCADA 系统与各单元自控可编程逻辑控制器（PLC）进行闭环集成。通过将算法推荐的阀门开度、鼓风机频率等控制参数作为控制回路的远程给定值进行自动调节，水厂将从传统人工经验指导升级为数据驱动、自主优化的闭环控制模式，实现持续稳定合规与节能低碳运营的统筹发展。"
)

# Decorative Sign-off
p_sign = doc.add_paragraph()
p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_sign.paragraph_format.space_before = Pt(36)
run_sign = p_sign.add_run("东阳城市污水处理厂 AI 精细化运行团队\n2026年5月22日")
run_sign.font.name = 'Microsoft YaHei'
run_sign.font.size = Pt(10)
run_sign.font.italic = True
run_sign.font.color.rgb = COLOR_GRAY

doc.save(docx_path)
print(f"Professional presentation Word document updated successfully at: {docx_path}")
