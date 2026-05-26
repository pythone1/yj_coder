from pathlib import Path
import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


BASE = Path(r"E:\PY\research\0526")
OUT_DIR = BASE / "output" / "doc"
md_candidates = sorted(OUT_DIR.glob("*漏水*.md"), key=lambda p: p.stat().st_mtime, reverse=True)
if not md_candidates:
    md_candidates = sorted(OUT_DIR.glob("*.md"), key=lambda p: p.stat().st_mtime, reverse=True)
if not md_candidates:
    raise SystemExit("No markdown report found")

md_path = md_candidates[0]
docx_path = OUT_DIR / (md_path.stem + ".docx")
lines = md_path.read_text(encoding="utf-8").splitlines()

highlight_terms = [
    "本项目聚焦“大连路隧道内部管网泄漏监测”",
    "大连路隧道公开资料显示其为黄浦江越江公路盾构隧道，不是综合管廊",
    "公开资料未直接披露",
    "这部分必须向业主或运维单位确认",
    "消防给水管网",
    "冲洗/生产给水管网",
    "隧道排水管网",
    "引道雨水管网",
    "结构渗漏收集/导排系统",
    "电缆通道排水/水浸监测",
    "泡沫/水喷雾/细水雾系统",
    "隧道给水采用生产、生活和消防分开的给水系统",
    "冲洗废水、结构渗漏水、消防废水及引道雨水",
    "全长约 2526.88 m",
    "泥水盾构施工",
    "排水监测",
    "确认本次纳入范围的管网",
    "索要管网竣工图/BIM/CAD",
    "索要传感器点表",
    "索要历史数据",
    "索要事件台账",
    "压力/流量",
    "液位",
    "泵状态",
]


def set_run_font(run, size=None, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Microsoft YaHei")
    fonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(fonts)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def split_highlight_chunks(text):
    chunks = [(text, False)]
    for term in sorted(highlight_terms, key=len, reverse=True):
        next_chunks = []
        for chunk, highlighted in chunks:
            if highlighted or term not in chunk:
                next_chunks.append((chunk, highlighted))
                continue
            parts = chunk.split(term)
            for idx, part in enumerate(parts):
                if part:
                    next_chunks.append((part, False))
                if idx < len(parts) - 1:
                    next_chunks.append((term, True))
        chunks = next_chunks
    return chunks


def add_text_run(paragraph, text, size=10.5, bold=False):
    for chunk, highlighted in split_highlight_chunks(text):
        run = paragraph.add_run(chunk)
        set_run_font(run, size=size, bold=bold)
        if highlighted:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def add_inline(paragraph, text, size=10.5, bold=False):
    text = text.replace("**", "")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    pos = 0
    pattern = re.compile(r"\[([^\]]+)\]\((https?://[^)]+)\)|<((?:https?://)[^>]+)>")
    for match in pattern.finditer(text):
        if match.start() > pos:
            add_text_run(paragraph, text[pos : match.start()], size=size, bold=bold)
        if match.group(1):
            add_hyperlink(paragraph, match.group(1), match.group(2))
        else:
            url = match.group(3)
            add_hyperlink(paragraph, url, url)
        pos = match.end()
    if pos < len(text):
        add_text_run(paragraph, text[pos:], size=size, bold=bold)


def is_table_start(index):
    if index + 1 >= len(lines):
        return False
    return lines[index].strip().startswith("|") and re.match(
        r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$",
        lines[index + 1],
    )


def parse_table_row(line):
    return [cell.strip().replace("<", "").replace(">", "") for cell in line.strip().strip("|").split("|")]


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

styles = doc.styles
styles["Normal"].font.name = "Microsoft YaHei"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(10.5)
for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
    styles[style_name].font.name = "Microsoft YaHei"
    styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles[style_name].font.size = Pt(size)
    styles[style_name].font.bold = True

first_title = True
idx = 0
while idx < len(lines):
    stripped = lines[idx].strip()
    if not stripped:
        idx += 1
        continue

    if is_table_start(idx):
        headers = parse_table_row(lines[idx])
        idx += 2
        rows = []
        while idx < len(lines) and lines[idx].strip().startswith("|"):
            rows.append(parse_table_row(lines[idx]))
            idx += 1

        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for col_idx, header in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.text = ""
            add_inline(paragraph, header, size=9, bold=True)

        for row in rows:
            cells = table.add_row().cells
            for col_idx in range(len(headers)):
                cell = cells[col_idx]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraph = cell.paragraphs[0]
                paragraph.text = ""
                add_inline(paragraph, row[col_idx] if col_idx < len(row) else "", size=8.5)
        doc.add_paragraph()
        continue

    heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
    if heading:
        level = len(heading.group(1))
        title = heading.group(2).strip()
        if level == 1 and first_title:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text_run(paragraph, title, size=18, bold=True)
            first_title = False
        else:
            paragraph = doc.add_heading(level=min(level, 3))
            paragraph.clear()
            add_text_run(paragraph, title, size={1: 16, 2: 14, 3: 12}.get(level, 11), bold=True)
        idx += 1
        continue

    numbered = re.match(r"^(\d+)[\.、]\s*(.*)$", stripped)
    if numbered:
        paragraph = doc.add_paragraph()
        add_inline(paragraph, f"{numbered.group(1)}、{numbered.group(2)}")
        idx += 1
        continue

    if stripped.startswith("- "):
        paragraph = doc.add_paragraph(style="List Bullet")
        add_inline(paragraph, stripped[2:])
        idx += 1
        continue

    paragraph = doc.add_paragraph()
    add_inline(paragraph, stripped)
    idx += 1

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("大连路隧道漏水监测项目调研报告")
set_run_font(run, size=9, color=(100, 100, 100))

doc.save(docx_path)
print(docx_path)
