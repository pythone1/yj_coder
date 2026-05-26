# -*- coding: utf-8 -*-
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell padding (in dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level, space_before=12, space_after=6):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(space_before)
    heading.paragraph_format.space_after = Pt(space_after)
    heading.paragraph_format.keep_with_next = True
    
    # Set colors and fonts
    run = heading.runs[0]
    run.font.name = 'Microsoft YaHei'
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(31, 78, 121) # Dark Blue
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(46, 117, 182) # Medium Blue
    else:
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return heading

def add_paragraph_styled(doc, text="", space_after=6, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    p.alignment = align
    if text:
        run = p.add_run(text)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(10.5)
        run.font.bold = bold
        run.font.italic = italic
    return p

def add_bullet_styled(doc, label, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    # Bold part
    if label:
        run_l = p.add_run(label)
        run_l.font.name = 'Microsoft YaHei'
        run_l.font.size = Pt(10.5)
        run_l.font.bold = True
        
    run_t = p.add_run(text)
    run_t.font.name = 'Microsoft YaHei'
    run_t.font.size = Pt(10.5)
    return p

def add_callout(doc, title, content_lines, color_hex="F2F5F8", border_hex="2E75B6"):
    """Create a shaded box for notes/alerts."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.0)
    set_cell_background(cell, color_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Set left border only (simulate callout box)
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    # Add content to cell
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run(title)
    run_title.font.name = 'Microsoft YaHei'
    run_title.font.bold = True
    run_title.font.size = Pt(10.5)
    run_title.font.color.rgb = RGBColor(46, 117, 182)
    
    for line in content_lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(10)

def main():
    doc = docx.Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Document Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(6)
    run_title = title.add_run("中学物理电学与力学综合题解析报告")
    run_title.font.name = 'Microsoft YaHei'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(31, 78, 121)
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    run_sub = subtitle.add_run("第 25、18、19、15、16 题详细分析与多方案解答")
    run_sub.font.name = 'Microsoft YaHei'
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(128, 128, 128)
    
    # Intro
    add_paragraph_styled(doc, "本报告针对您提供的物理题目中第 25、18、19、15、16 题进行了详细的受力分析、电路拓扑推导与计算验证。对于第 16 题（2），除了便签记录的解答外，另外设计了 3 种更易理解与直接的解题方案。")
    
    # ------------------ Q25 ------------------
    add_heading_styled(doc, "第 25 题：液体密度杠杆秤（力矩平衡）", level=1)
    add_paragraph_styled(doc, "本题利用杠杆平衡原理制作一个测量液体密度的仪器。相关参数如下：", space_after=4)
    add_bullet_styled(doc, "重物 A 质量：", "m_A = 10 g")
    add_bullet_styled(doc, "两空容器位置：", "固定在两侧的 10 cm 刻度线处。")
    add_bullet_styled(doc, "水及待测液体体积：", "V = 30 mL = 30 cm³")
    
    add_heading_styled(doc, "详细解答步骤：", level=2)
    add_bullet_styled(doc, "(1) 调平方向：", "图甲中杠杆左端下沉（左重右轻），应将平衡螺母向右（即图中 → 方向）调节，使杠杆在水平位置平衡。")
    
    add_bullet_styled(doc, "(2) 待测液体密度计算：", "由于两空容器挂在杠杆两侧等长的位置（均为 10 cm），且在实验前调平，因此空容器的重力力矩在两侧相互抵消。")
    add_paragraph_styled(doc, "向左侧倒入水，向右侧倒入待测液体，并在右侧距离支点 l_A = 9 cm 处挂重物 A，杠杆再次平衡：", space_after=4)
    add_paragraph_styled(doc, "G_水 × L_左 = G_液 × L_右 + G_A × l_A", space_after=4, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_styled(doc, "两边消去重力加速度 g 转化为质量关系：", space_after=4)
    add_paragraph_styled(doc, "m_水 × 10 = m_液 × 10 + m_A × l_A", space_after=4, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_styled(doc, "已知 m_水 = ρ_水 × V = 1.0 g/cm³ × 30 cm³ = 30 g，m_A = 10 g，l_A = 9 cm。代入数据：", space_after=4)
    add_paragraph_styled(doc, "30 × 10 = m_液 × 10 + 10 × 9\n300 = 10 × m_液 + 90  =>  m_液 = 21 g", space_after=4, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_styled(doc, "因此液体的密度为：", space_after=4)
    add_paragraph_styled(doc, "ρ_液 = m_液 / V = 21 g / 30 cm³ = 0.7 g/cm³", space_after=6, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    add_bullet_styled(doc, "(3) 函数关系图线选择：", "将质量代入力矩方程 ρ_水·V · 10 = ρ_液·V · 10 + m_A · l，整理得到 ρ_液 关于位置 l 的函数关系：")
    add_paragraph_styled(doc, "ρ_液 = ρ_水 - [ m_A / (10·V) ] × l", space_after=4, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_styled(doc, "由于斜率 k = -m_A / (10·V) < 0，这是一条斜率为负的直线。图丙中，a 斜率为正，c 为曲线，只有 b 符合直线且斜率为负的要求。因此选择 b。", space_after=6)
    
    add_bullet_styled(doc, "(4) 提高测量精确度（分度值与灵敏度分析）：", "“提高测量精确度”在杠杆秤中意味着使刻度盘上的刻度分布更稀疏、分度值更小（即当液体密度发生微小的改变时，重物 A 需要移动更大的距离，这样在尺子上读数时受估读误差的影响更小，读数更精确）。")
    add_paragraph_styled(doc, "根据力矩平衡得出的函数关系式：ρ_液 = ρ_水 - [ m_A / (10·V) ] × l_A", space_after=4, bold=True)
    add_paragraph_styled(doc, "若密度改变 Δρ_液，对应的重物 A 位移变化量 Δl_A 为：Δl_A = [ 10V / m_A ] × Δρ_液。定义秤的灵敏度 S = Δl_A / Δρ_液 = 10V / m_A。由于液体的体积 V = 30 mL 是恒定不变的，因此要增大灵敏度 S 以提高测量精确度，必须减小重物 A 的质量 m_A。")
    add_paragraph_styled(doc, "例如，已知 V = 30 cm³，若要标出密度变化 Δρ = 0.1 g/cm³ 的刻度线：\n"
                            "1. 若 m_A 较大（如 15 g），重物移动距离仅为 Δl_A = (10 × 30 / 15) × 0.1 = 2 cm，刻度密集，估读非常困难；\n"
                            "2. 若 m_A 较小（如 5 g），重物移动距离增大为 Δl_A = (10 × 30 / 5) × 0.1 = 6 cm，刻度十分稀疏，很容易读取出更小的密度值。\n"
                            "物理直观上：重物 A 越轻，就需要移动越长的力臂才能平衡两侧微小的重量差，因而刻度更稀疏、精确度更高。因此，应换用质量更小的重物 A。")
    
    # ------------------ Q18 ------------------
    add_heading_styled(doc, "第 18 题：多开关控制电路（定值电阻与灯泡并/串联）", level=1)
    add_paragraph_styled(doc, "小灯泡 L 标有“6V 3W”字样，其额定参数为：", space_after=4)
    add_bullet_styled(doc, "额定电流：", "I_L额 = P_L额 / U_L额 = 3W / 6V = 0.5 A")
    add_bullet_styled(doc, "灯丝电阻：", "R_L = U_L额² / P_L额 = 36 / 3 = 12 Ω")
    
    add_heading_styled(doc, "电路状态分析：", level=2)
    add_bullet_styled(doc, "状态 1（S1, S2, S3 均闭合）：", "S3 闭合导致灯泡 L 被短路，不发光。定值电阻 R1 与 R2 并联在电源 U 两端。")
    add_paragraph_styled(doc, "电流表 A2 测 R1 支路电流，I1 = 1 A。干路总电流 I_总 = 1.5 A。\n所以流过 R2 的电流为：I2 = I_总 - I1 = 1.5 A - 1 A = 0.5 A。\n电源电压关系式为：", space_after=4)
    add_paragraph_styled(doc, "U = I1 × R1 = 1.0 × R1  =>  U = R1  --- (式1)\nU = I2 × R2 = 0.5 × R2  =>  R2 = 2U  --- (式2)", space_after=6, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    add_bullet_styled(doc, "状态 2（仅闭合 S1）：", "S2 和 S3 断开，R2 支路开路，灯泡 L 与 R1 串联接入电路。")
    add_paragraph_styled(doc, "灯泡 L 正常发光，此时电路中电流为 I = I_L额 = 0.5 A，灯两端电压为 U_L = 6 V。电源电压为：", space_after=4)
    add_paragraph_styled(doc, "U = U_L + I × R1 = 6 + 0.5 × R1  --- (式3)", space_after=6, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    add_heading_styled(doc, "参数求解与热量比：", level=2)
    add_paragraph_styled(doc, "将（式1）U = R1 代入（式3）可得：", space_after=4)
    add_paragraph_styled(doc, "R1 = 6 + 0.5 × R1  =>  0.5 × R1 = 6  =>  R1 = 12 Ω\n电源电压：U = 12 V\nR2 阻值：R2 = 2U = 24 Ω", space_after=6, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_styled(doc, "前后两次电路中相同时间内 R1 产生的热量之比为：", space_after=4)
    add_paragraph_styled(doc, "Q1 : Q2 = I1² × R1 × t : I² × R1 × t = 1.0² : 0.5² = 1 : 0.25 = 4 : 1", space_after=6, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    # ------------------ Q19 ------------------
    add_heading_styled(doc, "第 19 题：电饭锅双档功耗（串联电路阻值调整）", level=1)
    add_paragraph_styled(doc, "高温档功率 P_高 = 1100 W，保温档功率 P_保 = 22 W，电源电压 U = 220 V。电路中，接 1 时为高温档，接 2 时 R1 与 R2 串联为保温档。", space_after=4)
    
    add_heading_styled(doc, "详细解答：", level=2)
    add_bullet_styled(doc, "(1) 保温时电流：", "I_保 = P_保 / U = 22 W / 220 V = 0.1 A")
    add_bullet_styled(doc, "(2) R2 阻值计算：", "高温档仅 R1 工作：R1 = U² / P_高 = 220² / 1100 = 44 Ω。\n保温档总电阻：R_总 = U² / P_保 = 220² / 22 = 2200 Ω。\n定值电阻 R2 = R_总 - R1 = 2200 - 44 = 2156 Ω。")
    add_bullet_styled(doc, "(3) 电能表闪烁次数：", "高温档工作 10 min 消耗电能：\nW = P_高 × t = 1.1 kW × (10/60) h = 11/60 kWh。\n指示灯闪烁次数：N = (11/60 kWh) × 1200 imp/kWh = 220 次。")
    add_bullet_styled(doc, "(4) 功率均提升 10% 后阻值之比：", "为使两档功率均提升 10%，设新的功率为 P_高' = 1.1 P_高，P_保' = 1.1 P_保。\n新阻值为 R1' = R1 / 1.1，总阻值 R1' + R2' = (R1 + R2) / 1.1，从而可得 R2' = R2 / 1.1。\n调整后 R1' 和 R2' 的阻值之比保持 不变。")
    
    # ------------------ Q15 ------------------
    add_heading_styled(doc, "第 15 题：恒流源驱动并联电路安全范围", level=1)
    add_paragraph_styled(doc, "恒流电源输出恒定电流 I_总 = 1 A。定值电阻 R1 = 5 Ω，滑动变阻器 R2 最大阻值 25 Ω（允许最大电流 1 A）。电压表量程 0~3 V，电流表量程 0~0.6 A。", space_after=4)
    add_paragraph_styled(doc, "通过分析 U = 2.6 V 时，R1 电流为 0.52 A，流过 R2 支路电流为 0.48 A，验证电流表串联在 R2 支路上。", space_after=6)
    
    add_heading_styled(doc, "安全范围推导：", level=2)
    add_paragraph_styled(doc, "恒流电路的总功率为 P_总 = U × I_总 = U × 1 A = U (W)。求总功率范围即求并联两端电压 U 的允许变化范围。", space_after=4)
    add_bullet_styled(doc, "电压上限约束：", "由电压表量程 0~3 V 限制，U_max = 3 V。此时最大总功率 P_max = 3 W。")
    add_bullet_styled(doc, "电压下限约束：", "由电流表量程限制，流过 R2 的电流 I2 = 1 - U/R1 = 1 - U/5 ≤ 0.6 A。\n从而解得 U ≥ 2 V，即最小电压 U_min = 2 V。此时最小总功率 P_min = 2 W。")
    add_paragraph_styled(doc, "因此，保证电路安全的前提下，电路总功率的范围是 2 - 3 W。", space_after=6, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    # ------------------ Q16 ------------------
    add_heading_styled(doc, "第 16 题：电路状态切换与灯泡额定参数推导", level=1)
    add_heading_styled(doc, "第一部分：总功率之比 P1 : P2", level=2)
    add_paragraph_styled(doc, "已知 R1 = 2R2。 S 断开时，R1 与 R2 串联，R_串 = 3R2，P1 = U² / (3R2)。\n当开关与电源对调且 S 闭合时，R1 与 R2 并联，R_并 = (2/3)R2，P2 = U² / [(2/3)R2] = 1.5 U² / R2。\n因此 P1 : P2 = (1/3) : (3/2) = 2 : 9。", space_after=6)
    
    add_heading_styled(doc, "第二部分：灯泡额定电压求取方案对比", level=2)
    add_paragraph_styled(doc, "已知：R1 = 8 Ω，电源电压 U = 12 V，灯泡额定功率 P_L = 4 W。闭合开关后灯泡正常发光且电阻比 R1 大（即 R_L > 8 Ω）。求小灯泡的额定电压 U_L。", space_after=6)
    
    # Callout for original method
    add_callout(doc, "【原便签记录方案】（以电流 I 为元变量）", [
        "设电流为 x。由电压分配关系：8x + 4/x = 12",
        "解一元二次方程得：x = 1 A 或 x = 0.5 A。",
        "若 x = 1 A，电阻 R_L = P/I² = 4 Ω < 8 Ω（舍去）。",
        "若 x = 0.5 A，电阻 R_L = P/I² = 16 Ω > 8 Ω（符合要求）。",
        "计算额定电压：U_L = P_L / I = 4 W / 0.5 A = 8 V。"
    ])
    doc.add_paragraph().paragraph_format.space_after = Pt(6) # Spacing
    
    # Alternative 1
    add_heading_styled(doc, "替选方案 1：直接以额定电压 U_L 为变量（最直观、推荐使用）", level=3)
    add_paragraph_styled(doc, "直接设所求物理量小灯泡额定电压为 U_L，免去中间换算：", space_after=4)
    add_bullet_styled(doc, "1. 表达电流：", "由于 L 与 R1 串联，R1 两端电压为 U1 = 12 - U_L，电路电流为 I = (12 - U_L) / 8。")
    add_bullet_styled(doc, "2. 建立功率方程：", "P_L = U_L × I = U_L × (12 - U_L) / 8 = 4 W  =>  12 U_L - U_L² = 32")
    add_bullet_styled(doc, "3. 整理得一元二次方程：", "U_L² - 12 U_L + 32 = 0  =>  (U_L - 8)(U_L - 4) = 0")
    add_bullet_styled(doc, "4. 求解与筛选：", "解得 U_L = 8 V 或 U_L = 4 V。\n当 U_L = 4 V 时，灯泡电阻 R_L = U_L²/P_L = 4 Ω < 8 Ω（舍去）。\n当 U_L = 8 V 时，灯泡电阻 R_L = 16 Ω > 8 Ω（符合题意）。\n因此小灯泡额定电压为 8 V。")
    
    # Alternative 2
    add_heading_styled(doc, "替选方案 2：直接以灯泡电阻 R_L 为变量", level=3)
    add_paragraph_styled(doc, "由于限制条件是关于电阻大小的（R_L > R1 = 8 Ω），以 R_L 为变量可以最方便地应用限制条件：", space_after=4)
    add_bullet_styled(doc, "1. 表达总电流：", "I = U / (R1 + R_L) = 12 / (8 + R_L)")
    add_bullet_styled(doc, "2. 建立功率方程：", "P_L = I² × R_L = 144 R_L / (8 + R_L)² = 4 W  =>  36 R_L = (8 + R_L)²")
    add_bullet_styled(doc, "3. 展开并整理：", "R_L² - 20 R_L + 64 = 0  =>  (R_L - 16)(R_L - 4) = 0")
    add_bullet_styled(doc, "4. 筛选与求解：", "解得 R_L = 16 Ω 或 R_L = 4 Ω。\n因为要求 R_L > 8 Ω，故 R_L = 16 Ω。\n计算额定电压：U_L = sqrt(P_L × R_L) = sqrt(4 × 16) = 8 V。")
    
    # Alternative 3
    add_heading_styled(doc, "替选方案 3：定性与特值估算结合法（考场极速秒杀法）", level=3)
    add_paragraph_styled(doc, "在填空题中，可利用定性分压规律快速心算得出答案：", space_after=4)
    add_bullet_styled(doc, "1. 范围定性：", "因为灯泡电阻 R_L 超过定值电阻 R1 = 8 Ω，在串联电路中，电阻大分压多，所以灯泡两端电压必然占总电压 12 V 的一半以上，即 U_L > 6 V。")
    add_bullet_styled(doc, "2. 尝试特值：", "在大于 6 V 的常见电压中，代入 U_L = 8 V 进行尝试：\nR1 两端电压为 12 - 8 = 4 V；\n电路电流 I = 4 V / 8 Ω = 0.5 A；\n此时灯泡实际功率 P_L = U_L × I = 8 V × 0.5 A = 4 W，与题设 4 W 完全相符，故直接锁定答案 8 V。")

    # Save docx
    filename = "物理综合题解析报告.docx"
    doc.save(filename)
    print(f"Document saved successfully as {filename}")

if __name__ == "__main__":
    main()
