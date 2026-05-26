# -*- coding: utf-8 -*-
import os
import re
from collections import defaultdict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


BASE = r"D:\Users\xwechat_files\wxid_4668346683612_4126\msg\file\2026-04"
SRC_NAME = "参考文献.docx"
OUT_NAME = "参考文献_按年份梳理.docx"


def setup_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    return doc


def add_text(doc, text, size=11, bold=False, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(size)
    r.bold = bold
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    return p


def read_entries():
    path = os.path.join(BASE, SRC_NAME)
    doc = Document(path)
    raw = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if t == "参考文献":
            continue
        if t.startswith("[1] Sergio Manzettiab"):
            continue
        if "学位论文的撰写应本着严谨求实的科学态度" in t:
            break
        if any(t.startswith(prefix) for prefix in ["引用文献的作者不超过3位", "连续出版物:", "专 译 著:", "论 文 集:", "学位论文:", "专    利:", "技术标准:"]):
            continue
        if t.startswith("[2]***"):
            continue
        raw.append(t)
    return raw


def normalize_key(entry):
    e = entry
    e = re.sub(r"\s+", "", e)
    e = e.replace("．", ".").replace("：", ":").replace("，", ",").replace("；", ";")
    e = e.replace("“", "").replace("”", "").replace('"', "")
    # normalize common type marker glitches
    e = e.replace("[[M]", "[M]").replace("[s]", "[S]")
    m = re.match(r"(.+?)\.(.+?)(\[[A-Za-z]\])", e)
    if m:
        author = m.group(1)
        title = m.group(2)
        typ = m.group(3)
        return f"{author}.{title}{typ}".lower()
    return e.lower()


def clean_entry(entry):
    e = entry.strip()
    e = e.replace("[[M]", "[M]").replace("[s]", "[S]")
    e = e.replace("．", ".").replace("，", ",").replace("：", ":")
    e = re.sub(r"\s+", " ", e)
    e = e.replace(" .", ".").strip()
    return e


def type_of(entry):
    m = re.search(r"\[([A-Za-z])\]", entry)
    return m.group(1).upper() if m else ""


def year_of(entry):
    years = re.findall(r"(19|20)\d{2}", entry)
    if years:
        # finditer to preserve full year
        full = re.findall(r"(?:19|20)\d{2}", entry)
        return int(full[-1])
    return 9999


def is_chinese(entry):
    letters = re.findall(r"[A-Za-z]", entry)
    chinese = re.findall(r"[\u4e00-\u9fff]", entry)
    if chinese and len(chinese) >= len(letters):
        return True
    # translated books with Chinese title/publication should still be Chinese
    if re.search(r"[\u4e00-\u9fff]", entry) and re.search(r"北京|上海|济南|郑州|华东师范大学出版社|人民教育出版社|中华书局", entry):
        return True
    return False


def category_group(entry):
    t = type_of(entry)
    zh = is_chinese(entry)
    if zh:
        if t == "M":
            return ("中文", "专著类")
        if t == "J":
            return ("中文", "期刊文章类")
        if t == "D":
            return ("中文", "学术论文类")
        if t == "S":
            return ("中文", "课程标准类")
        return ("中文", "期刊文章类")
    else:
        if t == "M":
            return ("外文", "专著类")
        return ("外文", "期刊论文类")


def dedupe(entries):
    best = {}
    for entry in entries:
        c = clean_entry(entry)
        key = normalize_key(c)
        # keep the longer/more complete one
        if key not in best or len(c) > len(best[key]):
            best[key] = c
    return list(best.values())


def sort_entries(entries):
    return sorted(entries, key=lambda x: (year_of(x), x))


def main():
    entries = read_entries()
    entries = dedupe(entries)

    grouped = {
        ("中文", "专著类"): [],
        ("中文", "期刊文章类"): [],
        ("中文", "学术论文类"): [],
        ("中文", "课程标准类"): [],
        ("外文", "专著类"): [],
        ("外文", "期刊论文类"): [],
    }
    for e in entries:
        grouped[category_group(e)].append(e)
    for k in grouped:
        grouped[k] = sort_entries(grouped[k])

    doc = setup_doc()
    add_text(doc, "参考文献", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_text(doc, "（一）中文文献", size=13, bold=True)
    add_text(doc, "1.专著类", size=12, bold=True)
    for e in grouped[("中文", "专著类")]:
        add_text(doc, e)
    add_text(doc, "2.期刊文章类", size=12, bold=True)
    for e in grouped[("中文", "期刊文章类")]:
        add_text(doc, e)
    add_text(doc, "3.学术论文类", size=12, bold=True)
    for e in grouped[("中文", "学术论文类")]:
        add_text(doc, e)
    add_text(doc, "4.课程标准类", size=12, bold=True)
    for e in grouped[("中文", "课程标准类")]:
        add_text(doc, e)

    add_text(doc, "（二）外文文献", size=13, bold=True)
    add_text(doc, "1.专著类", size=12, bold=True)
    for e in grouped[("外文", "专著类")]:
        add_text(doc, e)
    add_text(doc, "2.期刊论文类", size=12, bold=True)
    for e in grouped[("外文", "期刊论文类")]:
        add_text(doc, e)

    out_path = os.path.join(BASE, OUT_NAME)
    doc.save(out_path)
    print(out_path)
    for k, v in grouped.items():
        print(k, len(v))


if __name__ == "__main__":
    main()
