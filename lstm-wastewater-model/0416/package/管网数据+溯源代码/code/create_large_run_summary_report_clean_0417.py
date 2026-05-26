from __future__ import annotations

import heapq
import json
import math
import sys
from pathlib import Path

import matplotlib.font_manager as fm
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from config_0416 import (  # noqa: E402
    CANDIDATE_NODES,
    MODEL_1D_INP,
    MONITOR_NODES,
    OUTFALL_NODE,
    DATA_SUMMARY_JSON,
    RESULT_DIR,
    TRUTH_INJECTION_NODES,
)


RUN_DIR = RESULT_DIR / "large_run"
ANALYSIS_DIR = RUN_DIR / "analysis_large_0417"
REPORT_DOCX = ANALYSIS_DIR / "0417大参数版本结果总结与布局优化依据.docx"
REPORT_DOCX_FALLBACK = ANALYSIS_DIR / "0417大参数版本结果总结与布局优化依据_修正版.docx"
REPORT_MD = ANALYSIS_DIR / "0417大参数版本结果总结与布局优化依据.md"

FONT_PATH = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_PROP = fm.FontProperties(fname=str(FONT_PATH))
FONT_BOLD = fm.FontProperties(fname=str(Path(r"C:\Windows\Fonts\msyhbd.ttc")))

FIGURES = {
    "ga_convergence": ANALYSIS_DIR / "clean_fig1_GA收敛过程.png",
    "solution_bar": ANALYSIS_DIR / "clean_fig2_关键解份额对比.png",
    "posterior_interval": ANALYSIS_DIR / "clean_fig3_后验权重区间.png",
    "monitor_nse": ANALYSIS_DIR / "clean_fig4_监测点NSE对比.png",
    "truth_map": ANALYSIS_DIR / "clean_fig5_真值注入空间图.png",
    "ga_map": ANALYSIS_DIR / "clean_fig6_GA最优空间图.png",
    "am_map": ANALYSIS_DIR / "clean_fig7_AM后验最优空间图.png",
    "median_map": ANALYSIS_DIR / "clean_fig8_后验中位数空间图.png",
    "mean_map": ANALYSIS_DIR / "clean_fig9_后验均值空间图.png",
}


def setup_font() -> None:
    fm.fontManager.addfont(str(FONT_PATH))
    plt.rcParams["font.family"] = FONT_PROP.get_name()
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.dpi"] = 150
    plt.rcParams["savefig.dpi"] = 180


def read_section(inp_path: Path, section_name: str) -> list[list[str]]:
    rows: list[list[str]] = []
    section = ""
    for raw in inp_path.read_text(encoding="gbk", errors="ignore").splitlines():
        line = raw.split(";", 1)[0].strip()
        if not line:
            continue
        if line.startswith("[") and line.endswith("]"):
            section = line[1:-1].upper()
            continue
        if section == section_name.upper():
            rows.append(line.split())
    return rows


def parse_network() -> tuple[dict[str, tuple[float, float]], list[tuple[str, str, str, float]], set[str]]:
    coords: dict[str, tuple[float, float]] = {}
    for row in read_section(MODEL_1D_INP, "COORDINATES"):
        if len(row) >= 3:
            coords[row[0]] = (float(row[1]), float(row[2]))

    conduits: list[tuple[str, str, str, float]] = []
    for row in read_section(MODEL_1D_INP, "CONDUITS"):
        if len(row) >= 4:
            try:
                length = float(row[3])
            except ValueError:
                length = 1.0
            conduits.append((row[0], row[1], row[2], length))

    outfalls = {row[0] for row in read_section(MODEL_1D_INP, "OUTFALLS") if row}
    return coords, conduits, outfalls


def truth_values() -> dict[str, float]:
    each = 1.0 / len(TRUTH_INJECTION_NODES)
    return {node: each if node in TRUTH_INJECTION_NODES else 0.0 for node in CANDIDATE_NODES}


def row_values(row: pd.Series) -> dict[str, float]:
    return {node: float(row.get(node, 0.0)) for node in CANDIDATE_NODES}


def top_items(values: dict[str, float], limit: int = 8) -> list[tuple[str, float]]:
    return sorted(values.items(), key=lambda kv: kv[1], reverse=True)[:limit]


def nearest_truth_distances(conduits: list[tuple[str, str, str, float]]) -> dict[str, tuple[str, float]]:
    graph: dict[str, list[tuple[str, float]]] = {}
    for _, start, end, length in conduits:
        graph.setdefault(start, []).append((end, length))
        graph.setdefault(end, []).append((start, length))

    def dijkstra(start: str) -> dict[str, float]:
        dist = {start: 0.0}
        heap = [(0.0, start)]
        while heap:
            now_dist, node = heapq.heappop(heap)
            if now_dist != dist[node]:
                continue
            for nxt, weight in graph.get(node, []):
                new_dist = now_dist + weight
                if new_dist < dist.get(nxt, float("inf")):
                    dist[nxt] = new_dist
                    heapq.heappush(heap, (new_dist, nxt))
        return dist

    truth_dists = {truth: dijkstra(truth) for truth in TRUTH_INJECTION_NODES}
    nearest: dict[str, tuple[str, float]] = {}
    for node in CANDIDATE_NODES:
        best_truth = ""
        best_dist = float("inf")
        for truth in TRUTH_INJECTION_NODES:
            value = truth_dists[truth].get(node, float("inf"))
            if value < best_dist:
                best_dist = value
                best_truth = truth
        nearest[node] = (best_truth, best_dist)
    return nearest


def fig_title(ax, title: str, subtitle: str | None = None) -> None:
    ax.set_title(title, fontproperties=FONT_BOLD, fontsize=15, loc="left", pad=13)
    if subtitle:
        ax.text(
            0,
            1.01,
            subtitle,
            transform=ax.transAxes,
            fontproperties=FONT_PROP,
            fontsize=9.5,
            color="#5d6b63",
            va="bottom",
        )


def make_ga_convergence(analysis: dict) -> None:
    gen = pd.DataFrame(analysis["large"]["gen_stats"])
    fig, ax = plt.subplots(figsize=(11.2, 5.2))
    fig.patch.set_facecolor("#f7f3e8")
    ax.set_facecolor("#fffaf0")
    ax.plot(gen["generation"], gen["best_nse"], marker="o", lw=2.4, color="#0f766e", label="每代最优")
    ax.plot(gen["generation"], gen["q90_nse"], marker="s", lw=2.0, color="#f97316", label="每代90分位")
    ax.plot(gen["generation"], gen["median_nse"], marker="^", lw=2.0, color="#64748b", label="每代中位数")
    ax.axhline(0.99, color="#b91c1c", ls="--", lw=1.2, label="NSE=0.99")
    fig_title(ax, "GA 收敛过程", "前几代快速出现高分解，后续主要是在高分代偿区域内微调。")
    ax.set_xlabel("迭代代数", fontproperties=FONT_PROP)
    ax.set_ylabel("mean NSE", fontproperties=FONT_PROP)
    ax.grid(True, color="#e6ddcc", lw=0.8)
    ax.legend(prop=FONT_PROP, frameon=True)
    fig.tight_layout()
    fig.savefig(FIGURES["ga_convergence"], bbox_inches="tight")
    plt.close(fig)


def make_solution_bar(solution_shares: pd.DataFrame) -> None:
    candidates = list(CANDIDATE_NODES)
    data = pd.DataFrame(
        {
            "真值": truth_values(),
            "GA最优": row_values(solution_shares.loc["ga_best"]),
            "AM最优": row_values(solution_shares.loc["posterior_best_map"]),
            "后验中位数": row_values(solution_shares.loc["posterior_median_summary"]),
        }
    ).T[candidates]

    fig, ax = plt.subplots(figsize=(13.4, 5.8))
    fig.patch.set_facecolor("#f7f3e8")
    ax.set_facecolor("#fffaf0")
    colors = ["#111827", "#2563eb", "#ea580c", "#16a34a"]
    x = range(len(candidates))
    width = 0.19
    for idx, (name, color) in enumerate(zip(data.index, colors)):
        offset = (idx - 1.5) * width
        ax.bar([i + offset for i in x], data.loc[name], width=width, label=name, color=color, alpha=0.88)
    fig_title(ax, "关键解份额对比", "单独比较真值、GA最优、AM最优和后验中位数，避免把不同阶段混在一起解读。")
    ax.set_xticks(list(x))
    ax.set_xticklabels(candidates, rotation=45, ha="right", fontproperties=FONT_PROP)
    ax.set_ylabel("注入份额", fontproperties=FONT_PROP)
    ax.grid(axis="y", color="#e6ddcc", lw=0.8)
    ax.legend(prop=FONT_PROP, ncol=4, frameon=True)
    fig.tight_layout()
    fig.savefig(FIGURES["solution_bar"], bbox_inches="tight")
    plt.close(fig)


def make_posterior_interval(posterior: pd.DataFrame) -> None:
    df = posterior.copy()
    df = df[df["node"].isin(CANDIDATE_NODES)].sort_values("posterior_mean", ascending=False).head(12)
    fig, ax = plt.subplots(figsize=(10.8, 6.2))
    fig.patch.set_facecolor("#f7f3e8")
    ax.set_facecolor("#fffaf0")
    y = range(len(df))
    x = df["posterior_mean"]
    left = df["posterior_mean"] - df["p05"]
    right = df["p95"] - df["posterior_mean"]
    ax.errorbar(x, y, xerr=[left, right], fmt="o", color="#0f766e", ecolor="#94a3b8", elinewidth=2.0, capsize=4)
    ax.scatter(df["posterior_median"], y, marker="s", color="#f97316", label="后验中位数", zorder=3)
    ax.set_yticks(list(y))
    ax.set_yticklabels(df["node"], fontproperties=FONT_PROP)
    ax.invert_yaxis()
    fig_title(ax, "后验权重区间", "区间越宽，说明该候选井在后验样本中越不稳定。")
    ax.set_xlabel("注入份额", fontproperties=FONT_PROP)
    ax.grid(axis="x", color="#e6ddcc", lw=0.8)
    ax.legend(prop=FONT_PROP, frameon=True)
    fig.tight_layout()
    fig.savefig(FIGURES["posterior_interval"], bbox_inches="tight")
    plt.close(fig)


def make_monitor_nse(per_monitor: pd.DataFrame) -> None:
    pivot = per_monitor.pivot(index="monitor", columns="solution", values="nse").loc[list(MONITOR_NODES)]
    names = {
        "ga_best": "GA最优",
        "posterior_best_map": "AM最优",
        "posterior_median_summary": "后验中位数",
    }
    fig, ax = plt.subplots(figsize=(12.5, 5.4))
    fig.patch.set_facecolor("#f7f3e8")
    ax.set_facecolor("#fffaf0")
    colors = {"ga_best": "#2563eb", "posterior_best_map": "#ea580c", "posterior_median_summary": "#16a34a"}
    for col in ["ga_best", "posterior_best_map", "posterior_median_summary"]:
        ax.plot(pivot.index, pivot[col], marker="o", lw=2.2, color=colors[col], label=names[col])
    fig_title(ax, "各监测点拟合 NSE 对比", "AM最优在各监测点均接近1，但高拟合不等于点位唯一。")
    ax.set_xlabel("监测点", fontproperties=FONT_PROP)
    ax.set_ylabel("NSE", fontproperties=FONT_PROP)
    ax.tick_params(axis="x", rotation=35)
    ax.grid(True, color="#e6ddcc", lw=0.8)
    ax.legend(prop=FONT_PROP, frameon=True)
    fig.tight_layout()
    fig.savefig(FIGURES["monitor_nse"], bbox_inches="tight")
    plt.close(fig)


def draw_network_solution(
    filename: Path,
    values: dict[str, float],
    title: str,
    subtitle: str,
    score_text: str,
    nearest: dict[str, tuple[str, float]],
) -> None:
    coords, conduits, outfalls = parse_network()
    fig = plt.figure(figsize=(14.2, 7.4))
    fig.patch.set_facecolor("#f7f3e8")
    gs = fig.add_gridspec(1, 3, width_ratios=[5.0, 0.16, 1.35], wspace=0.08)
    ax = fig.add_subplot(gs[0, 0])
    cax = fig.add_subplot(gs[0, 1])
    side = fig.add_subplot(gs[0, 2])
    ax.set_facecolor("#fffaf0")
    side.set_facecolor("#fffaf0")

    for _, start, end, _ in conduits:
        if start in coords and end in coords:
            ax.plot(
                [coords[start][0], coords[end][0]],
                [coords[start][1], coords[end][1]],
                color="#8c9790",
                lw=1.45,
                alpha=0.45,
                zorder=1,
            )

    special = set(CANDIDATE_NODES) | set(MONITOR_NODES) | set(outfalls)
    normal_nodes = [node for node in coords if node not in special]
    ax.scatter(
        [coords[n][0] for n in normal_nodes],
        [coords[n][1] for n in normal_nodes],
        s=12,
        color="#c8d0ca",
        alpha=0.78,
        zorder=2,
    )

    candidate_values = [values.get(node, 0.0) for node in CANDIDATE_NODES]
    sizes = [65 + 1800 * max(v, 0.0) for v in candidate_values]
    sc = ax.scatter(
        [coords[n][0] for n in CANDIDATE_NODES],
        [coords[n][1] for n in CANDIDATE_NODES],
        c=candidate_values,
        s=sizes,
        cmap="YlOrRd",
        vmin=0.0,
        vmax=1.0 / 3.0,
        edgecolors="#26352f",
        linewidths=0.9,
        alpha=0.9,
        zorder=4,
    )

    monitor_nodes = [node for node in MONITOR_NODES if node in coords]
    ax.scatter(
        [coords[n][0] for n in monitor_nodes],
        [coords[n][1] for n in monitor_nodes],
        marker="s",
        s=64,
        color="#1d4ed8",
        edgecolors="white",
        linewidths=0.9,
        zorder=5,
    )

    truth_nodes = [node for node in TRUTH_INJECTION_NODES if node in coords]
    ax.scatter(
        [coords[n][0] for n in truth_nodes],
        [coords[n][1] for n in truth_nodes],
        s=520,
        facecolors="none",
        edgecolors="#b91c1c",
        linewidths=2.0,
        zorder=6,
    )

    outfall_nodes = [node for node in outfalls if node in coords]
    ax.scatter(
        [coords[n][0] for n in outfall_nodes],
        [coords[n][1] for n in outfall_nodes],
        marker="D",
        s=95,
        color="#111827",
        zorder=6,
    )

    label_nodes = {node for node, value in top_items(values, 8) if value > 0.005}
    label_nodes |= set(TRUTH_INJECTION_NODES)
    label_nodes |= {node for node in ["J1", "J2", "J21", "J84", "J72", "J86"] if values.get(node, 0) > 0.015}
    for node in sorted(label_nodes):
        if node not in coords:
            continue
        x, y = coords[node]
        ax.text(
            x + 10,
            y + 10,
            f"{node}\n{values.get(node, 0):.3f}",
            fontproperties=FONT_BOLD,
            fontsize=8.2,
            color="#1f2a24",
            zorder=7,
            bbox=dict(boxstyle="round,pad=0.18", fc="#fffaf0", ec="#d6cdb9", alpha=0.85),
        )

    fig_title(ax, title, subtitle)
    ax.set_xlabel("模型 X 坐标", fontproperties=FONT_PROP)
    ax.set_ylabel("模型 Y 坐标", fontproperties=FONT_PROP)
    ax.grid(True, color="#e8dfcf", lw=0.8)
    ax.set_aspect("equal", adjustable="box")

    cbar = fig.colorbar(sc, cax=cax)
    cbar.set_label("")
    cax.set_title("份额", fontproperties=FONT_PROP, fontsize=9, pad=8)

    side.axis("off")
    side.text(0.02, 0.98, "主要份额节点", fontproperties=FONT_BOLD, fontsize=14, va="top", color="#20302b", transform=side.transAxes)
    side.text(0.02, 0.91, score_text, fontproperties=FONT_PROP, fontsize=9.6, va="top", color="#5d6b63", wrap=True, transform=side.transAxes)
    y0 = 0.79
    for idx, (node, value) in enumerate(top_items(values, 8)):
        if value <= 0:
            continue
        y = y0 - idx * 0.087
        truth, dist = nearest.get(node, ("-", float("inf")))
        label = "真值点" if node in TRUTH_INJECTION_NODES else f"近邻 {truth}，距约 {dist:.0f}"
        color = "#b91c1c" if node in TRUTH_INJECTION_NODES else "#b45309"
        side.text(0.02, y, f"{node}  {value:.4f}", fontproperties=FONT_BOLD, fontsize=10.5, color=color, va="top", transform=side.transAxes)
        side.text(0.02, y - 0.035, label, fontproperties=FONT_PROP, fontsize=8.8, color="#66756f", va="top", transform=side.transAxes)
        side.plot([0.02, min(0.95, 0.02 + value / 0.34 * 0.88)], [y - 0.061, y - 0.061], color=color, lw=4.5, solid_capstyle="round", transform=side.transAxes)

    fig.savefig(filename, bbox_inches="tight")
    plt.close(fig)


def make_network_maps(solution_shares: pd.DataFrame, posterior: pd.DataFrame) -> None:
    _, conduits, _ = parse_network()
    nearest = nearest_truth_distances(conduits)
    posterior_mean = {row["node"]: float(row["posterior_mean"]) for _, row in posterior.iterrows() if row["node"] in CANDIDATE_NODES}

    draw_network_solution(
        FIGURES["truth_map"],
        truth_values(),
        "真值注入空间分布",
        "J20、J48、J11 为三处真实注入点，各占三分之一。",
        "基准答案；用于对照识别结果。",
        nearest,
    )
    draw_network_solution(
        FIGURES["ga_map"],
        row_values(solution_shares.loc["ga_best"]),
        "GA 最优解空间分布",
        "GA 已取得高拟合分数，但 J48 被 J84、J20 附近被 J21 等节点代偿。",
        "mean NSE=0.997718；SSE=0.00802651。",
        nearest,
    )
    draw_network_solution(
        FIGURES["am_map"],
        row_values(solution_shares.loc["posterior_best_map"]),
        "AM 后验最优 MAP 空间分布",
        "当前最优结果：J11 与 J48 基本恢复，J20 仍有少量邻近代偿。",
        "mean NSE=0.999866；SSE=0.000248442。",
        nearest,
    )
    draw_network_solution(
        FIGURES["median_map"],
        row_values(solution_shares.loc["posterior_median_summary"]),
        "后验中位数空间分布",
        "中位数用于表达不确定性；多峰情况下会把多个模式混合，不能作为唯一最终解。",
        "mean NSE=0.987021；SSE=0.029131。",
        nearest,
    )
    draw_network_solution(
        FIGURES["mean_map"],
        posterior_mean,
        "后验均值空间分布",
        "均值反映后验总体质量中心，可观察哪些区域长期承担不确定性。",
        "后验统计图；不是单次最优方案。",
        nearest,
    )


def make_all_figures(summary: dict, analysis: dict, solution_shares: pd.DataFrame, posterior: pd.DataFrame, per_monitor: pd.DataFrame) -> None:
    setup_font()
    make_ga_convergence(analysis)
    make_solution_bar(solution_shares)
    make_posterior_interval(posterior)
    make_monitor_nse(per_monitor)
    make_network_maps(solution_shares, posterior)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        set_cell_shading(table.rows[0].cells[i], "DDEFEA")
        if widths:
            table.rows[0].cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Cm(widths[i])
    document.add_paragraph()


def add_picture(document: Document, path: Path, caption: str, width_cm: float = 15.8) -> None:
    document.add_picture(str(path), width=Cm(width_cm))
    p = document.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(90, 90, 90)


def setup_doc(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    document.styles["Normal"].font.size = Pt(10.5)


def add_para(document: Document, text: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)


def percent(value: float) -> str:
    return f"{value * 100:.1f}%"


def top_text(values: dict[str, float], limit: int = 6) -> str:
    return "；".join(f"{node}={value:.4f}" for node, value in top_items(values, limit) if value > 1e-8)


def build_docx(summary: dict, data_summary: dict, analysis: dict, solution_shares: pd.DataFrame, posterior: pd.DataFrame) -> None:
    large = analysis["large"]
    config = summary["config"]
    doc = Document()
    setup_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("0417 大参数版本结果总结与布局优化依据")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("优化版：分阶段单独出图，修正中文乱码，并调整局部代偿解释口径")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(90, 90, 90)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    doc.add_heading("一、结论摘要", level=1)
    add_para(
        doc,
        "本轮大参数版本在曲线拟合层面已经达到较高水平。AM 后验最优 MAP 的 mean NSE 为 "
        f"{summary['solution_scores']['posterior_best_map']['mean_nse']:.6f}，SSE 为 "
        f"{summary['solution_scores']['posterior_best_map']['sse']:.6g}；GA 全局最优 mean NSE 为 "
        f"{summary['solution_scores']['ga_best']['mean_nse']:.6f}，SSE 为 "
        f"{summary['solution_scores']['ga_best']['sse']:.6g}。"
    )
    add_para(
        doc,
        "从定位上看，J11 最稳定，J48 在 AM MAP 中恢复较好，J20 仍被 J2/J1 等邻近节点分走少量份额。"
        "需要调整的判断是：靠得很近、属于同一局部水力区域的代偿，在工程溯源中可以视为区域定位有效；"
        "但如果下一阶段目标是精确到单井，就必须补充能区分相邻井的局部监测点。"
    )
    add_para(
        doc,
        "因此下一轮优化重点不是继续单纯加大参数规模，而是围绕 J20 邻近组和 J48 邻近组优化监测布局，"
        "把原来只看远端拟合的约束，改成同时强调局部区分性。"
    )

    doc.add_heading("二、数据与参数配置", level=1)
    add_table(
        doc,
        ["项目", "本轮设置"],
        [
            ["基线模型", str(data_summary["baseline_inp"])],
            ["事件模型", str(data_summary["truth_event_inp"])],
            ["时间尺度", f"{data_summary['rows']} 个 5 分钟样本，覆盖约 48 小时"],
            ["总入流量", f"事件排口总量 - 旱天排口总量 = {data_summary['outfall_delta_total_volume_m3']:.3f} m3"],
            ["真值注入", f"J20、J48、J11 三点，各 {data_summary['truth_scaled_volumes_m3']['J20']:.3f} m3"],
            ["候选井 / 监测点", f"{len(CANDIDATE_NODES)} 个候选井；{len(MONITOR_NODES)} 个监测点"],
            ["真值回放", f"mean NSE={large['truth']['mean_nse']:.6f}，SSE={large['truth']['sse']:.3g}"],
        ],
        widths=[4.0, 12.5],
    )
    add_table(
        doc,
        ["阶段", "参数规模", "说明"],
        [
            ["GA", f"{config['ga_population_count']} 个种群 x {config['ga_population_size']} 个体 x {config['ga_generations']} 代", f"共 {large['ga_rows']} 个样本，按 mean NSE 排序。"],
            ["initial PPD", f"{large['initial_rows']} 个样本", "作为 AM 起点池和协方差估计基础。"],
            ["AM", f"{config['am_chain_count']} 条链 x {config['am_samples_per_chain']} 步，warmup {config['am_warmup']} 步", f"prior 不进入接受率，平均接受率 {large['am_accept_mean']:.3f}。"],
            ["后验验证", f"{summary['posterior_validation_sample_count']} 个样本", f"90% 覆盖率均值 {summary['posterior_coverage_mean']:.3f}。"],
        ],
        widths=[3.2, 5.4, 7.8],
    )
    add_table(
        doc,
        ["阶段", "主指标", "代码执行含义", "报告解释口径"],
        [
            [
                "GA 迭代",
                "mean NSE 最大化",
                "种群排序、精英保留、交叉变异、竞争迁移均按 mean NSE 进行；SSE 只记录，部分同分排序时越小越好。",
                "GA 结果主要看 mean NSE，同时列出 SSE 作为误差记录；不能说 GA 是按 SSE 搜索。",
            ],
            [
                "initial PPD",
                "GA mean NSE 排序池",
                "从 GA 最后一代按 mean NSE 为主、SSE 为辅筛选，再用 rank 权重抽样并去重。",
                "这是 AM 起点池，不是最终后验结论。",
            ],
            [
                "AM 接受率",
                "SSE likelihood",
                "log_like = -0.5 x SSE / sigma^2；接受率由 proposal_log_like - current_log_like 决定。",
                "AM 过程主要看 SSE/log_like，mean NSE 只作为重新评价曲线拟合效果的记录值。",
            ],
            [
                "AM MAP",
                "log_like 最大",
                "posterior_best_map 从 AM 样本中按 log_like 最大选出，等价于 SSE 最小。",
                "AM 最优不是按 NSE 选出的，报告应写 AM MAP = SSE 最小样本。",
            ],
            [
                "后验中位数",
                "统计汇总",
                "逐节点取后验中位数再归一化，不参与 GA 或 AM 搜索目标。",
                "只能解释不确定性，不能作为单一最优解。",
            ],
        ],
        widths=[2.5, 3.0, 6.5, 5.0],
    )

    doc.add_heading("三、分阶段结果", level=1)
    rows = []
    for item in large["solutions"]:
        values = row_values(solution_shares.loc[item["solution"]])
        rows.append([
            item["solution"],
            f"{item['mean_nse']:.6f}",
            f"{item['sse']:.6g}",
            percent(item["truth_mass"]),
            percent(item["outside_mass"]),
            top_text(values),
        ])
    add_table(
        doc,
        ["方案", "mean NSE", "SSE", "真值质量", "非真值质量", "前六份额节点"],
        rows,
        widths=[3.6, 2.1, 2.4, 2.2, 2.2, 5.8],
    )
    add_para(
        doc,
        "GA 阶段能快速找到高拟合区域，但 GA 最优仍包含 J84、J21、J91、J86 等代偿份额。"
        "GA 的搜索主指标就是 mean NSE，因此这里的高分指 GA 指标高；这不是 AM 的接受指标。"
    )
    add_para(
        doc,
        "AM 阶段进一步降低 SSE，并把真值质量提高到 90.3%。其中 J11=0.3302、J48=0.3293、J20=0.2435，"
        "主要非真值份额集中在 J2=0.0559、J1=0.0373。AM MAP 是按 log_like 最大选出的，也就是 SSE 最小样本，"
        "不是按 NSE 选出的。"
    )
    add_para(
        doc,
        "后验中位数不应作为唯一最终方案。它的作用是描述不确定性；当后验存在多个相邻模式时，逐节点中位数会把模式混合，"
        "所以其拟合分数低于 AM MAP。"
    )
    add_picture(doc, FIGURES["ga_convergence"], "图 1  GA 收敛过程。")
    add_picture(doc, FIGURES["solution_bar"], "图 2  真值、GA最优、AM最优、后验中位数的份额对比。")
    add_picture(doc, FIGURES["posterior_interval"], "图 3  后验权重区间，反映各候选井不确定性。")
    add_picture(doc, FIGURES["monitor_nse"], "图 4  各监测点 NSE 对比。")

    doc.add_heading("四、单独空间图分析", level=1)
    add_para(
        doc,
        "下面将不同阶段结果分别画在管网结构上。红色空心圈是真值注入点，蓝色方块是监测点，圆点大小和颜色表示该方案下的注入份额。"
        "这种图比混合热点图更适合判断代偿是否发生在局部邻近区域。"
    )
    add_picture(doc, FIGURES["truth_map"], "图 5  真值注入空间图。")
    add_picture(doc, FIGURES["ga_map"], "图 6  GA 最优空间图：高分但仍存在较明显代偿。")
    add_picture(doc, FIGURES["am_map"], "图 7  AM 后验最优 MAP 空间图：当前最优，代偿主要收缩到 J20 邻近区域。")
    add_picture(doc, FIGURES["median_map"], "图 8  后验中位数空间图：用于看不确定性，不作为唯一最终解。")
    add_picture(doc, FIGURES["mean_map"], "图 9  后验均值空间图：用于看后验质量中心。")

    doc.add_heading("五、代偿解释口径", level=1)
    add_table(
        doc,
        ["GA 高分阈值", "样本数", "真值质量中位数", "非真值质量中位数", "解释"],
        [
            ["NSE >= 0.95", str(large["high_ga"]["0.95"]["count"]), f"{large['high_ga']['0.95']['truth_mass_median']:.3f}", f"{large['high_ga']['0.95']['outside_median']:.3f}", "高分样本数量多，但位置仍分散。"],
            ["NSE >= 0.98", str(large["high_ga"]["0.98"]["count"]), f"{large['high_ga']['0.98']['truth_mass_median']:.3f}", f"{large['high_ga']['0.98']['outside_median']:.3f}", "拟合更高，但仍有明显非真值质量。"],
            ["NSE >= 0.99", str(large["high_ga"]["0.99"]["count"]), f"{large['high_ga']['0.99']['truth_mass_median']:.3f}", f"{large['high_ga']['0.99']['outside_median']:.3f}", "说明高分不等于唯一定位。"],
            ["NSE >= 0.995", str(large["high_ga"]["0.995"]["count"]), f"{large['high_ga']['0.995']['truth_mass_median']:.3f}", f"{large['high_ga']['0.995']['outside_median']:.3f}", "极高分样本也可能是局部代偿。"],
        ],
        widths=[2.8, 1.7, 2.8, 2.8, 5.8],
    )
    add_para(
        doc,
        "本轮对代偿的解释应分级：第一类是局部邻近代偿，例如 J20 与 J2/J1/J21 这类同一片区或相邻支路上的替代，"
        "在工程排查中可以接受为区域定位结果；第二类是远距离或跨支路代偿，这类会误导现场排查，需要通过监测布局或候选集约束压制。"
    )
    add_para(
        doc,
        "当前 AM MAP 的主要非真值份额为 J2 和 J1，属于 J20 附近代偿，严重程度低于 GA 阶段的 J84/J21 等大比例代偿。"
        "所以本轮相比 GA 和中参数版本已经有实质改善。"
    )

    doc.add_heading("六、下一轮布局优化方向", level=1)
    add_table(
        doc,
        ["区域", "本轮表现", "下一轮目标"],
        [
            ["J20 - J2 - J1 - J21", "AM MAP 中 J20=0.2435，J2+J1=0.0932。", "加能区分 J20 与 J2/J1/J21 的局部监测点。"],
            ["J48 - J84 - J72 - J86", "GA 和后验中仍能看到该组不确定性，但 AM MAP 已恢复 J48=0.3293。", "保留关键下游监测，验证 J48 与 J84/J72/J86 的响应差异。"],
            ["J11 周边", "J11 稳定，AM MAP 为 0.3302。", "无需过度增加监测，保留校核即可。"],
            ["远端主干", "对总体拟合有帮助，但对局部点位区分不足。", "保留少量总量约束点，新增资源优先给局部支路。"],
        ],
        widths=[3.8, 6.0, 6.2],
    )
    add_para(
        doc,
        "下一轮建议只改监测布局，暂不改候选井数量和真值注入方案。这样可以清楚判断：识别改善是否来自监测点位优化，"
        "而不是参数规模、候选集或数据构造同时变化导致无法归因。"
    )

    try:
        doc.save(REPORT_DOCX)
    except PermissionError:
        doc.save(REPORT_DOCX_FALLBACK)


def make_markdown(summary: dict, data_summary: dict, analysis: dict) -> str:
    large = analysis["large"]
    return "\n".join(
        [
            "# 0417 大参数版本结果总结与布局优化依据",
            "",
            "## 结论",
            "",
            f"AM 后验最优 MAP：mean NSE={summary['solution_scores']['posterior_best_map']['mean_nse']:.6f}，SSE={summary['solution_scores']['posterior_best_map']['sse']:.6g}。",
            f"GA 全局最优：mean NSE={summary['solution_scores']['ga_best']['mean_nse']:.6f}，SSE={summary['solution_scores']['ga_best']['sse']:.6g}。",
            "GA 的搜索主指标是 mean NSE；AM 的接受主指标是 SSE likelihood，即 log_like=-0.5*SSE/sigma^2。",
            "AM MAP 是 AM 样本中 log_like 最大的解，等价于 SSE 最小；不是按 NSE 选出的。",
            "J11 稳定，J48 在 AM MAP 中恢复较好，J20 仍有 J2/J1 邻近代偿。",
            "靠得很近的局部代偿在工程溯源中可接受为区域定位；若要精确到井，需要补局部监测。",
            "",
            "## 数据",
            "",
            f"总入流量：{data_summary['outfall_delta_total_volume_m3']:.3f} m3。",
            f"真值注入：J20、J48、J11 各 {data_summary['truth_scaled_volumes_m3']['J20']:.3f} m3。",
            f"候选井：{len(CANDIDATE_NODES)} 个；监测点：{len(MONITOR_NODES)} 个。",
            "",
            "## 代偿",
            "",
            f"GA 中 NSE>=0.99 的样本有 {large['high_ga']['0.99']['count']} 个，真值质量中位数 {large['high_ga']['0.99']['truth_mass_median']:.3f}，非真值质量中位数 {large['high_ga']['0.99']['outside_median']:.3f}。",
            "下一轮优化重点是 J20-J2-J1-J21 与 J48-J84-J72-J86 两组局部区分。",
            "",
        ]
    )


def main() -> None:
    ANALYSIS_DIR.mkdir(parents=True, exist_ok=True)
    summary = json.loads((RUN_DIR / "0417_summary.json").read_text(encoding="utf-8"))
    data_summary = json.loads(DATA_SUMMARY_JSON.read_text(encoding="utf-8"))
    analysis = json.loads((ANALYSIS_DIR / "large_analysis_summary.json").read_text(encoding="utf-8"))
    solution_shares = pd.read_csv(RUN_DIR / "0417_solution_shares.csv").set_index("solution")
    posterior = pd.read_csv(RUN_DIR / "0417_posterior_node_weights.csv")
    per_monitor = pd.read_csv(ANALYSIS_DIR / "per_monitor_nse_large.csv")

    make_all_figures(summary, analysis, solution_shares, posterior, per_monitor)
    build_docx(summary, data_summary, analysis, solution_shares, posterior)
    REPORT_MD.write_text(make_markdown(summary, data_summary, analysis), encoding="utf-8")

    print(REPORT_DOCX)
    for path in FIGURES.values():
        print(path)


if __name__ == "__main__":
    main()
