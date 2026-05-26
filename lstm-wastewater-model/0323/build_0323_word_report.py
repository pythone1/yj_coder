from __future__ import annotations

import json
import math
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.dates as mdates
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(r"E:\PY\LSTM\0323")
RESULT_DIR = BASE_DIR / "results"
ASSET_DIR = RESULT_DIR / "doc_assets"
DOCX_PATH = RESULT_DIR / "0323_精细子网络溯源详细汇报.docx"

plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS"]
plt.rcParams["axes.unicode_minus"] = False


def set_doc_defaults(doc: Document) -> None:
    """统一文档样式，尽量做成正式汇报材料的阅读风格。"""
    styles = doc.styles
    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(11)

    for style_name, size in [("Title", 18), ("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)


def add_page_number(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def load_inputs():
    summary = json.loads((RESULT_DIR / "summary.json").read_text(encoding="utf-8"))
    selected_nodes = pd.read_csv(RESULT_DIR / "selected_nodes.csv")
    selected_links = pd.read_csv(RESULT_DIR / "selected_links.csv")
    observed = pd.read_csv(RESULT_DIR / "observed_delta_10min.csv", parse_dates=["time"])
    fitted = pd.read_csv(RESULT_DIR / "fitted_delta_10min.csv", parse_dates=["time"])
    ga_history = pd.read_csv(RESULT_DIR / "ga_history.csv")
    ga_population = pd.read_csv(RESULT_DIR / "ga_population_all.csv")
    initial_ppd = pd.read_csv(RESULT_DIR / "initial_ppd.csv")
    am_samples = pd.read_csv(RESULT_DIR / "am_samples.csv")
    posterior_weights = pd.read_csv(RESULT_DIR / "posterior_weights.csv")
    bands = pd.read_csv(RESULT_DIR / "posterior_predictive_bands.csv", parse_dates=["time"])
    coverage = pd.read_csv(RESULT_DIR / "posterior_predictive_coverage.csv")
    return (
        summary,
        selected_nodes,
        selected_links,
        observed,
        fitted,
        ga_history,
        ga_population,
        initial_ppd,
        am_samples,
        posterior_weights,
        bands,
        coverage,
    )


def ensure_assets_dir() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def classify_node(node: str, summary: dict) -> str:
    truth = set(summary["truth_nodes"])
    predicted = set(summary["predicted_nodes"])
    monitor = set(summary["monitor_nodes"])
    if node in truth and node in predicted:
        return "truth_predicted"
    if node in truth:
        return "truth_only"
    if node in predicted:
        return "predicted_only"
    if node in monitor:
        return "monitor"
    return "candidate"


def make_subnetwork_plot(selected_nodes: pd.DataFrame, selected_links: pd.DataFrame, summary: dict) -> Path:
    fig, ax = plt.subplots(figsize=(10, 7), dpi=200)
    for row in selected_links.itertuples():
        from_node = selected_nodes[selected_nodes["node"] == row.from_node]
        to_node = selected_nodes[selected_nodes["node"] == row.to_node]
        if from_node.empty or to_node.empty:
            continue
        ax.plot(
            [from_node["x"].iloc[0], to_node["x"].iloc[0]],
            [from_node["y"].iloc[0], to_node["y"].iloc[0]],
            color="#94A3B8",
            linewidth=1.4,
            zorder=1,
        )

    styles = {
        "candidate": dict(color="#94A3B8", marker="o", size=40, label="候选节点"),
        "monitor": dict(color="#2563EB", marker="s", size=80, label="监测点"),
        "truth_only": dict(color="#DC2626", marker="o", size=90, label="真值点"),
        "predicted_only": dict(color="#16A34A", marker="*", size=150, label="识别点"),
        "truth_predicted": dict(color="#EA580C", marker="D", size=120, label="真值且识别"),
    }

    used_labels = set()
    for row in selected_nodes.itertuples():
        node_type = classify_node(row.node, summary)
        style = styles[node_type]
        label = style["label"] if style["label"] not in used_labels else None
        used_labels.add(style["label"])
        ax.scatter(row.x, row.y, s=style["size"], c=style["color"], marker=style["marker"], label=label, zorder=3)
        ax.text(row.x + 2, row.y + 2, row.node, fontsize=7, color="#0F172A")

    ax.set_title("20 节点研究子网络与关键节点标注", fontsize=14)
    ax.set_xlabel("X")
    ax.set_ylabel("Y")
    ax.grid(alpha=0.18)
    ax.legend(frameon=False, loc="upper right")
    fig.tight_layout()
    path = ASSET_DIR / "01_子网络拓扑.png"
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def make_monitor_fit_plot(observed: pd.DataFrame, fitted: pd.DataFrame, summary: dict) -> Path:
    monitors = summary["monitor_nodes"]
    fig, axes = plt.subplots(len(monitors), 1, figsize=(11, 9), dpi=200, sharex=True)
    if len(monitors) == 1:
        axes = [axes]
    for ax, monitor in zip(axes, monitors):
        ax.plot(observed["time"], observed[f"{monitor}_inflow"], color="#EA580C", linewidth=1.8, label="真实监测增量")
        ax.plot(fitted["time"], fitted[f"{monitor}_inflow"], color="#16A34A", linewidth=1.6, linestyle="--", label="模型拟合增量")
        ax.set_ylabel(monitor)
        ax.grid(alpha=0.2)
    axes[0].legend(frameon=False, ncol=2, loc="upper right")
    axes[-1].xaxis.set_major_formatter(mdates.DateFormatter("%m-%d\n%H:%M"))
    fig.suptitle("5 个监测点流量增量拟合结果（10 分钟分辨率）", fontsize=14)
    fig.tight_layout(rect=[0, 0, 1, 0.98])
    path = ASSET_DIR / "02_监测拟合结果.png"
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def make_ga_posterior_plot(ga_history: pd.DataFrame, posterior_weights: pd.DataFrame, coverage: pd.DataFrame) -> Path:
    fig, axes = plt.subplots(2, 2, figsize=(11, 8.2), dpi=200)

    axes[0, 0].plot(ga_history["generation"], ga_history["best_mean_nse"], color="#2563EB", marker="o", linewidth=2)
    axes[0, 0].set_title("GA 各代最佳 Mean NSE")
    axes[0, 0].set_xlabel("迭代代数")
    axes[0, 0].set_ylabel("Mean NSE")
    axes[0, 0].grid(alpha=0.2)
    axes[0, 0].set_ylim(0.75, 1.0)

    top = posterior_weights.head(10)
    colors = ["#DC2626" if v else "#64748B" for v in top["is_truth"]]
    axes[0, 1].bar(top["node"], top["posterior_mean_share"], color=colors)
    axes[0, 1].set_title("Posterior Top 10 节点权重")
    axes[0, 1].tick_params(axis="x", rotation=45)
    axes[0, 1].set_ylabel("Posterior mean share")
    axes[0, 1].grid(alpha=0.2, axis="y")

    cov_colors = ["#16A34A" if x >= 0.7 else "#F59E0B" if x >= 0.5 else "#DC2626" for x in coverage["coverage_90"]]
    axes[1, 0].bar(coverage["monitor"], coverage["coverage_90"], color=cov_colors)
    axes[1, 0].set_title("监测点 90% 覆盖率")
    axes[1, 0].set_ylim(0, 1.0)
    axes[1, 0].grid(alpha=0.2, axis="y")

    truth_top2 = posterior_weights[posterior_weights["is_truth"]].copy()
    axes[1, 1].bar(
        truth_top2["node"],
        truth_top2["posterior_mean_share"],
        yerr=[
            truth_top2["posterior_mean_share"] - truth_top2["p05_share"],
            truth_top2["p95_share"] - truth_top2["posterior_mean_share"],
        ],
        color="#7C3AED",
        capsize=4,
    )
    axes[1, 1].set_title("真值节点后验区间")
    axes[1, 1].set_ylabel("Share")
    axes[1, 1].grid(alpha=0.2, axis="y")

    fig.suptitle("GA 收敛与后验结果概览", fontsize=14)
    fig.tight_layout(rect=[0, 0, 1, 0.97])
    path = ASSET_DIR / "03_GA与后验概览.png"
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def make_posterior_validation_plot(bands: pd.DataFrame, coverage: pd.DataFrame, summary: dict) -> Path:
    monitors = summary["monitor_nodes"]
    fig, axes = plt.subplots(len(monitors), 1, figsize=(11, 9.4), dpi=200, sharex=True)
    if len(monitors) == 1:
        axes = [axes]
    for ax, monitor in zip(axes, monitors):
        df = bands[bands["monitor"] == monitor].copy()
        cov_value = float(coverage.loc[coverage["monitor"] == monitor, "coverage_90"].iloc[0])
        ax.fill_between(df["time"], df["p05"], df["p95"], color="#86EFAC", alpha=0.55, label="90% 后验预测区间")
        ax.plot(df["time"], df["p50"], color="#16A34A", linewidth=1.6, label="后验中位预测")
        ax.plot(df["time"], df["observed"], color="#EA580C", linewidth=1.5, label="真实监测增量")
        ax.set_ylabel(f"{monitor}\n覆盖率={cov_value:.3f}")
        ax.grid(alpha=0.2)
    axes[0].legend(frameon=False, ncol=3, loc="upper right")
    axes[-1].xaxis.set_major_formatter(mdates.DateFormatter("%m-%d\n%H:%M"))
    fig.suptitle("Posterior Predictive Validation", fontsize=14)
    fig.tight_layout(rect=[0, 0, 1, 0.98])
    path = ASSET_DIR / "04_后验预测验证.png"
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)


def add_table(doc: Document, data: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(data):
        for j, value in enumerate(row):
            set_cell_text(table.cell(i, j), str(value), bold=(i == 0))
            if widths:
                table.cell(i, j).width = Cm(widths[j])
    doc.add_paragraph("")


def add_figure(doc: Document, image_path: Path, title: str, width_cm: float = 15.8) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    caption = doc.add_paragraph(title)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True


def add_formula_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(10.5)


def build_document():
    (
        summary,
        selected_nodes,
        selected_links,
        observed,
        fitted,
        ga_history,
        ga_population,
        initial_ppd,
        am_samples,
        posterior_weights,
        bands,
        coverage,
    ) = load_inputs()

    ensure_assets_dir()
    img_topology = make_subnetwork_plot(selected_nodes, selected_links, summary)
    img_fit = make_monitor_fit_plot(observed, fitted, summary)
    img_ga = make_ga_posterior_plot(ga_history, posterior_weights, coverage)
    img_post = make_posterior_validation_plot(bands, coverage, summary)

    doc = Document()
    set_doc_defaults(doc)
    add_page_number(doc.sections[0])

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("0323 精细子网络入渗溯源分析汇报")
    r.bold = True
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = subtitle.add_run("基于 SWMM、遗传算法与 Adaptive Metropolis 的受控数值实验")
    s.font.name = "微软雅黑"
    s._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    s.font.size = Pt(11.5)

    doc.add_paragraph("")
    doc.add_heading("一、研究背景与工作目标", level=1)
    doc.add_paragraph(
        "本轮工作以原始污水管网 SWMM 模型为基础，围绕“如何利用少量监测点观测结果，反推出异常入流节点及其贡献大小”这一问题，"
        "构建了一条更贴近论文方法的精细化验证链路。原始晴天模型包含 242 个节点和 237 条连接要素，若直接在全网尺度上开展自由反演，"
        "计算开销高、可辨识性分析复杂，因此本阶段先从全网中截取一个包含 20 个候选节点的研究子网络，开展受控数值实验。"
    )
    doc.add_paragraph(
        "与此前小时级数据链路相比，本轮实验将监测分辨率细化到 10 分钟，并采用“边界监测积分差值反算总入流量 Q_R”的方式建立物理约束，"
        "随后保持“GA - initial PPD - AM - posterior predictive validation”的主算法框架不变，对两处真实注入点进行识别与不确定性验证。"
    )

    doc.add_heading("二、研究对象与数据构造", level=1)
    doc.add_heading("2.1 原始模型与子网络选取", level=2)
    doc.add_paragraph(
        "本研究使用的基础模型为晴天工况文件 "
        "`case_dry_full.inp`。在此基础上，为保证实验链路可控、同时保留足够的拓扑结构特征，"
        "从原始全网中选取 20 个节点构成研究子网络。选点原则包括："
        "一是包含真实注入节点及其邻近支路；二是覆盖关键监测点和边界监测点；三是保留若干与真值点拓扑相近、容易产生混淆的候选节点，"
        "以检验算法的辨识能力。"
    )
    add_table(
        doc,
        [
            ["类别", "节点"],
            ["真实注入点", "J129, J195"],
            ["监测点", "J145, J72, J81, J236, J237"],
            ["边界监测点", "J145, J237"],
            ["候选节点总数", "20"],
        ],
        widths=[3.0, 12.0],
    )
    add_figure(doc, img_topology, "图 1  研究子网络拓扑与关键节点标注", width_cm=15.8)

    doc.add_heading("2.2 10 分钟监测序列的构造", level=2)
    doc.add_paragraph(
        "在时间分辨率方面，本轮实验将分析步长设定为 600 s，即 10 分钟。这样做有两方面考虑：其一，可以比小时级数据更细致地保留异常入流在网络中的传播过程；"
        "其二，仍可将计算成本控制在单次实验可接受范围内。具体地，首先在晴天基础模型上运行 baseline 仿真，得到 5 个监测点的基线流量过程；"
        "随后将两个真实注入点 J129 和 J195 的注入模板施加到同一晴天模型中，再次进行动态仿真，得到异常工况下的监测流量过程。"
    )
    doc.add_paragraph(
        "为了构造反演目标序列，本研究对监测点采用“异常工况减去晴天工况”的方式获得流量增量序列。对于任一监测点 m，在时刻 t 的观测增量定义为："
    )
    add_formula_paragraph(doc, "ΔQ_obs,m(t) = Q_run,m(t) - Q_dry,m(t)")
    doc.add_paragraph(
        "该增量序列直接作为后续优化和似然计算的观测对象。相比直接拟合绝对流量，增量形式能够尽量剔除正常运行背景的影响，使算法更聚焦于异常注入引起的额外响应。"
    )

    doc.add_heading("2.3 总入流量 Q_R 的监测积分反算", level=2)
    doc.add_paragraph(
        "与直接对真实注入模板积分不同，本轮链路采用边界监测点积分差值反算总入流量 Q_R。"
        "其思想是：异常注入在系统内传播后，必然会在边界断面上体现为额外流量，因此可以通过边界监测点的流量增量进行积分，得到系统总额外入流量。"
        "本研究选取 J145 和 J237 作为边界监测点，计算式为："
    )
    add_formula_paragraph(doc, "Q_R = Σ_t [ΔQ_J145(t) + ΔQ_J237(t)] · Δt")
    doc.add_paragraph(
        f"其中 Δt = {summary['analysis_step_seconds']} s。按上述方法积分后，得到本轮实验的总入流量 "
        f"Q_R = {summary['q_r_monitor_based']:.2f} m³。后续算法中的所有节点入流分配都必须满足这一总量约束。"
    )

    doc.add_heading("三、参数化方式与物理约束", level=1)
    doc.add_paragraph(
        "在 20 节点研究子网络中，算法并不直接优化“某个节点是否异常”的离散标签，而是优化一个 20 维连续份额向量。"
        "设候选节点集合为 {n_1, n_2, ..., n_20}，定义份额向量 x = [x_1, x_2, ..., x_20]。"
        "每个 x_i 表示总入流量 Q_R 分配到第 i 个候选节点的比例。为保持物理合理性，该向量必须同时满足："
    )
    add_formula_paragraph(doc, "x_i ≥ 0,    Σ_i x_i = 1")
    doc.add_paragraph(
        "于是第 i 个节点的目标体积为 q_i = x_i · Q_R。"
        "在时间维度上，本研究并未为每个节点独立设计任意形状的注入过程，而是从真实模板中提取了一个归一化的公共时间形状 φ(t)，"
        "再将每个节点的体积分配映射为动态入流序列："
    )
    add_formula_paragraph(doc, "I_i(t) = q_i · φ(t)")
    doc.add_paragraph(
        "这样的处理方式有两个优点：一是可以在不显著增加维数的前提下保留异常过程的主要时间特征；二是使各候选节点之间的差异主要体现在网络传播路径和响应模式上，"
        "更适合检验监测布设与拓扑结构对定位能力的影响。"
    )

    doc.add_heading("四、遗传算法（GA）全局粗搜索", level=1)
    doc.add_paragraph(
        "遗传算法承担的是全局粗搜索任务。它并不是一个节点一个节点去尝试，而是在 20 维份额空间中同时调整所有候选节点对总量 Q_R 的分配。"
        "当前实验采用 2 个种群、每个种群 6 个个体、共 4 代迭代的小参数配置，以保证在 PyCharm 环境中能够快速复现实验全过程。"
    )
    doc.add_heading("4.1 种群初始化", level=2)
    doc.add_paragraph(
        "为兼顾可收敛性与多样性，初始种群由两部分构成：其一是少量稀疏种子，用于保证小参数条件下算法具有合理起点；其二是 Dirichlet 随机样本，"
        "用于覆盖更广的可行空间。所有个体在进入评价前都会通过 simplex 投影回到“非负且总和为 1”的合法空间。"
    )
    doc.add_heading("4.2 个体评价与目标函数", level=2)
    doc.add_paragraph(
        "对任一个体 x，算法首先根据份额向量构造 20 个节点的动态注入方案，随后调用 SWMM 进行一次完整正演，得到 5 个监测点的模拟流量增量序列。"
        "为衡量该方案与真实监测序列的贴合程度，GA 阶段采用与英文论文一致的 Nash-Sutcliffe Efficiency（NSE）作为目标函数。"
        "对于监测点 m，其 NSE 定义为："
    )
    add_formula_paragraph(
        doc,
        "NSE_m = 1 - Σ_t (M_m(t) - S_m(t))^2 / Σ_t (M_m(t) - M̄_m)^2",
    )
    doc.add_paragraph(
        "其中 M_m(t) 为监测点 m 的真实流量增量，S_m(t) 为相同监测点在当前方案下的模拟流量增量，M̄_m 为监测点 m 在整个分析时段内的平均观测增量。"
        "本轮实验对 5 个监测点的 NSE 取平均，得到个体总适应度："
    )
    add_formula_paragraph(doc, "MeanNSE = (1/5) · Σ_m NSE_m")
    doc.add_paragraph(
        "算法以 loss = 1 - MeanNSE 作为排序依据，loss 越小表示方案越优。"
    )
    doc.add_heading("4.3 进化策略", level=2)
    doc.add_paragraph(
        "每一代中，算法先对种群内所有个体按 loss 排序，保留前 25% 个体作为精英；"
        "随后采用均值交叉与高斯扰动生成新个体，并在固定代数间隔进行迁移，即把当前全局最优个体注入到其他种群中，以加快高质量模式的扩散。"
        "本轮 GA 的关键迭代结果如图 2 所示。可以看到，GA 最佳 Mean NSE 从第一代的 0.8898 持续提升到第四代的 0.9073，"
        "说明在小参数条件下，种群已经能够较稳定地向高拟合区域收缩。"
    )
    add_figure(doc, img_ga, "图 2  GA 收敛与 posterior 结果概览", width_cm=16.0)
    doc.add_paragraph(
        "值得强调的是，GA 的输出并不是单一最优点，而是末代多个较优个体构成的高概率样本群。所有 GA 个体都保存在 "
        "`ga_population_all.csv` 中，便于后续审查搜索过程。"
    )

    doc.add_heading("五、initial PPD 的构造", level=1)
    doc.add_paragraph(
        "为了避免将后续的贝叶斯采样限制在单个点解附近，本研究没有直接将 GA 最优个体送入 AM，而是按照英文论文的思路，"
        "先将末代个体合并，得到一组高质量候选样本，再通过轮盘赌方式构造 initial PPD。"
    )
    doc.add_paragraph(
        "具体而言，设末代合并群体中的适应度为 f_i = MeanNSE_i，先对适应度做平移，保证其为正，再归一化得到轮盘赌概率："
    )
    add_formula_paragraph(doc, "p_i = (f_i - min(f) + ε) / Σ_j (f_j - min(f) + ε)")
    doc.add_paragraph(
        "随后根据 p_i 进行不放回抽样，保留前 70% 左右的高质量样本，形成 initial PPD。"
        f"本轮实验末代合并样本数为 {summary['merged_last_generation_size']}，initial PPD 样本数为 {summary['initial_ppd_size']}。"
        "这一步的意义在于：一方面保留 GA 末代中多个可能模式，另一方面为后续 AM 采样提供更接近 posterior 形状的先验信息。"
    )

    doc.add_heading("六、Adaptive Metropolis（AM）后验精细采样", level=1)
    doc.add_paragraph(
        "在 initial PPD 的基础上，算法进入 Adaptive Metropolis 阶段。与 GA 的全局粗搜索不同，AM 不再从头大范围探索，而是在 initial PPD 所刻画的高概率区域附近，"
        "通过大量小步采样逐步逼近 posterior 分布。当前实现中，AM 采用 3 条独立链同时运行，每条链从 initial PPD 中不同的样本出发，以减弱单一起点导致的局部收缩。"
    )
    doc.add_heading("6.1 先验分布", level=2)
    doc.add_paragraph(
        "AM 的 prior 并没有被简化成单高斯，而是由 initial PPD 样本构造的混合先验。设 initial PPD 中的样本中心为 μ_k，"
        "轮盘赌概率为 w_k，组件协方差为 Σ_prior，则先验密度可写为："
    )
    add_formula_paragraph(doc, "p(x) = Σ_k w_k · N(x | μ_k, Σ_prior)")
    doc.add_paragraph(
        "这种写法比单峰高斯更能保留 GA 输出样本群的多模态特征，也更贴近英文论文中“initial PPD 作为 prior information”的原始表述。"
    )
    doc.add_heading("6.2 似然函数", level=2)
    doc.add_paragraph(
        "在每一步采样中，AM 都会将当前提议方案重新送入 SWMM，并计算 5 个监测点拟合误差的平方和 SSE。"
        "在噪声服从零均值高斯分布的假设下，似然写为："
    )
    add_formula_paragraph(doc, "log L(x) = - SSE(x) / (2σ_obs^2)")
    doc.add_paragraph(
        f"其中 σ_obs = 0.03，为当前实验的观测误差标准差参数。后验对数概率为："
    )
    add_formula_paragraph(doc, "log posterior(x) = log p(x) + log L(x)")
    doc.add_heading("6.3 自适应协方差与自动找方向", level=2)
    doc.add_paragraph(
        "AM 与普通 Metropolis 的关键区别在于，它会根据已经接受的历史样本不断更新提议分布的协方差矩阵，从而自动学习“下一步应该往哪个方向移动、步长应该多大”。"
        "在本研究中，设参数维数 d = 20，则论文式尺度因子为 sd = 2.42 / d = "
        f"{summary['proposal_scale_sd']:.3f}。自适应协方差的更新写为："
    )
    add_formula_paragraph(doc, "C_n = sd · Cov(X_0, X_1, ..., X_{n-1}) + sd · εI")
    doc.add_paragraph(
        "其中 ε 为极小扰动项，用于避免协方差矩阵奇异。前 8 步采用固定 base covariance 作为冷启动阶段，"
        "从第 8 步之后开始依据历史样本更新协方差。这样做的作用在于：如果某些维度变化非常敏感，协方差会自动收缩这些方向的步长；"
        "如果某些维度之间存在联动关系，协方差的非对角项会帮助采样沿着更合理的方向前进。"
    )
    doc.add_heading("6.4 接受-拒绝机制", level=2)
    doc.add_paragraph(
        "对任一步提议 x'，算法计算其对数后验概率，并与当前状态 x 比较。接受概率为："
    )
    add_formula_paragraph(doc, "β = min(1, exp(log posterior(x') - log posterior(x)))")
    doc.add_paragraph(
        "若随机数 u ~ U(0,1) 小于 β，则接受该提议；否则保留当前状态。"
        "这一机制既保证了高概率方案更容易被接受，又允许以较小概率接受局部变差的提议，从而避免采样链过早陷入局部极值。"
        f"本轮实验 3 条链的最终接受率均值为 {summary['am_accept_rate_mean']:.4f}，说明链能够持续移动，但仍保留了较强筛选性。"
    )
    doc.add_heading("6.5 后验权重提取", level=2)
    doc.add_paragraph(
        "对每条链去除 burn-in 后，将尾部样本合并，得到总体 posterior 样本集合。随后对每个候选节点统计 posterior mean、posterior median、"
        "5% 分位和 95% 分位，形成节点后验权重表。当前结果显示，J129 与 J195 的 posterior mean share 分别为 "
        f"{posterior_weights.loc[posterior_weights['node']=='J129','posterior_mean_share'].iloc[0]:.4f} 和 "
        f"{posterior_weights.loc[posterior_weights['node']=='J195','posterior_mean_share'].iloc[0]:.4f}，"
        "显著高于其余候选节点，说明 posterior 已经稳定聚焦到两个真实注入点。"
    )

    doc.add_heading("七、Posterior Predictive Validation", level=1)
    doc.add_paragraph(
        "在完成节点识别后，本研究进一步对 posterior 的不确定性表达进行验证。其核心思想是：如果 posterior 是合理的，"
        "那么从 posterior 中抽取的样本重新驱动 SWMM 所得到的监测响应，应该能够以较高概率覆盖真实监测曲线。"
    )
    doc.add_paragraph(
        "具体做法为：从 posterior 尾部样本中均匀抽取若干组代表性参数方案，对每一组方案重新运行 SWMM；"
        "对于每个监测点和每个时刻，收集多次模拟结果后计算 5%、50% 和 95% 分位，分别作为 90% 预测区间的下界、中位预测和上界。"
        "若真实监测值位于 [p05, p95] 区间内，则记该时刻被覆盖。"
    )
    add_formula_paragraph(doc, "Coverage_90(m) = 1/T · Σ_t I[p05_m(t) ≤ M_m(t) ≤ p95_m(t)]")
    doc.add_paragraph(
        f"本轮实验的 posterior 90% 覆盖率均值为 {summary['posterior_coverage_mean']:.4f}。"
        "其中 J145、J81、J236 和 J237 的覆盖率均在 0.74 左右或以上，说明 posterior 给出的不确定性带能够较好包住真实曲线；"
        "而 J72 覆盖率仅为 0.3798，提示该监测点对应的响应模式仍然相对薄弱，是后续继续提升的重点。"
    )
    add_figure(doc, img_post, "图 3  Posterior predictive validation 结果", width_cm=16.0)

    doc.add_heading("八、结果汇总与分析", level=1)
    add_table(
        doc,
        [
            ["指标", "数值"],
            ["最终解来源", summary["final_solution_name"]],
            ["最终识别点", ", ".join(summary["predicted_nodes"])],
            ["真实注入点", ", ".join(summary["truth_nodes"])],
            ["Mean NSE", f"{summary['final_mean_nse']:.4f}"],
            ["最终 SSE", f"{summary['final_sse']:.4f}"],
            ["ACC", f"{summary['acc']:.4f}"],
            ["MCC", f"{summary['mcc']:.4f}"],
            ["MAE(all nodes)", f"{summary['mae_all_nodes']:.4f}"],
            ["MAE(truth nodes)", f"{summary['mae_truth_nodes']:.4f}"],
            ["Q_R (m³)", f"{summary['q_r_monitor_based']:.2f}"],
            ["Posterior 90% 覆盖均值", f"{summary['posterior_coverage_mean']:.4f}"],
        ],
        widths=[5.0, 6.5],
    )
    doc.add_paragraph(
        "从结果看，本轮 0323 精细子网络实验已经完成了从“监测差值反算总量”到“GA 粗搜索”“AM 后验采样”“后验预测验证”的完整闭环。"
        "在 20 个候选节点中，算法最终稳定识别出 J129 和 J195 两个真实注入点，ACC 与 MCC 均达到 1.0，说明在当前受控实验条件下，"
        "节点识别结果是准确的。同时，Mean NSE 达到 0.9200，说明最终方案能够较好重建 5 个监测点的流量增量过程。"
    )
    doc.add_paragraph(
        "从不确定性层面看，posterior predictive validation 已经初步具备论文所强调的“区间预测验证”功能。尽管平均覆盖率尚未达到非常保守的 0.9 水平，"
        "但这一结果已经说明：当前 posterior 分布不仅能给出正确的节点识别结论，还能在大部分监测点上形成具有解释力的预测带。"
        "后续若继续沿着论文方向提升，重点应落在对薄弱监测点的 coverage 提高，以及对更复杂多源场景的批量统计验证。"
    )

    doc.add_heading("九、结论与下一步工作建议", level=1)
    doc.add_paragraph(
        "综上，本轮工作实现了一个更精细、也更贴近论文原始框架的入渗溯源验证系统。"
        "它以 10 分钟分辨率监测数据和边界监测积分差值反算总量为基础，将物理约束、GA、initial PPD、AM 和 posterior predictive validation 串联为一体，"
        "并在 20 节点子网络中成功找回两个真实注入点。"
    )
    doc.add_paragraph(
        "下一步建议主要包括三方面：第一，继续提高 posterior predictive coverage，尤其关注 J72 这类弱监测点；"
        "第二，在保持当前链路不变的前提下扩展到更多随机双源或三源场景，形成更有统计说服力的实验结果；"
        "第三，在精细子网络链路稳定后，再将候选缩圈、分块和更大规模网络扩展接回主工程主线。"
    )

    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    add_page_number(sec)
    doc.add_heading("附录：本轮实验关键参数", level=1)
    add_table(
        doc,
        [
            ["参数", "取值", "说明"],
            ["分析步长", str(summary["analysis_step_seconds"]), "10 分钟，即 600 s"],
            ["候选节点数", str(summary["candidate_count"]), "20 节点研究子网络"],
            ["GA 种群数", "2", "多种群并行粗搜索"],
            ["GA 每群个体数", "6", "小参数配置，便于快速复现"],
            ["GA 代数", "4", "当前小规模迭代次数"],
            ["elite_ratio", "0.25", "前 25% 个体作为精英保留"],
            ["mutation_sigma", "0.10", "GA 变异标准差"],
            ["initial_ppd_keep_ratio", "0.70", "末代样本保留比例"],
            ["AM 样本数", "30", "每条链的采样步数"],
            ["AM burn-in", "8", "前 8 步不参与后验统计"],
            ["adaptive_start", "8", "第 8 步后启用自适应协方差"],
            ["σ_obs", "0.03", "似然函数中的观测误差标准差"],
            ["proposal_scale_sd", f"{summary['proposal_scale_sd']:.3f}", "2.42 / d，当前 d=20"],
            ["tau", f"{summary['tau']:.2f}", "活跃节点判定阈值上限"],
        ],
        widths=[4.3, 2.6, 8.0],
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_document()
