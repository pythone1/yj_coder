from pathlib import Path
import textwrap

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.section import WD_ORIENT


BASE = Path(r"E:\PY\research")
WORK = BASE / "0424suzhou"
OUT = BASE / "output" / "doc"
ASSET = BASE / "output" / "assets_online_inspection"
OUT.mkdir(parents=True, exist_ok=True)
ASSET.mkdir(parents=True, exist_ok=True)

REPORT = OUT / "在线智能检测板块调研报告.docx"
XLSX = WORK / "调研需求梳理_发送.xlsx"
LSTM_TEMPLATE = WORK / "lstm_predictive_maintenance_template.py"


def set_font(run, size=None, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para_font(paragraph, size=10.5, color=None):
    for run in paragraph.runs:
        set_font(run, size=size, color=color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_width(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for i, width in enumerate(widths_cm):
            cell = row.cells[i]
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 567)))
            tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None, header_fill="1F4E79"):
    for row in rows:
        p = doc.add_paragraph()
        for i, value in enumerate(row):
            if i:
                sep = p.add_run("  |  ")
                set_font(sep, size=8.5, color="808080")
            label = p.add_run(f"{headers[i]}: ")
            set_font(label, size=9, bold=True, color=header_fill)
            val = p.add_run(str(value))
            set_font(val, size=9)
        p.paragraph_format.space_after = Pt(3)
    doc.add_paragraph()
    return None


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.add_run(item)
        set_para_font(p, 10)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, size=16 if level == 1 else 13 if level == 2 else 11, bold=True, color="1F4E79")
    return p


def add_paragraph(doc, text, size=10.5, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, size=size, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r, size=size)
    else:
        r = p.add_run(text)
        set_font(r, size=size)
    return p


def soften_long_text(text):
    text = str(text)
    for token in ["\\", "/", "_", "-", "."]:
        text = text.replace(token, token + "\u200b")
    return text


def clean_text(value):
    if pd.isna(value):
        return ""
    return str(value).replace("\n", " ").strip()


def load_demands():
    df = pd.read_excel(XLSX, sheet_name="Sheet1").fillna("")
    df.columns = [clean_text(c) for c in df.columns]
    df["需求描述"] = df["需求描述"].map(clean_text)
    online = df[df["场景名称"].eq("在线智能检测")].copy()
    related = df[
        df["场景名称"].isin(["设备运行监控与维护", "数字基础设施建设"])
        & df["需求描述"].str.contains("设备|视觉|检测|知识库|MES|图片", regex=True)
    ].copy()
    return df, online, related


sources = [
    ("S1", "IPC-9716 Requirements for Automated Optical Inspection Process Control", "IPC", "AOI过程控制要求，用于PCBA/电子装联AOI验收与过程管理。", "https://shop.ipc.org/ipc-9716"),
    ("S2", "IPC AI in AOI white paper", "IPC", "讨论AI在AOI中的应用边界、数据/过程控制与质量风险。", "https://www.ipc.org/news-release/new-ipc-white-paper-focuses-use-artificial-intelligence-automated-optical-inspection"),
    ("S3", "IPC-A-610 Acceptability of Electronic Assemblies", "IPC", "电子组件可接受性标准，作为PCBA AOI缺陷判定口径的重要依据。", "https://shop.ipc.org/ipc-a-610"),
    ("S4", "OPC UA overview", "OPC Foundation", "工业互操作与设备数据建模标准，适合作为开放设备数据接入层。", "https://opcfoundation.org/about/opc-technologies/opc-ua/"),
    ("S5", "MTConnect standard", "MTConnect Institute", "面向制造设备的开放、只读、结构化设备数据标准。", "https://www.mtconnect.org/standard"),
    ("S6", "ISA-95 / IEC 62264 overview", "ISA", "企业系统与制造运营系统集成标准，支撑ERP/MES/控制层边界划分。", "https://www.isa.org/standards-and-publications/isa-standards/isa-95"),
    ("S7", "Smart Manufacturing Operations Planning and Control Program", "NIST", "智能制造数据互操作、操作管理和系统集成相关研究。", "https://www.nist.gov/programs-projects/smart-manufacturing-operations-planning-and-control-program"),
    ("S8", "PatchCore: Towards Total Recall in Industrial Anomaly Detection", "Amazon/University of Tuebingen", "基于局部特征记忆库的工业异常检测方法，在MVTec AD上达到接近99% AUROC。", "https://arxiv.org/abs/2106.08265"),
    ("S9", "Segment Anything", "Meta AI", "通用分割基础模型，可辅助工业视觉中的标注、ROI生成和缺陷候选区域切分。", "https://arxiv.org/abs/2304.02643"),
    ("S10", "SAHI: Slicing Aided Hyper Inference", "OBSS", "大图/小目标检测的切片推理方案，可用于高分辨率工业图片。", "https://github.com/obss/sahi"),
    ("S11", "Recent Advances in Deep Learning Based Industrial Visual Anomaly Detection", "Survey", "工业视觉异常检测综述，覆盖重构、嵌入、蒸馏、扩散等主流路线。", "https://arxiv.org/abs/2401.16269"),
    ("S12", "Cognex In-Sight ViDi", "Cognex", "工业深度学习视觉检测产品线，面向缺陷检测、定位、分类和OCR。", "https://www.cognex.com/products/deep-learning/in-sight-vidi"),
    ("S13", "KEYENCE AI Vision", "KEYENCE", "面向现场易用性的AI视觉系统与检测工具。", "https://www.keyence.com/products/vision/vision-sys/"),
    ("S14", "LandingAI Visual AI platform", "LandingAI", "面向少样本视觉检测和模型迭代的工业AI视觉平台。", "https://landing.ai/platform"),
    ("S15", "Siemens Senseye Predictive Maintenance", "Siemens", "工业预测性维护SaaS/平台方案。", "https://www.siemens.com/global/en/products/services/digital-enterprise-services/senseye-predictive-maintenance.html"),
    ("S16", "IBM Maximo Application Suite", "IBM", "EAM/资产管理与AI设备维护平台。", "https://www.ibm.com/products/maximo"),
    ("S17", "Applied Materials AIx", "Applied Materials", "半导体设备与工艺优化AI平台，强调数据、模型和专家知识结合。", "https://www.appliedmaterials.com/us/en/semiconductor/technologies/aix.html"),
    ("S18", "Synopsys.ai", "Synopsys", "AI驱动芯片设计、验证和制造相关分析能力。", "https://www.synopsys.com/ai.html"),
    ("S19", "Deloitte: Generative AI in semiconductor", "Deloitte", "生成式AI在半导体研发、制造、供应链和知识管理中的应用分析。", "https://www2.deloitte.com/us/en/insights/industry/technology/generative-ai-semiconductor-industry.html"),
    ("S20", "Mordor Intelligence Machine Vision Market", "Mordor Intelligence", "机器视觉市场规模、增速和供应商格局参考。", "https://www.mordorintelligence.com/industry-reports/machine-vision-systems-market"),
]


FONT_PATH = next((p for p in [
    r"C:\Windows\Fonts\msyh.ttc",
    r"C:\Windows\Fonts\simhei.ttf",
    r"C:\Windows\Fonts\simsun.ttc",
] if Path(p).exists()), None)


def font(size, bold=False):
    if FONT_PATH:
        return ImageFont.truetype(FONT_PATH, size)
    return ImageFont.load_default()


def wrap_cn(text, max_chars):
    lines = []
    for para in str(text).split("\n"):
        para = para.strip()
        while len(para) > max_chars:
            lines.append(para[:max_chars])
            para = para[max_chars:]
        if para:
            lines.append(para)
    return "\n".join(lines)


def draw_center(draw, box_xy, text, fnt, fill="#17202A"):
    x1, y1, x2, y2 = box_xy
    lines = wrap_cn(text, max(4, int((x2 - x1) / (fnt.size * 0.95)))).split("\n")
    heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        widths.append(bbox[2] - bbox[0])
        heights.append(bbox[3] - bbox[1])
    total_h = sum(heights) + max(0, len(lines) - 1) * 6
    y = y1 + (y2 - y1 - total_h) / 2
    for line, w, h in zip(lines, widths, heights):
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += h + 6


def rect(draw, xy, fill, outline="#2F4050", width=3):
    draw.rounded_rectangle(xy, radius=8, fill=fill, outline=outline, width=width)


def arrow_line(draw, start, end, fill="#2F4050", width=4):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 12
    p1 = (x2, y2)
    p2 = (x2 - ux * size + px * size * 0.55, y2 - uy * size + py * size * 0.55)
    p3 = (x2 - ux * size - px * size * 0.55, y2 - uy * size - py * size * 0.55)
    draw.polygon([p1, p2, p3], fill=fill)


def interp_color(v, vmin=1, vmax=5, c1=(232, 246, 239), c2=(31, 78, 121)):
    t = (v - vmin) / (vmax - vmin)
    return tuple(int(c1[i] * (1 - t) + c2[i] * t) for i in range(3))


def save_heatmap(path):
    labels = ["AOI错漏装", "SOP步骤检测", "预测性维护", "芯片良率分析", "全流程过程检测", "高分辨率图片优化"]
    metrics = ["业务价值", "技术成熟度", "数据准备度", "落地复杂度", "优先级"]
    data = [
        [5, 4, 3, 3, 5],
        [4, 3, 2, 4, 4],
        [5, 4, 3, 4, 5],
        [4, 3, 2, 4, 3],
        [5, 3, 2, 5, 4],
        [4, 4, 2, 3, 4],
    ]
    img = Image.new("RGB", (2100, 1180), "white")
    draw = ImageDraw.Draw(img)
    title_f = font(42)
    head_f = font(28)
    body_f = font(28)
    draw_center(draw, (0, 35, 2100, 100), "在线智能检测需求优先级热力图（1低-5高）", title_f, "#1F4E79")
    left, top = 360, 210
    cw, ch = 260, 120
    for j, m in enumerate(metrics):
        draw_center(draw, (left + j * cw, top - 85, left + (j + 1) * cw, top - 10), m, head_f, "#17202A")
    for i, label in enumerate(labels):
        draw_center(draw, (30, top + i * ch, left - 25, top + (i + 1) * ch), label, head_f, "#17202A")
        for j, v in enumerate(data[i]):
            xy = (left + j * cw, top + i * ch, left + (j + 1) * cw - 4, top + (i + 1) * ch - 4)
            fill = interp_color(v)
            draw.rectangle(xy, fill=fill, outline="#FFFFFF", width=3)
            draw_center(draw, xy, str(v), body_f, "white" if v >= 4 else "#0B2239")
    draw.text((1680, 1000), "评分为调研侧综合判断", font=font(22), fill="#666666")
    img.save(path)


def save_architecture(path):
    img = Image.new("RGB", (2200, 1240), "white")
    draw = ImageDraw.Draw(img)
    draw_center(draw, (0, 35, 2200, 100), "在线智能检测总体技术架构", font(42), "#1F4E79")
    boxes = [
        ((60, 230, 420, 390), "视觉相机/光源/工位视频", "#DDEBF7"),
        ((60, 500, 420, 660), "MES/SCADA/PLC/设备开放数据", "#DDEBF7"),
        ((60, 770, 420, 930), "维修工单/专家经验/SOP文档", "#DDEBF7"),
        ((560, 230, 900, 390), "边缘推理与质量网关", "#E2F0D9"),
        ((560, 500, 900, 660), "时序特征与批次对齐", "#E2F0D9"),
        ((560, 770, 900, 930), "知识抽取与向量化", "#E2F0D9"),
        ((1040, 230, 1400, 390), "AOI/缺陷/错漏装模型", "#FFF2CC"),
        ((1040, 500, 1400, 660), "预测性维护/RUL/异常检测", "#FFF2CC"),
        ((1040, 770, 1400, 930), "RAG维修问答/经验库", "#FFF2CC"),
        ((1540, 360, 1880, 520), "规则+模型融合判定", "#FCE4D6"),
        ((1540, 650, 1880, 810), "告警、工单、报告生成", "#FCE4D6"),
        ((1980, 500, 2160, 690), "闭环反馈与持续学习", "#EADCF8"),
    ]
    for xy, text, fc in boxes:
        rect(draw, xy, fc)
        draw_center(draw, xy, text, font(26))
    for y in [310, 580, 850]:
        arrow_line(draw, (420, y), (560, y))
        arrow_line(draw, (900, y), (1040, y))
    arrow_line(draw, (1400, 310), (1540, 440))
    arrow_line(draw, (1400, 580), (1540, 730))
    arrow_line(draw, (1400, 850), (1540, 730))
    arrow_line(draw, (1880, 440), (1980, 590))
    arrow_line(draw, (1880, 730), (1980, 590))
    draw_center(draw, (1760, 960, 2160, 1110), "样本回流、阈值校准、工程师确认、模型再训练", font(24), "#34495E")
    img.save(path)


def save_roadmap(path):
    stages = [
        ("0-2个月\n数据/工位盘点", "点位字典、设备台账、相机/光源评估、缺陷/SOP标签规范"),
        ("2-4个月\nAOI试点", "错装漏装检测、SOP步骤识别、98%目标验证、边缘部署"),
        ("4-7个月\n预测性维护", "关键设备健康评分、异常检测、维修知识库和RAG问答"),
        ("7-12个月\n芯片/复杂工艺分析", "良率归因、经验库自动沉淀、辅助分析报告"),
        ("12个月+\n全流程闭环", "过程+结果融合、工单闭环、模型持续学习、跨产线复制"),
    ]
    img = Image.new("RGB", (2200, 960), "white")
    draw = ImageDraw.Draw(img)
    draw_center(draw, (0, 35, 2200, 100), "落地路线图：从单点AOI到全流程在线智能检测", font(42), "#1F4E79")
    y = 435
    draw.line([(170, y), (2030, y)], fill="#1F4E79", width=8)
    xs = [180 + i * 465 for i in range(len(stages))]
    colors = ["#DDEBF7", "#E2F0D9", "#FFF2CC", "#FCE4D6", "#EADCF8"]
    for x, (title, desc), color in zip(xs, stages, colors):
        draw.ellipse((x - 24, y - 24, x + 24, y + 24), fill="#1F4E79")
        xy = (x - 175, 165, x + 175, 300)
        rect(draw, xy, color)
        draw_center(draw, xy, title, font(24))
        arrow_line(draw, (x, 300), (x, y - 28))
        draw_center(draw, (x - 190, 530, x + 190, 760), desc, font(22), "#17202A")
    img.save(path)


def save_vendor_matrix(path):
    vendors = ["Cognex", "KEYENCE", "LandingAI", "Siemens Senseye", "IBM Maximo", "Applied AIx"]
    dims = ["视觉能力", "现场易用", "设备运维", "知识/报告", "生态集成"]
    data = [
        [5, 4, 2, 2, 4],
        [4, 5, 2, 1, 3],
        [4, 4, 1, 3, 3],
        [1, 3, 5, 3, 4],
        [1, 3, 5, 4, 5],
        [3, 3, 3, 5, 4],
    ]
    img = Image.new("RGB", (2100, 1080), "white")
    draw = ImageDraw.Draw(img)
    draw_center(draw, (0, 35, 2100, 100), "核心供应商能力矩阵（调研侧评估）", font(42), "#1F4E79")
    left, top = 380, 210
    cw, ch = 255, 115
    for j, d in enumerate(dims):
        draw_center(draw, (left + j * cw, top - 82, left + (j + 1) * cw, top - 10), d, font(27), "#17202A")
    for i, vendor_name in enumerate(vendors):
        draw_center(draw, (35, top + i * ch, left - 25, top + (i + 1) * ch), vendor_name, font(27), "#17202A")
        for j, v in enumerate(data[i]):
            xy = (left + j * cw, top + i * ch, left + (j + 1) * cw - 4, top + (i + 1) * ch - 4)
            fill = interp_color(v, c1=(238, 248, 244), c2=(71, 145, 110))
            draw.rectangle(xy, fill=fill, outline="#FFFFFF", width=3)
            draw_center(draw, xy, str(v), font(28), "white" if v >= 4 else "#0B2239")
    img.save(path)


def save_evolution(path):
    img = Image.new("RGB", (2200, 1040), "white")
    draw = ImageDraw.Draw(img)
    draw_center(draw, (0, 35, 2200, 100), "工业视觉检测演进路径", font(42), "#1F4E79")
    stages = [
        ("结果检测", "成品拍照\nOK/NG判定\n缺陷分类"),
        ("关键步骤检测", "工位动作\nSOP顺序\n缺件/错件即时提示"),
        ("过程检测", "参数+图像+视频\n在制品状态\n异常趋势预警"),
        ("全流程检测", "批次追溯\n设备健康\n质量/工艺/运维闭环"),
    ]
    for idx, (name, desc) in enumerate(stages):
        x = 110 + idx * 515
        xy = (x, 330, x + 360, 520)
        rect(draw, xy, "#DDEBF7" if idx < 2 else "#E2F0D9")
        draw_center(draw, xy, name, font(32), "#17202A")
        draw_center(draw, (x - 20, 585, x + 380, 790), desc, font(26), "#17202A")
        if idx < len(stages) - 1:
            arrow_line(draw, (x + 360, 425), (x + 500, 425))
    draw_center(draw, (0, 875, 2200, 940), "能力跃迁：单帧分类 -> 时序识别 -> 多模态融合 -> 闭环优化", font(30), "#1F4E79")
    img.save(path)


def save_lstm_flow(path):
    img = Image.new("RGB", (2200, 960), "white")
    draw = ImageDraw.Draw(img)
    draw_center(draw, (0, 35, 2200, 100), "LSTM预测性维护建模流程", font(42), "#1F4E79")
    steps = [
        ("数据接入", "MES节拍/批次\nSCADA/PLC点位\n维修工单"),
        ("窗口化", "按设备ID排序\n滑动窗口60-240步\n预测未来故障/健康分"),
        ("LSTM模型", "Masking\n2层LSTM/Dropout\nDense输出"),
        ("评估", "AUC/Recall\n提前预警时间\n误报率"),
        ("上线闭环", "边缘/服务推理\n告警工单\n维修反馈回流"),
    ]
    xs = [100, 520, 940, 1360, 1780]
    colors = ["#DDEBF7", "#E2F0D9", "#FFF2CC", "#FCE4D6", "#EADCF8"]
    for idx, ((title, desc), x, color) in enumerate(zip(steps, xs, colors)):
        xy = (x, 250, x + 320, 455)
        rect(draw, xy, color)
        draw_center(draw, (x, 265, x + 320, 325), title, font(31), "#17202A")
        draw_center(draw, (x + 15, 330, x + 305, 445), desc, font(24), "#17202A")
        draw_center(draw, (x - 10, 560, x + 330, 740), [
            "统一设备ID、时间戳、点位口径",
            "只用历史窗口特征，避免未来数据泄漏",
            "先二分类/健康评分，再做RUL",
            "以漏报和提前量为核心验收",
            "工程师确认结果作为新标签",
        ][idx], font(22), "#34495E")
        if idx < len(steps) - 1:
            arrow_line(draw, (x + 320, 355), (x + 410, 355))
    draw_center(draw, (0, 825, 2200, 895), "约束：当前Workspace无真实设备时序数据，报告提供建模方案与代码模板，不生成虚假模型指标。", font(27), "#A64D00")
    img.save(path)


def write_lstm_template():
    code = r'''"""
LSTM predictive maintenance template.

Input CSV requirements:
- equipment_id: machine identifier
- timestamp: sortable event/sample time
- target: 0/1 label, e.g. failure within forecast horizon
- feature columns: numeric sensor/process features from MES/SCADA/PLC

Run:
    conda run -n LSTM python lstm_predictive_maintenance_template.py data.csv
"""
from __future__ import annotations

import sys
from pathlib import Path

import numpy as np
import pandas as pd
from sklearn.metrics import classification_report, roc_auc_score
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler
from tensorflow import keras
from tensorflow.keras import layers


WINDOW = 120
STEP = 5
FORECAST_TARGET = "target"
ID_COL = "equipment_id"
TIME_COL = "timestamp"


def build_windows(df: pd.DataFrame, feature_cols: list[str]) -> tuple[np.ndarray, np.ndarray]:
    xs, ys = [], []
    for _, g in df.sort_values([ID_COL, TIME_COL]).groupby(ID_COL):
        values = g[feature_cols].to_numpy(dtype="float32")
        labels = g[FORECAST_TARGET].to_numpy(dtype="int32")
        if len(g) < WINDOW:
            continue
        for end in range(WINDOW, len(g) + 1, STEP):
            xs.append(values[end - WINDOW:end])
            ys.append(labels[end - 1])
    if not xs:
        raise ValueError("No windows built. Check data length per equipment_id and WINDOW.")
    return np.stack(xs), np.array(ys)


def main(csv_path: str):
    df = pd.read_csv(csv_path)
    required = {ID_COL, TIME_COL, FORECAST_TARGET}
    missing = required - set(df.columns)
    if missing:
        raise ValueError(f"Missing columns: {sorted(missing)}")

    feature_cols = [
        c for c in df.columns
        if c not in {ID_COL, TIME_COL, FORECAST_TARGET}
        and pd.api.types.is_numeric_dtype(df[c])
    ]
    if not feature_cols:
        raise ValueError("No numeric feature columns found.")

    df = df.sort_values([ID_COL, TIME_COL]).copy()
    scaler = StandardScaler()
    df[feature_cols] = scaler.fit_transform(df[feature_cols].fillna(method="ffill").fillna(0))

    x, y = build_windows(df, feature_cols)
    x_train, x_test, y_train, y_test = train_test_split(
        x, y, test_size=0.2, random_state=42, stratify=y if len(set(y)) > 1 else None
    )

    model = keras.Sequential([
        layers.Input(shape=(WINDOW, len(feature_cols))),
        layers.Masking(mask_value=0.0),
        layers.LSTM(64, return_sequences=True),
        layers.Dropout(0.2),
        layers.LSTM(32),
        layers.Dropout(0.2),
        layers.Dense(16, activation="relu"),
        layers.Dense(1, activation="sigmoid"),
    ])
    model.compile(
        optimizer=keras.optimizers.Adam(learning_rate=1e-3),
        loss="binary_crossentropy",
        metrics=[keras.metrics.AUC(name="auc"), keras.metrics.Recall(name="recall")],
    )
    callbacks = [
        keras.callbacks.EarlyStopping(monitor="val_auc", mode="max", patience=8, restore_best_weights=True),
        keras.callbacks.ModelCheckpoint("lstm_predictive_maintenance.keras", monitor="val_auc", mode="max", save_best_only=True),
    ]
    model.fit(
        x_train, y_train,
        validation_split=0.2,
        epochs=80,
        batch_size=128,
        callbacks=callbacks,
        verbose=2,
    )

    pred = model.predict(x_test).ravel()
    print("AUC:", roc_auc_score(y_test, pred) if len(set(y_test)) > 1 else "undefined")
    print(classification_report(y_test, (pred >= 0.5).astype(int), digits=4))
    print("Saved: lstm_predictive_maintenance.keras")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        raise SystemExit("Usage: python lstm_predictive_maintenance_template.py data.csv")
    main(sys.argv[1])
'''
    LSTM_TEMPLATE.write_text(code, encoding="utf-8")


def create_report():
    write_lstm_template()
    heatmap = ASSET / "demand_heatmap.png"
    arch = ASSET / "architecture.png"
    roadmap = ASSET / "roadmap.png"
    vendor = ASSET / "vendor_matrix.png"
    evolution = ASSET / "evolution.png"
    lstm_flow = ASSET / "lstm_flow.png"
    save_heatmap(heatmap)
    save_architecture(arch)
    save_roadmap(roadmap)
    save_vendor_matrix(vendor)
    save_evolution(evolution)
    save_lstm_flow(lstm_flow)

    df, online, related = load_demands()

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("在线智能检测板块调研报告")
    set_font(r, 24, True, "1F4E79")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("需求梳理、技术路径、市场方案与实施建议")
    set_font(r, 13, False, "4F5B62")
    doc.add_paragraph()
    meta = [
        ("项目目录", soften_long_text(str(WORK))),
        ("源文件", soften_long_text(XLSX.name)),
        ("检索日期", "2026-04-24"),
        ("输出", soften_long_text(REPORT.name)),
    ]
    add_table(doc, ["项目", "内容"], meta, widths=[4.0, 21.0], header_fill="5B9BD5")
    doc.add_picture(str(arch), width=Cm(24.5))
    cap = doc.add_paragraph("图 1 在线智能检测总体技术架构")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_font(cap, 8.5, "666666")
    doc.add_page_break()

    add_heading(doc, "一、结论摘要", 1)
    add_bullets(doc, [
        "在线智能检测应拆成三条主线：AOI视觉质检与SOP步骤检测、MES/开放设备数据驱动的预测性维护、芯片制造良率分析与经验库/报告自动化。",
        "98%-99%准确率不是单靠大模型得到的指标，关键在于成像稳定、缺陷/步骤标准化、样本闭环、阈值校准、误检漏检分层管理和上线后的SPC/MSA治理。",
        "预测性维护先从关键设备无监督/半监督异常检测做MVP，再逐步叠加RUL预测、工单推荐、备件策略和RAG维修问答。",
        "芯片制造场景短期适合做良率分析助手、工艺/缺陷经验库、分析报告草稿生成；自动决策和闭环调参必须保留工程师确认。",
        "视觉检测演进路径是：结果检测 -> 关键步骤检测 -> 过程检测 -> 过程+结果全流程检测；数据底座和质量闭环比模型选型更先决。",
    ])
    add_paragraph(doc, "本报告基于源Excel、Workspace历史文档/表格、公开标准/论文/供应商资料和行业报告整理。未在Workspace中发现邮件或聊天记录文件；发现的内部资料主要是既有AI平台建设报告和综合评估表。")
    doc.add_picture(str(heatmap), width=Cm(24.5))
    cap = doc.add_paragraph("图 2 需求优先级热力图")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_font(cap, 8.5, "666666")

    add_heading(doc, "二、源Excel需求汇总", 1)
    demand_rows = []
    for _, row in online.iterrows():
        demand_rows.append([row["需求名称"], row["需求描述"]])
    add_table(doc, ["在线智能检测需求", "原始需求摘要"], demand_rows, widths=[5.0, 20.0])

    add_heading(doc, "相邻支撑需求", 2)
    support_rows = []
    for _, row in related.iterrows():
        support_rows.append([row["场景名称"], row["需求名称"], row["需求描述"]])
    add_table(doc, ["场景", "需求", "关联点"], support_rows, widths=[4.0, 4.5, 16.5], header_fill="70AD47")

    add_heading(doc, "需求场景细分", 2)
    scenario_rows = [
        ["AOI质量检测", "装配完成品错装、漏装、反装、缺件、型号混用", "相机/光源/治具、检测标准、缺陷样本、工位节拍", "边缘视觉模型+规则复核+人工复判闭环"],
        ["SOP步骤检测", "工人操作顺序、关键动作、工具/物料使用、漏步骤提示", "标准步骤视频、动作拆解、时间窗、工位布局", "视频动作识别+目标检测+时序状态机"],
        ["设备预测性维护", "减少故障停机，输出维护方案，维修知识库问答", "MES、SCADA/PLC、传感器、EAM/工单、故障标签", "异常检测/RUL+知识图谱/RAG问答"],
        ["芯片智能检测优化", "良品/良率分析、经验库、辅助分析报告", "工艺参数、量测数据、缺陷图、批次/wafer追溯", "多模态数据湖+根因分析+LLM报告助手"],
        ["生产关键步骤监测", "从结果检测扩展到过程检测与全流程检测", "跨工序ID、批次追踪、过程图像/参数", "过程状态识别+质量预测+闭环工单"],
        ["高分辨率图像优化", "图片分辨率高、体积大，需提升识别能力", "ROI标注、金字塔/切片策略、GPU资源", "切片推理、Patch/Anomaly模型、主动学习"],
    ]
    add_table(doc, ["子场景", "业务目标", "数据条件", "技术组合"], scenario_rows, widths=[4.0, 7.0, 7.0, 7.0], header_fill="5B9BD5")

    add_heading(doc, "三、AOI工业质检与SOP检测", 1)
    add_paragraph(doc, "技术标准方面，PCBA/电子装联AOI可参考IPC-9716、IPC-A-610等标准，并用企业内部缺陷样本库固化判定口径。非电子装配场景没有单一通用AOI标准，应把产品BOM、装配SOP、质量控制计划、抽检规则和复判流程转化为模型标签体系。")
    aoi_rows = [
        ["装配错漏装", "YOLO/RT-DETR/Mask R-CNN/ViT检测；模板匹配和几何校验；OCR/码识别", "强光源和治具固定；BOM-工位-相机位映射；每个关键件建立OK/NG样本；小目标采用高分辨率切片。"],
        ["外观缺陷", "CNN/ViT分类；PatchCore/PaDiM/FastFlow等异常检测；SAM辅助标注", "缺陷稀缺时先做异常检测和少样本学习；用人工复判回流扩充长尾缺陷。"],
        ["SOP步骤检测", "目标检测+人体姿态/手部动作识别+时序模型；状态机约束步骤顺序", "先由工程师拆解标准步骤，定义可观测动作、时间窗、允许偏差和提示策略。"],
        ["98%-99%准确率", "模型准确率+工位准确率+复判准确率分开考核", "独立测试集覆盖班次、人员、光照、产品批次；重点看漏检率、误报率、单工位节拍和漂移。"],
    ]
    add_table(doc, ["任务", "主流算法架构", "实现高准确率的关键"], aoi_rows, widths=[4.0, 9.0, 12.0])
    add_paragraph(doc, "落地建议：先选择装配错漏装和2-3个关键SOP动作做试点，不直接覆盖全部流程。上线验收建议采用分层指标：缺陷级召回率>=99%、误报可由复判降低；工位级OK/NG准确率>=98%；SOP关键步骤识别准确率>=98%，并记录无法判定比例。")

    add_heading(doc, "四、预测性维护与维修知识库", 1)
    add_paragraph(doc, "数据架构建议采用ISA-95划分企业层/MES层/控制层；设备开放数据优先用OPC UA、MTConnect、MQTT/Sparkplug等标准接入，维修历史来自EAM/CMMS/工单系统。先统一设备台账、点位字典、故障编码、停机时段和维修动作。")
    doc.add_picture(str(lstm_flow), width=Cm(24.5))
    cap = doc.add_paragraph("图 3 LSTM预测性维护建模流程")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_font(cap, 8.5, "666666")
    pm_rows = [
        ["数据接入", "PLC/SCADA/传感器、MES批次/节拍、EAM工单、备件、巡检记录", "边缘网关+时序库+湖仓；按设备ID和时间对齐"],
        ["特征工程", "温度、振动、电流、压力、节拍、停机、报警、维修动作", "滑窗统计、频域特征、趋势/斜率、工况分段、异常上下文"],
        ["模型路线", "无监督异常检测 -> LSTM健康评分/故障预警 -> RUL预测", "故障标签不足时先做异常检测；有维修标签后用LSTM滑动窗口预测未来故障或健康分"],
        ["知识库", "维修手册、故障现象、原因、处理步骤、备件、专家经验", "结构化故障树+向量检索RAG；回答必须带来源、适用设备和置信提示"],
        ["业务闭环", "告警、工单建议、备件建议、维修反馈", "工程师确认后回写EAM，作为标签持续训练"],
    ]
    add_table(doc, ["模块", "内容", "技术实现"], pm_rows, widths=[4.0, 9.0, 12.0], header_fill="70AD47")

    add_heading(doc, "LSTM建模细化", 2)
    lstm_rows = [
        ["输入特征", "MES节拍/批次、PLC/SCADA点位、振动/电流/温度/压力、报警码、开停机、维修动作"],
        ["样本构造", "按设备ID和时间排序，使用60-240个采样点滑动窗口；标签为未来N小时/天内故障、停机或维修事件。"],
        ["模型结构", "Masking -> LSTM(64, return_sequences) -> Dropout -> LSTM(32) -> Dense -> Sigmoid/回归输出。"],
        ["训练策略", "按设备/时间切分训练集与测试集，避免同一故障片段泄漏；类别不平衡时使用class_weight或过采样。"],
        ["验收指标", "召回率、误报率、AUC、提前预警时间、每台设备月告警数、工单命中率。"],
        ["代码模板", soften_long_text(str(LSTM_TEMPLATE))],
    ]
    add_table(doc, ["项目", "方案"], lstm_rows, widths=[4.0, 21.0], header_fill="70AD47")

    add_heading(doc, "五、芯片制造AI大模型应用", 1)
    add_paragraph(doc, "半导体场景的数据链更复杂：wafer map、缺陷图、量测、设备recipe、批次路径、SPC报警、工程师分析记录均需按批次/wafer/工艺段关联。大模型适合做知识检索、根因候选生成、分析报告草稿和经验库沉淀；良率结论仍应由统计模型、工艺规则和工程师审核共同确认。")
    chip_rows = [
        ["良率分析", "把wafer map、缺陷分类、量测数据、recipe和SPC事件关联，输出根因候选", "统计过程控制、因果/相关分析、图模型、AutoML、LLM解释层"],
        ["经验库自动构建", "从工程师报告、8D、FA记录和邮件文档中抽取问题-原因-措施", "信息抽取+知识图谱+向量索引+审核工作流"],
        ["辅助报告生成", "自动生成批次异常摘要、影响范围、相似历史案例、下一步实验建议", "RAG+模板化报告；关键数值来自可追溯数据表，不由模型编造"],
        ["成功案例方向", "设备商/EDA厂商已将AI用于工艺优化、设计/验证、制造分析平台", "可参考Applied Materials AIx、Synopsys.ai和生成式AI半导体行业分析"],
    ]
    add_table(doc, ["应用", "目标", "实现路径"], chip_rows, widths=[4.0, 10.0, 11.0], header_fill="8064A2")

    add_heading(doc, "六、从结果检测到全流程检测", 1)
    doc.add_picture(str(evolution), width=Cm(24.5))
    cap = doc.add_paragraph("图 4 工业视觉检测演进路径")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_font(cap, 8.5, "666666")
    img_rows = [
        ["切片推理", "SAHI/滑窗/金字塔；对小缺陷和大图有效", "需处理切片重叠、NMS融合和边缘缺陷"],
        ["ROI先验", "先定位产品/关键区域，再局部高分辨率识别", "要求工位相机位和治具稳定"],
        ["异常检测", "PatchCore/PaDiM/FastFlow等少缺陷样本路线", "适合缺陷样本少，但需要稳定OK样本"],
        ["多尺度Transformer", "ViT/Swin/RT-DETR等兼顾全局和局部", "训练成本高，需数据增强和蒸馏"],
        ["主动学习", "优先标注模型不确定样本和误检漏检样本", "能降低标注成本，提升长尾缺陷覆盖"],
        ["边缘压缩", "蒸馏、量化、TensorRT/OpenVINO加速", "保障节拍，避免只在离线环境准确"],
    ]
    add_table(doc, ["策略", "用途", "注意点"], img_rows, widths=[4.0, 10.0, 11.0], header_fill="5B9BD5")

    add_heading(doc, "七、市场趋势与供应商方案", 1)
    add_paragraph(doc, "市场趋势：机器视觉和AOI仍处于增长阶段，增长驱动来自新能源、电子、半导体、汽车、医疗器械等行业的自动化质检需求；AI视觉从单一深度学习工具转向数据闭环平台。不同市场报告口径差异较大，本报告只将其作为方向性参考。")
    doc.add_picture(str(vendor), width=Cm(24.5))
    cap = doc.add_paragraph("图 5 核心供应商能力矩阵")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_font(cap, 8.5, "666666")
    vendor_rows = [
        ["Cognex", "AOI、深度学习视觉、工业相机/读码生态成熟", "工业落地强、工具链成熟", "复杂过程检测和跨系统知识问答需集成"],
        ["KEYENCE", "现场易用、硬件和软件一体化", "部署快、适合单点检测", "开放性和深度定制空间相对有限"],
        ["LandingAI", "少样本视觉检测、数据中心化管理", "适合快速迭代和缺陷样本少场景", "需与产线硬件/MES做集成"],
        ["Siemens Senseye", "预测性维护平台", "适合设备健康与停机预警", "视觉AOI能力需外部组合"],
        ["IBM Maximo", "资产管理、工单、AI维护", "EAM/CMMS闭环强", "模型和现场数据接入需项目化"],
        ["Applied AIx/Synopsys.ai", "半导体工艺/EDA/制造AI", "贴近芯片行业数据和流程", "通用制造迁移需评估成本和适配"],
    ]
    add_table(doc, ["供应商", "方案定位", "优势", "局限"], vendor_rows, widths=[4.0, 7.0, 7.0, 7.0], header_fill="70AD47")

    add_heading(doc, "八、Workspace历史资料检索结论", 1)
    local_rows = [
        ["邮件/聊天记录", "未发现 .eml/.msg 或聊天记录文件", "Unknown：仅能确认当前Workspace未检索到此类文件。"],
        ["既有DOCX", "AI平台建设与生产模块报告", "强调数据接入层需覆盖PLC/SCADA/DCS、MES、ERP、WMS、QMS、LIMS，先统一点位、设备、批次、质量口径。"],
        ["既有XLSX", "十二AI场景综合评估表", "设备预测性维护建议选包装线、泵、输送、AGV、空压/锅炉等关键设备做MVP；故障样本少时先做无监督/半监督预警。"],
        ["当前Excel", "在线智能检测7条直接需求", "AOI、SOP、设备维护、芯片良率、智能质量/工艺预测和全流程检测是核心。"],
    ]
    add_table(doc, ["资料类型", "检索结果", "可复用结论"], local_rows, widths=[4.0, 7.0, 14.0], header_fill="A5A5A5")

    add_heading(doc, "九、落地实施建议", 1)
    doc.add_picture(str(roadmap), width=Cm(24.5))
    cap = doc.add_paragraph("图 6 实施路线图")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_font(cap, 8.5, "666666")
    implement_rows = [
        ["P0", "数据与标准准备", "点位字典、设备台账、缺陷/SOP标签、相机光源评估、质量验收指标", "2个月"],
        ["P1", "AOI错漏装+SOP试点", "选1条装配线和2-3个关键动作；边缘推理；人工复判闭环", "2-4个月"],
        ["P1", "设备预测性维护MVP", "选3-5类关键设备；健康评分、异常预警、维修知识库问答", "4-7个月"],
        ["P2", "芯片/复杂工艺分析助手", "良率/质量分析经验库、报告模板、相似案例检索", "7-12个月"],
        ["P2", "全流程过程检测", "跨工序批次追踪，过程+结果融合，告警工单闭环", "12个月+"],
    ]
    add_table(doc, ["优先级", "工作包", "交付物", "周期"], implement_rows, widths=[3.0, 6.0, 13.0, 3.0], header_fill="5B9BD5")
    add_heading(doc, "关键挑战与控制措施", 2)
    add_bullets(doc, [
        "成像不稳定：先解决相机、光源、治具、节拍和遮挡，再谈模型。",
        "样本和标签不足：建立OK样本库、缺陷分层、主动学习和复判回流。",
        "系统割裂：用设备ID、批次ID、工艺段ID打通MES、SCADA、EAM、QMS。",
        "大模型幻觉：所有维修/良率问答必须引用知识库来源，关键结论禁止无来源生成。",
        "指标误读：准确率之外必须管理漏检率、误检率、无法判定率、节拍、稳定性和漂移。",
    ])

    add_heading(doc, "十、详细调研报告大纲", 1)
    outline_rows = [
        ["1. 项目背景与范围", "在线智能检测边界、业务目标、数据范围、验收指标"],
        ["2. 需求场景细分", "AOI错漏装、SOP、预测性维护、芯片良率、过程检测、高分辨率图片"],
        ["3. 技术标准与质量体系", "IPC/企业质量标准、SOP标准化、数据治理、模型验证"],
        ["4. 技术实现路径", "视觉模型、时序模型、知识库/RAG、多模态数据融合、边缘部署"],
        ["5. 市场与供应商", "AOI/视觉、预测性维护、EAM、半导体AI平台供应商比较"],
        ["6. 风险与挑战", "数据、现场、算法、系统集成、安全合规、人员流程"],
        ["7. 实施路线与预算假设", "MVP范围、阶段交付、资源角色、采购/自研建议"],
        ["8. 验收与持续运营", "指标体系、复判机制、模型监控、知识库更新、闭环工单"],
    ]
    add_table(doc, ["章节", "内容要点"], outline_rows, widths=[6.0, 19.0], header_fill="1F4E79")

    add_heading(doc, "参考资料", 1)
    ref_rows = [[sid, title, org, note, soften_long_text(url)] for sid, title, org, note, url in sources]
    add_table(doc, ["编号", "资料", "机构/来源", "用途", "链接"], ref_rows, widths=[1.4, 6.0, 3.2, 8.0, 6.4], header_fill="404040")

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "在线智能检测板块调研报告 | 2026-04-24"
        set_para_font(footer, 8, "666666")

    doc.save(REPORT)
    return REPORT


if __name__ == "__main__":
    print(create_report())
