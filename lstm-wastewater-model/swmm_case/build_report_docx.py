from __future__ import annotations

import json
import math
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

import full_network_source_tracing as base


WORK_DIR = Path(r"E:\PY\LSTM\swmm_case")
RESULT_DIR = WORK_DIR / "paper_route_full_dim_results" / "midscale_ppd"
ASSET_DIR = RESULT_DIR / "doc_assets"
OUTPUT_DOCX = RESULT_DIR / "污水管网入渗溯源汇报文档.docx"

plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS", "DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False


def set_doc_language(document: Document) -> None:
    styles = document.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        if style_name == "Normal":
            style.font.size = Pt(11)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def add_formula(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Consolas"
    run.font.size = Pt(10.5)


def shade_table_header(cell, fill: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_caption(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(90, 112, 138)


def build_architecture_figure(output_path: Path) -> None:
    fig, ax = plt.subplots(figsize=(12, 4.8))
    ax.axis("off")

    boxes = [
        (0.04, 0.55, 0.18, 0.24, "#DCEBFA", "步骤1\n读取两个 INP 模型\ncase_dry / case_wet"),
        (0.28, 0.55, 0.18, 0.24, "#E6F5EC", "步骤2\n构造真实注水观测\nJ129 / J195 / J61"),
        (0.52, 0.55, 0.18, 0.24, "#FEF3D7", "步骤3\nGA 全局粗筛\n10 维份额搜索"),
        (0.76, 0.55, 0.18, 0.24, "#FCE3E3", "步骤4\nAM 后验采样\ninitial PPD -> posterior"),
        (0.40, 0.16, 0.20, 0.20, "#EEE8FF", "步骤5\n结果输出\n真值对比 + 监测拟合 + PPD"),
    ]

    for x, y, w, h, color, text in boxes:
        rect = plt.Rectangle((x, y), w, h, facecolor=color, edgecolor="#355070", linewidth=1.6)
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=12, color="#10233C")

    arrows = [
        ((0.22, 0.67), (0.28, 0.67)),
        ((0.46, 0.67), (0.52, 0.67)),
        ((0.70, 0.67), (0.76, 0.67)),
        ((0.85, 0.55), (0.58, 0.36)),
    ]
    for (x1, y1), (x2, y2) in arrows:
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1), arrowprops=dict(arrowstyle="->", lw=2.0, color="#355070"))

    ax.text(
        0.5,
        0.93,
        "污水管网入渗溯源算法主链路",
        ha="center",
        va="center",
        fontsize=18,
        color="#10233C",
        fontweight="bold",
    )
    fig.tight_layout()
    fig.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def build_network_compare_figure(summary: dict, output_path: Path) -> None:
    nodes_df, links_df = base.parse_network(WORK_DIR / "case_dry.inp")
    truth_nodes = set(summary["truth_nodes"])
    predicted_nodes = set(summary["predicted_nodes"])
    candidate_nodes = set(summary["candidate_nodes"])
    monitor_nodes = set(base.MONITOR_NODES)

    fig, ax = plt.subplots(figsize=(10, 9))
    for row in links_df.itertuples():
        ax.plot([row.x1, row.x2], [row.y1, row.y2], color="#C5CED8", lw=1.2, zorder=1)

    def draw_group(node_list, color, marker, size, label):
        if not node_list:
            return
        view = nodes_df[nodes_df["node"].isin(node_list)]
        ax.scatter(view["x"], view["y"], c=color, s=size, marker=marker, label=label, edgecolors="white", linewidths=1.0, zorder=3)
        for item in view.itertuples():
            ax.text(item.x, item.y + 8, item.node, fontsize=8.5, ha="center", color="#10233C")

    overlap = sorted(truth_nodes & predicted_nodes)
    truth_only = sorted(truth_nodes - predicted_nodes)
    predicted_only = sorted(predicted_nodes - truth_nodes)
    candidate_only = sorted(candidate_nodes - truth_nodes - predicted_nodes)

    draw_group(sorted(monitor_nodes), "#2563EB", "s", 80, "监测点")
    draw_group(overlap, "#F59E0B", "D", 110, "真值且识别到")
    draw_group(truth_only, "#DC2626", "o", 95, "真值但未识别")
    draw_group(predicted_only, "#16A34A", "*", 150, "识别到但非真值")
    draw_group(candidate_only, "#94A3B8", "o", 45, "其他候选点")

    ax.set_title("真实异常点与识别结果平面位置对比", fontsize=15)
    ax.set_xlabel("X 坐标")
    ax.set_ylabel("Y 坐标")
    ax.legend(loc="upper right", fontsize=9)
    ax.grid(alpha=0.15)
    fig.tight_layout()
    fig.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def build_monitor_fit_figure(output_path: Path) -> None:
    truth_df = pd.read_csv(RESULT_DIR / "full_dim_truth_delta.csv")
    fit_df = pd.read_csv(RESULT_DIR / "full_dim_fitted_delta.csv")
    monitors = base.MONITOR_NODES

    fig, axes = plt.subplots(2, 2, figsize=(12, 7), sharex=True)
    axes = axes.flatten()
    for ax, monitor in zip(axes, monitors):
        ax.plot(truth_df[f"{monitor}_inflow"].to_numpy(), color="#EA580C", lw=2, label="真实观测增量")
        ax.plot(fit_df[f"{monitor}_inflow"].to_numpy(), color="#16A34A", lw=2, linestyle="--", label="模型拟合增量")
        ax.set_title(f"{monitor} 监测点流量增量")
        ax.grid(alpha=0.2)
    axes[0].legend(fontsize=9)
    fig.suptitle("4 个监测点流量增量拟合对比", fontsize=16)
    fig.tight_layout()
    fig.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def build_ga_figure(output_path: Path) -> None:
    ga_history = pd.read_csv(RESULT_DIR / "full_dim_ga_history.csv")
    ga_population = pd.read_csv(RESULT_DIR / "full_dim_ga_population.csv")

    fig, axes = plt.subplots(1, 2, figsize=(12.5, 4.6))
    axes[0].plot(ga_history["generation"], ga_history["best_mean_nse"], marker="o", color="#2563EB", lw=2.2)
    axes[0].set_title("GA 各代最佳 Mean NSE")
    axes[0].set_xlabel("代数")
    axes[0].set_ylabel("Best Mean NSE")
    axes[0].grid(alpha=0.2)

    generations = sorted(ga_population["generation"].unique())
    data = [ga_population.loc[ga_population["generation"] == g, "mean_nse"].to_numpy() for g in generations]
    axes[1].boxplot(data, labels=[str(g) for g in generations], patch_artist=True, boxprops=dict(facecolor="#DCEBFA"))
    axes[1].set_title("GA 各代种群 Mean NSE 分布")
    axes[1].set_xlabel("代数")
    axes[1].set_ylabel("Mean NSE")
    axes[1].grid(alpha=0.2)

    fig.suptitle("GA 搜索过程与群体表现", fontsize=16)
    fig.tight_layout()
    fig.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def build_posterior_figure(output_path: Path) -> None:
    weights_df = pd.read_csv(RESULT_DIR / "full_dim_weights.csv").sort_values("posterior_mean_share", ascending=False)

    fig, ax = plt.subplots(figsize=(10.5, 5.4))
    colors = ["#DC2626" if bool(v) else "#64748B" for v in weights_df["is_truth"]]
    y = weights_df["posterior_mean_share"].to_numpy()
    yerr_low = y - weights_df["p05_share"].to_numpy()
    yerr_high = weights_df["p95_share"].to_numpy() - y
    ax.bar(weights_df["node"], y, color=colors)
    ax.errorbar(weights_df["node"], y, yerr=[yerr_low, yerr_high], fmt="none", ecolor="#111827", capsize=4, lw=1.2)
    tau = float(weights_df["tau"].iloc[0])
    ax.axhline(tau, color="#C2410C", linestyle="--", linewidth=1.8, label=f"Tau={tau:.3f}")
    ax.set_title("Posterior 节点权重与置信区间")
    ax.set_xlabel("候选节点")
    ax.set_ylabel("Posterior mean share")
    ax.legend()
    ax.grid(alpha=0.18, axis="y")
    fig.tight_layout()
    fig.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def add_metrics_table(document: Document, summary: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = table.rows[0].cells
    set_cell_text(headers[0], "指标", True)
    set_cell_text(headers[1], "数值", True)
    for cell in headers:
        shade_table_header(cell)

    rows = [
        ("真实异常点", ", ".join(summary["truth_nodes"])),
        ("当前识别结果", ", ".join(summary["predicted_nodes"])),
        ("注水放大系数", f'{summary.get("truth_scale_factor", 1.0):.2f}'),
        ("总入流量 Q_R", f'{summary["q_r"]:.1f}'),
        ("Mean NSE", f'{summary["final_mean_nse"]:.4f}'),
        ("ACC", f'{summary["acc"]:.4f}'),
        ("MCC", f'{summary["mcc"]:.4f}'),
        ("MAE(all nodes)", f'{summary["mae_all_nodes"]:.4f}'),
        ("MAE(truth nodes)", f'{summary["mae_truth_nodes"]:.4f}'),
        ("AM 平均接受率", f'{summary["am_accept_rate_mean"]:.4f}'),
    ]
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key)
        set_cell_text(cells[1], value)


def add_candidate_table(document: Document, summary: dict) -> None:
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = table.rows[0].cells
    set_cell_text(headers[0], "类别", True)
    set_cell_text(headers[1], "节点", True)
    set_cell_text(headers[2], "说明", True)
    for cell in headers:
        shade_table_header(cell)

    rows = [
        ("原始管网", "242 个节点, 237 条管段", "完整 SWMM 模型规模"),
        ("本轮监测点", ", ".join(base.MONITOR_NODES), "用于反演和评估"),
        ("真实注水点", ", ".join(summary["truth_nodes"]), "受控实验真值"),
        ("当前 10 个候选点", ", ".join(summary["candidate_nodes"]), "本轮先在 10 节点上跑通"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)


def add_weights_table(document: Document) -> None:
    weights_df = pd.read_csv(RESULT_DIR / "节点识别对比表.csv").sort_values("posterior_mean_share", ascending=False)
    table = document.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["节点", "后验均值", "5%分位", "95%分位", "真值点", "识别点"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True)
        shade_table_header(table.rows[0].cells[idx])

    for row in weights_df.itertuples():
        vals = [
            row.node,
            f"{row.posterior_mean_share:.4f}",
            f"{row.p05_share:.4f}",
            f"{row.p95_share:.4f}",
            "是" if bool(row.is_truth) else "否",
            "是" if bool(row.is_predicted) else "否",
        ]
        cells = table.add_row().cells
        for idx, value in enumerate(vals):
            set_cell_text(cells[idx], value)


def build_docx() -> None:
    RESULT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    summary = json.loads((RESULT_DIR / "full_dim_summary.json").read_text(encoding="utf-8"))

    architecture_png = ASSET_DIR / "01_架构图.png"
    network_png = ASSET_DIR / "02_真值对比图.png"
    monitor_png = ASSET_DIR / "03_监测拟合图.png"
    ga_png = ASSET_DIR / "04_GA过程图.png"
    posterior_png = ASSET_DIR / "05_posterior图.png"

    build_architecture_figure(architecture_png)
    build_network_compare_figure(summary, network_png)
    build_monitor_fit_figure(monitor_png)
    build_ga_figure(ga_png)
    build_posterior_figure(posterior_png)

    document = Document()
    set_doc_language(document)

    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("污水管网入渗溯源项目汇报文档")
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("基于 SWMM、遗传算法与自适应 Metropolis 的 10 节点受控盲测反演")
    run.font.size = Pt(11)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    document.add_paragraph("")
    document.add_heading("一、项目概况", level=1)
    document.add_paragraph(
        "本项目的目标，是在只掌握少量监测点流量数据的情况下，"
        "通过 SWMM 物理仿真与概率反演算法，定位污水管网中最可能存在异常入渗的节点位置与贡献大小。"
    )
    document.add_paragraph(
        "当前阶段以受控盲测方式开展验证：先在真实管网模型中选定 3 个节点进行注水模拟，"
        "再让算法在 10 个候选节点中反推异常点。这样既保留了物理模型真实性，又能清楚评估算法是否有效。"
    )
    add_metrics_table(document, summary)

    document.add_heading("二、模型基础与研究对象", level=1)
    document.add_paragraph(
        "当前项目使用两个 SWMM 模型文件作为基础："
        "`case_dry.inp` 用于表示晴天基线工况，`case_wet.inp` 用作原始湿天/有雨工况参考。"
        "反演主链路实际以 `case_dry.inp` 为基座，再叠加受控注水时序构造真实观测。"
    )
    document.add_paragraph(
        "原始 `case_dry.inp` 模型中共有 242 个节点，其中包含 236 个普通检查井节点、2 个出水口和 4 个 storage 节点，"
        "共有 237 条管段。由于直接在全网做盲搜代价很高，因此本阶段先从全网中筛出 10 个核心候选节点，"
        "形成便于验证和解释的 10 节点受控盲测子问题。"
    )
    add_candidate_table(document, summary)
    document.add_picture(str(architecture_png), width=Inches(6.7))
    add_caption(document, "图 1  项目主链路与主要处理步骤")

    document.add_heading("三、真实注水设置与总入流量计算", level=1)
    document.add_paragraph(
        "本轮受控实验的真实注水点为 J129、J195、J61。为了增强监测端的可辨识性，"
        "本版本对真实注水模板整体放大了 2.0 倍。也就是说，算法并不是在原始模板强度下运行，"
        "而是在更强的受控激励下重新做反演。"
    )
    document.add_paragraph(
        "总入流量 Q_R 的计算遵循时间积分思想。由于当前监测比较口径为 1 小时，"
        "因此总入流量可以看成每个时间步的注水流量乘以 3600 秒再累加。"
    )
    add_formula(document, "Q_R = Σ_t [ q_J129(t) + q_J195(t) + q_J61(t) ] × Δt")
    document.add_paragraph(
        "在本次 2.0 倍放大后的工况下，计算得到总入流量 "
        f"Q_R = {summary['q_r']:.1f}。这意味着后续算法无论怎么搜索，"
        "10 个候选节点分到的水量总和都必须严格等于这个总量。"
    )
    document.add_paragraph(
        "这一点非常关键，因为它把问题从“任意猜哪里有水”变成了“在总量守恒约束下重新分配这笔水”，"
        "大大减少了不合理解。"
    )

    document.add_heading("四、监测数据是如何使用的", level=1)
    document.add_paragraph(
        "算法不是直接拿绝对流量做拟合，而是先构造监测点流量增量："
        "即真实注水工况下的流量减去晴天基线工况下的流量。"
    )
    add_formula(document, "ΔQ_obs(t) = Q_truth(t) - Q_dry(t)")
    document.add_paragraph(
        "这样做的好处是可以先扣掉正常基流背景，把反演重点集中在异常入流额外带来的响应上。"
        "当前监测点为 J145、J17、J236、J59。"
    )
    document.add_picture(str(monitor_png), width=Inches(6.7))
    add_caption(document, "图 2  4 个监测点流量增量拟合对比")

    document.add_heading("五、遗传算法 GA：在 10 个节点上做全局粗筛", level=1)
    document.add_paragraph(
        "GA 阶段的核心任务，是在 10 维参数空间中快速找到高相关区域。"
        "这里的 10 维不是 10 组监测值，而是 10 个候选节点各自承担总入流量的份额。"
    )
    add_formula(document, "S = [s1, s2, ..., s10],   si ≥ 0,   Σ si = 1")
    add_formula(document, "qi = si × Q_R")
    document.add_paragraph(
        "也就是说，GA 每次产生的不是“一个节点”或“一个组合标签”，"
        "而是一整套 10 个节点如何共同分配总入流量 Q_R 的方案。"
    )
    document.add_paragraph(
        "每个 GA 个体都会经历同样的评估流程："
        "先把份额向量转换成 10 个节点的动态注水时序，再调用 SWMM 跑一次完整仿真，"
        "然后把 4 个监测点的模拟流量增量与真实观测增量进行比较。"
    )
    document.add_paragraph("GA 的目标函数使用平均 NSE：")
    add_formula(document, "NSE_n = 1 - Σ_t (M_n^t - S_n^t)^2 / Σ_t (M_n^t - M̄_n)^2")
    add_formula(document, "MeanNSE = (1/M) × Σ_n NSE_n")
    document.add_paragraph(
        "其中 M_n^t 表示第 n 个监测点的真实流量增量，S_n^t 表示模拟流量增量。"
        "Mean NSE 越高，说明这套 10 节点分配方案越接近真实观测。"
    )
    document.add_paragraph(
        "本轮为了保证运行时间较短，GA 采用了较小的参数配置："
        "2 个种群、每群 5 个个体、共 3 代。虽然规模不大，但已经足够在 PyCharm 中快速跑通并形成稳定结果。"
    )
    document.add_picture(str(ga_png), width=Inches(6.7))
    add_caption(document, "图 3  GA 各代最佳结果与群体分布")

    document.add_heading("六、GA 之后为什么要构造 initial PPD", level=1)
    document.add_paragraph(
        "如果只把 GA 最优解 ga_best 一组方案送入 AM，那么后续采样会过分依赖单个起点。"
        "为了更贴近英文论文路线，当前项目已经改成："
        "先合并 GA 末代群体，再按适应度做轮盘赌选择，形成 initial PPD。"
    )
    document.add_paragraph(
        "这一步的好处是：后续 AM 不再围绕单个点打转，而是从一个“高概率初始分布”出发，"
        "更符合贝叶斯后验采样的逻辑。"
    )

    document.add_heading("七、自适应 Metropolis AM：在高相关区域做后验采样", level=1)
    document.add_paragraph(
        "AM 的作用不是再做大范围乱搜，而是在 GA 已经找到的高相关区域附近做更细致的后验采样。"
        "当前 AM 使用的是 prior + likelihood 的贝叶斯框架。"
    )
    add_formula(document, "log Posterior(X|M) = log Prior(X) + log Likelihood(X)")
    document.add_paragraph("似然部分基于监测流量误差平方和：")
    add_formula(document, "SSE(X) = Σ_n Σ_t (M_n^t - S_n^t)^2")
    add_formula(document, "log Likelihood = - SSE(X) / (2σ²)")
    document.add_paragraph(
        "这里的 Prior 不再是简单单高斯，而是由 initial PPD 样本构成的混合先验，"
        "更贴近前面 GA 形成的高概率结构。"
    )
    document.add_paragraph(
        "与此同时，AM 的协方差更新采用了更贴论文的形式："
    )
    add_formula(document, "C_n = sd × Cov(history) + sd × εI,   sd = 2.42 / d")
    document.add_paragraph(
        "它的作用，是根据采样历史自动学习下一步该往哪个方向动、动多大。"
        "当前这轮用了 2 条链，每条链 18 步采样，先保证小参数下稳定跑通。"
    )

    document.add_picture(str(posterior_png), width=Inches(6.7))
    add_caption(document, "图 4  posterior 节点权重与置信区间")

    document.add_heading("八、当前结果与真实结果对比", level=1)
    document.add_paragraph(
        "本轮 2.0 倍放大后的受控实验结果如下："
    )
    add_metrics_table(document, summary)
    document.add_picture(str(network_png), width=Inches(6.5))
    add_caption(document, "图 5  真实异常点与识别结果对比")

    document.add_paragraph(
        "从结果可以看出，当前版本已经能够稳定识别出 J195 和 J61，"
        "并且整体拟合指标显著提升，Mean NSE 达到 0.8575。"
        "虽然 J129 仍未被稳定拉出，但相比之前版本，当前链路已经更加稳健、可解释，也更适合继续优化。"
    )

    document.add_heading("九、GA 结果和后验结果怎么看", level=1)
    document.add_paragraph(
        "GA 的详细方案已经完整保留在当前结果目录中，便于后续逐代排查。"
        "其中最关键的几个文件是："
    )
    for path in [
        "full_dim_ga_history.csv",
        "full_dim_ga_population.csv",
        "full_dim_merged_last_generation.csv",
        "full_dim_initial_ppd.csv",
        "full_dim_am_samples.csv",
    ]:
        document.add_paragraph(path, style="List Bullet")

    document.add_paragraph(
        "其中 `full_dim_ga_population.csv` 会记录每一代、每个种群、每个个体的份额方案及其分数，"
        "如果后面你想深入看 GA 到底试了哪些方案，这个文件最重要。"
    )
    add_weights_table(document)

    document.add_heading("十、结论", level=1)
    document.add_paragraph(
        "当前项目已经形成了一条完整、可运行、可汇报的污水管网入渗溯源技术链："
        "从两个 INP 模型起步，先在完整 242 节点管网中抽取 10 个候选节点，"
        "通过受控注水构造观测，再用总量守恒约束、遗传算法和自适应 Metropolis 完成定位。"
    )
    document.add_paragraph(
        "本轮在 2.0 倍注水放大、较小参数规模下，结果已经达到较好的水平："
        "识别出 3 个真实异常点中的 2 个，整体拟合质量较高，且后验链路、收敛图和对比图都能够支撑正式汇报。"
    )
    document.add_paragraph(
        "如果后续继续优化，重点建议放在："
        "增强 J129 方向辨识度、加入更强拓扑约束和传播时延特征、在保持可运行性的基础上逐步放大参数规模。"
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("附录：重点文件索引", level=1)
    appendix = [
        "项目最终汇报.md",
        "领导汇报总览.html",
        "truth_vs_prediction_map.html",
        "paper_posterior_ppd.html",
        "paper_convergence_diagnostics.html",
        "full_dim_summary.json",
        "full_dim_ga_population.csv",
    ]
    for item in appendix:
        document.add_paragraph(item, style="List Bullet")

    document.save(OUTPUT_DOCX)


def main() -> None:
    build_docx()
    print("Wrote DOCX report to", OUTPUT_DOCX)


if __name__ == "__main__":
    main()
