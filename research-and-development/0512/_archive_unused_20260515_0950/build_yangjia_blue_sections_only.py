from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "output" / "doc"
OUT.mkdir(parents=True, exist_ok=True)

RED = RGBColor(192, 0, 0)
BLACK = RGBColor(0, 0, 0)


def set_run(run, red=True, bold=False, size=10.5, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RED if red else BLACK


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def setup(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    for name in ["Normal", "List Bullet"]:
        st = doc.styles[name]
        st.font.name = "宋体"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        st.font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        st = doc.styles[name]
        st.font.name = "黑体"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.font.color.rgb = BLACK


def title(doc, text, sub):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run(r, False, True, 20, "黑体")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(sub)
    set_run(r, True, False, 11)
    doc.add_paragraph()


def h(doc, text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        set_run(r, False, True, 15 if level == 1 else 13 if level == 2 else 11.5, "黑体")


def p(doc, text, red=True):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(21)
    para.paragraph_format.line_spacing = 1.25
    r = para.add_run(text)
    set_run(r, red)


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.line_spacing = 1.15
    r = para.add_run(text)
    set_run(r, True)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, head in enumerate(headers):
        c = t.rows[0].cells[i]
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(c, "D9EAF7")
        pp = c.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pp.add_run(head)
        set_run(r, False, True, 9, "黑体")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            pp = cells[i].paragraphs[0]
            pp.paragraph_format.line_spacing = 1.05
            r = pp.add_run(val)
            set_run(r, True, False, 8.5)
    doc.add_paragraph()


def main():
    doc = Document()
    setup(doc)
    title(
        doc,
        "AI模型在供水管网DMA系统漏损检测中的应用",
        "杨佳蓝色部分专项扩写版：仅保留模型相关内容；红色为补充写入内容"
    )

    p(doc, "写作范围说明：本稿只写你负责的“模型相关”部分。原始Word的蓝色标记在程序中无法稳定识别，因此本稿按你口述的“蓝色杨佳部分=AI模型、算法、DMA数据建模、模型应用和模型落地”来收敛，不再展开政策背景、传统机理模型发展史、泛泛组织管理等非模型内容。")

    h(doc, "1.3 供水管网AI模型应用的发展（详细）", 1)
    h(doc, "1.3.1 从阈值报警到动态基线", 2)
    p(doc, "供水管网漏损检测的AI化，并不是一开始就进入深度学习或大模型阶段。早期DMA漏损管控主要依赖固定阈值，例如入口流量超过上限、夜间最小流量超过经验值、压力低于设定值即报警。这类方法简单透明，但容易出现两个问题：一是不同DMA的用户结构、地形高差、压力制度和夜间用水习惯差异很大，统一阈值难以适配；二是节假日、天气、二供补水、商业夜间用水等正常扰动，容易被误判为漏损。")
    p(doc, "AI模型的第一步价值，是把固定阈值升级为动态基线。所谓动态基线，是指模型根据每个DMA自身的历史流量、压力、日期、节假日、天气和用户结构，学习“正常情况下此时应该是什么样”。当实际曲线持续偏离正常基线，且偏离方向和持续时间符合漏损特征时，系统再触发复核或派单建议。")

    h(doc, "1.3.2 从单点异常到多源联动", 2)
    p(doc, "传统报警常常只看单点流量或单点压力，但真实漏损通常表现为多源数据的联动变化。入口流量可能抬升，末梢压力可能下降，夜间底流可能持续变高，历史维修高风险管段附近可能出现重复异常。AI模型可以把这些变量组合起来，形成更接近工程实际的判断。")
    p(doc, "在DMA场景中，建议把模型输入分为五类：运行时序数据，包括入口流量、压力点、泵阀状态；时间上下文数据，包括小时、星期、节假日、季节、天气；空间资产数据，包括管径、材质、管龄、高程、道路等级；历史事件数据，包括爆管、维修、投诉、巡检；业务状态数据，包括阀门操作、施工计划、设备维护和抄表周期。")

    h(doc, "1.3.3 从算法模型到智能体协同", 2)
    p(doc, "AI在供水漏损检测中的应用可以分为三层。底层是专业检测模型，例如LSTM/GRU、孤立森林、自编码器、随机森林、GBDT和GNN，负责预测、识别、排序和定位。中层是数据和业务平台，负责SCADA、GIS、营收、工单、巡检和维修结果的打通。上层是大模型智能体，负责把模型结果转化为管理语言，生成日报、异常解释、巡检建议和复盘材料。")
    p(doc, "需要强调的是，大模型智能体不是底层漏损检测器。它不能替代时序模型、异常检测模型和水力模型。更合理的定位是“业务协同层”：它读取专业模型输出和业务系统数据，帮助调度人员理解异常、组织派单、形成报告和沉淀案例。")

    h(doc, "1.3.4 供水漏损AI模型的技术演进路线", 2)
    table(doc, ["阶段", "主要方法", "应用特点", "局限"], [
        ["经验阈值阶段", "固定上下限、MNF经验阈值、压力低限报警", "简单、透明、易实施", "误报多，难适配不同DMA"],
        ["统计分析阶段", "移动平均、同比环比、控制图、季节分解", "能识别历史偏离和周期变化", "对复杂非线性和多源联动不足"],
        ["机器学习阶段", "孤立森林、DBSCAN、随机森林、GBDT", "适合少标签异常和管段风险排序", "依赖特征质量和标签口径"],
        ["深度学习阶段", "LSTM、GRU、CNN-LSTM、自编码器、GNN", "适合高频时序、多变量耦合和拓扑学习", "需要更多数据和验证，解释成本更高"],
        ["智能体协同阶段", "知识库、大模型、工具调用、工单联动", "适合解释、问答、报告和派单协同", "必须保留专业模型和人工确认"],
    ])

    h(doc, "2.1 智慧水务场景下的AI技术核心认知", 1)
    h(doc, "2.1.1 AI不是替代人工，而是压缩排查范围", 2)
    p(doc, "供水漏损检测具有明显的工程属性，AI模型不能直接替代现场复核。模型更现实的价值是压缩排查范围、提高异常发现效率、减少无效巡检、沉淀经验标签。对运维人员而言，模型输出不应只是一个“异常分数”，而应提供异常时段、异常证据、可能原因、候选管段、复核方式和优先级。")
    p(doc, "因此，AI系统的设计目标不是“全自动开挖”，而是“更早发现、更准排序、更少盲查、更快复盘”。这个定位对外演讲时很重要，可以避免听众把AI理解成脱离现场经验的黑箱系统。")

    h(doc, "2.1.2 DMA是AI建模的基本单元", 2)
    p(doc, "每个DMA都有自己的用水节奏、压力制度、管龄结构、用户类型和历史漏损特征。一个全市统一模型往往难以覆盖所有分区差异。实际落地时，可以采用“分区模型+全局模型”的组合：分区模型学习本DMA的动态基线，全局模型学习跨DMA的共性风险特征。")
    p(doc, "例如，居民区DMA夜间合法用水低，MNF对漏损更敏感；商业区或医院学校附近夜间用水可能较高，模型需要学习其特殊模式；老旧管网DMA需要更多关注管龄、材质和维修历史；压力波动大的DMA则需要把泵阀工况和压力控制策略纳入特征。")

    h(doc, "2.1.3 模型建设必须先于算法选择明确业务问题", 2)
    p(doc, "漏损AI不是单一问题，而是一组业务任务。若目标是提前预警，应优先做时序预测和异常检测；若目标是减少巡检盲目性，应做候选区域收敛；若目标是年度改造计划，应做管段风险排序；若目标是调压降漏，应做压力优化和需水预测。不同任务对应不同模型，不能用一个模型解决所有问题。")

    h(doc, "2.1.4 数据质量决定模型上限", 2)
    p(doc, "供水管网AI项目的核心风险往往不在算法，而在数据。常见问题包括：流量和压力时间戳不一致，设备编码在SCADA和GIS中不统一，工单只有文字描述没有管段坐标，阀门状态长期未更新，DMA边界存在隐性连通。若这些问题不解决，复杂模型只会把错误数据包装成复杂结论。")

    h(doc, "2.2 适配管网漏损检测的核心AI技术及业务价值", 1)
    h(doc, "2.2.1 时序预测模型：LSTM/GRU/CNN-LSTM", 2)
    p(doc, "LSTM通过门控结构处理长期依赖，适合学习供水流量和压力的日周期、周周期、季节性和短期扰动。GRU结构更轻、参数更少，适合算力有限或需要快速迭代的场景。CNN-LSTM可以先用一维卷积提取多传感器局部特征，再用LSTM处理时间依赖，适合多个压力点和入口流量联合建模。")
    p(doc, "在DMA漏损检测中，LSTM/GRU更适合作为“正常曲线预测器”。模型输入可以是过去24小时、48小时或7天的入口流量、关键压力点、天气、星期、节假日等特征；输出可以是未来1小时、3小时或24小时的正常预测值和预测区间。当实际值持续高于预测区间，且主要发生在夜间低用水窗口，就形成疑似漏损残差信号。")
    p(doc, "对外讲解时，不建议重点讲输入门、遗忘门、输出门等结构细节，而应讲清楚业务逻辑：模型先学习正常节奏，再看实际曲线是否持续偏离。这样非算法听众更容易理解。")

    h(doc, "2.2.2 无监督异常检测：孤立森林、DBSCAN、自编码器", 2)
    p(doc, "很多水司缺少高质量漏损标签，直接训练监督模型容易过拟合。孤立森林、DBSCAN和自编码器更适合在早期作为“异常候选发现工具”。它们不需要大量准确漏损标签，可以先学习多数正常样本的分布，再把少数异常时段筛出来交给人工复核。")
    p(doc, "孤立森林适合多维运行特征的快速异常筛查。它的直观解释是：异常样本少且不同，因此更容易被随机切分隔离。DBSCAN适合识别低密度离群点或异常簇，适用于夜间流量-压力特征空间。自编码器适合多变量联合异常识别，它先学习正常状态下多传感器之间的组合关系，当异常模式出现时，重构误差会变大。")
    p(doc, "这类模型的输出应定位为“复核清单”，不能直接等同于漏损结论。复核结果需要回填为真实漏损、合法用水、设备故障、边界异常、施工扰动或未知原因，为后续监督模型积累标签。")

    h(doc, "2.2.3 监督学习：随机森林、GBDT、HGB", 2)
    p(doc, "当资产台账和历史工单较完整时，可以引入监督学习做管段风险排序和候选区域识别。随机森林、GBDT、HGB等树模型适合处理管龄、材质、口径、压力等级、历史维修、道路等级、周边施工、投诉频次等结构化变量。")
    p(doc, "这类模型的优势是结果相对可解释。它不仅能输出风险分，还可以说明主要影响因素，例如管龄偏长、维修频次高、压力波动大、道路施工多。对管理人员而言，这比单纯展示一个黑箱概率更有价值。")

    h(doc, "2.2.4 机理模型与AI融合：水力仿真+机器学习", 2)
    p(doc, "水力模型与AI融合是漏点定位的重要方向。常见方法是利用水力模型模拟不同漏点位置、漏量和工况下的压力响应，生成仿真样本库；再用机器学习模型学习压力响应模式，在线监测时将真实压力响应与样本库匹配，输出候选管段。")
    p(doc, "融合方法的优势是能够缓解真实漏损标签不足的问题。但前提是水力模型本身可靠，如果拓扑、阀门状态、粗糙系数或需水模式不准确，仿真样本会把错误规律带给AI模型。因此上线前应使用已知漏损事件、压力试验或现场验证数据校准模型。")

    h(doc, "2.2.5 图神经网络与知识图谱", 2)
    p(doc, "供水管网天然是图结构，节点可以是水源、泵站、阀门、压力点、用户和水池，边是管段。GNN可以把拓扑关系纳入模型，使异常识别不只看单点曲线，还看上下游传播关系。漏点发生后，不同位置会形成不同压力响应模式，图模型有机会学习这种空间传播。")
    p(doc, "知识图谱则更适合业务关联，把管段、设备、DMA、工单、维修、投诉、材质、管龄等对象连接起来。它可以服务大模型智能体，让系统回答“这个DMA为什么异常”“过去相似事件如何处置”“哪些管段优先巡检”等问题。")

    h(doc, "2.2.6 算法选型表", 2)
    table(doc, ["业务任务", "推荐模型", "输入数据", "输出结果", "注意事项"], [
        ["动态基线预测", "LSTM、GRU、CNN-LSTM、移动基线", "入口流量、压力点、天气、日历", "预测值、预测区间、残差", "按DMA训练，采用时间切分验证"],
        ["少标签异常发现", "孤立森林、DBSCAN、自编码器", "MNF特征、残差、多维时序", "异常分数、异常时段", "作为复核清单，不直接确认漏损"],
        ["管段风险排序", "随机森林、GBDT、HGB", "管龄、材质、维修、压力、道路、投诉", "风险分、优先级", "需要结构化工单和资产台账"],
        ["漏点定位收敛", "水力仿真+RF/GBDT、GNN", "拓扑、多点压力、仿真样本", "候选片区、候选管段TopN", "依赖水力模型和压力点质量"],
        ["压力优化调度", "遗传算法、贝叶斯优化、强化学习", "泵阀状态、压力约束、需水预测", "调压策略、泵阀建议", "必须满足供水安全和服务压力"],
        ["业务解释协同", "知识图谱、大模型智能体", "模型输出、GIS、工单、知识库", "问答、报告、派单建议", "保留人工确认和权限控制"],
    ])

    h(doc, "3. 核心应用：模型相关场景补充", 1)
    h(doc, "3.1 事前防控：基于AI的DMA异常预警", 2)
    p(doc, "AI预警的核心是从“是否超过阈值”升级为“是否偏离动态正常模式”。系统可为每个DMA建立日内基线、周周期基线和季节基线，持续比较实际值与预测值的残差。若残差持续扩大，且出现在MNF窗口或低用水时段，同时伴随压力点响应异常，则可提高漏损风险等级。")
    p(doc, "预警输出应包含异常时间、异常持续时长、残差幅度、压力响应、相似历史事件、疑似原因和建议动作。这样调度人员才能判断是漏损、合法用水、设备故障还是边界异常。")

    h(doc, "3.2 事中处置：漏点候选区域收敛", 2)
    p(doc, "事中处置的目标不是让模型直接给出唯一漏点，而是把排查范围从整个DMA收敛到若干候选管段。入口流量异常说明区域异常，多点压力响应和水力模型可以帮助判断异常可能发生在上游、下游、末梢或某些关键分支。")
    p(doc, "建议模型输出TopN候选管段，每个候选项包含管段编号、位置、置信度、主要证据、历史维修次数、影响用户数和建议复核方式。这样更符合工程实际，也能降低模型误判带来的风险。")

    h(doc, "3.3 事后复盘：模型持续学习", 2)
    p(doc, "每一次报警和现场核查都应成为训练样本。复盘字段至少包括DMA编号、报警时间、异常特征、派单时间、现场结果、漏点位置、漏损类型、维修完成时间、估算漏量、误报原因和模型版本。没有复盘标签，AI系统就无法越用越准。")
    p(doc, "复盘还能反向优化特征和阈值。若某类误报频繁出现，例如二供补水、商业夜间用水或流量计漂移，应把这些场景转化为特征或规则，减少下次误报。若某类漏损未被发现，则应检查压力点布设、数据采样、模型窗口和阈值策略。")

    h(doc, "3.4 调度类模型应用：压力优化与需水预测", 2)
    p(doc, "压力优化是模型应用的重要拓展。漏损通常与压力相关，高压会增加背景漏损和爆管风险。AI模型可以预测不同时段的需水量，结合压力点和水力模型，给出分时压力控制、泵站启停和阀门调节建议。")
    p(doc, "但压力优化必须满足最不利点服务压力、消防保障、二供补水和用户体验约束。因此，对外表述应是“在满足供水安全前提下寻找降漏和节能空间”，而不是简单说“AI自动降压”。")

    h(doc, "3.5 DMA规划类模型应用：分区和监测点优化", 2)
    p(doc, "AI可以参与DMA分区方案和监测点布设优化。分区方案可根据拓扑、地形、用户数量、压力制度、入口数量、边界阀数量和施工成本进行多目标评价。监测点优化可通过水力仿真不同漏点情景，比较不同压力点组合对定位效果的贡献。")

    h(doc, "4. 模型建设与落地路径", 1)
    h(doc, "4.1 数据治理：模型上线前的底座", 2)
    p(doc, "模型建设前应先完成数据治理，至少形成设备表、管网表、时序表、工单表和标签表。设备表记录流量计、压力计、RTU和边缘网关；管网表记录管段、阀门、DMA归属和资产属性；时序表记录采样时间、流量压力、缺失标记和清洗版本；工单表记录报警、核查、维修和结果；标签表记录真实漏损、误报原因和模型版本。")

    h(doc, "4.2 模型训练：从基线模型到融合模型", 2)
    p(doc, "建议按四步建设模型。第一步建立MNF和移动基线，形成最小可用预警能力。第二步训练LSTM/GRU动态基线模型，识别持续残差。第三步引入孤立森林、自编码器等异常检测模型，形成多模型交叉验证。第四步结合水力模型、资产数据和工单标签，建立候选管段定位和风险排序模型。")

    h(doc, "4.3 模型验证：算法指标与业务指标并重", 2)
    p(doc, "算法指标包括MAE、RMSE、精确率、召回率、F1、AUC等；业务指标包括提前预警时间、误报工单比例、平均排查范围、TopN命中率、维修闭环时间、维修后MNF回落和复盘回填率。对外演讲中，业务指标比单纯准确率更有说服力。")

    h(doc, "4.4 模型运营：漂移监控与再训练", 2)
    p(doc, "DMA模型上线后会面临漂移。季节变化、用户结构变化、管网改造、阀门调整、传感器更换和压力制度改变都会影响数据分布。建议每月复盘误报漏报，每季度检查模型漂移，每半年结合新增工单进行再训练，重大管网变更后重新评估模型基线。")

    h(doc, "5. PPT拆页建议（仅模型相关）", 1)
    for item in [
        "AI发展页：从阈值报警、统计分析、机器学习、深度学习到智能体协同。",
        "DMA建模页：说明每个DMA是独立建模单元，模型要学习分区自己的正常曲线。",
        "LSTM页：用预测曲线和实际曲线残差解释动态基线。",
        "异常检测页：用孤立森林、DBSCAN、自编码器解释少标签异常发现。",
        "风险排序页：用随机森林/GBDT解释管段风险和改造优先级。",
        "机理融合页：水力仿真提供物理约束，AI输出候选管段TopN。",
        "GNN/知识图谱页：管网是图结构，资产、传感器、工单可形成关联知识网络。",
        "数据治理页：设备表、管网表、时序表、工单表、标签表。",
        "模型闭环页：报警、派单、现场复核、维修回填、模型再训练。",
    ]:
        bullet(doc, item)

    h(doc, "6. 可直接用于演讲的总结表述", 1)
    for item in [
        "AI模型不是替代现场人员，而是把异常发现和排查排序提前做细。",
        "DMA是AI漏损检测的基本单元，没有边界和计量，模型无法可靠判断。",
        "LSTM/GRU负责学习正常曲线，异常检测模型负责筛出可疑时段，水力模型负责帮助定位。",
        "文献中的准确率只能说明研究条件下的表现，项目效果必须用本地DMA数据验证。",
        "没有结构化工单和复盘标签，就没有持续变准的AI模型。",
    ]:
        bullet(doc, item)

    h(doc, "参考资料", 1)
    for ref in [
        "US EPA. EPANET: Application for Modeling Drinking Water Distribution Systems.",
        "Water Network Partitioning into District Metered Areas: A State-Of-The-Art Review. Water, 2020.",
        "Artificial Intelligence in Water Distribution Networks: A Systematic Review of Models, Input Variables, Databases, and Output Strategies for Leak Detection. Smart Cities, 2026.",
        "Hochreiter S., Schmidhuber J. Long Short-Term Memory. Neural Computation, 1997.",
        "《AI供水管网漏损检测研究.docx》作为算法素材库使用，未直接采用其中未经核验的效果百分比作为项目承诺。",
    ]:
        p(doc, ref)

    out = OUT / "AI模型在供水管网DMA系统漏损检测中的应用_杨佳蓝色模型部分专项扩写版.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
