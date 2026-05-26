"""
项目名称: drainage-network-source-tracking
技术领域: 04-smart-water-systems
模块说明: analyze_0520_model_data.py - 核心业务算法实现
作者: 杨佳 (资深 AI 算法与遥感工程师)
"""

from __future__ import annotations

from collections import Counter, defaultdict, deque
from dataclasses import asdict, dataclass
from html import escape
from pathlib import Path
import json
import math
import re

import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"
OUT_DIR = ROOT / "analysis_0520"
OUT_DIR.mkdir(parents=True, exist_ok=True)

INP_PATH = max(DATA_DIR.glob("*.inp"), key=lambda p: p.stat().st_mtime)
RPT_PATH = max(DATA_DIR.glob("*.rpt"), key=lambda p: p.stat().st_mtime)
DOCX_PATH = max(DATA_DIR.glob("*.docx"), key=lambda p: p.stat().st_mtime)

def section_rows(inp_path: Path) -> dict[str, list[list[str]]]:
    sections: dict[str, list[list[str]]] = defaultdict(list)
    section = ""
    for raw in inp_path.read_text(encoding="gbk", errors="ignore").splitlines():
        stripped = raw.strip()
        if stripped.startswith("[") and stripped.endswith("]"):
            section = stripped[1:-1].upper()
            continue
        if not section or not stripped or stripped.startswith(";"):
            continue
        sections[section].append(stripped.split())
    return sections


def as_float(value: str | None, default: float = 0.0) -> float:
    if value is None:
        return default
    try:
        return float(value)
    except Exception:
        return default


def parse_docx(path: Path) -> dict[str, object]:
    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        payload: list[list[str]] = []
        for row in table.rows:
            payload.append([cell.text.strip().replace("\n", " / ") for cell in row.cells])
        tables.append(payload)
    return {
        "paragraphs": paragraphs,
        "table_count": len(doc.tables),
        "image_count": len(doc.inline_shapes),
        "tables": tables,
    }


def parse_options(rows: list[list[str]]) -> dict[str, str]:
    return {row[0]: " ".join(row[1:]) for row in rows if len(row) >= 2}


def parse_nodes(sections: dict[str, list[list[str]]]) -> pd.DataFrame:
    rows: list[dict[str, object]] = []
    coords = {r[0]: (as_float(r[1]), as_float(r[2])) for r in sections.get("COORDINATES", []) if len(r) >= 3}
    for r in sections.get("JUNCTIONS", []):
        name = r[0]
        x, y = coords.get(name, (math.nan, math.nan))
        rows.append(
            {
                "node": name,
                "type": "junction",
                "invert_elev_m": as_float(r[1] if len(r) > 1 else None),
                "max_depth_m": as_float(r[2] if len(r) > 2 else None),
                "init_depth_m": as_float(r[3] if len(r) > 3 else None),
                "surcharge_depth_m": as_float(r[4] if len(r) > 4 else None),
                "ponded_area_m2": as_float(r[5] if len(r) > 5 else None),
                "x": x,
                "y": y,
            }
        )
    for r in sections.get("STORAGE", []):
        name = r[0]
        x, y = coords.get(name, (math.nan, math.nan))
        rows.append(
            {
                "node": name,
                "type": "storage",
                "invert_elev_m": as_float(r[1] if len(r) > 1 else None),
                "max_depth_m": as_float(r[2] if len(r) > 2 else None),
                "init_depth_m": as_float(r[3] if len(r) > 3 else None),
                "surcharge_depth_m": math.nan,
                "ponded_area_m2": 0.0,
                "x": x,
                "y": y,
            }
        )
    for r in sections.get("OUTFALLS", []):
        name = r[0]
        x, y = coords.get(name, (math.nan, math.nan))
        rows.append(
            {
                "node": name,
                "type": "outfall",
                "invert_elev_m": as_float(r[1] if len(r) > 1 else None),
                "max_depth_m": math.nan,
                "init_depth_m": math.nan,
                "surcharge_depth_m": math.nan,
                "ponded_area_m2": 0.0,
                "x": x,
                "y": y,
            }
        )
    return pd.DataFrame(rows)


def parse_links(sections: dict[str, list[list[str]]]) -> pd.DataFrame:
    xsections = {r[0]: r for r in sections.get("XSECTIONS", [])}
    rows: list[dict[str, object]] = []
    for r in sections.get("CONDUITS", []):
        xs = xsections.get(r[0], [])
        rows.append(
            {
                "link": r[0],
                "type": "conduit",
                "from_node": r[1] if len(r) > 1 else "",
                "to_node": r[2] if len(r) > 2 else "",
                "length_m": as_float(r[3] if len(r) > 3 else None),
                "roughness": as_float(r[4] if len(r) > 4 else None),
                "shape": xs[1] if len(xs) > 1 else "",
                "diameter_or_height_m": as_float(xs[2] if len(xs) > 2 else None),
            }
        )
    for r in sections.get("PUMPS", []):
        rows.append(
            {
                "link": r[0],
                "type": "pump",
                "from_node": r[1] if len(r) > 1 else "",
                "to_node": r[2] if len(r) > 2 else "",
                "length_m": 0.0,
                "roughness": math.nan,
                "shape": "PUMP",
                "diameter_or_height_m": math.nan,
            }
        )
    return pd.DataFrame(rows)


def parse_subcatchments(sections: dict[str, list[list[str]]]) -> pd.DataFrame:
    rows = []
    for r in sections.get("SUBCATCHMENTS", []):
        rows.append(
            {
                "subcatchment": r[0],
                "raingage": r[1] if len(r) > 1 else "",
                "outlet": r[2] if len(r) > 2 else "",
                "area_ha": as_float(r[3] if len(r) > 3 else None),
                "imperv_pct": as_float(r[4] if len(r) > 4 else None),
                "width_m": as_float(r[5] if len(r) > 5 else None),
                "slope_pct": as_float(r[6] if len(r) > 6 else None),
            }
        )
    return pd.DataFrame(rows)


def parse_timeseries(sections: dict[str, list[list[str]]]) -> pd.DataFrame:
    rows = []
    for r in sections.get("TIMESERIES", []):
        if len(r) >= 3:
            rows.append({"series": r[0], "time": r[1], "value": as_float(r[2])})
    return pd.DataFrame(rows)


def parse_inflows(sections: dict[str, list[list[str]]]) -> pd.DataFrame:
    rows = []
    for r in sections.get("INFLOWS", []):
        rows.append(
            {
                "node": r[0] if len(r) > 0 else "",
                "constituent": r[1] if len(r) > 1 else "",
                "timeseries": r[2] if len(r) > 2 else "",
                "type": r[3] if len(r) > 3 else "",
                "m_factor": as_float(r[4] if len(r) > 4 else None),
                "s_factor": as_float(r[5] if len(r) > 5 else None),
                "baseline": as_float(r[6] if len(r) > 6 else None),
                "pattern": r[7] if len(r) > 7 else "",
            }
        )
    return pd.DataFrame(rows)


def rpt_section_text(text: str, title: str, next_titles: tuple[str, ...] = ()) -> str:
    idx = text.find(title)
    if idx < 0:
        return ""
    end = len(text)
    for next_title in next_titles:
        next_idx = text.find(next_title, idx + len(title))
        if next_idx > idx:
            end = min(end, next_idx)
    return text[idx:end]


def parse_rpt(path: Path) -> dict[str, pd.DataFrame | dict[str, float | str]]:
    text = path.read_text(encoding="gbk", errors="ignore")
    continuity: dict[str, float | str] = {}
    flow_continuity_text = rpt_section_text(text, "Flow Routing Continuity", ("Quality Routing Continuity",))
    for key in [
        "Wet Weather Inflow",
        "External Outflow",
        "Flooding Loss",
        "Final Stored Volume",
        "Continuity Error (%)",
    ]:
        for line in flow_continuity_text.splitlines():
            if key in line:
                nums = re.findall(r"[-+]?\d+(?:\.\d+)?", line)
                if nums:
                    continuity[key] = as_float(nums[-1])
                break
    node_flooding = []
    section = rpt_section_text(text, "Node Flooding Summary", ("Storage Volume Summary",))
    for line in section.splitlines():
        match = re.match(r"^\s*(\S+)\s+([0-9.]+)\s+([0-9.]+)\s+(\d+)\s+(\d+:\d+)\s+([0-9.]+)\s+([0-9.]+)", line)
        if match:
            node_flooding.append(
                {
                    "node": match.group(1),
                    "hours_flooded": as_float(match.group(2)),
                    "max_rate_cms": as_float(match.group(3)),
                    "time_of_max": f"{match.group(4)} {match.group(5)}",
                    "total_flood_10^6_ltr": as_float(match.group(6)),
                    "max_ponded_depth_m": as_float(match.group(7)),
                }
            )
    outfalls = []
    section = rpt_section_text(text, "Outfall Loading Summary", ("Link Flow Summary",))
    for line in section.splitlines():
        match = re.match(r"^\s*(\S+)\s+([0-9.]+)\s+([0-9.]+)\s+([0-9.]+)\s+([0-9.]+)\s+([0-9.]+)\s+([0-9.]+)", line)
        if match and match.group(1) not in {"Outfall", "System"}:
            outfalls.append(
                {
                    "outfall": match.group(1),
                    "flow_freq_pct": as_float(match.group(2)),
                    "avg_flow_cms": as_float(match.group(3)),
                    "max_flow_cms": as_float(match.group(4)),
                    "total_volume_10^6_ltr": as_float(match.group(5)),
                    "total_COD_kg": as_float(match.group(6)),
                    "total_NH3_kg": as_float(match.group(7)),
                }
            )
    link_flow = []
    section = rpt_section_text(text, "Link Flow Summary", ("Flow Classification Summary", "Conduit Surcharge Summary"))
    for line in section.splitlines():
        match = re.match(r"^\s*(\S+)\s+(CONDUIT|PUMP)\s+([0-9.]+)\s+(\d+)\s+(\d+:\d+)\s+([0-9.]+)\s+([0-9.]+)\s+([0-9.]+)", line)
        if match:
            link_flow.append(
                {
                    "link": match.group(1),
                    "type": match.group(2).lower(),
                    "max_abs_flow_cms": as_float(match.group(3)),
                    "time_of_max": f"{match.group(4)} {match.group(5)}",
                    "max_velocity_mps": as_float(match.group(6)),
                    "max_full_flow_ratio": as_float(match.group(7)),
                    "max_full_depth_ratio": as_float(match.group(8)),
                }
            )
    node_depth = []
    section = rpt_section_text(text, "Node Depth Summary", ("Node Inflow Summary", "Node Surcharge Summary"))
    for line in section.splitlines():
        match = re.match(r"^\s*(\S+)\s+(JUNCTION|STORAGE|OUTFALL)\s+([0-9.]+)\s+([0-9.]+)\s+([0-9.]+)\s+(\d+)\s+(\d+:\d+)\s+([0-9.]+)", line)
        if match:
            node_depth.append(
                {
                    "node": match.group(1),
                    "type": match.group(2).lower(),
                    "avg_depth_m": as_float(match.group(3)),
                    "max_depth_m": as_float(match.group(4)),
                    "max_hgl_m": as_float(match.group(5)),
                    "time_of_max": f"{match.group(6)} {match.group(7)}",
                    "reported_max_depth_m": as_float(match.group(8)),
                }
            )
    storage = []
    section = rpt_section_text(text, "Storage Volume Summary", ("Outfall Loading Summary",))
    for line in section.splitlines():
        match = re.match(r"^\s*(\S+)\s+([0-9.]+)\s+([0-9.]+)\s+([0-9.]+)\s+([0-9.]+)\s+([0-9.]+)\s+([0-9.]+)\s+(\d+)\s+(\d+:\d+)\s+([0-9.]+)", line)
        if match:
            storage.append(
                {
                    "storage": match.group(1),
                    "avg_volume_1000m3": as_float(match.group(2)),
                    "avg_pct_full": as_float(match.group(3)),
                    "max_volume_1000m3": as_float(match.group(6)),
                    "max_pct_full": as_float(match.group(7)),
                    "time_of_max": f"{match.group(8)} {match.group(9)}",
                    "max_outflow_cms": as_float(match.group(10)),
                }
            )
    return {
        "continuity": continuity,
        "node_flooding": pd.DataFrame(node_flooding),
        "outfall_loading": pd.DataFrame(outfalls),
        "link_flow": pd.DataFrame(link_flow),
        "node_depth": pd.DataFrame(node_depth),
        "storage_volume": pd.DataFrame(storage),
    }


def downstream_reachability(nodes: pd.DataFrame, links: pd.DataFrame, outfalls: set[str]) -> pd.DataFrame:
    reverse_graph: dict[str, list[str]] = defaultdict(list)
    for _, row in links.iterrows():
        reverse_graph[str(row["to_node"])].append(str(row["from_node"]))
    reachable = set(outfalls)
    queue = deque(outfalls)
    while queue:
        node = queue.popleft()
        for upstream in reverse_graph.get(node, []):
            if upstream not in reachable:
                reachable.add(upstream)
                queue.append(upstream)
    result = nodes[["node", "type"]].copy()
    result["can_reach_outfall"] = result["node"].isin(reachable)
    return result


def svg_network(nodes: pd.DataFrame, links: pd.DataFrame, subcatchments: pd.DataFrame, flooding: pd.DataFrame) -> str:
    node_map = nodes.dropna(subset=["x", "y"]).set_index("node")[["x", "y", "type"]].to_dict("index")
    xs = [v["x"] for v in node_map.values()]
    ys = [v["y"] for v in node_map.values()]
    if not xs or not ys:
        return "<p>模型缺少坐标，无法绘制空间图。</p>"
    minx, maxx, miny, maxy = min(xs), max(xs), min(ys), max(ys)
    width, height = 1180, 760

    def sx(x: float) -> float:
        return 40 + (x - minx) / (maxx - minx or 1) * (width - 80)

    def sy(y: float) -> float:
        return height - 40 - (y - miny) / (maxy - miny or 1) * (height - 80)

    flood_map = {}
    if not flooding.empty:
        flood_map = flooding.set_index("node")["total_flood_10^6_ltr"].to_dict()
    area_by_node = subcatchments.groupby("outlet")["area_ha"].sum().to_dict() if not subcatchments.empty else {}

    parts = [f"<svg viewBox='0 0 {width} {height}' class='network'>"]
    for _, row in links.iterrows():
        a, b = str(row["from_node"]), str(row["to_node"])
        if a not in node_map or b not in node_map:
            continue
        color = "#7b8fa2" if row["type"] == "conduit" else "#8b5cf6"
        stroke = 1.1 if row["type"] == "conduit" else 2.4
        parts.append(
            f"<line x1='{sx(node_map[a]['x']):.1f}' y1='{sy(node_map[a]['y']):.1f}' "
            f"x2='{sx(node_map[b]['x']):.1f}' y2='{sy(node_map[b]['y']):.1f}' "
            f"stroke='{color}' stroke-width='{stroke}' opacity='.62' />"
        )
    for node, area in area_by_node.items():
        if node not in node_map or area <= 0:
            continue
        r = min(18, 3 + math.sqrt(area) * 2.2)
        parts.append(
            f"<circle cx='{sx(node_map[node]['x']):.1f}' cy='{sy(node_map[node]['y']):.1f}' r='{r:.1f}' "
            "fill='#fbbf24' opacity='.18' stroke='none' />"
        )
    for node, value in flood_map.items():
        if node not in node_map:
            continue
        r = min(55, 10 + math.sqrt(max(value, 0)) * 115)
        parts.append(
            f"<circle cx='{sx(node_map[node]['x']):.1f}' cy='{sy(node_map[node]['y']):.1f}' r='{r:.1f}' "
            "fill='#ef4444' opacity='.28' stroke='#dc2626' stroke-width='1.2' />"
        )
    color_map = {"junction": "#2f80ed", "storage": "#f59e0b", "outfall": "#111827"}
    label_nodes = set(nodes[nodes["type"].isin(["storage", "outfall"])]["node"].astype(str))
    label_nodes |= set(flood_map)
    for _, row in nodes.dropna(subset=["x", "y"]).iterrows():
        node = str(row["node"])
        color = color_map.get(str(row["type"]), "#64748b")
        r = 3.2 if row["type"] == "junction" else 6.5
        parts.append(
            f"<circle cx='{sx(row['x']):.1f}' cy='{sy(row['y']):.1f}' r='{r}' fill='{color}' "
            "stroke='white' stroke-width='1' />"
        )
        if node in label_nodes:
            parts.append(
                f"<text x='{sx(row['x']) + 7:.1f}' y='{sy(row['y']) - 7:.1f}' class='svg-label'>{escape(node)}</text>"
            )
    parts.append("</svg>")
    return "\n".join(parts)


def bar_svg(items: list[tuple[str, float]], title: str, unit: str = "") -> str:
    if not items:
        return "<p>无数据。</p>"
    width, height = 700, 260
    left, top, bottom = 125, 35, 28
    maxv = max(v for _, v in items) or 1
    row_h = (height - top - bottom) / len(items)
    parts = [f"<svg viewBox='0 0 {width} {height}' class='bar'><text x='0' y='18' class='chart-title'>{escape(title)}</text>"]
    for i, (name, value) in enumerate(items):
        y = top + i * row_h
        w = (width - left - 70) * value / maxv
        parts.append(f"<text x='0' y='{y + row_h * 0.65:.1f}' class='bar-label'>{escape(str(name))}</text>")
        parts.append(f"<rect x='{left}' y='{y + 3:.1f}' width='{w:.1f}' height='{max(6, row_h - 6):.1f}' rx='3' fill='#2563eb' opacity='.78'/>")
        parts.append(f"<text x='{left + w + 8:.1f}' y='{y + row_h * 0.65:.1f}' class='bar-value'>{value:.3g}{escape(unit)}</text>")
    parts.append("</svg>")
    return "\n".join(parts)


def html_table(df: pd.DataFrame, max_rows: int = 20) -> str:
    if df.empty:
        return "<p>无数据。</p>"
    view = df.head(max_rows).copy()
    return view.to_html(index=False, classes="table", escape=True)


def main() -> None:
    doc = parse_docx(DOCX_PATH)
    sections = section_rows(INP_PATH)
    options = parse_options(sections.get("OPTIONS", []))
    nodes = parse_nodes(sections)
    links = parse_links(sections)
    subcatchments = parse_subcatchments(sections)
    timeseries = parse_timeseries(sections)
    inflows = parse_inflows(sections)
    rpt = parse_rpt(RPT_PATH)

    reachability = downstream_reachability(nodes, links, set(nodes[nodes["type"] == "outfall"]["node"].astype(str)))
    node_depth = rpt["node_depth"] if isinstance(rpt["node_depth"], pd.DataFrame) else pd.DataFrame()
    flooding = rpt["node_flooding"] if isinstance(rpt["node_flooding"], pd.DataFrame) else pd.DataFrame()
    link_flow = rpt["link_flow"] if isinstance(rpt["link_flow"], pd.DataFrame) else pd.DataFrame()

    enriched_nodes = nodes.merge(reachability[["node", "can_reach_outfall"]], on="node", how="left")
    if not node_depth.empty:
        enriched_nodes = enriched_nodes.merge(node_depth[["node", "avg_depth_m", "max_depth_m", "reported_max_depth_m"]], on="node", how="left")
    if not flooding.empty:
        enriched_nodes = enriched_nodes.merge(flooding, on="node", how="left")

    enriched_links = links.copy()
    if not link_flow.empty:
        enriched_links = enriched_links.merge(link_flow, on=["link", "type"], how="left")

    ts_summary = (
        timeseries.groupby("series")
        .agg(point_count=("value", "count"), min_value=("value", "min"), max_value=("value", "max"), mean_value=("value", "mean"), total_value=("value", "sum"))
        .reset_index()
        if not timeseries.empty
        else pd.DataFrame()
    )
    sub_by_outlet = (
        subcatchments.groupby("outlet").agg(subcatchment_count=("subcatchment", "count"), total_area_ha=("area_ha", "sum")).reset_index().sort_values("total_area_ha", ascending=False)
        if not subcatchments.empty
        else pd.DataFrame()
    )

    files = {
        "nodes": OUT_DIR / "0520_nodes_classified.csv",
        "links": OUT_DIR / "0520_links_classified.csv",
        "subcatchments": OUT_DIR / "0520_subcatchments.csv",
        "timeseries": OUT_DIR / "0520_timeseries_summary.csv",
        "inflows": OUT_DIR / "0520_inflows.csv",
        "flooding": OUT_DIR / "0520_node_flooding.csv",
        "summary": OUT_DIR / "0520_model_data_summary.json",
        "dashboard": OUT_DIR / "0520_model_data_dashboard.html",
    }
    enriched_nodes.to_csv(files["nodes"], index=False, encoding="utf-8-sig")
    enriched_links.to_csv(files["links"], index=False, encoding="utf-8-sig")
    subcatchments.to_csv(files["subcatchments"], index=False, encoding="utf-8-sig")
    ts_summary.to_csv(files["timeseries"], index=False, encoding="utf-8-sig")
    inflows.to_csv(files["inflows"], index=False, encoding="utf-8-sig")
    flooding.to_csv(files["flooding"], index=False, encoding="utf-8-sig")

    node_type_counts = nodes["type"].value_counts().to_dict()
    diameter_counts = links[links["type"] == "conduit"]["diameter_or_height_m"].round(3).value_counts().sort_index().to_dict()
    summary = {
        "source_files": {"inp": str(INP_PATH), "rpt": str(RPT_PATH), "docx": str(DOCX_PATH)},
        "docx_key_points": doc["paragraphs"],
        "options": options,
        "section_counts": {k: len(v) for k, v in sorted(sections.items())},
        "node_type_counts": node_type_counts,
        "link_type_counts": links["type"].value_counts().to_dict(),
        "conduit_total_length_m": float(links[links["type"] == "conduit"]["length_m"].sum()),
        "diameter_counts": {str(k): int(v) for k, v in diameter_counts.items()},
        "subcatchment_count": int(len(subcatchments)),
        "subcatchment_total_area_ha": float(subcatchments["area_ha"].sum()) if not subcatchments.empty else 0.0,
        "timeseries_count": int(ts_summary["series"].nunique()) if not ts_summary.empty else 0,
        "inflow_count": int(len(inflows)),
        "flooding_node_count": int(len(flooding)),
        "continuity": rpt["continuity"],
        "outfall_loading": (rpt["outfall_loading"].to_dict("records") if isinstance(rpt["outfall_loading"], pd.DataFrame) else []),
    }
    files["summary"].write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")

    flood_items = []
    if not flooding.empty:
        flood_items = [(str(r["node"]), float(r["total_flood_10^6_ltr"]) * 1000.0) for _, r in flooding.sort_values("total_flood_10^6_ltr", ascending=False).head(10).iterrows()]
    area_items = []
    if not sub_by_outlet.empty:
        area_items = [(str(r["outlet"]), float(r["total_area_ha"])) for _, r in sub_by_outlet.head(10).iterrows()]
    diameter_items = [(str(k), float(v)) for k, v in diameter_counts.items()]

    cards = [
        ("检查井", int(node_type_counts.get("junction", 0))),
        ("调蓄/泵站节点", int(node_type_counts.get("storage", 0))),
        ("排口", int(node_type_counts.get("outfall", 0))),
        ("管段", int((links["type"] == "conduit").sum())),
        ("泵", int((links["type"] == "pump").sum())),
        ("汇水区", int(len(subcatchments))),
        ("总面积ha", f"{summary['subcatchment_total_area_ha']:.2f}"),
        ("溢流节点", int(len(flooding))),
    ]
    card_html = "\n".join(f"<div class='card'><span>{escape(str(k))}</span><b>{escape(str(v))}</b></div>" for k, v in cards)
    option_rows = pd.DataFrame([{"参数": k, "值": v} for k, v in options.items()])

    html = f"""<!doctype html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>0520 新模型数据解析与分类可视化</title>
<style>
body{{margin:0;background:#f5f7fb;color:#172033;font-family:'Microsoft YaHei','SimHei',Arial,sans-serif;}}
header{{padding:28px 42px;background:#12304a;color:white;}}
main{{padding:24px 42px 42px;}}
h1{{margin:0;font-size:28px;}} h2{{margin:0 0 14px;font-size:20px;color:#12304a;}}
.sub{{margin-top:8px;color:#d9e6f2;}}
.grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(170px,1fr));gap:14px;margin:18px 0;}}
.card{{background:white;border-radius:10px;padding:14px 16px;box-shadow:0 8px 24px #12304a12;}}
.card span{{display:block;color:#64748b;font-size:13px;}} .card b{{font-size:28px;color:#0f766e;}}
.panel{{background:white;border-radius:12px;padding:18px;margin:18px 0;box-shadow:0 8px 24px #12304a10;}}
.network{{width:100%;height:auto;background:#fbfdff;border:1px solid #d9e2ec;border-radius:10px;}}
.svg-label{{font-size:12px;font-weight:700;fill:#172033;}}
.table{{border-collapse:collapse;width:100%;font-size:13px;}} .table th,.table td{{border-bottom:1px solid #e5e7eb;padding:7px 9px;text-align:left;}}
.table th{{background:#eff6ff;color:#1e3a8a;}}
.cols{{display:grid;grid-template-columns:1fr 1fr;gap:18px;}}
.chart-title{{font-size:15px;font-weight:700;fill:#12304a;}} .bar-label,.bar-value{{font-size:12px;fill:#334155;}}
.note{{background:#fff7ed;border-left:5px solid #f97316;padding:12px 14px;border-radius:8px;margin:10px 0;line-height:1.7;}}
@media(max-width:900px){{.cols{{grid-template-columns:1fr;}} main{{padding:18px;}}}}
</style>
</head>
<body>
<header><h1>0520 新模型数据解析与分类可视化</h1><div class="sub">模型：{escape(INP_PATH.name)}；说明文档：{escape(DOCX_PATH.name)}</div></header>
<main>
<section class="grid">{card_html}</section>
<section class="panel"><h2>一、说明文档关键信息</h2>
<div class="note">{escape('；'.join(str(x) for x in doc['paragraphs'][:10]))}</div>
{html_table(pd.DataFrame(doc['tables'][0][1:], columns=doc['tables'][0][0]) if doc['tables'] else pd.DataFrame(), 10)}
</section>
<section class="panel"><h2>二、模型空间结构分类图</h2>
{svg_network(nodes, links, subcatchments, flooding)}
<p>蓝色为检查井，橙色为调蓄/泵站相关节点，黑色为排口，紫色线为泵，红色半透明圈为 RPT 中出现溢流的节点，黄色淡圈表示接入该节点的汇水区面积。</p>
</section>
<section class="cols">
<section class="panel"><h2>三、溢流节点排序</h2>{bar_svg(flood_items, '节点溢流体积 Top10', ' m3')}</section>
<section class="panel"><h2>四、汇水区接入节点面积</h2>{bar_svg(area_items, '接入节点面积 Top10', ' ha')}</section>
</section>
<section class="cols">
<section class="panel"><h2>五、管径分类</h2>{bar_svg(diameter_items, '圆管直径数量分布', ' 条')}</section>
<section class="panel"><h2>六、时间序列</h2>{html_table(ts_summary, 20)}</section>
</section>
<section class="panel"><h2>七、模型时间与计算设置</h2>{html_table(option_rows, 60)}</section>
<section class="panel"><h2>八、入流与污染物设置</h2>
<div class="note">模型使用 RAINGAGE + SUBCATCHMENTS 表达 48h 污水量过程；INFLOWS 中仅检测到污染物外部输入记录，未检测到节点水量外部入流。</div>
{html_table(inflows, 20)}
</section>
<section class="panel"><h2>九、RPT 溢流与排口结果</h2>
{html_table(flooding, 20)}
{html_table(rpt['outfall_loading'] if isinstance(rpt['outfall_loading'], pd.DataFrame) else pd.DataFrame(), 20)}
</section>
</main></body></html>"""
    files["dashboard"].write_text(html, encoding="utf-8")
    print(json.dumps({k: str(v) for k, v in files.items()}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
