from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE = Path(r"E:\PY\research")
XLSX = next((BASE / "output" / "spreadsheet").glob("*在线智能检测模块调研汇总表_汇报版.xlsx"))
OUT = BASE / "output" / "doc" / "在线智能检测板块调研报告_汇报版.docx"


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    r = p.add_run(text)
    set_font(r, 16 if level == 1 else 13, True, "1F4E79")
    return p


def add_para(doc, text, size=10.5, color=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(str(text))
    set_font(r, size, False, color)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_label_para(doc, label, text):
    p = doc.add_paragraph()
    r = p.add_run(label + "：")
    set_font(r, 10.5, True, "1F4E79")
    r = p.add_run(str(text))
    set_font(r, 10.5)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r, 10.5)
        p.paragraph_format.space_after = Pt(2)


def split_links(text):
    return [x.strip() for x in str(text).split("|") if x.strip()]


def main():
    df = pd.read_excel(XLSX, sheet_name="模块总表", header=1).fillna("")
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.6)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("在线智能检测板块调研报告（汇报版）")
    set_font(r, 22, True, "1F4E79")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("按模块说明：做什么、怎么做、用什么数据、收益和风险")
    set_font(r, 12, False, "666666")
    doc.add_paragraph()

    add_heading(doc, "一、整体结论", 1)
    add_bullets(doc, [
        "在线智能检测不是单一视觉识别项目，应分成视觉质检、操作步骤检测、设备预测性维护、维修知识库、良率分析和数据打通几类模块。",
        "先做能快速验证价值的模块：AOI错装漏装检测、设备预测性维护、数据准备和系统打通。",
        "LSTM适合预测设备状态变化，但前提是有连续设备数据和维修标签。",
        "大模型适合查资料、整理问答和生成报告草稿，不能直接替代质量、工艺或维修结论。",
        "98%-99%这类目标要放到真实工位验证，重点看漏检、误报、节拍和人工复核结果。"
    ])

    add_heading(doc, "二、模块汇总", 1)
    for _, row in df.iterrows():
        module = row.get("模块", "")
        if not module or str(module).startswith("在线智能检测"):
            continue
        add_heading(doc, str(module), 2)
        add_label_para(doc, "要解决什么", row.get("要解决什么", ""))
        add_label_para(doc, "具体怎么做", row.get("具体怎么做", ""))
        add_label_para(doc, "需要哪些数据", row.get("需要哪些数据", ""))
        add_label_para(doc, "用什么算法/技术", row.get("用什么算法/技术", ""))
        add_label_para(doc, "系统输出", row.get("系统输出", ""))
        add_label_para(doc, "能带来什么收益", row.get("能带来什么收益", ""))
        add_label_para(doc, "建议看哪些指标", row.get("建议看哪些指标", ""))
        add_label_para(doc, "主要风险/前置条件", row.get("主要风险/前置条件", ""))
        links = split_links(row.get("调研链接/案例", ""))
        if links:
            p = doc.add_paragraph()
            r = p.add_run("贴合场景的调研链接：")
            set_font(r, 10.5, True, "1F4E79")
            p.paragraph_format.space_after = Pt(2)
            for link in links:
                add_para(doc, link, size=9, color="444444", space_after=1)

    add_heading(doc, "三、建议汇报口径", 1)
    add_bullets(doc, [
        "讲AOI时，不说“模型能力很强”，改说“相机固定、缺陷标准明确、样本回流后，系统能稳定检查错装漏装”。",
        "讲SOP时，不说“动作识别平台”，改说“系统判断工人做到哪一步，发现漏步骤就提醒”。",
        "讲LSTM时，不说“时序深度学习”，改说“用设备过去一段时间的运行状态，预测接下来是否容易出问题”。",
        "讲知识库时，不说“RAG智能体”，改说“先查企业自己的维修资料，再生成带来源的回答”。",
        "讲芯片良率时，不说“大模型自动分析良率”，改说“统计模型找线索，大模型帮工程师整理证据和报告草稿”。",
    ])

    add_heading(doc, "四、需要补充的数据", 1)
    add_bullets(doc, [
        "每条产线当前人工检测人数、单件检测时间、漏检/误检记录。",
        "设备停机成本、故障记录、维修工单、备件费用和维修响应时间。",
        "关键产品的缺陷样本、合格样本、工位图片和复判结果。",
        "芯片或复杂工艺场景下的批次、工艺参数、量测、缺陷图和工程师分析记录。",
        "如果要测算ROI，还需要产量、人工成本、返工成本、报废成本和停机损失。"
    ])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("在线智能检测板块调研报告（汇报版）")
    set_font(r, 8, False, "666666")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
